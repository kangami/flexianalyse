import os
import json
import requests
from flask import Flask, request, jsonify, Response, stream_with_context
from flask.helpers import make_response
from flask_cors import CORS
from dotenv import load_dotenv
from docx import Document as DocxDocument
import PyPDF2
from io import BytesIO
from openai import OpenAI
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from typing import Dict, List, Optional, Any
import aiocache
import asyncio
import logging
import time
import re
from datetime import datetime
from auth import register_auth_routes, init_database

# Imports des modules refactorisés
from config import (
    AuthConfig, FlaskConfig, AIConfig,
    MODEL_CONFIG, DEFAULT_MODEL, MISTRAL_MODEL, OLLAMA_MODELS, OLLAMA_API_URL
)
from utils.file_utils import extract_text_from_docx, extract_text_from_pdf
from utils.translations import translations
from services.api_clients import (
    call_openai_api, call_mistral_api, call_ollama_api, call_gemini_api,
    stream_response, openai_client, get_model_config
)
from services.analysis_service import analyze_file_content, save_file_description
from services.search_service import perform_online_search, search_serpapi, rerank_documents_with_llm
from services.hybrid_retrieval import hybrid_retrieve_documents
from services.vector_store_service import (
    vector_stores, embeddings, get_vector_store, 
    create_vector_store, add_documents_to_vector_store
)

# Load environment variables
load_dotenv()

# Debug: Vérifier le chargement des variables d'environnement
print(f"GOOGLE_CLIENT_ID chargé: {os.getenv('GOOGLE_CLIENT_ID')[:20] + '...' if os.getenv('GOOGLE_CLIENT_ID') else 'NON CHARGÉ'}")
print(f"OPENAI_API_KEY chargé: {'✅' if os.getenv('OPENAI_API_KEY') else '❌'}")

# Variable globale pour stocker les vector stores par session
vector_stores = {}

# Logger setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration des modèles importée depuis config.models
# MODEL_CONFIG, DEFAULT_MODEL, MISTRAL_MODEL, OLLAMA_MODELS, OLLAMA_API_URL sont maintenant importés

# Flask setup
app = Flask(__name__)
CORS(app, resources={r"/*": {
    "origins": ["http://flexianalyse.com", "http://localhost:5173", "https://flexianalyse.com"],
    "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    "allow_headers": ["Content-Type", "Authorization", "Session-ID"]
}})

# Ajouter un handler pour les requêtes OPTIONS (CORS preflight)
@app.before_request
def handle_preflight():
    if request.method == "OPTIONS":
        response = make_response()
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add('Access-Control-Allow-Headers', "Content-Type,Authorization,Session-ID")
        response.headers.add('Access-Control-Allow-Methods', "GET,PUT,POST,DELETE,OPTIONS")
        return response

# Clients et composants initialisés dans les modules
# openai_client, embeddings sont maintenant importés depuis services.api_clients et services.vector_store_service
vector_store = None  # Gardé pour compatibilité avec le code existant

# Cache setup
cache = aiocache.SimpleMemoryCache()

# Description file path
DESCRIPTIONS_FILE = os.path.join(os.path.dirname(__file__), "file_descriptions.json")
if not os.path.exists(DESCRIPTIONS_FILE):
    with open(DESCRIPTIONS_FILE, 'w') as f:
        json.dump([], f)

# Traductions importées depuis utils.translations
# translations est maintenant importé
_OLD_translations = {
    'en': {
        # Existing keys
        'analyze': 'Analyze',
        'content_of_file': "the content of the file",
        'and_provide': "and provide a description of its purpose.",
        'content': "Content",
        'description': "Description",
        'template': "I want an answer like: The file text.js aims to .....",
        'project_structure': 'Project structure',
        'file': 'File',
        'other_file': 'Other file',
        'question': 'Question',
        'emoji': 'Add some relevant emojis when it necessary to make it pleasant to read 😊📄✨',
        
        # New keys for local query mode
        'local_analysis_mode': '🔒 LOCAL ANALYSIS MODE - Analyze ONLY the context provided below.',
        'no_external_search': 'Do NOT perform external searches or reference information not provided.',
        'main_file': 'Main file',
        'file_content': 'File content',
        'directory_context': 'Directory context',
        'no_other_files': 'No other files in context.',
        'instructions': 'Instructions',
        'base_response_only': '- Base your response ONLY on the code/content provided above',
        'missing_info_clarify': '- If information is not in the provided context, state it clearly',
        'no_speculation': '- Do not speculate on elements not present in the files',
        'focus_local_analysis': '- Focus on analyzing the local code/content',
        'error_local_analysis': 'Error during local analysis',
        'cached_result': 'Using cached result',
        'analyzing': 'Analyzing',
        'with_context': 'with context',
        'files_in_directory': 'files in directory',
        
        # Query analysis keys
        'analyze_query_prompt': 'Analyze this question and determine if it requires recent/current information that your model might not have in its training data.',
        'query_label': 'Question',
        'json_response_format': 'Respond in JSON format ONLY',
        'needs_search_field': 'needs_search',
        'reason_field': 'reason',
        'search_keywords_field': 'search_keywords',
        'cutoff_relevance_field': 'estimated_cutoff_relevance',
        'short_explanation': 'short explanation',
        'keyword_or_null': 'or null',
        'examples_need_search': 'Examples of questions requiring search',
        'current_prices_stocks': '- Current prices, recent stock prices',
        'recent_events_news': '- Recent events, news',
        'new_software_versions': '- New software/technology versions',
        'recent_statistics_data': '- Recent statistics, government data',
        'recent_people_companies': '- Information about recent people/companies',
        'examples_no_search': 'Examples of questions NOT requiring search',
        'general_concepts': '- General concepts, established theories',
        'programming_syntax': '- Programming, language syntax',
        'history_facts': '- History, historical facts',
        'math_science': '- Mathematics, fundamental sciences',
        'analysis_error': 'Analysis error',
        'query_analysis_error': 'Error during query analysis',
        'automatic_analysis_failed': 'Automatic analysis inconclusive',
        'error_analysis_fallback': 'Analysis error, searching for safety',
        'analyzing_query': 'Analyzing query',
        'search_needed': 'Search needed',
        'no_search_needed': 'No search needed',
        'json_parse_failed': 'Failed to parse JSON response',
        'fallback_analysis': 'Using fallback analysis',

        # NEW KEYS FOR ONLINE MODE
        'online_mode_title': '🤖 Answer this question using your training knowledge.',
        'recent_info_mention': 'If you think more recent information could enrich your response, mention it at the end.',
        'online_instructions_title': 'Instructions:',
        'give_best_answer': '- Give your best answer based on your knowledge',
        'be_precise_dates': '- Be precise about dates/versions you know',
        'mention_recent_useful': '- Mention if more recent info would be useful',
        'enrichment_title': '🔄 RESPONSE ENRICHMENT',
        'initial_response': 'Your initial response:',
        'new_info_found': 'New information found:',
        'enrichment_instructions': 'Instructions:',
        'combine_intelligently': '- Intelligently combine your initial response with new info',
        'update_obsolete': '- Update obsolete parts if necessary',
        'distinguish_info': '- Clearly distinguish basic info from recent info',
        'cite_sources': '- Cite sources for recent information',
        'keep_structure': '- Keep the structure and style of your initial response',
        'prioritize_recent': '- If new info contradicts your knowledge, prioritize recent sources',
        'enriched_response': 'Enriched response:',
        'info_enriched': '💡 **Information enriched**: This response combines my basic knowledge with recent data found online.',
        'search_reason': '🔍 **Search reason**: {reason}',
        'default_search_reason': 'Potentially outdated information',
        'source_training': '🧠 **Source**: Response based on my training knowledge. No recent search was deemed necessary.',
        'online_processing_error': 'Error during online mode processing: {error}',
        'search_analysis_log': '🔍 Search analysis: {analysis}',
        'enriching_current_data': '🌐 Enriching with current data...',
        'no_search_necessary': '✅ No search necessary, response based on model knowledge',
        'error_online_mode': 'Error during online mode: {error}',
        'separator': '---'
    },
    'fr': {
        # Existing keys
        'analyze': 'Analysez',
        'content_of_file': "le contenu du fichier",
        'and_provide': "et fournissez une description de son objectif.",
        'content': "Contenu",
        'description': "Description",
        'template': "Je veux une réponse du genre: Le Fichier text.js a pour Objectif : .....",
        'project_structure': 'Structure du projet',
        'file': 'Fichier',
        'other_file': 'Autre fichier',
        'question': 'Question',
        'emoji': 'Ajoute quelques emojis pertinents quand c\'est nécessaire pour rendre la lecture agréable 😊📄✨',
        
        # New keys for local query mode
        'local_analysis_mode': '🔒 MODE ANALYSE LOCALE - Analysez UNIQUEMENT le contexte fourni ci-dessous.',
        'no_external_search': 'Ne faites PAS de recherche externe ou de référence à des informations non fournies.',
        'main_file': 'Fichier principal',
        'file_content': 'Contenu du fichier',
        'directory_context': 'Contexte du répertoire',
        'no_other_files': 'Aucun autre fichier dans le contexte.',
        'instructions': 'Instructions',
        'base_response_only': '- Basez votre réponse UNIQUEMENT sur le code/contenu fourni ci-dessus',
        'missing_info_clarify': '- Si l\'information n\'est pas dans le contexte fourni, dites-le clairement',
        'no_speculation': '- Ne spéculez pas sur des éléments non présents dans les fichiers',
        'focus_local_analysis': '- Concentrez-vous sur l\'analyse du code/contenu local',
        'error_local_analysis': 'Erreur lors de l\'analyse locale',
        'cached_result': 'Utilisation du résultat en cache',
        'analyzing': 'Analyse de',
        'with_context': 'avec contexte de',
        'files_in_directory': 'fichiers dans le répertoire',
        
        # Query analysis keys
        'analyze_query_prompt': 'Analysez cette question et déterminez si elle nécessite des informations récentes/actuelles que votre modèle pourrait ne pas avoir dans ses données d\'entraînement.',
        'query_label': 'Question',
        'json_response_format': 'Répondez au format JSON UNIQUEMENT',
        'needs_search_field': 'needs_search',
        'reason_field': 'reason',
        'search_keywords_field': 'search_keywords',
        'cutoff_relevance_field': 'estimated_cutoff_relevance',
        'short_explanation': 'explication courte',
        'keyword_or_null': 'ou null',
        'examples_need_search': 'Exemples de questions nécessitant une recherche',
        'current_prices_stocks': '- Prix actuels, cours de bourse récents',
        'recent_events_news': '- Événements récents, actualités',
        'new_software_versions': '- Nouvelles versions de logiciels/technologies',
        'recent_statistics_data': '- Statistiques récentes, données gouvernementales',
        'recent_people_companies': '- Informations sur des personnes/entreprises récentes',
        'examples_no_search': 'Exemples de questions NE nécessitant PAS de recherche',
        'general_concepts': '- Concepts généraux, théories établies',
        'programming_syntax': '- Programmation, syntaxe de langages',
        'history_facts': '- Histoire, faits historiques',
        'math_science': '- Mathématiques, sciences fondamentales',
        'analysis_error': 'Erreur d\'analyse',
        'query_analysis_error': 'Erreur lors de l\'analyse de la requête',
        'automatic_analysis_failed': 'Analyse automatique non concluante',
        'error_analysis_fallback': 'Erreur d\'analyse, recherche par sécurité',
        'analyzing_query': 'Analyse de la requête',
        'search_needed': 'Recherche nécessaire',
        'no_search_needed': 'Aucune recherche nécessaire',
        'json_parse_failed': 'Échec du parsing JSON',
        'fallback_analysis': 'Utilisation de l\'analyse de secours',

        # NEW KEYS FOR ONLINE MODE
        'online_mode_title': '🤖 Répondez à cette question en utilisant vos connaissances d\'entraînement.',
        'recent_info_mention': 'Si vous pensez que des informations plus récentes pourraient enrichir votre réponse, mentionnez-le à la fin.',
        'online_instructions_title': 'Instructions :',
        'give_best_answer': '- Donnez votre meilleure réponse basée sur vos connaissances',
        'be_precise_dates': '- Soyez précis sur les dates/versions que vous connaissez',
        'mention_recent_useful': '- Mentionnez si des infos plus récentes seraient utiles',
        'enrichment_title': '🔄 ENRICHISSEMENT DE RÉPONSE',
        'initial_response': 'Votre réponse initiale :',
        'new_info_found': 'Nouvelles informations trouvées :',
        'enrichment_instructions': 'Instructions :',
        'combine_intelligently': '- Combinez intelligemment votre réponse initiale avec les nouvelles infos',
        'update_obsolete': '- Mettez à jour les parties obsolètes si nécessaire',
        'distinguish_info': '- Distinguez clairement les infos de base des infos récentes',
        'cite_sources': '- Citez les sources pour les informations récentes',
        'keep_structure': '- Gardez la structure et le style de votre réponse initiale',
        'prioritize_recent': '- Si les nouvelles infos contredisent vos connaissances, privilégiez les sources récentes',
        'enriched_response': 'Réponse enrichie :',
        'info_enriched': '💡 **Informations enrichies** : Cette réponse combine mes connaissances de base avec des données récentes trouvées en ligne.',
        'search_reason': '🔍 **Raison de la recherche** : {reason}',
        'default_search_reason': 'Information potentiellement obsolète',
        'source_training': '🧠 **Source** : Réponse basée sur mes connaissances d\'entraînement. Aucune recherche récente n\'a été jugée nécessaire.',
        'online_processing_error': 'Erreur lors du traitement en mode online : {error}',
        'search_analysis_log': '🔍 Analyse de recherche : {analysis}',
        'enriching_current_data': '🌐 Enrichissement avec des données actuelles...',
        'no_search_necessary': '✅ Pas de recherche nécessaire, réponse basée sur les connaissances du modèle',
        'error_online_mode': 'Erreur lors du mode online : {error}',
        'separator': '---'
    },
    'es': {
        # Existing keys
        'analyze': 'Analiza',
        'content_of_file': "el contenido del archivo",
        'and_provide': "y proporciona una descripción de su propósito.",
        'content': "Contenido",
        'description': "Descripción",
        'template': "Quiero una respuesta como: El archivo text.js tiene como objetivo .....",
        'project_structure': 'Estructura del proyecto',
        'file': 'Archivo',
        'other_file': 'Otro archivo',
        'question': 'Pregunta',
        'emoji': 'Agrega algunos emojis relevantes cuando sea necesario para que sea agradable de leer 😊📄✨',
        
        # New keys for local query mode
        'local_analysis_mode': '🔒 MODO ANÁLISIS LOCAL - Analiza ÚNICAMENTE el contexto proporcionado a continuación.',
        'no_external_search': 'NO realices búsquedas externas o referencias a información no proporcionada.',
        'main_file': 'Archivo principal',
        'file_content': 'Contenido del archivo',
        'directory_context': 'Contexto del directorio',
        'no_other_files': 'Ningún otro archivo en el contexto.',
        'instructions': 'Instrucciones',
        'base_response_only': '- Basa tu respuesta ÚNICAMENTE en el código/contenido proporcionado arriba',
        'missing_info_clarify': '- Si la información no está en el contexto proporcionado, indícalo claramente',
        'no_speculation': '- No especules sobre elementos no presentes en los archivos',
        'focus_local_analysis': '- Concéntrate en analizar el código/contenido local',
        'error_local_analysis': 'Error durante el análisis local',
        'cached_result': 'Usando resultado en caché',
        'analyzing': 'Analizando',
        'with_context': 'con contexto de',
        'files_in_directory': 'archivos en el directorio',
        
        # Query analysis keys
        'analyze_query_prompt': 'Analiza esta pregunta y determina si requiere información reciente/actual que tu modelo podría no tener en sus datos de entrenamiento.',
        'query_label': 'Pregunta',
        'json_response_format': 'Responde en formato JSON ÚNICAMENTE',
        'needs_search_field': 'needs_search',
        'reason_field': 'reason',
        'search_keywords_field': 'search_keywords',
        'cutoff_relevance_field': 'estimated_cutoff_relevance',
        'short_explanation': 'explicación corta',
        'keyword_or_null': 'o null',
        'examples_need_search': 'Ejemplos de preguntas que requieren búsqueda',
        'current_prices_stocks': '- Precios actuales, precios de acciones recientes',
        'recent_events_news': '- Eventos recientes, noticias',
        'new_software_versions': '- Nuevas versiones de software/tecnologías',
        'recent_statistics_data': '- Estadísticas recientes, datos gubernamentales',
        'recent_people_companies': '- Información sobre personas/empresas recientes',
        'examples_no_search': 'Ejemplos de preguntas que NO requieren búsqueda',
        'general_concepts': '- Conceptos generales, teorías establecidas',
        'programming_syntax': '- Programación, sintaxis de lenguajes',
        'history_facts': '- Historia, hechos históricos',
        'math_science': '- Matemáticas, ciencias fundamentales',
        'analysis_error': 'Error de análisis',
        'query_analysis_error': 'Error durante el análisis de la consulta',
        'automatic_analysis_failed': 'Análisis automático no concluyente',
        'error_analysis_fallback': 'Error de análisis, búsqueda por seguridad',
        'analyzing_query': 'Analizando consulta',
        'search_needed': 'Búsqueda necesaria',
        'no_search_needed': 'No se necesita búsqueda',
        'json_parse_failed': 'Error al analizar respuesta JSON',
        'fallback_analysis': 'Usando análisis de respaldo',

        # NEW KEYS FOR ONLINE MODE
        'online_mode_title': '🤖 Responde esta pregunta usando tus conocimientos de entrenamiento.',
        'recent_info_mention': 'Si crees que información más reciente podría enriquecer tu respuesta, menciónalo al final.',
        'online_instructions_title': 'Instrucciones:',
        'give_best_answer': '- Da tu mejor respuesta basada en tus conocimientos',
        'be_precise_dates': '- Sé preciso sobre fechas/versiones que conoces',
        'mention_recent_useful': '- Menciona si información más reciente sería útil',
        'enrichment_title': '🔄 ENRIQUECIMIENTO DE RESPUESTA',
        'initial_response': 'Tu respuesta inicial:',
        'new_info_found': 'Nueva información encontrada:',
        'enrichment_instructions': 'Instrucciones:',
        'combine_intelligently': '- Combina inteligentemente tu respuesta inicial con la nueva info',
        'update_obsolete': '- Actualiza las partes obsoletas si es necesario',
        'distinguish_info': '- Distingue claramente la info básica de la info reciente',
        'cite_sources': '- Cita fuentes para la información reciente',
        'keep_structure': '- Mantén la estructura y estilo de tu respuesta inicial',
        'prioritize_recent': '- Si la nueva info contradice tus conocimientos, prioriza fuentes recientes',
        'enriched_response': 'Respuesta enriquecida:',
        'info_enriched': '💡 **Información enriquecida**: Esta respuesta combina mis conocimientos básicos con datos recientes encontrados en línea.',
        'search_reason': '🔍 **Razón de búsqueda**: {reason}',
        'default_search_reason': 'Información potencialmente obsoleta',
        'source_training': '🧠 **Fuente**: Respuesta basada en mis conocimientos de entrenamiento. No se consideró necesaria una búsqueda reciente.',
        'online_processing_error': 'Error durante el procesamiento en modo online: {error}',
        'search_analysis_log': '🔍 Análisis de búsqueda: {analysis}',
        'enriching_current_data': '🌐 Enriqueciendo con datos actuales...',
        'no_search_necessary': '✅ No es necesaria búsqueda, respuesta basada en conocimientos del modelo',
        'error_online_mode': 'Error durante modo online: {error}',
        'separator': '---'
    }
}

# Description template
description_template = "Je veux une reponse du genre: Le Fichier text.js a pour Objectif : ....."

# Fonctions utilitaires importées depuis les modules
# extract_text_from_docx, extract_text_from_pdf sont maintenant importés depuis utils.file_utils
# get_model_config est maintenant importé depuis services.api_clients

# call_openai_api est maintenant importé depuis services.api_clients
def _OLD_call_openai_api(prompt, selected_model="gpt-3.5-turbo", max_retries=3, max_tokens_override=None):
    """
    Enhanced OpenAI API call supporting both Chat Completions and Responses API
    max_tokens_override: override la limite de tokens de la config du modèle
    """
    model_config = get_model_config(selected_model)
    model_id = model_config["model_id"]
    api_type = model_config.get("api_type", "chat")
    
    for attempt in range(max_retries):
        try:
            logger.info(f"[OpenAI] Using model: {model_id} with API: {api_type}")
            
            # GPT-5 variants use the new Responses API
            if api_type == "responses":
                request_params = {
                    "model": model_id,
                    "input": prompt
                }
                
                # Add reasoning parameters for GPT-5
                if "reasoning_effort" in model_config:
                    request_params["reasoning"] = {
                        "effort": model_config["reasoning_effort"]
                    }
                
                # Add verbosity parameters for GPT-5
                if "verbosity" in model_config:
                    request_params["text"] = {
                        "verbosity": model_config["verbosity"]
                    }
                
                logger.info(f"[OpenAI] GPT-5 request params: reasoning={model_config.get('reasoning_effort')}, verbosity={model_config.get('verbosity')}")
                response = openai_client.responses.create(**request_params)
                
                # Extract text from responses API - Updated extraction logic
                try:
                    # The actual content is in output array, look for message type with content
                    if hasattr(response, 'output') and response.output:
                        for item in response.output:
                            # Look for message type items
                            if hasattr(item, 'type') and item.type == 'message':
                                if hasattr(item, 'content') and item.content:
                                    for content_item in item.content:
                                        if hasattr(content_item, 'text') and content_item.text:
                                            logger.info(f"[OpenAI] Successfully extracted GPT-5 response")
                                            return content_item.text.strip()
                    
                    # Fallback extraction methods
                    if hasattr(response, 'text') and hasattr(response.text, 'content'):
                        return response.text.content.strip()
                    elif hasattr(response, 'content'):
                        return response.content.strip()
                    else:
                        # Final fallback: try to extract from string representation
                        response_str = str(response)
                        logger.warning(f"[OpenAI] Using fallback extraction for GPT-5 response")
                        
                        # Try to extract text from the string representation
                        import re
                        text_match = re.search(r"text='([^']*)'", response_str)
                        if text_match:
                            extracted_text = text_match.group(1)
                            # Decode escape sequences
                            extracted_text = extracted_text.replace('\\n', '\n').replace("\\'", "'")
                            return extracted_text.strip()
                        
                        return "Error: Could not extract response content from GPT-5"
                        
                except Exception as extraction_error:
                    logger.error(f"[OpenAI] Error extracting GPT-5 response: {extraction_error}")
                    return f"Error extracting response: {str(extraction_error)}"
            
            # Standard models use Chat Completions API
            else:
                # Utiliser max_tokens_override si fourni, sinon la config du modèle
                max_tokens = max_tokens_override if max_tokens_override is not None else model_config.get("max_tokens", 500)
                request_params = {
                    "model": model_id,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.5,
                    "max_tokens": max_tokens
                }
                
                response = openai_client.chat.completions.create(**request_params)
                return response.choices[0].message.content.strip()
            
        except Exception as e:
            logger.warning(f"[OpenAI] Attempt {attempt+1}/{max_retries} failed with {model_id}: {str(e)}")
            
            # Smart fallback strategy
            error_str = str(e).lower()
            if any(keyword in error_str for keyword in ["model", "not found", "unsupported", "invalid"]):
                # GPT-5 fallback chain
                if model_id == "gpt-5":
                    logger.info("[OpenAI] GPT-5 not available, falling back to GPT-5-mini")
                    model_id = "gpt-5-mini"
                    model_config = get_model_config("gpt-5-mini")
                    api_type = model_config.get("api_type", "responses")
                    continue
                elif model_id == "gpt-5-mini":
                    logger.info("[OpenAI] GPT-5-mini not available, falling back to GPT-5-nano")
                    model_id = "gpt-5-nano"
                    model_config = get_model_config("gpt-5-nano")
                    api_type = model_config.get("api_type", "responses")
                    continue
                elif model_id == "gpt-5-nano":
                    logger.info("[OpenAI] GPT-5-nano not available, falling back to GPT-4o")
                    model_id = "gpt-4o"
                    model_config = get_model_config("gpt-4o")
                    api_type = model_config.get("api_type", "chat")
                    continue
                elif model_id == "gpt-4o":
                    logger.info("[OpenAI] GPT-4o not available, falling back to GPT-3.5-turbo")
                    model_id = "gpt-3.5-turbo"
                    model_config = get_model_config("gpt-3.5-turbo")
                    api_type = model_config.get("api_type", "chat")
                    continue
            
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logger.warning(f"[OpenAI] Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                continue
            else:
                logger.error(f"[OpenAI] Max retries exceeded: {str(e)}")
                raise e
    
    raise RuntimeError("[OpenAI] Max retries exceeded")

# stream_response est maintenant importé depuis services.api_clients
def _OLD_stream_response(prompt, selected_model="gpt-3.5-turbo"):
    """
    Génère une réponse en streaming pour n'importe quel modèle
    """
    model_config = get_model_config(selected_model)
    
    try:
        # Pour les modèles OpenAI avec support du streaming
        if selected_model.lower() in ["gpt-3.5-turbo", "gpt-4o", "openai"]:
            model_id = model_config["model_id"]
            
            stream = openai_client.chat.completions.create(
                model=model_id,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=model_config.get("max_tokens", 500),
                stream=True
            )
            
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
            
            yield f"data: {json.dumps({'done': True})}\n\n"
            
        # Pour Mistral - pas de streaming natif, on simule
        elif selected_model.lower() == "mistral":
            try:
                # Appel normal à Mistral (sans streaming)
                response = call_mistral_api(prompt)
                
                # Simuler le streaming en envoyant par morceaux
                words = response.split(' ')
                for i in range(0, len(words), 3):  # Envoyer 3 mots à la fois
                    chunk = ' '.join(words[i:i+3])
                    if i + 3 < len(words):
                        chunk += ' '
                    yield f"data: {json.dumps({'content': chunk})}\n\n"
                    time.sleep(0.05)  # Petit délai pour simuler le streaming
                
                yield f"data: {json.dumps({'done': True})}\n\n"
                
            except Exception as mistral_error:
                # Si Mistral échoue, fallback vers GPT-3.5
                logger.warning(f"Mistral failed: {str(mistral_error)}, falling back to GPT-3.5")
                yield f"data: {json.dumps({'warning': 'Mistral indisponible, utilisation de GPT-3.5'})}\n\n"
                
                # Utiliser GPT-3.5 en fallback
                stream = openai_client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5,
                    max_tokens=500,
                    stream=True
                )
                
                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
                
                yield f"data: {json.dumps({'done': True})}\n\n"
        
        # Pour Llama/Ollama - pas de streaming natif
        elif selected_model.lower() in ["llama3", "llama3.2"]:
            try:
                response = call_ollama_api(prompt, selected_model)
                
                # Simuler le streaming
                words = response.split(' ')
                for i in range(0, len(words), 3):
                    chunk = ' '.join(words[i:i+3])
                    if i + 3 < len(words):
                        chunk += ' '
                    yield f"data: {json.dumps({'content': chunk})}\n\n"
                    time.sleep(0.05)
                
                yield f"data: {json.dumps({'done': True})}\n\n"
                
            except Exception as ollama_error:
                logger.warning(f"Ollama failed: {str(ollama_error)}, falling back to GPT-3.5")
                yield f"data: {json.dumps({'warning': 'Ollama indisponible, utilisation de GPT-3.5'})}\n\n"
                
                # Fallback vers GPT-3.5
                stream = openai_client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5,
                    max_tokens=500,
                    stream=True
                )
                
                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
                
                yield f"data: {json.dumps({'done': True})}\n\n"
        
        else:
            # Modèle inconnu, utiliser GPT-3.5 par défaut
            logger.warning(f"Unknown model {selected_model}, using GPT-3.5")
            yield f"data: {json.dumps({'warning': f'Modèle {selected_model} inconnu, utilisation de GPT-3.5'})}\n\n"
            
            stream = openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=500,
                stream=True
            )
            
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
            
            yield f"data: {json.dumps({'done': True})}\n\n"
            
    except Exception as e:
        logger.error(f"Erreur streaming: {str(e)}")
        yield f"data: {json.dumps({'error': str(e)})}\n\n"

# call_mistral_api est maintenant importé depuis services.api_clients
def _OLD_call_mistral_api(prompt, max_retries=3):
    mistral_api_key = os.getenv("MISTRAL_API_KEY")
    if not mistral_api_key:
        raise ValueError("MISTRAL_API_KEY environment variable not set")

    url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {mistral_api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": MISTRAL_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 500,
        "temperature": 0.5
    }

    for attempt in range(max_retries):
        try:
            response = requests.post(url, headers=headers, json=payload)
            if response.status_code == 429:
                retry_after = int(response.headers.get("Retry-After", 1))
                logger.warning(f"[Mistral] Rate limited. Retrying in {retry_after} sec (attempt {attempt+1}/{max_retries})...")
                time.sleep(retry_after)
                continue

            response.raise_for_status()
            result = response.json()
            return result["choices"][0]["message"]["content"].strip()

        except requests.exceptions.HTTPError as e:
            if response.status_code == 429 and attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logger.warning(f"[Mistral] 429 Too Many Requests. Retrying in {wait_time} sec (attempt {attempt+1}/{max_retries})...")
                time.sleep(wait_time)
                continue
            else:
                logger.error(f"[Mistral] HTTP Error: {e}")
                raise e

        except Exception as e:
            logger.error(f"[Mistral] Unexpected Error: {e}")
            raise e

    raise RuntimeError("[Mistral] Max retries exceeded")

# call_ollama_api est maintenant importé depuis services.api_clients
def _OLD_call_ollama_api(prompt, selected_model="llama3", max_retries=3):
    """Enhanced Ollama API call"""
    model_config = get_model_config(selected_model)
    model_id = model_config["model_id"]
    
    for attempt in range(max_retries):
        try:
            response = requests.post(OLLAMA_API_URL, json={
                "model": model_id,
                "prompt": prompt,
                "stream": False
            })
            response.raise_for_status()
            data = response.json()
            return data.get("response") or data.get("message", {}).get("content", "No response").strip()
            
        except Exception as e:
            logger.warning(f"[Ollama] Attempt {attempt+1}/{max_retries} failed: {str(e)}")
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logger.warning(f"[Ollama] Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                continue
            else:
                logger.error(f"[Ollama] Max retries exceeded: {str(e)}")
                raise e
    
    raise RuntimeError("[Ollama] Max retries exceeded")

async def infer_corpus_actions(documents: List[Document], language: str = 'en') -> Dict[str, Any]:
    """
    Utilise un petit appel modèle pour deviner le type de corpus (CV, rapports annuels, etc.)
    et proposer des actions suggérées (boutons) adaptées.
    """
    try:
        # D'abord, détecter le type de document dominant dans le corpus (avec plus de contexte)
        document_types = {}
        document_details = {}  # Stocker plus d'infos par type
        
        for doc in documents[:30]:  # Analyser plus de documents
            doc_content = doc.page_content[:4000] if len(doc.page_content) > 4000 else doc.page_content  # Plus de contenu pour une meilleure détection
            doc_type = await detect_document_type(doc_content, doc.metadata.get('fileName', ''))
            document_types[doc_type] = document_types.get(doc_type, 0) + 1
            
            # Stocker les détails pour enrichir le prompt
            if doc_type not in document_details:
                document_details[doc_type] = []
            meta_name = doc.metadata.get('fileName') or doc.metadata.get('source') or 'document'
            # Prendre un échantillon plus intelligent (début + milieu si disponible)
            content_sample = doc_content[:1000]
            if len(doc_content) > 2000:
                mid_point = len(doc_content) // 2
                content_sample += " ... " + doc_content[mid_point:mid_point+500]
            document_details[doc_type].append({
                'name': meta_name,
                'snippet': content_sample.replace('\n', ' ').strip()[:1200]  # Plus de contexte
            })
        
        # Trouver le type de document le plus fréquent
        dominant_type = max(document_types.items(), key=lambda x: x[1])[0] if document_types else 'document_generique'
        
        # Construire un résumé enrichi du corpus pour le prompt
        sample_texts = []
        # Inclure tous les documents du type dominant (jusqu'à 10)
        dominant_docs = document_details.get(dominant_type, [])[:10]
        for doc_info in dominant_docs:
            sample_texts.append(f"- {doc_info['name']}: {doc_info['snippet']}")
        
        # Ajouter quelques exemples d'autres types si présents
        other_types = [dt for dt in document_types.keys() if dt != dominant_type]
        for other_type in other_types[:2]:  # Max 2 autres types
            for doc_info in document_details.get(other_type, [])[:2]:  # 2 exemples par type
                sample_texts.append(f"- [{other_type}] {doc_info['name']}: {doc_info['snippet'][:600]}")
        
        corpus_preview = "\n".join(sample_texts)
        
        # Informations additionnelles pour améliorer la précision
        type_distribution = ", ".join([f"{dt}: {count}" for dt, count in sorted(document_types.items(), key=lambda x: x[1], reverse=True)[:3]])

        # Actions spécifiques selon le type de document et la langue
        specific_actions_prompts = {
            'fr': {
                'contrat_location': """
Actions spécifiques pour un contrat de location (bail) :
1. "Vérifier les parties" - Identifie et liste toutes les parties (locataire, bailleur) avec leurs coordonnées complètes
2. "Vérifier les dates" - Extrait toutes les dates importantes : signature, début, fin, durée, préavis
3. "Vérifier les montants" - Liste tous les montants : loyer, caution, charges, indexation
4. "Analyser les clauses à risque" - Identifie les clauses potentiellement problématiques ou désavantageuses
5. "Vérifier les obligations" - Liste les obligations du locataire et du bailleur
6. "Vérifier le bien loué" - Détaille les caractéristiques du bien (adresse, superficie, type)
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'contrat_travail': """
Actions spécifiques pour un contrat de travail :
1. "Vérifier les parties" - Identifie l'employeur et l'employé avec leurs coordonnées
2. "Vérifier les dates" - Extrait les dates : signature, début, période d'essai, fin
3. "Vérifier la rémunération" - Détaille le salaire, primes, avantages, révisions
4. "Vérifier les obligations" - Liste les obligations de l'employé et de l'employeur
5. "Analyser les clauses à risque" - Identifie les clauses restrictives ou problématiques
6. "Vérifier les conditions" - Détaille les conditions de travail, horaires, lieu
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'contrat_vente': """
Actions spécifiques pour un contrat de vente :
1. "Vérifier les parties" - Identifie l'acheteur et le vendeur avec leurs coordonnées
2. "Vérifier les dates" - Extrait les dates : signature, livraison, paiement
3. "Vérifier les montants" - Détaille le prix, acompte, modalités de paiement
4. "Vérifier l'objet" - Décrit précisément l'objet de la vente
5. "Analyser les garanties" - Liste les garanties et conditions de garantie
6. "Vérifier les conditions" - Détaille les conditions de vente, délais, pénalités
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'contrat_generique': """
Actions spécifiques pour un contrat :
1. "Vérifier les parties" - Identifie toutes les parties avec leurs coordonnées
2. "Vérifier les dates importantes" - Extrait toutes les dates clés du contrat
3. "Vérifier les montants" - Liste tous les montants et modalités financières
4. "Analyser les clauses à risque" - Identifie les clauses potentiellement problématiques
5. "Vérifier les obligations" - Liste les obligations de chaque partie
6. "Vérifier l'objet" - Décrit précisément l'objet du contrat
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'testament': """
Actions spécifiques pour un testament :
1. "Vérifier le testateur" - Identifie le testateur et ses coordonnées
2. "Vérifier les bénéficiaires" - Liste tous les bénéficiaires et leurs parts
3. "Vérifier les dates" - Extrait les dates : rédaction, signature, modifications
4. "Vérifier les legs" - Détaille tous les legs et héritages
5. "Vérifier les conditions" - Liste les conditions et clauses particulières
6. "Vérifier l'exécuteur" - Identifie l'exécuteur testamentaire
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'acte_notarie': """
Actions spécifiques pour un acte notarié :
1. "Vérifier les parties" - Identifie toutes les parties impliquées
2. "Vérifier les dates" - Extrait toutes les dates importantes
3. "Vérifier les montants" - Liste tous les montants et transactions
4. "Vérifier l'objet" - Décrit précisément l'objet de l'acte
5. "Vérifier le notaire" - Identifie le notaire et son étude
6. "Vérifier les conditions" - Détaille les conditions et clauses
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'lettre': """
Actions spécifiques pour une lettre :
1. "Vérifier l'expéditeur" - Identifie l'expéditeur avec ses coordonnées
2. "Vérifier le destinataire" - Identifie le destinataire avec ses coordonnées
3. "Vérifier la date" - Extrait la date de la lettre
4. "Vérifier l'objet" - Décrit l'objet et le but de la lettre
5. "Vérifier infos clés" - Extrait les informations importantes (montants, références, engagements)
6. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'document_financier': """
Actions spécifiques pour un document financier :
1. "Vérifier les parties" - Identifie les parties concernées (employeur, employé, institution)
2. "Vérifier la période" - Extrait la période couverte par le document
3. "Vérifier les montants" - Liste tous les montants (revenus, déductions, impôts, totaux)
4. "Vérifier déductions" - Détaille toutes les déductions (impôts, cotisations)
5. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'cv_resume': """
Actions spécifiques pour un CV/Resume :
1. "Vérifier identité" - Identifie le nom complet, coordonnées et informations de contact
2. "Vérifier expérience" - Liste toutes les expériences professionnelles avec dates, postes et entreprises
3. "Vérifier formation" - Détaille les diplômes, formations et certifications avec dates et institutions
4. "Vérifier compétences" - Liste les compétences techniques, linguistiques et autres compétences
5. "Vérifier réalisations" - Extrait les réalisations, projets et accomplissements majeurs
6. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'facture_invoice': """
Actions spécifiques pour une facture/Invoice :
1. "Vérifier les parties" - Identifie l'émetteur (vendeur) et le client avec leurs coordonnées
2. "Vérifier les dates" - Extrait la date d'émission, date d'échéance et date de paiement
3. "Vérifier les montants" - Liste le montant HT, TVA, montant TTC et modalités de paiement
4. "Vérifier les articles" - Détaille tous les articles/lignes de facturation avec quantités et prix
5. "Vérifier références" - Extrait le numéro de facture, références client et numéros de commande
6. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'contrat_prenuptial': """
Actions spécifiques pour un contrat de mariage/prénuptial :
1. "Vérifier les parties" - Identifie les époux/futurs époux avec leurs coordonnées complètes
2. "Vérifier les dates" - Extrait la date de signature et la date de mariage prévue
3. "Vérifier le régime matrimonial" - Détaille le régime choisi (séparation de biens, communauté, etc.)
4. "Vérifier les biens" - Liste les biens propres et les biens communs avec leurs valeurs
5. "Vérifier les clauses particulières" - Identifie les clauses spécifiques (héritage, donation, etc.)
6. "Analyser les conditions" - Détaille les conditions de modification et de dissolution
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'procuration_poa': """
Actions spécifiques pour une procuration/Power Of Attorney :
1. "Vérifier le mandant" - Identifie le mandant (donneur de procuration) avec ses coordonnées
2. "Vérifier le mandataire" - Identifie le mandataire (représentant) avec ses coordonnées
3. "Vérifier les dates" - Extrait la date de signature, date de début et date d'expiration
4. "Vérifier les pouvoirs" - Détaille tous les pouvoirs accordés (signature, gestion, décisions)
5. "Vérifier les limitations" - Liste les limitations et restrictions des pouvoirs
6. "Vérifier les conditions" - Détaille les conditions de révocation et d'utilisation
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'accord_confidentialite_nda': """
Actions spécifiques pour un accord de confidentialité/NDA :
1. "Vérifier les parties" - Identifie toutes les parties (divulgateur et bénéficiaire) avec leurs coordonnées
2. "Vérifier les dates" - Extrait la date de signature et la durée de l'accord
3. "Vérifier les informations confidentielles" - Détaille la portée des informations couvertes
4. "Vérifier les obligations" - Liste les obligations de confidentialité et de non-divulgation
5. "Analyser les exceptions" - Identifie les exceptions autorisées (loi, ordre judiciaire, etc.)
6. "Vérifier les sanctions" - Détaille les pénalités et recours en cas de violation
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON""",
                'acte_propriete_immobiliere': """
Actions spécifiques pour un acte de propriété immobilière/Real Estate Deed :
1. "Vérifier le propriétaire" - Identifie le propriétaire actuel avec ses coordonnées
2. "Vérifier le bien" - Détaille les caractéristiques du bien (adresse, superficie, type, parcelle cadastrale)
3. "Vérifier les dates" - Extrait la date d'acquisition, date de l'acte et historique des transactions
4. "Vérifier les montants" - Liste le prix d'acquisition, taxes et frais associés
5. "Vérifier les charges" - Détaille les servitudes, hypothèques et autres charges
6. "Vérifier le bornage" - Identifie les limites et bornes de la propriété
7. "Extraire données structurées" - Extrait toutes les données dans un format structuré JSON"""
            },
            'en': {
                'contrat_location': """
Specific actions for a rental contract (lease):
1. "Verify parties" - Identifies and lists all parties (tenant, landlord) with their full contact details
2. "Verify dates" - Extracts all important dates: signature, start, end, duration, notice period
3. "Verify amounts" - Lists all amounts: rent, deposit, charges, indexation
4. "Analyze risky clauses" - Identifies potentially problematic or disadvantageous clauses
5. "Verify obligations" - Lists the obligations of the tenant and landlord
6. "Verify rented property" - Details the property characteristics (address, area, type)
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_travail': """
Specific actions for an employment contract:
1. "Verify parties" - Identifies the employer and employee with their contact details
2. "Verify dates" - Extracts dates: signature, start, probation period, end
3. "Verify remuneration" - Details salary, bonuses, benefits, revisions
4. "Verify obligations" - Lists the obligations of the employee and employer
5. "Analyze risky clauses" - Identifies restrictive or problematic clauses
6. "Verify conditions" - Details working conditions, hours, location
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_vente': """
Specific actions for a sale contract:
1. "Verify parties" - Identifies the buyer and seller with their contact details
2. "Verify dates" - Extracts dates: signature, delivery, payment
3. "Verify amounts" - Details price, down payment, payment terms
4. "Verify subject" - Precisely describes the subject of the sale
5. "Analyze guarantees" - Lists guarantees and warranty conditions
6. "Verify conditions" - Details sale conditions, deadlines, penalties
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_generique': """
Specific actions for a contract:
1. "Verify parties" - Identifies all parties with their contact details
2. "Verify important dates" - Extracts all key dates in the contract
3. "Verify amounts" - Lists all amounts and financial terms
4. "Analyze risky clauses" - Identifies potentially problematic clauses
5. "Verify obligations" - Lists the obligations of each party
6. "Verify subject" - Precisely describes the subject of the contract
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'testament': """
Specific actions for a will:
1. "Verify testator" - Identifies the testator and their contact details
2. "Verify beneficiaries" - Lists all beneficiaries and their shares
3. "Verify dates" - Extracts dates: drafting, signature, modifications
4. "Verify bequests" - Details all bequests and inheritances
5. "Verify conditions" - Lists conditions and particular clauses
6. "Verify executor" - Identifies the executor
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'acte_notarie': """
Specific actions for a notarial act:
1. "Verify parties" - Identifies all parties involved
2. "Verify dates" - Extracts all important dates
3. "Verify amounts" - Lists all amounts and transactions
4. "Verify subject" - Precisely describes the subject of the act
5. "Verify notary" - Identifies the notary and their office
6. "Verify conditions" - Details conditions and clauses
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'lettre': """
Specific actions for a letter:
1. "Verify sender" - Identifies the sender with their contact details
2. "Verify recipient" - Identifies the recipient with their contact details
3. "Verify date" - Extracts the letter date
4. "Verify subject" - Describes the subject and purpose of the letter
5. "Verify key info" - Extracts important information (amounts, references, commitments)
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'document_financier': """
Specific actions for a financial document:
1. "Verify parties" - Identifies parties concerned (employer, employee, institution)
2. "Verify period" - Extracts the period covered by the document
3. "Verify amounts" - Lists all amounts (income, deductions, taxes, totals)
4. "Verify deductions" - Details all deductions (taxes, contributions)
5. "Extract structured data" - Extracts all data in a structured JSON format""",
                'cv_resume': """
Specific actions for a CV/Resume:
1. "Verify identity" - Identifies full name, contact information and personal details
2. "Verify experience" - Lists all professional experiences with dates, positions and companies
3. "Verify education" - Details degrees, training and certifications with dates and institutions
4. "Verify skills" - Lists technical, language and other skills
5. "Verify achievements" - Extracts major achievements, projects and accomplishments
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'facture_invoice': """
Specific actions for an invoice:
1. "Verify parties" - Identifies the issuer (seller) and customer with their contact details
2. "Verify dates" - Extracts issue date, due date and payment date
3. "Verify amounts" - Lists net amount, tax, total amount and payment terms
4. "Verify items" - Details all invoice items/lines with quantities and prices
5. "Verify references" - Extracts invoice number, customer references and order numbers
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_prenuptial': """
Specific actions for a prenuptial agreement:
1. "Verify parties" - Identifies the spouses/future spouses with their full contact details
2. "Verify dates" - Extracts signature date and planned marriage date
3. "Verify marital regime" - Details the chosen regime (separation of property, community, etc.)
4. "Verify assets" - Lists separate and community property with their values
5. "Verify special clauses" - Identifies specific clauses (inheritance, donation, etc.)
6. "Analyze conditions" - Details conditions for modification and dissolution
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'procuration_poa': """
Specific actions for a Power Of Attorney:
1. "Verify principal" - Identifies the principal (grantor) with contact details
2. "Verify agent" - Identifies the agent (attorney-in-fact) with contact details
3. "Verify dates" - Extracts signature date, start date and expiration date
4. "Verify powers" - Details all granted powers (signing, management, decisions)
5. "Verify limitations" - Lists limitations and restrictions of powers
6. "Verify conditions" - Details conditions for revocation and use
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'accord_confidentialite_nda': """
Specific actions for a Non-Disclosure Agreement (NDA):
1. "Verify parties" - Identifies all parties (discloser and recipient) with their contact details
2. "Verify dates" - Extracts signature date and duration of the agreement
3. "Verify confidential information" - Details the scope of covered information
4. "Verify obligations" - Lists confidentiality and non-disclosure obligations
5. "Analyze exceptions" - Identifies authorized exceptions (law, court order, etc.)
6. "Verify penalties" - Details penalties and remedies in case of breach
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'acte_propriete_immobiliere': """
Specific actions for a Real Estate Deed:
1. "Verify owner" - Identifies the current owner with contact details
2. "Verify property" - Details property characteristics (address, area, type, cadastral lot)
3. "Verify dates" - Extracts acquisition date, deed date and transaction history
4. "Verify amounts" - Lists acquisition price, taxes and associated fees
5. "Verify encumbrances" - Details easements, mortgages and other encumbrances
6. "Verify boundaries" - Identifies property limits and boundaries
7. "Extract structured data" - Extracts all data in a structured JSON format"""
            },
            'es': {
                'contrat_location': """
Acciones específicas para un contrato de alquiler (arrendamiento):
1. "Verificar partes" - Identifica y enumera todas las partes (inquilino, arrendador) con sus datos de contacto completos
2. "Verificar fechas" - Extrae todas las fechas importantes: firma, inicio, fin, duración, preaviso
3. "Verificar montos" - Enumera todos los montos: alquiler, depósito, gastos, indexación
4. "Analizar cláusulas de riesgo" - Identifica cláusulas potencialmente problemáticas o desventajosas
5. "Verificar obligaciones" - Enumera las obligaciones del inquilino y del arrendador
6. "Verificar propiedad alquilada" - Detalla las características de la propiedad (dirección, superficie, tipo)
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_travail': """
Acciones específicas para un contrato de trabajo:
1. "Verificar partes" - Identifica al empleador y al empleado con sus datos de contacto
2. "Verificar fechas" - Extrae las fechas: firma, inicio, período de prueba, fin
3. "Verificar remuneración" - Detalla salario, bonos, beneficios, revisiones
4. "Verificar obligaciones" - Enumera las obligaciones del empleado y del empleador
5. "Analizar cláusulas de riesgo" - Identifica cláusulas restrictivas o problemáticas
6. "Verificar condiciones" - Detalla las condiciones de trabajo, horarios, lugar
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_vente': """
Acciones específicas para un contrato de venta:
1. "Verificar partes" - Identifica al comprador y al vendedor con sus datos de contacto
2. "Verificar fechas" - Extrae las fechas: firma, entrega, pago
3. "Verificar montos" - Detalla precio, anticipo, modalidades de pago
4. "Verificar objeto" - Describe precisamente el objeto de la venta
5. "Analizar garantías" - Enumera las garantías y condiciones de garantía
6. "Verificar condiciones" - Detalla las condiciones de venta, plazos, penalizaciones
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_generique': """
Acciones específicas para un contrato:
1. "Verificar partes" - Identifica todas las partes con sus datos de contacto
2. "Verificar fechas importantes" - Extrae todas las fechas clave del contrato
3. "Verificar montos" - Enumera todos los montos y modalidades financieras
4. "Analizar cláusulas de riesgo" - Identifica cláusulas potencialmente problemáticas
5. "Verificar obligaciones" - Enumera las obligaciones de cada parte
6. "Verificar objeto" - Describe precisamente el objeto del contrato
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'testament': """
Acciones específicas para un testamento:
1. "Verificar testador" - Identifica al testador y sus datos de contacto
2. "Verificar beneficiarios" - Enumera todos los beneficiarios y sus partes
3. "Verificar fechas" - Extrae las fechas: redacción, firma, modificaciones
4. "Verificar legados" - Detalla todos los legados y herencias
5. "Verificar condiciones" - Enumera las condiciones y cláusulas particulares
6. "Verificar ejecutor" - Identifica al albacea
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'acte_notarie': """
Acciones específicas para un acta notarial:
1. "Verificar partes" - Identifica todas las partes involucradas
2. "Verificar fechas" - Extrae todas las fechas importantes
3. "Verificar montos" - Enumera todos los montos y transacciones
4. "Verificar objeto" - Describe precisamente el objeto del acta
5. "Verificar notario" - Identifica al notario y su estudio
6. "Verificar condiciones" - Detalla las condiciones y cláusulas
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'lettre': """
Acciones específicas para una carta:
1. "Verificar remitente" - Identifica al remitente con sus datos de contacto
2. "Verificar destinatario" - Identifica al destinatario con sus datos de contacto
3. "Verificar fecha" - Extrae la fecha de la carta
4. "Verificar objeto" - Describe el objeto y el propósito de la carta
5. "Verificar información clave" - Extrae información importante (montos, referencias, compromisos)
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'document_financier': """
Acciones específicas para un documento financiero:
1. "Verificar partes" - Identifica las partes concernidas (empleador, empleado, institución)
2. "Verificar período" - Extrae el período cubierto por el documento
3. "Verificar montos" - Enumera todos los montos (ingresos, deducciones, impuestos, totales)
4. "Verificar deducciones" - Detalla todas las deducciones (impuestos, cotizaciones)
5. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'cv_resume': """
Acciones específicas para un CV/Resume:
1. "Verificar identidad" - Identifica el nombre completo, datos de contacto e información personal
2. "Verificar experiencia" - Enumera todas las experiencias profesionales con fechas, puestos y empresas
3. "Verificar formación" - Detalla títulos, formaciones y certificaciones con fechas e instituciones
4. "Verificar competencias" - Enumera las competencias técnicas, lingüísticas y otras habilidades
5. "Verificar logros" - Extrae los logros, proyectos y realizaciones principales
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'facture_invoice': """
Acciones específicas para una factura/Invoice:
1. "Verificar partes" - Identifica el emisor (vendedor) y el cliente con sus datos de contacto
2. "Verificar fechas" - Extrae la fecha de emisión, fecha de vencimiento y fecha de pago
3. "Verificar montos" - Enumera el importe sin IVA, IVA, importe total y modalidades de pago
4. "Verificar artículos" - Detalla todos los artículos/líneas de facturación con cantidades y precios
5. "Verificar referencias" - Extrae el número de factura, referencias del cliente y números de pedido
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_prenuptial': """
Acciones específicas para un contrato de matrimonio/prenupcial:
1. "Verificar partes" - Identifica los cónyuges/futuros cónyuges con sus datos de contacto completos
2. "Verificar fechas" - Extrae la fecha de firma y la fecha de matrimonio prevista
3. "Verificar régimen matrimonial" - Detalla el régimen elegido (separación de bienes, comunidad, etc.)
4. "Verificar bienes" - Enumera los bienes propios y los bienes comunes con sus valores
5. "Verificar cláusulas particulares" - Identifica las cláusulas específicas (herencia, donación, etc.)
6. "Analizar condiciones" - Detalla las condiciones de modificación y disolución
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'procuration_poa': """
Acciones específicas para una procuración/Power Of Attorney:
1. "Verificar mandante" - Identifica al mandante (otorgante) con sus datos de contacto
2. "Verificar mandatario" - Identifica al mandatario (representante) con sus datos de contacto
3. "Verificar fechas" - Extrae la fecha de firma, fecha de inicio y fecha de expiración
4. "Verificar poderes" - Detalla todos los poderes concedidos (firma, gestión, decisiones)
5. "Verificar limitaciones" - Enumera las limitaciones y restricciones de los poderes
6. "Verificar condiciones" - Detalla las condiciones de revocación y uso
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'accord_confidentialite_nda': """
Acciones específicas para un acuerdo de confidencialidad/NDA:
1. "Verificar partes" - Identifica todas las partes (divulgador y beneficiario) con sus datos de contacto
2. "Verificar fechas" - Extrae la fecha de firma y la duración del acuerdo
3. "Verificar información confidencial" - Detalla el alcance de la información cubierta
4. "Verificar obligaciones" - Enumera las obligaciones de confidencialidad y no divulgación
5. "Analizar excepciones" - Identifica las excepciones autorizadas (ley, orden judicial, etc.)
6. "Verificar sanciones" - Detalla las penalizaciones y recursos en caso de violación
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructur ado""",
                'acte_propriete_immobiliere': """
Acciones específicas para un acta de propiedad inmobiliaria/Real Estate Deed:
1. "Verificar propietario" - Identifica al propietario actual con sus datos de contacto
2. "Verificar propiedad" - Detalla las características de la propiedad (dirección, superficie, tipo, lote catastral)
3. "Verificar fechas" - Extrae la fecha de adquisición, fecha del acta e historial de transacciones
4. "Verificar montos" - Enumera el precio de adquisición, impuestos y honorarios asociados
5. "Verificar cargas" - Detalla las servidumbres, hipotecas y otras cargas
6. "Verificar límites" - Identifica los límites y linderos de la propiedad
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado"""
            }
        }

        # Sélectionner les prompts spécifiques selon la langue
        lang_prompts = specific_actions_prompts.get(language, specific_actions_prompts['en'])
        specific_prompt = lang_prompts.get(dominant_type, '')
        
        base_prompts = {
            'fr': f"""
Tu reçois un aperçu de plusieurs documents importés par un utilisateur.
Le type de document dominant détecté est : {dominant_type}

{specific_prompt}

À partir de ces textes UNIQUEMENT, propose jusqu'à 7 actions suggérées au format JSON (elles seront affichées comme des boutons dans l'interface).
Les actions doivent être PRATIQUES et correspondre à ce que les utilisateurs vérifient habituellement pour ce type de document.

Retourne du JSON STRICT avec exactement cette forme :
{{
  "domain": "etiquette_courte_du_domaine",
  "suggested_actions": [
    {{
      "id": "identifiant_machine",
      "title": "Label court du bouton (max 25 caractères)",
      "description": "Une phrase expliquant ce que fait cette action pour l'utilisateur.",
      "sample_prompt": "Prompt complet en langage naturel que l'app pourra envoyer à l'assistant quand l'utilisateur clique sur ce bouton."
    }}
  ]
}}

IMPORTANT :
- Les titres doivent être courts et clairs (max 25 caractères)
- Les actions doivent être spécifiques au type de document détecté
- Priorise les actions que les utilisateurs vérifient habituellement (parties, dates, montants, clauses, obligations)
- Inclus toujours "Extraire données structurées" comme dernière action
""",
            'en': f"""
You are a document analysis expert. You receive a preview of documents uploaded by a user.

CONTEXTUAL INFORMATION:
- Dominant document type detected: {dominant_type}
- Type distribution in corpus: {type_distribution}
- Total documents analyzed: {len(documents)}

{specific_prompt}

ANALYZE THE PROVIDED DOCUMENTS and propose up to 7 PRECISE and USEFUL suggested actions in JSON format.
Actions must be:
1. SPECIFIC to the actual content of the documents (analyze the content to propose relevant actions)
2. PRACTICAL (what users actually check in this document type)
3. ACTIONABLE (each action should allow extracting or verifying precise information)

DETAILED INSTRUCTIONS:
- First analyze the document content to identify key elements present (dates, amounts, parties, clauses, etc.)
- Propose actions that exactly match what is present in the documents
- Titles must be short and clear (max 25 chars), actionable and direct
- sample_prompt must be precise and request specific information findable in the documents
- Always prioritize: parties/identities → dates → amounts/finances → obligations/clauses → other details
- Always include "Extract structured data" as the last action

Return STRICT valid JSON with this exact shape:
{{
  "domain": "short_domain_label",
  "suggested_actions": [
    {{
      "id": "machine_readable_id",
      "title": "Short button label (max 25 chars)",
      "description": "One sentence explaining what this action does for the user.",
      "sample_prompt": "Complete, precise and actionable natural language prompt the app will send to the assistant when the user clicks this action. The prompt must request specific information present in the documents."
    }}
  ]
}}
""",
            'es': f"""
Eres un experto en análisis documental. Recibes una vista previa de documentos cargados por un usuario.

INFORMACIÓN CONTEXTUAL:
- Tipo de documento dominante detectado: {dominant_type}
- Distribución de tipos en el corpus: {type_distribution}
- Total de documentos analizados: {len(documents)}

{specific_prompt}

ANALIZA LOS DOCUMENTOS PROPORCIONADOS y propone hasta 7 acciones sugeridas PRECISAS y ÚTILES en formato JSON.
Las acciones deben ser:
1. ESPECÍFICAS al contenido real de los documentos (analiza el contenido para proponer acciones relevantes)
2. PRÁCTICAS (lo que los usuarios realmente verifican en este tipo de documento)
3. ACCIONABLES (cada acción debe permitir extraer o verificar información precisa)

INSTRUCCIONES DETALLADAS:
- Primero analiza el contenido del documento para identificar elementos clave presentes (fechas, montos, partes, cláusulas, etc.)
- Propone acciones que correspondan exactamente a lo que está presente en los documentos
- Los títulos deben ser cortos y claros (máx 25 caracteres), accionables y directos
- Los sample_prompt deben ser precisos y solicitar información específica encontrable en los documentos
- Prioriza siempre: partes/identidades → fechas → montos/finanzas → obligaciones/cláusulas → otros detalles
- Siempre incluye "Extraer datos estructurados" como última acción

Retorna JSON ESTRICTO con exactamente esta forma:
{{
  "domain": "etiqueta_corta_del_dominio",
  "suggested_actions": [
    {{
      "id": "identificador_legible_por_maquina",
      "title": "Etiqueta corta del botón (máx 25 caracteres)",
      "description": "Una oración que explica lo que hace esta acción para el usuario.",
      "sample_prompt": "Prompt completo, preciso y accionable en lenguaje natural que la app enviará al asistente cuando el usuario haga clic en esta acción. El prompt debe solicitar información específica presente en los documentos."
    }}
  ]
}}
"""
        }
        
        prompt = base_prompts.get(language, base_prompts['en']) + "\n\nCORPUS PREVIEW:\n" + corpus_preview

        # Use Mistral instead of OpenAI for suggested actions generation
        raw_response = call_mistral_api(prompt)

        # Essayer d'extraire un JSON valide
        json_match = re.search(r'\{.*\}', raw_response, re.DOTALL)
        if json_match:
            actions = json.loads(json_match.group())
        else:
            actions = json.loads(raw_response)

        # Validation minimale
        if not isinstance(actions, dict) or "suggested_actions" not in actions:
            raise ValueError("Invalid actions format")

        return actions
    except Exception as e:
        logger.warning(f"⚠️  Unable to infer corpus actions: {e}")
        
        # Détecter le type de document pour le fallback
        try:
            # Utiliser le premier document pour détecter le type
            if documents:
                first_doc = documents[0]
                doc_type = await detect_document_type(
                    first_doc.page_content[:2000] if first_doc.page_content else '',
                    first_doc.metadata.get('fileName', '')
                )
            else:
                doc_type = 'document_generique'
        except:
            doc_type = 'document_generique'
        
        # Fallback avec actions spécifiques selon le type de document et la langue
        fallback_actions = {
            'fr': {
                'contrat_location': [
                    {"id": "verify_parties", "title": "Vérifier les parties", "description": "Identifie toutes les parties (locataire, bailleur) avec leurs coordonnées", "sample_prompt": "Identifie toutes les parties de ce contrat de location : le locataire et le bailleur. Liste leurs noms complets, adresses, téléphones et emails si disponibles."},
                    {"id": "verify_dates", "title": "Vérifier les dates", "description": "Extrait toutes les dates importantes du contrat", "sample_prompt": "Extrais toutes les dates importantes de ce contrat de location : date de signature, date de début, date de fin, durée, préavis."},
                    {"id": "verify_amounts", "title": "Vérifier les montants", "description": "Liste tous les montants : loyer, caution, charges", "sample_prompt": "Liste tous les montants mentionnés dans ce contrat de location : le loyer mensuel, la caution, les charges, et toute indexation prévue."},
                    {"id": "analyze_risky_clauses", "title": "Analyser clauses à risque", "description": "Identifie les clauses potentiellement problématiques", "sample_prompt": "Analyse ce contrat de location et identifie les clauses potentiellement problématiques ou désavantageuses pour le locataire ou le bailleur."},
                    {"id": "verify_obligations", "title": "Vérifier les obligations", "description": "Liste les obligations du locataire et du bailleur", "sample_prompt": "Liste toutes les obligations du locataire et du bailleur mentionnées dans ce contrat de location."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de ce contrat de location : parties, dates, montants, bien loué, clauses importantes."}
                ],
                'contrat_travail': [
                    {"id": "verify_parties", "title": "Vérifier les parties", "description": "Identifie l'employeur et l'employé", "sample_prompt": "Identifie les parties de ce contrat de travail : l'employeur et l'employé. Liste leurs noms complets et coordonnées."},
                    {"id": "verify_dates", "title": "Vérifier les dates", "description": "Extrait les dates importantes du contrat", "sample_prompt": "Extrais toutes les dates importantes de ce contrat de travail : date de signature, date de début, période d'essai, date de fin si applicable."},
                    {"id": "verify_remuneration", "title": "Vérifier rémunération", "description": "Détaille le salaire, primes et avantages", "sample_prompt": "Détaille la rémunération dans ce contrat de travail : salaire de base, primes, avantages, révisions salariales prévues."},
                    {"id": "verify_obligations", "title": "Vérifier les obligations", "description": "Liste les obligations de l'employé et de l'employeur", "sample_prompt": "Liste toutes les obligations de l'employé et de l'employeur mentionnées dans ce contrat de travail."},
                    {"id": "analyze_risky_clauses", "title": "Analyser clauses à risque", "description": "Identifie les clauses restrictives ou problématiques", "sample_prompt": "Analyse ce contrat de travail et identifie les clauses potentiellement restrictives ou problématiques (clause de non-concurrence, clause d'exclusivité, etc.)."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de ce contrat de travail : parties, dates, rémunération, obligations, conditions de travail."}
                ],
                'contrat_vente': [
                    {"id": "verify_parties", "title": "Vérifier les parties", "description": "Identifie l'acheteur et le vendeur", "sample_prompt": "Identifie les parties de ce contrat de vente : l'acheteur et le vendeur. Liste leurs noms complets et coordonnées."},
                    {"id": "verify_dates", "title": "Vérifier les dates", "description": "Extrait les dates importantes", "sample_prompt": "Extrais toutes les dates importantes de ce contrat de vente : date de signature, date de livraison, dates de paiement."},
                    {"id": "verify_amounts", "title": "Vérifier les montants", "description": "Détaille le prix et modalités de paiement", "sample_prompt": "Détaille tous les montants de ce contrat de vente : prix total, acompte, modalités de paiement, échéances."},
                    {"id": "verify_object", "title": "Vérifier l'objet", "description": "Décrit précisément l'objet de la vente", "sample_prompt": "Décris précisément l'objet de cette vente : nature du bien, caractéristiques, quantité, état."},
                    {"id": "verify_guarantees", "title": "Vérifier garanties", "description": "Liste les garanties et conditions", "sample_prompt": "Liste toutes les garanties mentionnées dans ce contrat de vente et leurs conditions."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de ce contrat de vente : parties, dates, montants, objet, garanties, conditions."}
                ],
                'contrat_generique': [
                    {"id": "verify_parties", "title": "Vérifier les parties", "description": "Identifie toutes les parties", "sample_prompt": "Identifie toutes les parties de ce contrat avec leurs noms complets et coordonnées."},
                    {"id": "verify_dates", "title": "Vérifier les dates", "description": "Extrait toutes les dates importantes", "sample_prompt": "Extrais toutes les dates importantes de ce contrat : signature, échéances, dates de paiement."},
                    {"id": "verify_amounts", "title": "Vérifier les montants", "description": "Liste tous les montants et modalités", "sample_prompt": "Liste tous les montants mentionnés dans ce contrat et leurs modalités de paiement."},
                    {"id": "verify_object", "title": "Vérifier l'objet", "description": "Décrit précisément l'objet du contrat", "sample_prompt": "Décris précisément l'objet de ce contrat en une phrase claire."},
                    {"id": "analyze_risky_clauses", "title": "Analyser clauses à risque", "description": "Identifie les clauses problématiques", "sample_prompt": "Analyse ce contrat et identifie les clauses potentiellement problématiques ou désavantageuses."},
                    {"id": "verify_obligations", "title": "Vérifier les obligations", "description": "Liste les obligations de chaque partie", "sample_prompt": "Liste toutes les obligations de chaque partie mentionnées dans ce contrat."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de ce contrat : parties, dates, montants, objet, obligations, clauses importantes."}
                ],
                'testament': [
                    {"id": "verify_testator", "title": "Vérifier le testateur", "description": "Identifie le testateur", "sample_prompt": "Identifie le testateur de ce testament avec ses coordonnées complètes."},
                    {"id": "verify_beneficiaries", "title": "Vérifier bénéficiaires", "description": "Liste tous les bénéficiaires et leurs parts", "sample_prompt": "Liste tous les bénéficiaires de ce testament et leurs parts respectives."},
                    {"id": "verify_dates", "title": "Vérifier les dates", "description": "Extrait les dates importantes", "sample_prompt": "Extrais toutes les dates importantes de ce testament : date de rédaction, signature, modifications éventuelles."},
                    {"id": "verify_legacies", "title": "Vérifier les legs", "description": "Détaille tous les legs et héritages", "sample_prompt": "Détaille tous les legs et héritages mentionnés dans ce testament."},
                    {"id": "verify_executor", "title": "Vérifier l'exécuteur", "description": "Identifie l'exécuteur testamentaire", "sample_prompt": "Identifie l'exécuteur testamentaire mentionné dans ce testament."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de ce testament : testateur, bénéficiaires, legs, dates, exécuteur."}
                ],
                'acte_notarie': [
                    {"id": "verify_parties", "title": "Vérifier les parties", "description": "Identifie toutes les parties", "sample_prompt": "Identifie toutes les parties impliquées dans cet acte notarié avec leurs coordonnées."},
                    {"id": "verify_dates", "title": "Vérifier les dates", "description": "Extrait toutes les dates importantes", "sample_prompt": "Extrais toutes les dates importantes de cet acte notarié."},
                    {"id": "verify_amounts", "title": "Vérifier les montants", "description": "Liste tous les montants et transactions", "sample_prompt": "Liste tous les montants et transactions mentionnés dans cet acte notarié."},
                    {"id": "verify_object", "title": "Vérifier l'objet", "description": "Décrit précisément l'objet de l'acte", "sample_prompt": "Décris précisément l'objet de cet acte notarié."},
                    {"id": "verify_notary", "title": "Vérifier le notaire", "description": "Identifie le notaire et son étude", "sample_prompt": "Identifie le notaire qui a rédigé cet acte et son étude notariale."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de cet acte notarié : parties, dates, montants, objet, notaire."}
                ],
                'lettre': [
                    {"id": "verify_sender", "title": "Vérifier l'expéditeur", "description": "Identifie l'expéditeur de la lettre", "sample_prompt": "Identifie l'expéditeur de cette lettre : nom, fonction, organisation, coordonnées."},
                    {"id": "verify_recipient", "title": "Vérifier le destinataire", "description": "Identifie le destinataire de la lettre", "sample_prompt": "Identifie le destinataire de cette lettre : nom, fonction, organisation, coordonnées."},
                    {"id": "verify_date", "title": "Vérifier la date", "description": "Extrait la date de la lettre", "sample_prompt": "Extrais la date de cette lettre (date d'écriture, date d'envoi si mentionnée)."},
                    {"id": "verify_object", "title": "Vérifier l'objet", "description": "Décrit l'objet et le but de la lettre", "sample_prompt": "Décris l'objet et le but principal de cette lettre en une phrase claire."},
                    {"id": "verify_key_information", "title": "Vérifier infos clés", "description": "Extrait les informations importantes mentionnées", "sample_prompt": "Extrais toutes les informations importantes mentionnées dans cette lettre : montants, dates, références, engagements."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de cette lettre : expéditeur, destinataire, date, objet, informations clés."}
                ],
                'document_financier': [
                    {"id": "verify_parties", "title": "Vérifier les parties", "description": "Identifie les parties concernées", "sample_prompt": "Identifie toutes les parties concernées par ce document financier : employeur, employé, institution, etc."},
                    {"id": "verify_period", "title": "Vérifier la période", "description": "Extrait la période couverte", "sample_prompt": "Extrais la période couverte par ce document financier : dates de début et de fin."},
                    {"id": "verify_amounts", "title": "Vérifier les montants", "description": "Liste tous les montants et totaux", "sample_prompt": "Liste tous les montants mentionnés dans ce document financier : revenus, déductions, impôts, totaux."},
                    {"id": "verify_deductions", "title": "Vérifier déductions", "description": "Détaille toutes les déductions", "sample_prompt": "Détaille toutes les déductions mentionnées dans ce document financier : impôts, cotisations, autres déductions."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de ce document financier : parties, période, montants, déductions, totaux."}
                ],
                'default': [
                    {"id": "summarize_all", "title": "Résumer documents", "description": "Génère un résumé global des documents", "sample_prompt": "Fournis un résumé clair et structuré de tous les documents uploadés, en mettant en évidence les thèmes principaux et les informations importantes. Adapte le résumé au type de document : pour un CV, concentre-toi sur l'expérience professionnelle, les compétences et les réalisations ; pour un document financier, mentionne les montants et chiffres pertinents ; pour un contrat, mentionne les parties et dates importantes. Ne mentionne PAS d'informations qui ne sont pas présentes dans les documents (par exemple, ne mentionne pas d'informations financières si le document est un CV)."},
                    {"id": "extract_key_points", "title": "Extraire points clés", "description": "Liste les points clés et entités", "sample_prompt": "Extrais les points clés, décisions importantes et entités nommées (personnes, entreprises, lieux) de tous les documents uploadés et organise-les en puces."},
                    {"id": "extract_structured", "title": "Extraire données structurées", "description": "Extrait toutes les données dans un format structuré", "sample_prompt": "Extrait toutes les données structurées de ce document : parties, dates, montants, informations clés."}
                ]
            },
            'en': {
                'contrat_location': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies all parties (tenant, landlord) with their contact details", "sample_prompt": "Identify all parties in this rental contract: the tenant and the landlord. List their full names, addresses, phone numbers and emails if available."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts all important dates from the contract", "sample_prompt": "Extract all important dates from this rental contract: signing date, start date, end date, duration, notice period."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists all amounts: rent, deposit, charges", "sample_prompt": "List all amounts mentioned in this rental contract: monthly rent, deposit, charges, and any planned indexation."},
                    {"id": "analyze_risky_clauses", "title": "Analyze risky clauses", "description": "Identifies potentially problematic clauses", "sample_prompt": "Analyze this rental contract and identify potentially problematic or disadvantageous clauses for the tenant or landlord."},
                    {"id": "verify_obligations", "title": "Verify obligations", "description": "Lists obligations of tenant and landlord", "sample_prompt": "List all obligations of the tenant and landlord mentioned in this rental contract."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this rental contract: parties, dates, amounts, rented property, important clauses."}
                ],
                'contrat_travail': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies employer and employee", "sample_prompt": "Identify the parties in this employment contract: the employer and the employee. List their full names and contact details."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts important dates from the contract", "sample_prompt": "Extract all important dates from this employment contract: signing date, start date, probation period, end date if applicable."},
                    {"id": "verify_remuneration", "title": "Verify remuneration", "description": "Details salary, bonuses and benefits", "sample_prompt": "Detail the remuneration in this employment contract: base salary, bonuses, benefits, planned salary revisions."},
                    {"id": "verify_obligations", "title": "Verify obligations", "description": "Lists obligations of employee and employer", "sample_prompt": "List all obligations of the employee and employer mentioned in this employment contract."},
                    {"id": "analyze_risky_clauses", "title": "Analyze risky clauses", "description": "Identifies restrictive or problematic clauses", "sample_prompt": "Analyze this employment contract and identify potentially restrictive or problematic clauses (non-compete clause, exclusivity clause, etc.)."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this employment contract: parties, dates, remuneration, obligations, working conditions."}
                ],
                'contrat_vente': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies buyer and seller", "sample_prompt": "Identify the parties in this sale contract: the buyer and the seller. List their full names and contact details."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts important dates", "sample_prompt": "Extract all important dates from this sale contract: signing date, delivery date, payment dates."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Details price and payment terms", "sample_prompt": "Detail all amounts in this sale contract: total price, down payment, payment terms, deadlines."},
                    {"id": "verify_object", "title": "Verify subject", "description": "Precisely describes the subject of the sale", "sample_prompt": "Precisely describe the subject of this sale: nature of the property, characteristics, quantity, condition."},
                    {"id": "verify_guarantees", "title": "Verify guarantees", "description": "Lists guarantees and conditions", "sample_prompt": "List all guarantees mentioned in this sale contract and their conditions."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this sale contract: parties, dates, amounts, subject, guarantees, conditions."}
                ],
                'contrat_generique': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies all parties", "sample_prompt": "Identify all parties in this contract with their full names and contact details."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts all important dates", "sample_prompt": "Extract all important dates from this contract: signing, deadlines, payment dates."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists all amounts and terms", "sample_prompt": "List all amounts mentioned in this contract and their payment terms."},
                    {"id": "verify_object", "title": "Verify subject", "description": "Precisely describes the subject of the contract", "sample_prompt": "Precisely describe the subject of this contract in a clear sentence."},
                    {"id": "analyze_risky_clauses", "title": "Analyze risky clauses", "description": "Identifies problematic clauses", "sample_prompt": "Analyze this contract and identify potentially problematic or disadvantageous clauses."},
                    {"id": "verify_obligations", "title": "Verify obligations", "description": "Lists obligations of each party", "sample_prompt": "List all obligations of each party mentioned in this contract."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this contract: parties, dates, amounts, subject, obligations, important clauses."}
                ],
                'testament': [
                    {"id": "verify_testator", "title": "Verify testator", "description": "Identifies the testator", "sample_prompt": "Identify the testator of this will with their complete contact details."},
                    {"id": "verify_beneficiaries", "title": "Verify beneficiaries", "description": "Lists all beneficiaries and their shares", "sample_prompt": "List all beneficiaries of this will and their respective shares."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts important dates", "sample_prompt": "Extract all important dates from this will: drafting date, signature, any modifications."},
                    {"id": "verify_legacies", "title": "Verify bequests", "description": "Details all bequests and inheritances", "sample_prompt": "Detail all bequests and inheritances mentioned in this will."},
                    {"id": "verify_executor", "title": "Verify executor", "description": "Identifies the executor", "sample_prompt": "Identify the executor mentioned in this will."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this will: testator, beneficiaries, bequests, dates, executor."}
                ],
                'acte_notarie': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies all parties", "sample_prompt": "Identify all parties involved in this notarial act with their contact details."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts all important dates", "sample_prompt": "Extract all important dates from this notarial act."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists all amounts and transactions", "sample_prompt": "List all amounts and transactions mentioned in this notarial act."},
                    {"id": "verify_object", "title": "Verify subject", "description": "Precisely describes the subject of the act", "sample_prompt": "Precisely describe the subject of this notarial act."},
                    {"id": "verify_notary", "title": "Verify notary", "description": "Identifies the notary and their office", "sample_prompt": "Identify the notary who drafted this act and their notarial office."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this notarial act: parties, dates, amounts, subject, notary."}
                ],
                'lettre': [
                    {"id": "verify_sender", "title": "Verify sender", "description": "Identifies the letter sender", "sample_prompt": "Identify the sender of this letter: name, function, organization, contact details."},
                    {"id": "verify_recipient", "title": "Verify recipient", "description": "Identifies the letter recipient", "sample_prompt": "Identify the recipient of this letter: name, function, organization, contact details."},
                    {"id": "verify_date", "title": "Verify date", "description": "Extracts the letter date", "sample_prompt": "Extract the date of this letter (writing date, sending date if mentioned)."},
                    {"id": "verify_object", "title": "Verify subject", "description": "Describes the subject and purpose of the letter", "sample_prompt": "Describe the subject and main purpose of this letter in a clear sentence."},
                    {"id": "verify_key_information", "title": "Verify key info", "description": "Extracts important information mentioned", "sample_prompt": "Extract all important information mentioned in this letter: amounts, dates, references, commitments."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this letter: sender, recipient, date, subject, key information."}
                ],
                'document_financier': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies concerned parties", "sample_prompt": "Identify all parties concerned by this financial document: employer, employee, institution, etc."},
                    {"id": "verify_period", "title": "Verify period", "description": "Extracts the covered period", "sample_prompt": "Extract the period covered by this financial document: start and end dates."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists all amounts and totals", "sample_prompt": "List all amounts mentioned in this financial document: income, deductions, taxes, totals."},
                    {"id": "verify_deductions", "title": "Verify deductions", "description": "Details all deductions", "sample_prompt": "Detail all deductions mentioned in this financial document: taxes, contributions, other deductions."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this financial document: parties, period, amounts, deductions, totals."}
                ],
                'default': [
                    {"id": "summarize_all", "title": "Summarize documents", "description": "Generates a global summary of documents", "sample_prompt": "Provide a clear and structured summary of all uploaded documents, highlighting the main themes and important information. Adapt the summary to the document type: for a CV, focus on professional experience, skills and achievements; for a financial document, mention relevant amounts and figures; for a contract, mention important parties and dates. Do NOT mention information that is not present in the documents (for example, do not mention financial information if the document is a CV)."},
                    {"id": "extract_key_points", "title": "Extract key points", "description": "Lists key points and entities", "sample_prompt": "Extract key points, important decisions and named entities (people, companies, places) from all uploaded documents and organize them in bullet points."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this document: parties, dates, amounts, key information."}
                ]
            },
            'es': {
                'contrat_location': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica todas las partes (inquilino, arrendador) con sus datos de contacto", "sample_prompt": "Identifica todas las partes de este contrato de alquiler: el inquilino y el arrendador. Enumera sus nombres completos, direcciones, teléfonos y correos electrónicos si están disponibles."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae todas las fechas importantes del contrato", "sample_prompt": "Extrae todas las fechas importantes de este contrato de alquiler: fecha de firma, fecha de inicio, fecha de fin, duración, preaviso."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera todos los montos: alquiler, depósito, gastos", "sample_prompt": "Enumera todos los montos mencionados en este contrato de alquiler: el alquiler mensual, el depósito, los gastos y cualquier indexación prevista."},
                    {"id": "analyze_risky_clauses", "title": "Analizar cláusulas de riesgo", "description": "Identifica cláusulas potencialmente problemáticas", "sample_prompt": "Analiza este contrato de alquiler e identifica las cláusulas potencialmente problemáticas o desventajosas para el inquilino o el arrendador."},
                    {"id": "verify_obligations", "title": "Verificar obligaciones", "description": "Enumera las obligaciones del inquilino y del arrendador", "sample_prompt": "Enumera todas las obligaciones del inquilino y del arrendador mencionadas en este contrato de alquiler."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este contrato de alquiler: partes, fechas, montos, propiedad alquilada, cláusulas importantes."}
                ],
                'contrat_travail': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica al empleador y al empleado", "sample_prompt": "Identifica las partes de este contrato de trabajo: el empleador y el empleado. Enumera sus nombres completos y datos de contacto."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae las fechas importantes del contrato", "sample_prompt": "Extrae todas las fechas importantes de este contrato de trabajo: fecha de firma, fecha de inicio, período de prueba, fecha de fin si es aplicable."},
                    {"id": "verify_remuneration", "title": "Verificar remuneración", "description": "Detalla el salario, bonos y beneficios", "sample_prompt": "Detalla la remuneración en este contrato de trabajo: salario base, bonos, beneficios, revisiones salariales previstas."},
                    {"id": "verify_obligations", "title": "Verificar obligaciones", "description": "Enumera las obligaciones del empleado y del empleador", "sample_prompt": "Enumera todas las obligaciones del empleado y del empleador mencionadas en este contrato de trabajo."},
                    {"id": "analyze_risky_clauses", "title": "Analizar cláusulas de riesgo", "description": "Identifica cláusulas restrictivas o problemáticas", "sample_prompt": "Analiza este contrato de trabajo e identifica las cláusulas potencialmente restrictivas o problemáticas (cláusula de no competencia, cláusula de exclusividad, etc.)."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este contrato de trabajo: partes, fechas, remuneración, obligaciones, condiciones de trabajo."}
                ],
                'contrat_vente': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica al comprador y al vendedor", "sample_prompt": "Identifica las partes de este contrato de venta: el comprador y el vendedor. Enumera sus nombres completos y datos de contacto."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae las fechas importantes", "sample_prompt": "Extrae todas las fechas importantes de este contrato de venta: fecha de firma, fecha de entrega, fechas de pago."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Detalla el precio y modalidades de pago", "sample_prompt": "Detalla todos los montos de este contrato de venta: precio total, anticipo, modalidades de pago, plazos."},
                    {"id": "verify_object", "title": "Verificar objeto", "description": "Describe precisamente el objeto de la venta", "sample_prompt": "Describe precisamente el objeto de esta venta: naturaleza del bien, características, cantidad, estado."},
                    {"id": "verify_guarantees", "title": "Verificar garantías", "description": "Enumera las garantías y condiciones", "sample_prompt": "Enumera todas las garantías mencionadas en este contrato de venta y sus condiciones."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este contrato de venta: partes, fechas, montos, objeto, garantías, condiciones."}
                ],
                'contrat_generique': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica todas las partes", "sample_prompt": "Identifica todas las partes de este contrato con sus nombres completos y datos de contacto."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae todas las fechas importantes", "sample_prompt": "Extrae todas las fechas importantes de este contrato: firma, plazos, fechas de pago."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera todos los montos y modalidades", "sample_prompt": "Enumera todos los montos mencionados en este contrato y sus modalidades de pago."},
                    {"id": "verify_object", "title": "Verificar objeto", "description": "Describe precisamente el objeto del contrato", "sample_prompt": "Describe precisamente el objeto de este contrato en una oración clara."},
                    {"id": "analyze_risky_clauses", "title": "Analizar cláusulas de riesgo", "description": "Identifica cláusulas problemáticas", "sample_prompt": "Analiza este contrato e identifica las cláusulas potencialmente problemáticas o desventajosas."},
                    {"id": "verify_obligations", "title": "Verificar obligaciones", "description": "Enumera las obligaciones de cada parte", "sample_prompt": "Enumera todas las obligaciones de cada parte mencionadas en este contrato."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurudos de este contrato: partes, fechas, montos, objeto, obligaciones, cláusulas importantes."}
                ],
                'testament': [
                    {"id": "verify_testator", "title": "Verificar testador", "description": "Identifica al testador", "sample_prompt": "Identifica al testador de este testamento con sus datos de contacto completos."},
                    {"id": "verify_beneficiaries", "title": "Verificar beneficiarios", "description": "Enumera todos los beneficiarios y sus partes", "sample_prompt": "Enumera todos los beneficiarios de este testamento y sus partes respectivas."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae las fechas importantes", "sample_prompt": "Extrae todas las fechas importantes de este testamento: fecha de redacción, firma, modificaciones eventuales."},
                    {"id": "verify_legacies", "title": "Verificar legados", "description": "Detalla todos los legados y herencias", "sample_prompt": "Detalla todos los legados y herencias mencionados en este testamento."},
                    {"id": "verify_executor", "title": "Verificar ejecutor", "description": "Identifica al albacea", "sample_prompt": "Identifica al albacea mencionado en este testamento."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este testamento: testador, beneficiarios, legados, fechas, ejecutor."}
                ],
                'acte_notarie': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica todas las partes", "sample_prompt": "Identifica todas las partes involucradas en este acta notarial con sus datos de contacto."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae todas las fechas importantes", "sample_prompt": "Extrae todas las fechas importantes de este acta notarial."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera todos los montos y transacciones", "sample_prompt": "Enumera todos los montos y transacciones mencionados en este acta notarial."},
                    {"id": "verify_object", "title": "Verificar objeto", "description": "Describe precisamente el objeto del acta", "sample_prompt": "Describe precisamente el objeto de este acta notarial."},
                    {"id": "verify_notary", "title": "Verificar notario", "description": "Identifica al notario y su estudio", "sample_prompt": "Identifica al notario que redactó este acta y su estudio notarial."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este acta notarial: partes, fechas, montos, objeto, notario."}
                ],
                'lettre': [
                    {"id": "verify_sender", "title": "Verificar remitente", "description": "Identifica al remitente de la carta", "sample_prompt": "Identifica al remitente de esta carta: nombre, función, organización, datos de contacto."},
                    {"id": "verify_recipient", "title": "Verificar destinatario", "description": "Identifica al destinatario de la carta", "sample_prompt": "Identifica al destinatario de esta carta: nombre, función, organización, datos de contacto."},
                    {"id": "verify_date", "title": "Verificar fecha", "description": "Extrae la fecha de la carta", "sample_prompt": "Extrae la fecha de esta carta (fecha de escritura, fecha de envío si se menciona)."},
                    {"id": "verify_object", "title": "Verificar objeto", "description": "Describe el objeto y el propósito de la carta", "sample_prompt": "Describe el objeto y el propósito principal de esta carta en una oración clara."},
                    {"id": "verify_key_information", "title": "Verificar info clave", "description": "Extrae información importante mencionada", "sample_prompt": "Extrae toda la información importante mencionada en esta carta: montos, fechas, referencias, compromisos."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de esta carta: remitente, destinatario, fecha, objeto, información clave."}
                ],
                'document_financier': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica las partes concernidas", "sample_prompt": "Identifica todas las partes concernidas por este documento financiero: empleador, empleado, institución, etc."},
                    {"id": "verify_period", "title": "Verificar período", "description": "Extrae el período cubierto", "sample_prompt": "Extrae el período cubierto por este documento financiero: fechas de inicio y de fin."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera todos los montos y totales", "sample_prompt": "Enumera todos los montos mencionados en este documento financiero: ingresos, deducciones, impuestos, totales."},
                    {"id": "verify_deductions", "title": "Verificar deducciones", "description": "Detalla todas las deducciones", "sample_prompt": "Detalla todas las deducciones mencionadas en este documento financiero: impuestos, cotizaciones, otras deducciones."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este documento financiero: partes, período, montos, deducciones, totales."}
                ],
                'default': [
                    {"id": "summarize_all", "title": "Resumir documentos", "description": "Genera un resumen global de los documentos", "sample_prompt": "Proporciona un resumen claro y estructurado de todos los documentos cargados, destacando los temas principales y la información importante. Adapta el resumen al tipo de documento: para un CV, concéntrate en la experiencia profesional, habilidades y logros; para un documento financiero, menciona montos y cifras relevantes; para un contrato, menciona partes y fechas importantes. NO menciones información que no esté presente en los documentos (por ejemplo, no menciones información financiera si el documento es un CV)."},
                    {"id": "extract_key_points", "title": "Extraer puntos clave", "description": "Enumera los puntos clave y entidades", "sample_prompt": "Extrae los puntos clave, decisiones importantes y entidades nombradas (personas, empresas, lugares) de todos los documentos cargados y organízalos en viñetas."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este documento: partes, fechas, montos, información clave."}
                ]
            }
        }
        
        # Sélectionner les actions selon la langue et le type de document
        lang_actions = fallback_actions.get(language, fallback_actions['en'])
        actions_list = lang_actions.get(doc_type, lang_actions.get('default', []))
        
        return {
            "domain": doc_type,
            "suggested_actions": actions_list
        }

# save_file_description est maintenant importé depuis services.analysis_service
async def _OLD_save_file_description(file_name, description):
    with open(DESCRIPTIONS_FILE, 'r') as f:
        descriptions = json.load(f)
    descriptions.append({"file_name": file_name, "description": description})
    with open(DESCRIPTIONS_FILE, 'w') as f:
        json.dump(descriptions, f, indent=4)

async def query_model(file_name, file_content, directory_content, repo_structure, user_query, is_binary=False, selected_model=DEFAULT_MODEL, language='en'):
    cache_key = f"query_{file_name}_{user_query}_{selected_model}_{language}"
    cached = await cache.get(cache_key)
    if cached:
        return cached

    # fallback to English if unknown language
    t = translations.get(language, translations['en'])

    # Construire le résumé des autres fichiers
    directory_content_summary = ' '.join(
        [f"{t['other_file']}: {doc['fileName']} : {doc['content']}" for doc in directory_content]
    )

    # Construire le prompt multilangue
    prompt = (
        f"{t['project_structure']}:\n{repo_structure}\n\n"
        f"{t['file']}: {file_name}\n\n"
        f"{t['content']}:\n{file_content}\n\n"
        f"{directory_content_summary}\n\n"
        f"{t['question']}: {user_query}\n\n"
        f"{t['emoji']}"
    )

    try:
        # Enhanced model routing with GPT-5 variants
        if selected_model.lower() in ["gpt-3.5-turbo", "gpt-4o"]:
            result = call_openai_api(prompt, selected_model)
        elif selected_model.lower() in ["gpt-5", "gpt-5-mini", "gpt-5-nano"]:
            result = call_openai_api(prompt, selected_model)
        elif selected_model.lower() == "openai":
            # Legacy compatibility
            result = call_openai_api(prompt, "openai")
        elif selected_model.lower() == "mistral":
            result = call_mistral_api(prompt)
        elif selected_model.lower().startswith("gemini") or selected_model.lower() in ["gemini-3-flash", "gemini-pro"]:
            result = call_gemini_api(prompt, selected_model)
        elif selected_model.lower() in OLLAMA_MODELS or selected_model.lower() == "llama3":
            result = call_ollama_api(prompt, selected_model)
        else:
            logger.warning(f"Unknown model {selected_model}, falling back to default {DEFAULT_MODEL}")
            result = call_openai_api(prompt, DEFAULT_MODEL)
            
    except Exception as e:
        result = f"Error querying model: {str(e)}"

    await cache.set(cache_key, result, ttl=3600)
    return result

# Validation de la configuration
try:
    FlaskConfig.validate_env_vars()
    logger.info("✅ Configuration validée")
except ValueError as e:
    logger.error(f"❌ Erreur de configuration: {e}")

# Initialisation de la base de données d'authentification
init_database()

# Enregistrement des routes d'authentification
register_auth_routes(app)

# New endpoint to get available models with detailed info
@app.route('/models', methods=['GET'])
def get_available_models():
    """
    Returns a list of available models with detailed configuration
    """
    models = []
    for model_id, config in MODEL_CONFIG.items():
        model_info = {
            "id": model_id,
            "name": config["name"],
            "provider": config["provider"],
            "description": config["description"],
            "cost_tier": config["cost_tier"],
            "max_tokens": config["max_tokens"],
            "is_default": model_id == DEFAULT_MODEL
        }
        models.append(model_info)
    
    return jsonify({
        "models": models,
        "default_model": DEFAULT_MODEL,
        "total_count": len(models)
    }), 200

@app.route('/models/<model_id>/test', methods=['POST'])
def test_model(model_id):
    """
    Test if a specific model is available and working
    """
    if model_id not in MODEL_CONFIG:
        return jsonify({"error": "Model not found"}), 404
    
    test_prompt = "Hello, please respond with 'Model is working correctly.'"
    
    try:
        if model_id in ["gpt-3.5-turbo", "gpt-4o", "gpt-5", "gpt-5-mini", "gpt-5-nano", "openai"]:
            response = call_openai_api(test_prompt, model_id)
        elif model_id == "mistral":
            response = call_mistral_api(test_prompt)
        elif model_id.startswith("gemini") or model_id in ["gemini-3-flash", "gemini-pro"]:
            response = call_gemini_api(test_prompt, model_id)
        elif model_id in OLLAMA_MODELS or model_id == "llama3":
            response = call_ollama_api(test_prompt, model_id)
        else:
            return jsonify({"error": "Model not supported"}), 400
            
        return jsonify({
            "model_id": model_id,
            "status": "available",
            "test_response": response,
            "config": MODEL_CONFIG[model_id]
        }), 200
        
    except Exception as e:
        logger.error(f"Model test failed for {model_id}: {str(e)}")
        return jsonify({
            "model_id": model_id,
            "status": "unavailable",
            "error": str(e),
            "config": MODEL_CONFIG[model_id]
        }), 503

# Enhanced Flask Routes
@app.route('/upload', methods=['POST'])
async def upload_files():
    if 'files' not in request.files:
        return jsonify({"error": "No files provided"}), 400

    files = request.files.getlist('files')
    selected_model = request.form.get('model', DEFAULT_MODEL)  # Use GPT-3.5-Turbo as default
    language = request.form.get('language', 'en')
    results, texts = [], []

    # Validate model
    if selected_model not in MODEL_CONFIG:
        logger.warning(f"Unknown model {selected_model}, using default {DEFAULT_MODEL}")
        selected_model = DEFAULT_MODEL

    for file in files:
        file_name = file.filename
        extension = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''
        is_binary = extension in ['docx', 'pdf']

        file_content = (extract_text_from_docx(file) if extension == 'docx' else
                        extract_text_from_pdf(file) if extension == 'pdf' else
                        file.read().decode('utf-8', errors='ignore'))

        description = await analyze_file_content(file_content, file_name, is_binary, extension, selected_model, language)
        await save_file_description(file_name, description)
        if not description.startswith("Error"):
            texts.append(description)
        results.append({
            "file_name": file_name, 
            "description": description,
            "model_used": selected_model
        })

    if texts:
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        docs = [doc for text in texts for doc in splitter.create_documents([text])]
        global vector_store
        if vector_store is None:
            vector_store = FAISS.from_documents(docs, embeddings)
        else:
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, vector_store.add_documents, docs)

    return jsonify({
        "message": "Files processed successfully", 
        "results": results,
        "model_used": selected_model,
        "model_config": MODEL_CONFIG.get(selected_model, {})
    }), 200

# perform_online_search et search_serpapi sont maintenant importés depuis services.search_service
def _OLD_perform_online_search(query: str, language: str = 'en') -> str:
    return _OLD_search_serpapi(query, language)

def _OLD_search_serpapi(query: str, language: str = 'en') -> str:
    """
    SerpAPI - 100 recherches gratuites/mois
    Inscription: https://serpapi.com/
    """
    api_key = os.getenv("SERPAPI_KEY")
    if not api_key:
        return "Clé API manquante"
    
    try:
        url = "https://serpapi.com/search"
        params = {
            'api_key': api_key,
            'engine': 'google',
            'q': query,
            'hl': language,
            'num': 5,
            'no_cache': 'true'  # Force fresh results
        }
        
        response = requests.get(url, params=params, timeout=15)
        response.raise_for_status()
        data = response.json()
        
        results = []
        
        # Réponse directe (featured snippet)
        if data.get('answer_box'):
            answer = data['answer_box']
            if answer.get('answer'):
                results.append(f"**Réponse directe**: {answer['answer']}")
            if answer.get('snippet'):
                results.append(f"**Information**: {answer['snippet']}")
        
        # Résultats organiques
        if data.get('organic_results'):
            for result in data['organic_results'][:5]:
                title = result.get('title', '')
                snippet = result.get('snippet', '')
                link = result.get('link', '')
                
                if title and snippet:
                    results.append(f"**{title}**\n{snippet}\nSource: {link}")
        
        return "\n\n".join(results) if results else "Aucun résultat trouvé."
        
    except Exception as e:
        raise Exception(f"Erreur SerpAPI: {str(e)}")

# rerank_documents_with_llm est maintenant importé depuis services.search_service
async def _OLD_rerank_documents_with_llm(query: str, documents: List[Document], model: str = DEFAULT_MODEL) -> List[Document]:
    """
    Re-rank les documents avec un LLM pour améliorer la pertinence sémantique.
    Le LLM évalue chaque document par rapport à la requête et les trie par pertinence.
    """
    try:
        if len(documents) <= 5:
            return documents  # Pas besoin de re-ranking pour peu de documents
        
        # Construire le prompt de re-ranking
        docs_summary = []
        for i, doc in enumerate(documents):
            file_name = doc.metadata.get("fileName") or doc.metadata.get("file_name") or f"document_{i}"
            content_snippet = doc.page_content[:500]  # Premier 500 chars pour le prompt
            docs_summary.append(f"[{i}] Fichier: {file_name}\nContenu: {content_snippet}...")
        
        rerank_prompt = f"""Tu dois classer ces documents par ordre de pertinence pour répondre à cette question: "{query}"

Documents à classer:
{chr(10).join(docs_summary)}

Retourne UNIQUEMENT une liste JSON des indices (nombres entre crochets) dans l'ordre de pertinence décroissante.
Format: [3, 1, 5, 2, 0, 4, ...]

Réponds UNIQUEMENT avec le JSON, sans explication."""
        
        # Appel au modèle pour re-ranking
        rerank_response = call_openai_api(rerank_prompt, model)
        
        # Extraire les indices du JSON
        import json
        json_match = re.search(r'\[.*?\]', rerank_response)
        if json_match:
            ranked_indices = json.loads(json_match.group())
            # Vérifier que tous les indices sont valides
            valid_indices = [idx for idx in ranked_indices if 0 <= idx < len(documents)]
            if len(valid_indices) == len(documents):
                reranked = [documents[idx] for idx in valid_indices]
                logger.info(f"✅ Re-ranking réussi: {len(reranked)} documents reclassés")
                return reranked
        
        # Fallback: retourner l'ordre original si le parsing échoue
        logger.warning("Re-ranking JSON invalide, utilisation de l'ordre original")
        return documents
        
    except Exception as e:
        logger.warning(f"Erreur lors du re-ranking: {str(e)}, utilisation de l'ordre original")
        return documents

def search_semantic_documents_sync(vector_store, user_query: str, session_id: str = 'default', conversation_history: Optional[List[Dict]] = None) -> List[Dict]:
    """
    Fonction helper synchrone pour rechercher des documents pertinents dans le vector store.
    Retourne une liste de dictionnaires avec 'fileName' et 'content'.
    AMÉLIORATION: Détecte les noms de personnes et filtre strictement par nom.
    Utilise l'historique de conversation pour détecter les pronoms (he, she, his, her).
    """
    try:
        # Enterprise-style hybrid retrieval (deterministic): semantic + lexical reranking.
        # This is more precise than raw similarity_search, and avoids LLM reranking drift.
        result = hybrid_retrieve_documents(
            vector_store=vector_store,
            query=user_query,
            k_candidates=100,  # Increased for better recall
            k_final=25,  # Increased from 20 to 25 for more comprehensive results
            semantic_weight=0.55,  # Adjusted weights for better precision
            bm25_weight=0.35,
            exact_weight=0.10,
        )
        if isinstance(result, tuple) and len(result) == 2:
            docs, debug = result
            if not isinstance(debug, dict):
                debug = {}
        else:
            docs = []
            debug = {}
        if docs:
            logger.info(f"🔎 search_semantic_documents_sync hybrid selected {len(docs)} docs | top={debug.get('top', []) if isinstance(debug, dict) else []}")
            out: List[Dict] = []
            seen = set()
            for d in docs:
                fn = d.metadata.get("fileName") or d.metadata.get("file_name") or d.metadata.get("source") or "document_vectorstore"
                key = (str(fn).lower(), (d.page_content or "")[:120])
                if key in seen:
                    continue
                seen.add(key)
                
                # Extraire les informations de page depuis les métadonnées
                page_number = d.metadata.get("page_number")
                is_page_chunk = d.metadata.get("is_page_chunk", False)
                
                result_dict = {
                    "fileName": fn,
                    "content": (d.page_content or "")[:2500]
                }
                
                # Ajouter les informations de page si disponibles
                if page_number is not None:
                    result_dict["pageNumber"] = page_number
                    result_dict["isPageChunk"] = is_page_chunk
                
                # Si c'est un chunk de page, extraire le numéro de page du contenu ou des métadonnées
                if is_page_chunk and page_number is None:
                    # Essayer d'extraire de la métadonnée ou du contenu
                    content_preview = d.page_content or ""
                    if "[Page" in content_preview and "de" in content_preview:
                        # Extraire le numéro de page du format "[Page X de filename]"
                        import re
                        match = re.search(r'\[Page\s+(\d+)', content_preview)
                        if match:
                            result_dict["pageNumber"] = int(match.group(1))
                            result_dict["isPageChunk"] = True
                
                out.append(result_dict)
            return out

        # DÉTECTION DES NOMS DE PERSONNES dans la requête ET l'historique
        person_names = set()
        words = user_query.split()
        
        # Détecter les pronoms et chercher le nom dans l'historique
        query_lower = user_query.lower()
        has_pronoun = any(pronoun in query_lower for pronoun in ['he', 'she', 'his', 'her', 'him', 'they', 'their'])
        
        if has_pronoun and conversation_history:
            # Chercher le dernier nom mentionné dans l'historique
            for turn in reversed(conversation_history):
                content = (turn.get("content") or "").lower()
                # Chercher les noms dans l'historique
                for name in ['karim', 'dominique', 'essome', 'ngami']:
                    if name in content:
                        person_names.add(name)
                        logger.info(f"👤 Nom détecté depuis l'historique (pronoun détecté): {name}")
                        break
                if person_names:
                    break
        
        # Détecter les noms dans la requête actuelle
        for i, word in enumerate(words):
            clean_word = word.strip('.,!?;:()[]{}"\'').strip()
            # Détecter les noms propres (capitalisés) ou les mots après "his", "her", "their", etc.
            if clean_word and clean_word[0].isupper() and len(clean_word) > 2:
                person_names.add(clean_word.lower())
            # Détecter après des mots indicateurs
            if i > 0 and words[i-1].lower() in ['his', 'her', 'their', 'karim', 'dominique', 'about', 'for']:
                if len(clean_word) > 2:
                    person_names.add(clean_word.lower())
        
        # Ajouter aussi les noms communs dans les requêtes (karim, dominique, etc.)
        common_names = ['karim', 'dominique', 'essome', 'ngami']
        for name in common_names:
            if name in query_lower:
                person_names.add(name)
        
        logger.info(f"👤 Noms de personnes détectés (requête + historique): {person_names}")
        
        # Recherche sémantique
        # AMÉLIORATION: Recherche avec plus de résultats et recherche multiple pour meilleure précision
        # Recherche principale avec la requête complète
        search_results_with_scores = vector_store.similarity_search_with_score(user_query, k=50)
        
        # Recherche supplémentaire avec les mots-clés individuels pour capturer les noms de fonctions/méthodes
        query_words = [word.strip('.,!?;:()[]{}"\'').lower() for word in user_query.split() if len(word.strip('.,!?;:()[]{}"\'')) > 2]
        additional_results = []
        for word in query_words[:5]:  # Limiter à 5 mots pour éviter trop de résultats
            try:
                word_results = vector_store.similarity_search_with_score(word, k=10)
                additional_results.extend(word_results)
            except:
                pass
        
        # Combiner et dédupliquer les résultats
        all_search_results = {}
        for doc, score in search_results_with_scores:
            doc_id = id(doc)
            if doc_id not in all_search_results or score < all_search_results[doc_id][1]:
                all_search_results[doc_id] = (doc, score)
        
        for doc, score in additional_results:
            doc_id = id(doc)
            if doc_id not in all_search_results or score < all_search_results[doc_id][1]:
                all_search_results[doc_id] = (doc, score)
        
        search_results_with_scores = list(all_search_results.values())
        
        # Extraire les mots-clés (y compris noms propres courts)
        query_keywords = set()
        for word in user_query.split():
            clean_word = word.strip('.,!?;:()[]{}"\'').lower()
            if len(clean_word) > 2:
                query_keywords.add(clean_word)
        
        logger.info(f"🔑 Mots-clés extraits: {query_keywords}")
        
        # Recherche par nom de fichier - Essayer de récupérer tous les documents
        filename_matches = []
        try:
            # Méthode 1: Recherche avec une requête très large pour récupérer beaucoup de documents
            all_docs_from_store = vector_store.similarity_search("document file content", k=1000)
            
            # Méthode 2: Si on peut accéder au docstore directement, l'utiliser
            if hasattr(vector_store, 'docstore') and hasattr(vector_store.docstore, '_dict'):
                all_docs_dict = vector_store.docstore._dict
                logger.info(f"📦 Accès direct au docstore: {len(all_docs_dict)} documents disponibles")
                # Convertir les valeurs du dict en documents
                all_docs_from_store = list(all_docs_dict.values()) if all_docs_dict else all_docs_from_store
            
            logger.info(f"🔍 Recherche dans {len(all_docs_from_store)} documents pour correspondance de nom")
            
            for doc in all_docs_from_store:
                file_name_from_meta = (doc.metadata.get("fileName") or doc.metadata.get("file_name") or "").lower()
                for keyword in query_keywords:
                    if keyword in file_name_from_meta:
                        filename_matches.append(doc)
                        logger.info(f"📁 Fichier trouvé par nom: {doc.metadata.get('fileName', 'N/A')} (mot-clé: '{keyword}')")
                        break
        except Exception as e:
            logger.warning(f"Erreur lors de la récupération de tous les documents: {str(e)}")
        
        logger.info(f"📂 {len(filename_matches)} fichiers trouvés par correspondance de nom")
        
        # AMÉLIORATION: Recherche explicite par nom de fonction/méthode dans le contenu
        function_name_matches = []
        # Extraire les noms potentiels de fonctions/méthodes de la requête (mots avec underscore ou camelCase)
        potential_function_names = []
        query_words_for_func = user_query.split()
        for word in query_words_for_func:
            clean_word = word.strip('.,!?;:()[]{}"\'').strip()
            # Détecter les noms de fonctions (avec underscore ou camelCase)
            if '_' in clean_word or (clean_word and clean_word[0].islower() and any(c.isupper() for c in clean_word[1:])):
                potential_function_names.append(clean_word.lower())
                potential_function_names.append(clean_word)  # Garder aussi la version originale
        
        # Rechercher explicitement ces noms dans le contenu
        if potential_function_names:
            try:
                all_docs_for_function_search = vector_store.similarity_search("function method def class", k=500)
                for doc in all_docs_for_function_search:
                    content_lower = doc.page_content.lower()
                    for func_name in potential_function_names:
                        if func_name.lower() in content_lower:
                            function_name_matches.append(doc)
                            logger.info(f"🔧 Fonction/méthode trouvée par nom: '{func_name}' dans {doc.metadata.get('fileName', 'N/A')}")
                            break
            except Exception as e:
                logger.warning(f"Erreur lors de la recherche par nom de fonction: {str(e)}")
        
        logger.info(f"🔧 {len(function_name_matches)} documents trouvés par correspondance de nom de fonction/méthode")
        
        # Combiner les résultats avec priorité aux correspondances exactes
        all_candidate_docs = {}
        # Priorité 1: Correspondances de noms de fonctions (score très élevé)
        for doc in function_name_matches:
            all_candidate_docs[id(doc)] = (doc, 0.95)  # Score très élevé pour correspondances exactes
        
        # Priorité 2: Correspondances de noms de fichiers (score élevé)
        for doc in filename_matches:
            doc_id = id(doc)
            if doc_id not in all_candidate_docs:
                all_candidate_docs[doc_id] = (doc, 0.85)
        
        # Priorité 3: Résultats de recherche sémantique (score normalisé)
        for doc, score in search_results_with_scores:
            doc_id = id(doc)
            # Si déjà présent avec un meilleur score, garder le meilleur
            if doc_id not in all_candidate_docs:
                all_candidate_docs[doc_id] = (doc, 1 - score)
            elif all_candidate_docs[doc_id][1] < (1 - score):
                # Si le nouveau score est meilleur, le remplacer
                all_candidate_docs[doc_id] = (doc, 1 - score)
        
        # FILTRAGE STRICT PAR NOM DE PERSONNE (CRITIQUE)
        # Si un nom de personne est détecté, ne garder QUE les fichiers correspondant à ce nom
        query_words = [w.strip('.,!?;:()[]{}"\'').lower() for w in user_query.split() if len(w.strip('.,!?;:()[]{}"\'')) > 2]
        filename_match_ids = {id(doc) for doc in filename_matches}
        filtered_docs = []
        
        # Liste des noms de personnes connus pour exclusion
        known_person_names = {
            'karim': ['dominique', 'essome'],
            'dominique': ['karim', 'ngami'],
            'essome': ['karim', 'ngami'],
            'ngami': ['dominique', 'essome']
        }
        
        # Déterminer les noms à exclure
        names_to_exclude = set()
        for detected_name in person_names:
            if detected_name in known_person_names:
                names_to_exclude.update(known_person_names[detected_name])
        
        logger.info(f"🚫 Noms à exclure (autres personnes): {names_to_exclude}")
        
        for doc, similarity_score in all_candidate_docs.values():
            doc_id = id(doc)
            file_name_lower = (doc.metadata.get("fileName") or doc.metadata.get("file_name") or "").lower()
            doc_content_lower = doc.page_content[:500].lower()  # Vérifier aussi dans le contenu
            
            # EXCLUSION STRICTE: Si un nom de personne est détecté, exclure les fichiers d'autres personnes
            should_exclude = False
            if person_names:
                # RÈGLE 1: Exclure TOUJOURS les fichiers contenant un nom à exclure (même avec score élevé)
                for exclude_name in names_to_exclude:
                    if exclude_name in file_name_lower:
                        should_exclude = True
                        logger.info(f"❌ Fichier EXCLU (nom de fichier contient '{exclude_name}'): {doc.metadata.get('fileName', 'N/A')}")
                        break
                    # Vérifier aussi dans le contenu (premiers 200 caractères pour être plus strict)
                    if exclude_name in doc_content_lower[:200]:
                        should_exclude = True
                        logger.info(f"❌ Fichier EXCLU (contenu contient '{exclude_name}'): {doc.metadata.get('fileName', 'N/A')}")
                        break
                
                # RÈGLE 2: Si un nom de personne est détecté, ne garder QUE les fichiers qui contiennent ce nom
                if not should_exclude:
                    file_contains_person_name = any(name in file_name_lower for name in person_names)
                    # Vérifier aussi dans le contenu si pas dans le nom
                    if not file_contains_person_name:
                        file_contains_person_name = any(name in doc_content_lower[:300] for name in person_names)
                    
                    if not file_contains_person_name:
                        # Si aucun nom de personne n'est dans le fichier, EXCLURE TOUJOURS (même avec score élevé)
                        should_exclude = True
                        logger.info(f"⚠️ Fichier EXCLU (ne contient AUCUN nom de la personne recherchée): {doc.metadata.get('fileName', 'N/A')}")
            
            if should_exclude:
                continue  # Exclure ce document
            
            filename_keyword_matches = sum(1 for word in query_words if word in file_name_lower)
            
            # Score élevé pour les fichiers trouvés par nom de personne
            if doc_id in filename_match_ids:
                # Bonus supplémentaire si le fichier contient le nom de la personne détectée
                person_name_bonus = 0.0
                if person_names:
                    for name in person_names:
                        if name in file_name_lower:
                            person_name_bonus = 0.15
                            break
                filtered_docs.append((doc, 0.9 + person_name_bonus))
                logger.info(f"✅ Fichier inclus (correspond au nom): {doc.metadata.get('fileName', 'N/A')}")
            elif similarity_score >= 0.45 or filename_keyword_matches >= 1:
                filtered_docs.append((doc, similarity_score))
        
        filtered_docs.sort(key=lambda x: x[1], reverse=True)
        # AMÉLIORATION: Augmenter le nombre de documents récupérés pour améliorer la précision
        # Si un nom de personne est détecté, limiter aux top 20 fichiers les plus pertinents
        # Sinon, prendre jusqu'à 30 documents pour avoir plus de contexte
        max_docs = 20 if person_names else 30
        top_docs = [doc for doc, score in filtered_docs[:max_docs]]
        
        logger.info(f"📊 {len(top_docs)} documents finaux sélectionnés après filtrage strict par nom")
        
        # Construire la liste de résultats avec amélioration du contenu
        seen_docs = set()
        results = []
        for doc in top_docs:
            file_name_from_meta = doc.metadata.get("fileName") or doc.metadata.get("file_name") or "document_vectorstore"
            content_preview = doc.page_content[:100]
            doc_key = f"{file_name_from_meta}:{content_preview}"
            
            if doc_key not in seen_docs:
                seen_docs.add(doc_key)
                # AMÉLIORATION: Augmenter la taille du contenu et chercher le contexte autour des mots-clés
                content = doc.page_content
                query_lower = user_query.lower()
                query_words = [w.strip('.,!?;:()[]{}"\'').lower() for w in user_query.split() if len(w.strip('.,!?;:()[]{}"\'')) > 2]
                
                # Si on trouve un mot-clé dans le contenu, essayer d'inclure plus de contexte autour
                content_lower = content.lower()
                for word in query_words:
                    if word in content_lower:
                        # Trouver la position du mot et inclure plus de contexte
                        idx = content_lower.find(word)
                        if idx > 0:
                            # Inclure 500 caractères avant et après pour capturer toute la fonction
                            start = max(0, idx - 500)
                            end = min(len(content), idx + len(word) + 500)
                            # Si le contenu extrait est plus pertinent, l'utiliser
                            if end - start > len(content[:3000]):
                                content = content[start:end]
                                break
                
                results.append({
                    "fileName": file_name_from_meta,
                    "content": content[:3000]  # Augmenté de 2500 à 3000 pour plus de contexte
                })
        
        logger.info(f"📚 {len(results)} documents uniques récupérés pour la requête")
        return results
        
    except Exception as e:
        logger.error(f"Erreur lors de la recherche sémantique synchrone: {str(e)}")
        return []

async def detect_missing_information(response: str, user_query: str, language: str = 'en') -> bool:
    """
    Détecte si la réponse du modèle indique que l'information n'est pas disponible dans les documents
    ou si une recherche en ligne serait utile pour obtenir des informations en temps réel.
    Retourne True si l'information semble absente ou si une recherche en ligne serait bénéfique.
    """
    response_lower = response.lower()
    query_lower = user_query.lower()
    
    # Phrases indicatrices que l'information n'est pas disponible
    missing_indicators = [
        "n'apparaît pas",
        "n'est pas disponible",
        "n'est pas présent",
        "n'est pas trouvé",
        "n'est pas mentionné",
        "n'est pas indiqué",
        "n'est pas fourni",
        "n'est pas dans",
        "ne figure pas",
        "ne contient pas",
        "aucune information",
        "pas d'information",
        "information absente",
        "information non disponible",
        "je ne trouve pas",
        "je n'ai pas trouvé",
        "impossible de trouver",
        "ne peut pas être trouvé",
        "does not appear",
        "is not available",
        "is not present",
        "is not found",
        "is not mentioned",
        "is not indicated",
        "is not provided",
        "is not in",
        "no information",
        "information not available",
        "i cannot find",
        "i did not find",
        "unable to find",
        "cannot be found"
    ]
    
    # Indicateurs que des informations plus récentes seraient utiles
    needs_recent_info_indicators = [
        "informations plus récentes",
        "information plus récente",
        "données plus récentes",
        "vérifier les changements",
        "confirmer que",
        "pourrait être utile",
        "serait utile",
        "pourrait nécessiter",
        "nécessiterait",
        "more recent information",
        "recent information",
        "more recent data",
        "check for changes",
        "confirm that",
        "might be useful",
        "would be useful",
        "might require",
        "would require",
        "should verify",
        "devrait vérifier",
        "il est important de vérifier",
        "it is important to check"
    ]
    
    # Mots-clés dans la requête qui indiquent un besoin d'informations en temps réel
    real_time_query_keywords = [
        "heure actuelle",
        "heure maintenant",
        "quelle heure",
        "what time",
        "current time",
        "time now",
        "prix actuel",
        "prix maintenant",
        "current price",
        "price now",
        "taux actuel",
        "current rate",
        "taux de change",
        "exchange rate",
        "cours actuel",
        "current rate",
        "événements récents",
        "recent events",
        "nouvelles récentes",
        "recent news",
        "maintenant",
        "now",
        "actuel",
        "current",
        "aujourd'hui",
        "today",
        "en ce moment",
        "right now",
        "à l'instant",
        "at the moment"
    ]
    
    # Vérifier si la réponse contient des indicateurs d'absence
    has_missing_indicator = any(indicator in response_lower for indicator in missing_indicators)
    
    # Vérifier si la réponse suggère qu'une information plus récente serait utile
    needs_recent_info = any(indicator in response_lower for indicator in needs_recent_info_indicators)
    
    # Vérifier si la requête demande des informations en temps réel
    needs_realtime_info = any(keyword in query_lower for keyword in real_time_query_keywords)
    
    # Vérifier aussi la longueur de la réponse - si elle est très courte, c'est probablement qu'il n'y a pas d'info
    is_very_short = len(response.strip()) < 100
    
    # Vérifier si la réponse contient des phrases comme "d'après les documents" ou "dans les documents fournis"
    # Si oui, c'est probablement qu'il n'y a pas d'info ailleurs
    mentions_documents_only = any(phrase in response_lower for phrase in [
        "dans les documents fournis",
        "dans les documents",
        "dans le document",
        "d'après les documents",
        "selon les documents",
        "in the provided documents",
        "in the documents",
        "in the document",
        "according to the documents"
    ])
    
    # Déclencher une recherche en ligne si :
    # 1. La réponse indique que l'info est absente
    # 2. La réponse suggère qu'une info plus récente serait utile
    # 3. La requête demande des informations en temps réel
    # 4. La réponse est très courte ET mentionne seulement les documents
    should_search_online = (
        has_missing_indicator or 
        needs_recent_info or 
        needs_realtime_info or 
        (is_very_short and mentions_documents_only)
    )
    
    if should_search_online:
        reason = []
        if has_missing_indicator:
            reason.append("information absente")
        if needs_recent_info:
            reason.append("besoin d'informations récentes")
        if needs_realtime_info:
            reason.append("demande d'informations en temps réel")
        logger.info(f"🔍 Recherche en ligne déclenchée: {', '.join(reason)}")
    
    return should_search_online

async def query_model_local_mode(file_name: str, file_content: str, directory_content: List[Dict], 
                                repo_structure: str, user_query: str, is_binary: bool = False, 
                                selected_model: str = DEFAULT_MODEL, language: str = 'en',
                                conversation_history: Optional[List[Dict]] = None,
                                enable_auto_online_search: bool = True) -> tuple[str, str]:
    """
    Mode LOCAL: analyse STRICTEMENT les documents fournis (fichier principal + contexte).
    On réduit et structure le contexte pour améliorer la précision et limiter les hallucinations.
    IMPORTANT: pas de cache sur ce mode, pour toujours prendre en compte le contexte le plus récent
    (nouveaux fichiers indexés, corrections, etc.).
    """
    t = translations.get(language, translations['en'])
    
    # Amélioration: Organiser le contexte par pertinence et préserver les informations clés
    max_chunk_len = 2000  # Augmenté encore plus pour préserver le maximum de contexte
    contextual_docs: List[str] = []
    
    # Grouper les documents par fichier source pour éviter la duplication et préserver le contexte complet
    # Structure: {file_label: [(content, page_number), ...]}
    docs_by_file = {}
    for doc in directory_content or []:
        if not isinstance(doc, dict):
            continue
        file_label = doc.get('fileName') or doc.get('file_name') or "document_contextuel"
        raw_content = doc.get('content', '') or ''
        if raw_content is None:
            raw_content = ''
        page_number = doc.get('pageNumber')  # Extraire le numéro de page si disponible
        
        if file_label not in docs_by_file:
            docs_by_file[file_label] = []
        docs_by_file[file_label].append((raw_content, page_number))
    
    # Construire un résumé structuré par fichier (les premiers sont les plus pertinents)
    # IMPORTANT: Préserver le maximum de contenu pour capturer les informations comme les adresses
    for idx, (file_label, contents_with_pages) in enumerate(docs_by_file.items(), 1):
        # Extraire les contenus et pages
        contents = [c[0] for c in contents_with_pages]
        pages = [c[1] for c in contents_with_pages if c[1] is not None]
        
        # Combiner les chunks du même fichier avec un séparateur clair
        combined_content = "\n---\n".join(contents)
        
        # Créer un label avec les informations de page
        page_info = ""
        if pages:
            unique_pages = sorted(set(pages))
            if len(unique_pages) == 1:
                page_info = f" 📄 Page {unique_pages[0]}"
            else:
                page_info = f" 📄 Pages {', '.join(map(str, unique_pages))}"
        
        # Pour les fichiers pertinents (top 5), préserver encore plus de contenu
        extended_max_len = max_chunk_len * 1.5 if idx <= 5 else max_chunk_len
        
        # Tronquer intelligemment en préservant le début et la fin
        if len(combined_content) > extended_max_len:
            # Garder le début (souvent le plus important) et un peu de la fin
            # Augmenter la partie finale pour capturer les informations en fin de document
            content = combined_content[:int(extended_max_len - 300)] + "\n[... section tronquée ...]\n" + combined_content[-300:]
        else:
            content = combined_content
        
        # Marquer les documents les plus pertinents (premiers dans la liste)
        relevance_marker = "⭐" if idx <= 3 else "🔍" if idx <= 8 else ""
        # Ajouter les informations de page si disponibles (depuis les contenus)
        page_info = ""
        if pages:
            unique_pages = sorted(set(pages))
            if len(unique_pages) == 1:
                page_info = f" 📄 Page {unique_pages[0]}"
            else:
                page_info = f" 📄 Pages {', '.join(map(str, unique_pages))}"
        
        contextual_docs.append(f"{relevance_marker} [{idx}] {file_label}{page_info}:\n{content}")

    directory_content_summary = "\n\n" + "="*80 + "\n\n".join(contextual_docs) + "\n\n" + "="*80 if contextual_docs else t.get(
        'no_other_files', "Aucun autre fichier dans le contexte."
    )
    
    # Ajouter un header explicatif si on a du contexte
    if contextual_docs:
        directory_content_summary = (
            f"⚠️ ATTENTION: {len(docs_by_file)} fichiers ont été récupérés par recherche sémantique "
            f"car ils sont potentiellement pertinents pour ta question.\n"
            f"Les documents sont classés par pertinence:\n"
            f"  ⭐ = Très pertinent (priorité haute)\n"
            f"  🔍 = Pertinent (priorité moyenne)\n"
            f"  [numéro] = Autre document à analyser\n\n"
            f"TU DOIS ANALYSER TOUS CES FICHIERS, même si l'information semble absente du document principal.\n\n"
            f"{directory_content_summary}"
        )

    # On tronque aussi légèrement le contenu du fichier principal si nécessaire
    main_max_len = 4000
    if file_content is None:
        file_content = ""
    trimmed_main_content = file_content[:main_max_len] + ("..." if len(file_content) > main_max_len else "")

    # Construire un résumé compact de l'historique de conversation (si fourni)
    history_section = ""
    if conversation_history:
      history_lines: List[str] = []
      for turn in conversation_history[-10:]:
          role = turn.get("role", "user")
          content = (turn.get("content") or "")[:600]
          prefix = "Utilisateur" if role == "user" else "Assistant"
          history_lines.append(f"{prefix}: {content}")
      if history_lines:
          history_section = "=== HISTORIQUE DE LA CONVERSATION ===\n" + "\n".join(history_lines) + "\n\n"

    # Détection des noms de personnes dans la requête pour instructions strictes
    person_names_in_query = set()
    words = user_query.split()
    for i, word in enumerate(words):
        clean_word = word.strip('.,!?;:()[]{}"\'').strip()
        if clean_word and clean_word[0].isupper() and len(clean_word) > 2:
            person_names_in_query.add(clean_word.lower())
        if i > 0 and words[i-1].lower() in ['his', 'her', 'their', 'karim', 'dominique', 'about', 'for']:
            if len(clean_word) > 2:
                person_names_in_query.add(clean_word.lower())
    
    query_lower = user_query.lower()
    common_names = ['karim', 'dominique', 'essome', 'ngami']
    for name in common_names:
        if name in query_lower:
            person_names_in_query.add(name)
    
    # Instructions strictes si un nom de personne est détecté
    person_filter_instructions = ""
    if person_names_in_query:
        person_names_str = ", ".join([name.capitalize() for name in person_names_in_query])
        person_filter_instructions = (
            f"\n🚨 INSTRUCTION CRITIQUE - FILTRAGE PAR NOM DE PERSONNE:\n"
            f"La question concerne {person_names_str}. TU DOIS UNIQUEMENT utiliser les documents qui contiennent "
            f"le(s) nom(s) '{person_names_str}' dans leur nom de fichier OU dans leur contenu.\n"
            f"❌ EXCLUSION STRICTE: N'utilise JAMAIS les documents d'autres personnes (comme Dominique si on demande Karim, "
            f"ou Karim si on demande Dominique).\n"
            f"✅ PRIORITÉ: Les fichiers dont le nom contient '{person_names_str}' sont les SEULS documents pertinents.\n"
            f"⚠️ Si un document ne contient pas le nom '{person_names_str}', IGNORE-LE COMPLÈTEMENT, même s'il semble "
            f"sémantiquement similaire à la question.\n\n"
        )
    
    # Prompt structuré et très explicite avec instructions améliorées pour la recherche sémantique
    # IMPORTANT: Le modèle doit analyser TOUS les documents, pas seulement le fichier sélectionné
    num_other_docs = len(contextual_docs) if contextual_docs else 0
    total_docs_note = f"\n⚠️ ATTENTION: Tu as accès à {num_other_docs} autres documents en plus du document principal. " \
                      f"L'information demandée peut être dans N'IMPORTE LEQUEL de ces documents. " \
                      f"ANALYSE TOUS LES DOCUMENTS avant de conclure qu'une information est absente.\n\n" if num_other_docs > 0 else ""

    prompt = (
        f"{t['local_analysis_mode']}\n"
        f"{t['no_external_search']}\n\n"
        f"{history_section}"
        f"{total_docs_note}"
        f"{person_filter_instructions}"
        f"=== CONTEXTE DU PROJET ===\n"
        f"{t['project_structure']}:\n{repo_structure}\n\n"
        f"=== DOCUMENT PRINCIPAL (FICHIER SÉLECTIONNÉ) ===\n"
        f"{t['main_file']}: {file_name}\n"
        f"⚠️ NOTE: Ce fichier est celui actuellement sélectionné, mais l'information peut être dans d'autres fichiers.\n"
        f"{t['file_content']}:\n{trimmed_main_content}\n\n"
        f"=== AUTRES DOCUMENTS DU DOSSIER (RÉCUPÉRÉS PAR RECHERCHE SÉMANTIQUE) ===\n"
        f"⚠️ IMPORTANT: Ces documents ont été récupérés car ils sont potentiellement pertinents pour ta question. "
        f"ANALYSE-TOUS MÉTICULEUSEMENT, même si l'information semble absente du document principal.\n"
        f"{directory_content_summary}\n\n"
        f"=== QUESTION À TRAITER ===\n"
        f"{t['question']}: {user_query}\n\n"
        f"{t['instructions']}:\n"
        f"{t['base_response_only']}\n"
        f"🔍 PROCÉDURE DE RECHERCHE OBLIGATOIRE:\n"
    )
    
    # Construire la première étape de la procédure (éviter les backslashes dans f-string)
    if person_names_in_query:
        step1_text = "Si la question concerne une personne spécifique, vérifie d'abord que le document principal contient le nom de cette personne. Sinon, cherche dans les autres documents."
    else:
        step1_text = f"Commence par analyser le document principal ({file_name})."
    
    prompt += (
        f"1. {step1_text}\n"
        f"2. Si l'information n'est pas trouvée, PARCOURS SYSTÉMATIQUEMENT TOUS les autres documents listés ci-dessus.\n"
        f"3. Les documents marqués ⭐ sont les plus pertinents selon la recherche sémantique - commence par ceux-là.\n"
        f"4. Ne conclus JAMAIS qu'une information est absente avant d'avoir analysé TOUS les documents fournis.\n"
        f"5. Si tu trouves l'information dans un autre document, cite explicitement le nom du fichier source.\n"
        f"6. IMPORTANT: Si un document a un numéro de page (📄 Page X), cite aussi la page dans ta réponse.\n"
        f"   Format de citation: 'Document X, page Y' ou '[Document X, page Y]'.\n"
        f"7. Si l'information est dans plusieurs documents, mentionne tous les fichiers concernés avec leurs pages si disponibles.\n\n"
        f"- Les documents sont classés par pertinence sémantique: les premiers sont les plus liés à ta question.\n"
        f"- Mais même les documents moins pertinents peuvent contenir l'information recherchée - ne les ignore pas.\n"
        f"{t['missing_info_clarify']}\n"
        f"{t['no_speculation']}\n"
        f"{t['focus_local_analysis']}\n"
        f"- Tu peux faire des calculs (totaux, moyennes, salaire mensuel/annuel, etc.) "
        f"à partir des montants et périodes présents dans les documents, mais explique "
        f"clairement ton raisonnement et les formules utilisées.\n"
        f"- Si une information n'est pas directement présente mais peut être déduite "
        f"par un calcul simple à partir des données (par exemple, convertir un salaire "
        f"par paie en salaire mensuel), effectue le calcul et montre les étapes.\n"
        f"- Si une information ne peut ni être lue ni déduite des textes ci-dessus APRÈS AVOIR ANALYSÉ TOUS LES DOCUMENTS, "
        f"répond explicitement qu'elle n'apparaît pas dans les documents fournis.\n\n"
        f"{t['emoji']}"
    )

    try:
        # Utiliser execute_model_query_with_fallback pour garantir la pertinence
        result, model_used = await execute_model_query_with_fallback(prompt, selected_model, user_query)
        logger.info(f"✅ Modèle utilisé pour la réponse: {model_used}")
        
        # Détection automatique si l'information n'est pas disponible et recherche en ligne si activée
        if enable_auto_online_search:
            missing_info = await detect_missing_information(result, user_query, language)
            
            if missing_info:
                logger.info("🔍 Information absente détectée dans les documents. Recherche en ligne automatique...")
                t = translations.get(language, translations['en'])
                
                try:
                    # Effectuer une recherche en ligne
                    search_query = user_query
                    search_results = perform_online_search(search_query, language)
                    
                    if search_results and search_results != "Aucun résultat trouvé." and search_results != "Clé API manquante":
                        # Fusionner les résultats locaux avec les résultats en ligne
                        enrichment_prompt = (
                            f"{t.get('enrichment_title', '🌐 Enrichissement avec recherche en ligne')}\n\n"
                            f"**Réponse initiale basée sur les documents locaux:**\n{result}\n\n"
                            f"**Résultats de recherche en ligne:**\n{search_results}\n\n"
                            f"{t.get('enrichment_instructions', 'Instructions:')}\n"
                            f"- Combine intelligemment les informations des documents locaux avec les données trouvées en ligne.\n"
                            f"- Si l'information n'est pas dans les documents locaux mais est trouvée en ligne, utilise les données en ligne.\n"
                            f"- Distingue clairement les sources: indique ce qui vient des documents locaux vs. ce qui vient de la recherche en ligne.\n"
                            f"- Cite les sources en ligne si disponibles.\n"
                            f"- Garde la structure et le format de la réponse initiale si possible.\n"
                            f"- Priorise les informations récentes trouvées en ligne pour les données qui peuvent être obsolètes.\n\n"
                            f"**Réponse enrichie:**"
                        )
                        
                        enriched_result, _ = await execute_model_query_with_fallback(enrichment_prompt, selected_model, user_query)
                        
                        # Ajouter une note indiquant que la recherche en ligne a été utilisée
                        result = (
                            f"{enriched_result}\n\n"
                            f"---\n"
                            f"💡 **Note**: Cette réponse combine les informations des documents locaux avec des données trouvées en ligne, "
                            f"car certaines informations n'étaient pas disponibles dans les documents fournis.\n"
                        )
                        
                        logger.info("✅ Réponse enrichie avec des données en ligne")
                    else:
                        logger.info("⚠️ Aucun résultat trouvé en ligne ou clé API manquante")
                        # Garder la réponse originale si la recherche en ligne n'a rien donné
                except Exception as search_error:
                    logger.warning(f"⚠️ Erreur lors de la recherche en ligne automatique: {str(search_error)}")
                    # En cas d'erreur, garder la réponse originale
                    pass
    except Exception as e:
        result = f"Erreur lors de l'analyse locale: {str(e)}"
        model_used = selected_model

    return result, model_used

def is_response_relevant(response: str, user_query: str) -> bool:
    """
    Détermine si une réponse est pertinente en vérifiant:
    1. Si la réponse n'est pas vide ou trop courte
    2. Si la réponse ne contient pas de phrases indiquant que l'information est absente
    3. Si la réponse semble contenir une information réelle (pas juste "je ne sais pas")
    """
    if not response or len(response.strip()) < 20:
        logger.warning(f"❌ Réponse trop courte ou vide: {len(response.strip() if response else 0)} chars")
        return False
    
    response_lower = response.lower()
    user_query_lower = user_query.lower()
    
    # Phrases indiquant que l'information n'est pas trouvée (plus complètes)
    missing_phrases = [
        "n'apparaît pas dans les documents",
        "n'est pas présente dans",
        "n'est pas disponible dans",
        "n'est pas trouvé",
        "n'est pas mentionné",
        "n'est pas indiqué",
        "ne trouve pas",
        "ne peut pas trouver",
        "ne peut pas être trouvé",
        "impossible de trouver",
        "information absente",
        "donnée absente",
        "non disponible",
        "pas d'information",
        "aucune information",
        "je ne peux pas",
        "je ne trouve pas",
        "je n'ai pas trouvé",
        "dans les documents fournis, l'information",
        "l'information sur",
        "n'est pas directement mentionné",
        "does not appear in",
        "is not present in",
        "is not available in",
        "is not found",
        "is not mentioned",
        "cannot find",
        "unable to find",
        "information missing",
        "data not available",
        "no information",
        "i cannot",
        "i cannot find",
        "i did not find",
        "in the provided documents, the information",
        "the information about",
        "is not directly mentioned"
    ]
    
    # Détection stricte: si une phrase d'absence apparaît dans la réponse, c'est non pertinent
    # Vérifier particulièrement dans les premiers 300 caractères (début de réponse)
    response_start = response_lower[:300]
    for phrase in missing_phrases:
        if phrase in response_lower:
            # Si la phrase d'absence apparaît dans les 300 premiers caractères, c'est définitivement non pertinent
            if phrase in response_start:
                logger.warning(f"❌ Phrase d'absence détectée dans les premiers 300 chars: '{phrase}'")
                return False
            # Si la phrase apparaît plus loin mais que la réponse contient principalement cette information, c'est aussi non pertinent
            # Compter combien de fois les phrases d'absence apparaissent
            missing_count = sum(1 for p in missing_phrases if p in response_lower)
            if missing_count >= 2:  # Si 2+ phrases d'absence, c'est probablement non pertinent
                logger.warning(f"❌ Plusieurs phrases d'absence détectées ({missing_count})")
                return False
    
    # Vérifier que la réponse contient au moins quelques mots de la requête (pertinence sémantique)
    query_words = set(word.lower().strip('.,!?;:()[]{}"\'') for word in user_query_lower.split() if len(word) > 2)
    response_words = set(word.lower().strip('.,!?;:()[]{}"\'') for word in response_lower.split() if len(word) > 2)
    
    # Extraire les mots-clés importants de la requête (noms propres, mots significatifs)
    important_query_words = [w for w in query_words if w not in ['est', 'sont', 'être', 'avoir', 'was', 'is', 'are', 'the', 'a', 'an', 'le', 'la', 'les', 'un', 'une', 'des']]
    
    if important_query_words:
        # Vérifier si au moins un mot important de la requête est dans la réponse
        common_words = query_words & response_words
        if len(common_words) == 0:
            logger.warning(f"❌ Aucun mot en commun entre requête et réponse. Requête: {important_query_words[:5]}")
            return False
        
        # Vérifier que la réponse contient au moins 30% des mots importants
        important_common = [w for w in important_query_words if w in response_words]
        if len(important_common) / len(important_query_words) < 0.3:
            logger.warning(f"❌ Trop peu de mots importants trouvés ({len(important_common)}/{len(important_query_words)})")
            # Mais accepter si la réponse est très détaillée et contient au moins un mot important
            if len(response.strip()) < 400 or len(important_common) == 0:
                return False
    
    # Vérifier si la réponse est trop générique (ex: "veuillez consulter les documents")
    generic_phrases = [
        "consultez les documents",
        "veuillez consulter",
        "please consult",
        "refer to the documents",
        "see the documents"
    ]
    if any(phrase in response_lower for phrase in generic_phrases) and len(response.strip()) < 150:
        logger.warning(f"❌ Réponse trop générique")
        return False
    
    logger.info(f"✅ Réponse considérée comme pertinente ({len(response.strip())} chars, {len(query_words & response_words)} mots en commun)")
    return True

async def execute_model_query_with_fallback(prompt: str, selected_model: str, user_query: str = "") -> tuple[str, str]:
    """
    Exécute la requête sur le modèle sélectionné avec fallback en cascade si la réponse n'est pas pertinente.
    Retourne (response, model_used)
    
    Ordre de fallback si GPT-3.5-turbo ne donne pas de réponse pertinente:
    1. Mistral
    2. GPT-4o
    3. GPT-5-Nano
    4. GPT-5-Mini
    5. GPT-5
    """
    # Liste complète des modèles de fallback (dans l'ordre de qualité croissante)
    all_fallback_models = ["gpt-3.5-turbo", "mistral", "gpt-4o", "gpt-5-nano", "gpt-5-mini", "gpt-5"]
    
    # Construire la liste des modèles à essayer en commençant par le modèle sélectionné
    selected_lower = selected_model.lower()
    
    if selected_lower in all_fallback_models:
        # Trouver l'index du modèle sélectionné et prendre tous les modèles suivants
        model_idx = all_fallback_models.index(selected_lower)
        models_to_try = all_fallback_models[model_idx:]
        logger.info(f"📋 Ordre de fallback: {models_to_try}")
    else:
        # Modèle non dans la liste, utiliser tel quel puis fallback complet
        models_to_try = [selected_lower] + all_fallback_models
        logger.info(f"📋 Modèle inconnu '{selected_lower}', utilisation avec fallback complet: {models_to_try}")
    
    last_exception = None
    
    for model in models_to_try:
        try:
            logger.info(f"🔄 Tentative avec le modèle: {model}")
            
            if model in ["gpt-3.5-turbo", "gpt-4o", "gpt-5", "gpt-5-mini", "gpt-5-nano"]:
                result = call_openai_api(prompt, model)
            elif model == "mistral":
                result = call_mistral_api(prompt)
            elif model.startswith("gemini") or model in ["gemini-3-flash", "gemini-pro"]:
                result = call_gemini_api(prompt, model)
            elif model in OLLAMA_MODELS or model == "llama3":
                result = call_ollama_api(prompt, model)
            else:
                # Modèle inconnu, utiliser OpenAI par défaut
                result = call_openai_api(prompt, DEFAULT_MODEL)
                model = DEFAULT_MODEL
            
            # Vérifier la pertinence si user_query est fourni
            # Pour GPT-3.5-turbo, on vérifie toujours (plus susceptible d'échouer)
            # Pour les autres modèles moins puissants (mistral), on vérifie aussi
            # Pour les modèles plus puissants (gpt-4o, gpt-5*), on fait confiance mais on peut quand même vérifier
            should_check_relevance = user_query and (
                model == "gpt-3.5-turbo" or 
                model == "mistral" or
                model in ["gpt-5-nano", "gpt-5-mini"]  # Vérifier aussi pour les modèles moins puissants
            )
            
            if should_check_relevance:
                if is_response_relevant(result, user_query):
                    logger.info(f"✅ Réponse pertinente obtenue avec {model}")
                    return result, model
                else:
                    logger.warning(f"⚠️ Réponse non pertinente avec {model}, passage au modèle suivant...")
                    continue
            else:
                # Pour les modèles plus puissants (gpt-4o, gpt-5) ou si pas de user_query, accepter la réponse
                logger.info(f"✅ Réponse obtenue avec {model} (vérification de pertinence sautée)")
                return result, model
                
        except Exception as e:
            logger.warning(f"❌ Erreur avec le modèle {model}: {str(e)}, passage au modèle suivant...")
            last_exception = e
            continue
    
    # Si tous les modèles ont échoué, lever l'exception du dernier
    if last_exception:
        raise last_exception
    else:
        raise RuntimeError("Tous les modèles de fallback ont échoué")

async def execute_model_query(prompt: str, selected_model: str) -> str:
    """
    Exécute la requête sur le modèle sélectionné (version simple, sans fallback)
    """
    try:
        if selected_model.lower() in ["gpt-3.5-turbo", "gpt-4o"]:
            result = call_openai_api(prompt, selected_model)
        elif selected_model.lower() in ["gpt-5", "gpt-5-mini", "gpt-5-nano"]:
            result = call_openai_api(prompt, selected_model)
        elif selected_model.lower() == "openai":
            result = call_openai_api(prompt, "openai")
        elif selected_model.lower() == "mistral":
            result = call_mistral_api(prompt)
        elif selected_model.lower().startswith("gemini") or selected_model.lower() in ["gemini-3-flash", "gemini-pro"]:
            result = call_gemini_api(prompt, selected_model)
        elif selected_model.lower() in OLLAMA_MODELS or selected_model.lower() == "llama3":
            result = call_ollama_api(prompt, selected_model)
        else:
            logger.warning(f"Modèle inconnu {selected_model}, utilisation du modèle par défaut {DEFAULT_MODEL}")
            result = call_openai_api(prompt, DEFAULT_MODEL)
        
        return result
    except Exception as e:
        raise e

async def detect_document_type(file_content: str, file_name: str) -> str:
    """
    Détecte le type de document de manière précise en utilisant une approche hybride:
    1. Analyse basée sur des patterns (règles)
    2. Analyse LLM pour les cas complexes
    """
    content_lower = file_content[:5000].lower()  # Augmenter la fenêtre d'analyse
    file_name_lower = file_name.lower()
    content_sample = file_content[:1500]  # Échantillon pour l'analyse LLM si nécessaire
    
    # ============ DÉTECTION PAR PATTERNS (FAST PATH) ============
    
    # Détection CV/Resume (très spécifique)
    cv_keywords = ['curriculum vitae', 'cv', 'resume', 'résumé professionnel', 'professional summary']
    cv_indicators = ['expérience professionnelle', 'professional experience', 'compétences', 'skills', 
                     'formation', 'education', 'emploi', 'job', 'poste', 'position',
                     'langues', 'languages', 'certifications', 'réalisations', 'achievements']
    if any(word in file_name_lower for word in cv_keywords) or \
       (any(word in content_lower[:800] for word in cv_keywords) and 
        sum(1 for word in cv_indicators if word in content_lower) >= 3):
        return 'cv_resume'
    
    # Détection facture/Invoice (très spécifique)
    invoice_keywords = ['facture', 'invoice', 'bill', 'reçu', 'receipt', 'quittance']
    invoice_indicators = ['montant total', 'total amount', 'tva', 'tax', 'date d\'émission', 'issue date',
                          'numéro de facture', 'invoice number', 'client', 'customer', 'paiement', 'payment']
    if any(word in file_name_lower for word in invoice_keywords) or \
       (any(word in content_lower[:800] for word in invoice_keywords) and 
        sum(1 for word in invoice_indicators if word in content_lower) >= 2):
        return 'facture_invoice'
    
    # Détection des lettres (doit être fait tôt pour éviter les faux positifs)
    letter_keywords = ['lettre', 'letter', 'correspondance', 'correspondence', 'courrier', 'mail']
    letter_context = ['soutien', 'support', 'recommandation', 'recommendation', 'demande', 'request', 
                     'attestation', 'certificate', 'certificat', 'justificatif', 'justification',
                     'cher monsieur', 'dear sir', 'madame', 'madam', 'monsieur', 'sir']
    if any(word in file_name_lower for word in letter_keywords) or \
       (any(word in content_lower[:500] for word in letter_keywords) and 
        (any(word in content_lower for word in letter_context) or 
         'objet:' in content_lower[:300] or 'subject:' in content_lower[:300])):
        return 'lettre'
    
    # Détection des documents financiers/fiscaux (améliorée)
    financial_keywords = ['t4', 't-4', 't4a', 't4a-', 'relevé', 'statement', 'payroll', 'paie', 
                          'salaire', 'salary', 'revenu', 'income', 'impôt', 'tax', 'déduction', 'deduction',
                          'feuillet', 'slip', 'relevé fiscal', 'tax statement', 'avis de cotisation', 'notice of assessment']
    financial_indicators = ['revenus bruts', 'gross income', 'revenus nets', 'net income', 'impôt retenu', 'tax withheld',
                           'année', 'year', 'période', 'period', 'numéro d\'assurance sociale', 'social insurance number']
    if any(word in file_name_lower for word in financial_keywords) or \
       (any(word in content_lower[:800] for word in financial_keywords) and 
        sum(1 for word in financial_indicators if word in content_lower) >= 2):
        return 'document_financier'
    
    # Détection des contrats (améliorée avec plus de contexte)
    contract_keywords = ['contrat', 'contract', 'agreement', 'convention', 'accord']
    if any(word in content_lower or word in file_name_lower for word in contract_keywords):
        # Contrat de location
        rental_indicators = ['location', 'rental', 'bail', 'loyer', 'rent', 'locataire', 'tenant', 
                            'bailleur', 'landlord', 'propriétaire', 'owner', 'garant', 'guarantor',
                            'charges locatives', 'maintenance fees', 'caution', 'deposit']
        if sum(1 for word in rental_indicators if word in content_lower) >= 3:
            return 'contrat_location'
        
        # Contrat de travail
        employment_indicators = ['travail', 'employment', 'employé', 'employee', 'employeur', 'employer',
                                'salaire', 'salary', 'rémunération', 'remuneration', 'poste', 'position',
                                'période d\'essai', 'probation', 'cde', 'cdi', 'cdd', 'contract duration']
        if sum(1 for word in employment_indicators if word in content_lower) >= 3:
            return 'contrat_travail'
        
        # Contrat de vente
        sale_indicators = ['vente', 'sale', 'achat', 'purchase', 'acheteur', 'buyer', 'vendeur', 'seller',
                          'prix', 'price', 'livraison', 'delivery', 'garantie', 'warranty', 'garant', 'guarantee']
        if sum(1 for word in sale_indicators if word in content_lower) >= 3:
            return 'contrat_vente'
        
        # Contrat générique
        return 'contrat_generique'
    
    # Détection des testaments
    will_keywords = ['testament', 'will', 'testamentaire', 'testamentary']
    will_indicators = ['héritier', 'heir', 'bénéficiaire', 'beneficiary', 'legs', 'bequest', 'legataire', 'legatee',
                      'exécuteur', 'executor', 'succession', 'inheritance']
    if any(word in content_lower or word in file_name_lower for word in will_keywords) or \
       (any(word in will_keywords for word in file_name_lower) and 
        sum(1 for word in will_indicators if word in content_lower) >= 2):
        return 'testament'
    
    # Détection des contrats de mariage/prénuptiaux (Prenuptial Agreement)
    prenuptial_keywords = ['contrat de mariage', 'prenuptial agreement', 'prenup', 'contrat prénuptial', 
                          'marriage contract', 'convention matrimoniale', 'marital agreement']
    prenuptial_indicators = ['régime matrimonial', 'matrimonial regime', 'biens', 'property', 'séparation de biens',
                            'separation of property', 'communauté', 'community', 'époux', 'spouse', 'mariage', 'marriage']
    if any(phrase in content_lower or phrase in file_name_lower for phrase in prenuptial_keywords) or \
       (any(word in file_name_lower for word in ['prenup', 'prenuptial', 'marriage contract', 'contrat mariage']) and 
        sum(1 for word in prenuptial_indicators if word in content_lower) >= 2):
        return 'contrat_prenuptial'
    
    # Détection des procurations (Power Of Attorney)
    poa_keywords = ['procuration', 'power of attorney', 'power-of-attorney', 'mandat', 'mandate', 'poa']
    poa_indicators = ['mandant', 'principal', 'mandataire', 'agent', 'attorney-in-fact', 'pouvoir', 'authority',
                     'représenter', 'represent', 'agir au nom', 'act on behalf', 'signer', 'sign']
    if any(word in content_lower or word in file_name_lower for word in poa_keywords) or \
       (any(word in file_name_lower for word in ['poa', 'power attorney', 'procuration']) and 
        sum(1 for word in poa_indicators if word in content_lower) >= 2):
        return 'procuration_poa'
    
    # Détection des accords de confidentialité (NDA)
    nda_keywords = ['accord de confidentialité', 'non-disclosure agreement', 'nda', 'n.d.a.', 
                   'confidentiality agreement', 'accord de non-divulgation']
    nda_indicators = ['confidentiel', 'confidential', 'secret', 'secret information', 'proprietary', 'propriétaire',
                     'divulgation', 'disclosure', 'révéler', 'reveal', 'informations confidentielles']
    if any(phrase in content_lower or phrase in file_name_lower for phrase in nda_keywords) or \
       (any(word in file_name_lower for word in ['nda', 'non-disclosure', 'confidentiality']) and 
        sum(1 for word in nda_indicators if word in content_lower) >= 2):
        return 'accord_confidentialite_nda'
    
    # Détection des actes de propriété immobilière (Real Estate Deed)
    deed_keywords = ['acte de propriété', 'real estate deed', 'property deed', 'acte de vente immobilière',
                    'deed of sale', 'acte notarié', 'notarial deed', 'title deed', 'titre de propriété']
    deed_indicators = ['propriétaire', 'owner', 'propriété immobilière', 'real estate', 'bien immobilier',
                      'property', 'parcelle', 'lot', 'cadastre', 'cadastral', 'superficie', 'area',
                      'adresse', 'address', 'bornage', 'boundary', 'hypothèque', 'mortgage']
    if any(phrase in content_lower or phrase in file_name_lower for phrase in deed_keywords) or \
       (any(word in file_name_lower for word in ['deed', 'acte propriété', 'titre propriété']) and 
        sum(1 for word in deed_indicators if word in content_lower) >= 3):
        return 'acte_propriete_immobiliere'
    
    # Détection des actes notariés (améliorée)
    notary_keywords = ['acte notarié', 'notarial act', 'acte authentique', 'authentic act', 'acte sous seing privé']
    notary_context = ['notaire', 'notary', 'étude notariale', 'notary office', 'minute', 
                     'authentification', 'authentification', 'signature authentique', 'répertoire des minutes']
    
    has_acte = any(word in content_lower or word in file_name_lower for word in ['acte', 'deed'])
    has_notary = any(word in content_lower for word in notary_context) or \
                 any(phrase in content_lower for phrase in notary_keywords)
    
    if has_acte and has_notary:
        return 'acte_notarie'
    
    # Détection bail/lease (séparé des contrats de location)
    if any(word in content_lower or word in file_name_lower for word in ['bail', 'lease']) and \
       'contrat' not in content_lower[:500]:  # Si c'est mentionné comme bail mais pas comme contrat
        return 'contrat_location'  # Normaliser vers contrat_location
    
    # ============ DÉTECTION PAR LLM POUR CAS COMPLEXES ============
    # Si aucun pattern n'a matché, utiliser un LLM pour une détection plus fine
    try:
        detection_prompt = f"""Analyze this document sample and identify its type. Return ONLY one of these exact types:
- cv_resume
- facture_invoice
- contrat_location
- contrat_travail
- contrat_vente
- contrat_generique
- contrat_prenuptial
- procuration_poa
- accord_confidentialite_nda
- acte_propriete_immobiliere
- testament
- acte_notarie
- lettre
- document_financier
- document_generique

File name: {file_name}
Content sample (first 1500 chars): {content_sample}

Return ONLY the type identifier, nothing else:"""
        
        detected_type = call_mistral_api(detection_prompt).strip().lower()
        
        # Valider que le type détecté est valide
        valid_types = ['cv_resume', 'facture_invoice', 'contrat_location', 'contrat_travail', 
                      'contrat_vente', 'contrat_generique', 'contrat_prenuptial', 'procuration_poa',
                      'accord_confidentialite_nda', 'acte_propriete_immobiliere', 'testament', 
                      'acte_notarie', 'lettre', 'document_financier', 'document_generique']
        if detected_type in valid_types:
            return detected_type
    except Exception as e:
        logger.warning(f"LLM document type detection failed: {e}")
    
    # Fallback
    return 'document_generique'

async def extract_structured_data(file_content: str, file_name: str, document_type: str, 
                                  selected_model: str = DEFAULT_MODEL, language: str = 'fr') -> Dict[str, Any]:
    """
    Extrait les données structurées d'un document selon son type
    Retourne un dictionnaire JSON structuré
    """
    # Prompts spécialisés par type de document
    prompts = {
        'contrat_location': """Tu dois extraire toutes les informations structurées de ce contrat de location.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Contrat de location",
  "parties": [
    {
      "nom": "nom complet",
      "role": "LOCATAIRE ou BAILLEUR",
      "adresse": "adresse complète",
      "telephone": "numéro si disponible",
      "email": "email si disponible"
    }
  ],
  "dates_importantes": [
    {
      "type": "Date de signature",
      "valeur": "date au format JJ/MM/AAAA"
    },
    {
      "type": "Date de début",
      "valeur": "date"
    },
    {
      "type": "Date de fin",
      "valeur": "date si disponible"
    },
    {
      "type": "Durée",
      "valeur": "durée (ex: 3 ans)"
    }
  ],
  "montants": [
    {
      "type": "Loyer mensuel",
      "valeur": nombre,
      "devise": "EUR"
    },
    {
      "type": "Caution",
      "valeur": nombre,
      "devise": "EUR"
    },
    {
      "type": "Charges",
      "valeur": nombre si disponible,
      "devise": "EUR"
    }
  ],
  "bien_loue": {
    "adresse": "adresse complète du bien",
    "superficie": "superficie si disponible",
    "type": "appartement, maison, etc."
  },
  "clauses_cles": [
    "liste des clauses importantes (renouvellement, indexation, etc.)"
  ],
  "conditions": [
    "conditions particulières mentionnées"
  ]
}
Extrait TOUTES les informations disponibles. Si une information n'est pas présente, utilise null ou omets le champ.""",

        'contrat_travail': """Tu dois extraire toutes les informations structurées de ce contrat de travail.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Contrat de travail",
  "employeur": {
    "nom": "nom de l'entreprise",
    "adresse": "adresse complète",
    "siret": "SIRET si disponible"
  },
  "employe": {
    "nom": "nom complet",
    "adresse": "adresse complète",
    "poste": "intitulé du poste",
    "fonctions": "description des fonctions"
  },
  "dates_importantes": [
    {
      "type": "Date de signature",
      "valeur": "date"
    },
    {
      "type": "Date d'embauche",
      "valeur": "date"
    },
    {
      "type": "Période d'essai",
      "valeur": "durée"
    }
  ],
  "remuneration": {
    "salaire_brut_mensuel": nombre,
    "salaire_net_mensuel": nombre si disponible,
    "devise": "EUR",
    "avantages": ["liste des avantages"]
  },
  "duree": {
    "type": "CDI, CDD, etc.",
    "duree": "durée si CDD"
  },
  "clauses_cles": [
    "clause de non-concurrence si présente",
    "clause de confidentialité si présente",
    "autres clauses importantes"
  ]
}""",

        'testament': """Tu dois extraire toutes les informations structurées de ce testament.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Testament",
  "testateur": {
    "nom": "nom complet",
    "adresse": "adresse complète",
    "date_naissance": "date si disponible"
  },
  "beneficiaires": [
    {
      "nom": "nom complet",
      "relation": "relation avec le testateur",
      "legs": "description du legs",
      "conditions": "conditions éventuelles"
    }
  ],
  "executeur_testamentaire": {
    "nom": "nom si mentionné",
    "fonctions": "fonctions"
  },
  "biens_legues": [
    {
      "description": "description du bien",
      "beneficiaire": "nom du bénéficiaire",
      "valeur_estimee": "valeur si mentionnée"
    }
  ],
  "dates_importantes": [
    {
      "type": "Date de rédaction",
      "valeur": "date"
    },
    {
      "type": "Date de signature",
      "valeur": "date"
    }
  ],
  "conditions_particulieres": [
    "conditions ou clauses particulières"
  ]
}""",

        'contrat_generique': """Tu dois extraire toutes les informations structurées de ce contrat.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Contrat",
  "parties": [
    {
      "nom": "nom complet",
      "role": "rôle dans le contrat",
      "adresse": "adresse si disponible"
    }
  ],
  "objet": "objet du contrat en une phrase",
  "dates_importantes": [
    {
      "type": "type de date (signature, échéance, etc.)",
      "valeur": "date"
    }
  ],
  "montants": [
    {
      "type": "type de montant",
      "valeur": nombre,
      "devise": "EUR"
    }
  ],
  "clauses_cles": [
    "liste des clauses importantes"
  ],
  "obligations": [
    "obligations de chaque partie"
  ],
  "conditions": [
    "conditions particulières"
  ]
}""",

        'document_generique': """Tu dois extraire les informations structurées de ce document.
Retourne UNIQUEMENT un JSON valide avec cette structure:
{
  "type_document": "Type de document détecté",
  "parties_mentionnees": [
    {
      "nom": "nom",
      "role": "rôle si identifiable"
    }
  ],
  "dates_importantes": [
    {
      "type": "type de date",
      "valeur": "date"
    }
  ],
  "montants": [
    {
      "type": "type de montant",
      "valeur": nombre,
      "devise": "EUR"
    }
  ],
  "informations_cles": [
    "points clés du document"
  ]
}"""
    }
    
    # Sélectionner le prompt approprié
    prompt_template = prompts.get(document_type, prompts['document_generique'])
    
    # Construire le prompt final
    extraction_prompt = f"""{prompt_template}

DOCUMENT À ANALYSER:
Nom du fichier: {file_name}

Contenu:
{file_content[:8000]}

IMPORTANT: 
- Retourne UNIQUEMENT le JSON, sans texte avant ou après
- Le JSON DOIT être complet et valide (toutes les accolades et crochets doivent être fermés)
- Utilise des valeurs null pour les champs manquants
- Sois précis et exhaustif
- Extrais TOUTES les informations disponibles dans le document
- Assure-toi que le JSON est bien formé et complet avant de répondre"""
    
    try:
        # Appeler le modèle pour l'extraction avec une limite de tokens plus élevée (2000 tokens pour les JSON structurés)
        # On utilise directement call_openai_api avec max_tokens_override
        raw_response = call_openai_api(extraction_prompt, selected_model, max_retries=3, max_tokens_override=2000)
        
        # Nettoyer la réponse : enlever les markdown code blocks si présents
        cleaned_response = raw_response.strip()
        if cleaned_response.startswith('```json'):
            cleaned_response = cleaned_response[7:]  # Enlever ```json
        if cleaned_response.startswith('```'):
            cleaned_response = cleaned_response[3:]  # Enlever ```
        if cleaned_response.endswith('```'):
            cleaned_response = cleaned_response[:-3]  # Enlever ``` à la fin
        cleaned_response = cleaned_response.strip()
        
        # Fonction pour réparer un JSON incomplet
        def repair_json(json_str: str) -> str:
            """Tente de réparer un JSON incomplet en fermant les structures ouvertes et les chaînes"""
            json_str = json_str.strip()
            if not json_str.startswith('{'):
                return json_str
            
            result = json_str
            
            # Réparer les chaînes tronquées (ex: "devise": "E -> "devise": "EUR")
            # Chercher les guillemets non fermés à la fin
            # On cherche le dernier guillemet ouvrant qui n'est pas suivi d'un guillemet fermant avant la fin
            # Pattern: trouver "..." qui n'est pas fermé
            # On va simplement fermer les chaînes ouvertes à la fin
            lines = result.split('\n')
            if lines:
                last_line = lines[-1].strip()
                # Si la dernière ligne se termine par un guillemet ouvrant ou une chaîne incomplète
                if last_line.count('"') % 2 != 0:  # Nombre impair de guillemets = chaîne non fermée
                    # Trouver le dernier guillemet ouvrant
                    last_quote_pos = last_line.rfind('"')
                    if last_quote_pos >= 0:
                        # Vérifier s'il y a un guillemet fermant après
                        after_quote = last_line[last_quote_pos + 1:]
                        if '"' not in after_quote:
                            # La chaîne est tronquée, on la ferme avec une valeur par défaut ou null
                            # On va plutôt essayer de compléter intelligemment
                            # Pour l'instant, on ferme simplement la chaîne
                            if last_line.endswith('"'):
                                # Déjà un guillemet, on ajoute juste la fermeture
                                pass
                            else:
                                # Tronqué au milieu, on ferme avec null
                                # Trouver où commence la valeur
                                if ':' in last_line:
                                    key_part = last_line[:last_line.rfind(':') + 1]
                                    # Remplacer la valeur tronquée par null
                                    lines[-1] = key_part + ' null'
                                else:
                                    # Juste fermer la chaîne
                                    lines[-1] = last_line + '"'
                        result = '\n'.join(lines)
            
            # Compter les accolades et crochets ouverts
            open_braces = result.count('{')
            close_braces = result.count('}')
            open_brackets = result.count('[')
            close_brackets = result.count(']')
            
            # Fermer les structures ouvertes
            if open_braces > close_braces:
                # Fermer les objets
                for _ in range(open_braces - close_braces):
                    # Trouver la dernière virgule ou le dernier caractère et fermer
                    if result.rstrip().endswith(','):
                        result = result.rstrip()[:-1]  # Enlever la virgule
                    # Nettoyer les virgules avant de fermer
                    result = re.sub(r',(\s*)$', r'\1', result)
                    result += '\n}'
            
            if open_brackets > close_brackets:
                # Fermer les tableaux
                for _ in range(open_brackets - close_brackets):
                    if result.rstrip().endswith(','):
                        result = result.rstrip()[:-1]  # Enlever la virgule
                    # Nettoyer les virgules avant de fermer
                    result = re.sub(r',(\s*)$', r'\1', result)
                    result += '\n]'
            
            # Nettoyer les virgules finales avant les fermetures
            result = re.sub(r',(\s*[}\]])', r'\1', result)
            
            # Nettoyer les chaînes mal fermées restantes
            # Si on a encore des problèmes, on remplace les valeurs tronquées par null
            # Chercher les patterns comme "key": "incomplete_string (sans guillemet fermant)
            result = re.sub(r':\s*"([^"]*?)(?<!")\s*$', r': null', result, flags=re.MULTILINE)
            # Aussi pour les dernières lignes qui se terminent par une chaîne incomplète
            lines = result.split('\n')
            if lines:
                last_line = lines[-1].strip()
                # Si la ligne se termine par un guillemet ouvrant sans fermeture
                if last_line.count('"') % 2 != 0 and ':' in last_line:
                    # Extraire la clé et remplacer la valeur par null
                    key_match = re.search(r'(\s*"[^"]+"\s*):\s*"[^"]*$', last_line)
                    if key_match:
                        lines[-1] = key_match.group(1) + ': null'
                    result = '\n'.join(lines)
            
            return result
        
        # Extraire le JSON de la réponse
        json_match = re.search(r'\{.*', cleaned_response, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            
            # Essayer de parser directement
            try:
                structured_data = json.loads(json_str)
            except json.JSONDecodeError:
                # Essayer de réparer le JSON
                logger.warning("⚠️ JSON incomplet détecté, tentative de réparation...")
                repaired_json = repair_json(json_str)
                try:
                    structured_data = json.loads(repaired_json)
                    logger.info("✅ JSON réparé avec succès")
                except json.JSONDecodeError as repair_error:
                    logger.error(f"❌ Impossible de réparer le JSON: {str(repair_error)}")
                    # Essayer d'extraire au moins les parties valides
                    try:
                        # Extraire juste le début du JSON jusqu'à la première erreur
                        # Trouver la position de l'erreur
                        error_pos = int(str(repair_error).split('char ')[1].split(')')[0]) if 'char' in str(repair_error) else len(json_str)
                        partial_json = json_str[:error_pos]
                        # Fermer proprement
                        partial_json = repair_json(partial_json)
                        structured_data = json.loads(partial_json)
                        logger.warning("⚠️ JSON partiel extrait (certaines données peuvent être manquantes)")
                    except:
                        raise repair_error
        else:
            # Essayer de parser directement
            structured_data = json.loads(cleaned_response)
        
        # Ajouter des métadonnées
        structured_data['metadata'] = {
            'file_name': file_name,
            'document_type': document_type,
            'extraction_date': datetime.utcnow().isoformat()
        }
        
        logger.info(f"✅ Extraction structurée réussie pour {file_name} (type: {document_type})")
        return structured_data
        
    except json.JSONDecodeError as e:
        logger.error(f"❌ Erreur de parsing JSON lors de l'extraction: {str(e)}")
        logger.error(f"Réponse brute (premiers 2000 caractères): {raw_response[:2000]}")
        logger.error(f"Longueur totale de la réponse: {len(raw_response)} caractères")
        # Retourner une structure minimale avec l'erreur
        return {
            "error": "Erreur lors de l'extraction des données structurées - JSON invalide ou incomplet",
            "error_details": str(e),
            "raw_response_preview": raw_response[:2000],
            "response_length": len(raw_response),
            "metadata": {
                "file_name": file_name,
                "document_type": document_type
            }
        }
    except Exception as e:
        logger.error(f"❌ Erreur lors de l'extraction structurée: {str(e)}")
        return {
            "error": f"Erreur lors de l'extraction: {str(e)}",
            "metadata": {
                "file_name": file_name,
                "document_type": document_type
            }
        }
    
async def analyze_query_need_for_search(user_query: str, selected_model: str, language: str = 'en') -> Dict:
    """
    Analyse si la query nécessite des informations actuelles
    Amélioré pour détecter les questions nécessitant des informations en temps réel
    """
    t = translations.get(language, translations['en'])
    
    # Détection rapide basée sur des mots-clés pour les questions en temps réel
    query_lower = user_query.lower()
    realtime_keywords = [
        "heure actuelle", "heure maintenant", "quelle heure", "what time", "current time", "time now",
        "prix actuel", "prix maintenant", "current price", "price now",
        "taux actuel", "current rate", "taux de change", "exchange rate",
        "cours actuel", "current rate",
        "maintenant", "now", "actuel", "current", "aujourd'hui", "today",
        "en ce moment", "right now", "à l'instant", "at the moment",
        "météo", "weather", "température", "temperature",
        "actualité", "news", "événements récents", "recent events"
    ]
    
    # Si la question contient des mots-clés de temps réel, déclencher directement la recherche
    needs_realtime = any(keyword in query_lower for keyword in realtime_keywords)
    
    if needs_realtime:
        logger.info(f"🔍 Question en temps réel détectée: {user_query[:50]}...")
        return {
            "needs_search": True,
            "reason": "Question nécessitant des informations en temps réel",
            "search_keywords": [user_query],
            "estimated_cutoff_relevance": "high"
        }
    
    needs_search_field = t['needs_search_field']
    reason_field = t['reason_field']
    search_keywords_field = t['search_keywords_field']
    cutoff_relevance_field = t['cutoff_relevance_field']
    
    analysis_prompt = (
        f"{t['analyze_query_prompt']}\n\n"
        f"{t['query_label']}: {user_query}\n\n"
        f"{t['json_response_format']}:\n"
        f"{{\n"
        f'  "{needs_search_field}": true/false,\n'
        f'  "{reason_field}": "{t["short_explanation"]}",\n'
        f'  "{search_keywords_field}": ["mot-clé1", "mot-clé2"] {t["keyword_or_null"]},\n'
        f'  "{cutoff_relevance_field}": "high/medium/low"\n'
        f"}}\n\n"
        f"{t['examples_need_search']}:\n"
        f"{t['current_prices_stocks']}\n"
        f"{t['recent_events_news']}\n"
        f"{t['new_software_versions']}\n"
        f"{t['recent_statistics_data']}\n"
        f"{t['recent_people_companies']}\n"
        f"- Questions sur l'heure actuelle dans un pays ou une ville\n"
        f"- Questions sur les prix actuels, taux de change, cours boursiers\n"
        f"- Questions sur la météo actuelle\n"
        f"- Questions sur les événements récents ou l'actualité\n\n"
        f"{t['examples_no_search']}:\n"
        f"{t['general_concepts']}\n"
        f"{t['programming_syntax']}\n"
        f"{t['history_facts']}\n"
        f"{t['math_science']}"
    )
    
    try:
        analysis_result = await execute_model_query(analysis_prompt, selected_model)
        # Nettoyer la réponse pour extraire le JSON
        json_match = re.search(r'\{.*\}', analysis_result, re.DOTALL)
        if json_match:
            analysis_data = json.loads(json_match.group())
            return analysis_data
        else:
            # Fallback si le JSON n'est pas parsable
            return {
                "needs_search": True,
                "reason": "Analyse automatique non concluante",
                "search_keywords": [user_query],
                "estimated_cutoff_relevance": "medium"
            }
    except Exception as e:
        logger.warning(f"Erreur lors de l'analyse de la query: {str(e)}")
        return {
            "needs_search": True,
            "reason": "Erreur d'analyse, recherche par sécurité",
            "search_keywords": [user_query],
            "estimated_cutoff_relevance": "medium"
        }
    
async def query_model_online_mode(user_query: str, selected_model: str = DEFAULT_MODEL, 
                                 language: str = 'en', enable_auto_online_search: bool = True) -> str:
    """
    Mode ONLINE: Réponse du modèle enrichie avec des données actuelles si nécessaire (comme SearchGPT)
    Avec recherche automatique si l'information n'est pas trouvée dans les connaissances du modèle
    """
    cache_key = f"online_query_{user_query}_{selected_model}_{language}"
    cached = await cache.get(cache_key)
    if cached:
        return cached

    t = translations.get(language, translations['en'])

    # ÉTAPE 1: Réponse initiale du modèle avec ses connaissances
    initial_prompt = (
        f"{t['online_mode_title']}\n"
        f"{t['recent_info_mention']}\n\n"
        f"{t['question']}: {user_query}\n\n"
        f"{t['online_instructions_title']}\n"
        f"{t['give_best_answer']}\n"
        f"{t['be_precise_dates']}\n"
        f"{t['mention_recent_useful']}\n\n"
        f"{t['emoji']}"
    )

    try:
        # Réponse initiale du modèle
        initial_response = await execute_model_query(initial_prompt, selected_model)
        
        # ÉTAPE 2: Analyse automatique du besoin de recherche
        search_analysis = await analyze_query_need_for_search(user_query, selected_model, language)
        
        logger.info(f"🔍 Analyse de recherche: {search_analysis}")
        
        # ÉTAPE 3: Recherche et enrichissement si nécessaire
        should_search = search_analysis.get("needs_search", False)
        
        # NOUVEAU: Vérifier aussi si la réponse indique que l'information n'est pas disponible
        if enable_auto_online_search and not should_search:
            missing_info = await detect_missing_information(initial_response, user_query, language)
            if missing_info:
                logger.info("🔍 Information absente détectée dans les connaissances du modèle. Recherche en ligne automatique...")
                should_search = True
        
        if should_search:
            logger.info("🌐 Enrichissement avec des données actuelles...")
            
            # Utiliser les mots-clés optimisés ou la query originale
            search_keywords = search_analysis.get("search_keywords", [user_query])
            search_query = " ".join(search_keywords) if isinstance(search_keywords, list) else user_query
            
            # Recherche en ligne
            try:
                search_results = perform_online_search(search_query, language)
                
                if search_results and search_results != "Aucun résultat trouvé." and search_results != "Clé API manquante":
                    # ÉTAPE 4: Fusion intelligente des informations avec formatage amélioré
                    enrichment_prompt = (
                        f"{t['enrichment_title']}\n\n"
                        f"{t['initial_response']}\n{initial_response}\n\n"
                        f"{t['new_info_found']}\n{search_results}\n\n"
                        f"{t['enrichment_instructions']}\n"
                        f"{t['combine_intelligently']}\n"
                        f"{t['update_obsolete']}\n"
                        f"{t['distinguish_info']}\n"
                        f"{t['cite_sources']}\n"
                        f"{t['keep_structure']}\n"
                        f"{t['prioritize_recent']}\n\n"
                        f"**FORMATAGE SPÉCIAL POUR CERTAINES RÉPONSES:**\n"
                        f"- Pour les questions sur l'heure: Réponds directement avec l'heure actuelle au format 'Il est XXhXX à [lieu]' ou 'It is XX:XX in [place]'\n"
                        f"- Pour les prix (cryptomonnaies, actions, devises): Formate comme une carte avec le prix en gras, le changement en pourcentage, et une description claire\n"
                        f"- Pour les taux de change: Affiche le taux actuel de manière claire et structurée\n"
                        f"- Pour la météo: Formate avec température, conditions, et lieu\n"
                        f"- Utilise des emojis appropriés (🕐 pour l'heure, 💰 pour les prix, 🌤️ pour la météo, etc.)\n"
                        f"- Sois concis et précis, va droit au but\n"
                        f"- Si les données de recherche sont claires, utilise-les directement sans trop d'explications\n\n"
                        f"{t['enriched_response']}"
                    )
                    
                    enriched_response, _ = await execute_model_query_with_fallback(enrichment_prompt, selected_model, user_query)
                    
                    # Ajout des métadonnées de recherche
                    final_response = (
                        f"{enriched_response}\n\n"
                        f"---\n"
                        f"💡 **Informations enrichies**: Cette réponse combine mes connaissances de base "
                        f"avec des données récentes trouvées en ligne.\n \n"
                        f"🔍 **Model**: {selected_model}"
                    )
                else:
                    # Pas de résultats en ligne, garder la réponse initiale
                    logger.info("⚠️ Aucun résultat trouvé en ligne ou clé API manquante")
                    final_response = (
                        f"{initial_response}\n\n"
                        f"{t['separator']}\n"
                        f"{t['source_training']}"
                    )
            except Exception as search_error:
                logger.warning(f"⚠️ Erreur lors de la recherche en ligne: {str(search_error)}")
                # En cas d'erreur, garder la réponse initiale
                final_response = (
                    f"{initial_response}\n\n"
                    f"{t['separator']}\n"
                    f"{t['source_training']}"
                )
            
        else:
            logger.info("✅ Pas de recherche nécessaire, réponse basée sur les connaissances du modèle")
            final_response = (
                f"{initial_response}\n\n"
                f"{t['separator']}\n"
                f"{t['source_training']}"
            )

    except Exception as e:
        logger.error(f"Erreur lors du mode online: {str(e)}")
        final_response = f"Erreur lors du traitement en mode online: {str(e)}"

    await cache.set(cache_key, final_response, ttl=1800)
    return final_response

@app.route('/index-directory', methods=['POST'])
async def index_directory():
    """
    Indexe les fichiers d'un repertoire avec OpenAI embeddings
    """
    try:
        data = request.get_json()
        files = data.get('files', [])
        session_id = request.headers.get('Session-ID', 'default')
        language = data.get('language', 'en')

        if not files:
            return jsonify({
                "error": "Aucun fichier fourni pour l'indexation",
                'files_received': len(files)
            }), 400
        
        logger.info(f"📂 Indexation de {len(files)} fichiers pour la session {session_id}")
        documents = []
        images_processed_count = 0
        ocr_enabled = os.getenv('ENABLE_OCR', 'false').lower() == 'true'
        
        for file_data in files:
            if not file_data.get('content') or not file_data.get('fileName'):
                logger.warning(f"Fichier ignoré (contenu ou nom manquant): {file_data.get('fileName', 'inconnu')}")
                continue

            # Extraire le contenu texte principal
            main_content = file_data['content']
            has_images = file_data.get('hasImages', False)
            images = file_data.get('images', [])
            metadata_info = file_data.get('metadata', {})
            
            # Si le fichier contient des images, essayer d'extraire le texte des images (OCR optionnel)
            image_texts = []
            file_images_processed = 0
            if has_images and images and ocr_enabled:
                logger.info(f"🖼️ Traitement de {len(images)} image(s) pour {file_data['fileName']}")
                try:
                    # Importer les bibliothèques OCR si disponibles
                    try:
                        from PIL import Image
                        import io
                        import base64
                        
                        # Essayer d'utiliser TrOCR ou pytesseract pour l'OCR
                        ocr_available = False
                        processor = None
                        model = None
                        
                        try:
                            from transformers import TrOCRProcessor, VisionEncoderDecoderModel
                            processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-printed")
                            model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-printed")
                            ocr_available = True
                            logger.info("✅ TrOCR disponible pour l'extraction de texte des images")
                        except (ImportError, Exception) as e:
                            ocr_available = False
                            logger.warning(f"⚠️ TrOCR non disponible, OCR désactivé pour les images: {str(e)}")
                        
                        if ocr_available and processor and model:
                            for img_idx, img_data in enumerate(images):
                                try:
                                    if img_data.get('dataUri'):
                                        # Extraire les données base64
                                        data_uri = img_data['dataUri']
                                        if data_uri.startswith('data:'):
                                            base64_data = data_uri.split(',')[1]
                                            image_bytes = base64.b64decode(base64_data)
                                            image = Image.open(io.BytesIO(image_bytes))
                                            
                                            # Utiliser TrOCR pour extraire le texte
                                            pixel_values = processor(image, return_tensors="pt").pixel_values
                                            generated_ids = model.generate(pixel_values)
                                            generated_text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
                                            
                                            if generated_text.strip():
                                                image_texts.append(f"[Image {img_idx + 1}: {generated_text.strip()}]")
                                                file_images_processed += 1
                                                images_processed_count += 1
                                                logger.info(f"✅ Texte extrait de l'image {img_idx + 1}: {generated_text[:50]}...")
                                except Exception as img_error:
                                    logger.warning(f"⚠️ Erreur lors du traitement de l'image {img_idx + 1}: {str(img_error)}")
                                    continue
                    except ImportError:
                        logger.warning("⚠️ PIL/Pillow non disponible, OCR désactivé pour les images")
                        
                except Exception as ocr_error:
                    logger.warning(f"⚠️ Erreur lors de l'initialisation OCR: {str(ocr_error)}")
            
            # Combiner le contenu principal avec les textes extraits des images
            if image_texts:
                main_content += "\n\n=== TEXTE EXTRAIT DES IMAGES ===\n" + "\n".join(image_texts)
            
            # Construire les métadonnées enrichies
            document_metadata = {
                    'fileName': file_data['fileName'],
                'file_size': metadata_info.get('file_size', len(main_content)),
                'word_count': metadata_info.get('wordCount'),
                'page_count': metadata_info.get('pageCount'),
                'has_scanned_content': metadata_info.get('hasScannedContent', False),
                'scanned_pages': metadata_info.get('scannedPages'),  # Pages scannées (PDF)
                'invoice_pages': metadata_info.get('invoicePages'),  # Pages de factures (PDF)
                'has_images': has_images,
                'images_count': len(images) if images else 0,
                'images_ocr_processed': file_images_processed,
                    'indexed_at': datetime.utcnow().isoformat()
                }
            
            # Ajouter des métadonnées sur les images si disponibles
            if images:
                document_metadata['image_descriptions'] = [
                    img.get('description', f"Image {idx + 1}") 
                    for idx, img in enumerate(images)
                ]
                # Ajouter les informations de page pour les images
                images_with_pages = [img for img in images if img.get('pageNumber')]
                if images_with_pages:
                    images_by_page_dict = {}
                    for img in images_with_pages:
                        page_num = str(img.get('pageNumber'))
                        if page_num not in images_by_page_dict:
                            images_by_page_dict[page_num] = []
                        images_by_page_dict[page_num].append({
                            'id': img.get('id'),
                            'description': img.get('description')
                        })
                    document_metadata['images_by_page'] = images_by_page_dict
            
            # Ajouter des informations sur les pages si disponibles
            pages_info = file_data.get('pages', [])
            if pages_info:
                document_metadata['pages_info'] = [
                    {
                        'page_number': page.get('pageNumber'),
                        'word_count': page.get('wordCount'),
                        'char_count': page.get('charCount'),
                        'has_images': page.get('hasImages', False),
                        'images_count': len(page.get('images', [])) if page.get('images') else 0
                    }
                    for page in pages_info
                ]
                
                # Ajouter des chunks par page pour permettre la recherche granulaire par page
                # Optionnel: permet de rechercher et référencer une page spécifique
                if len(pages_info) > 1:
                    logger.info(f"📄 Document {file_data['fileName']} divisé en {len(pages_info)} pages pour indexation granulaire")
                    
                    # Créer des documents additionnels par page pour recherche granulaire
                    # Le document principal reste inchangé, mais on ajoute des références par page
                    for page in pages_info:
                        page_content = page.get('content', '')
                        if page_content:
                            # Ajouter un préfixe pour identifier la page dans le contenu
                            page_document = Document(
                                page_content=f"[Page {page.get('pageNumber')} de {file_data['fileName']}]\n\n{page_content}",
                                metadata={
                                    **document_metadata,
                                    'page_number': page.get('pageNumber'),
                                    'is_page_chunk': True,
                                    'original_file': file_data['fileName']
                                }
                            )
                            documents.append(page_document)

            documents.append(Document(
                page_content=main_content,
                metadata=document_metadata
            ))
        if not documents:
            return jsonify({
                "error": "Aucun fichier valide à indexer",
                'files_processed':0
            }), 400
        
        logger.info(f"🗂️ {len(documents)} documents prêts pour l'indexation")

        # Diviser les documents en chunks avec stratégie améliorée pour préserver le contexte
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # Augmenté pour préserver plus de contexte
            chunk_overlap=150,  # Plus de chevauchement pour préserver les phrases complètes
            length_function=len,
            separators=["\n\n\n", "\n\n", "\n", ". ", " ", ""]  # Meilleure préservation des paragraphes
        )
        split_docs = text_splitter.split_documents(documents)
        logger.info(f"✂️ Documents divisés en {len(split_docs)} chunks (chunk_size=2000, overlap=300)")
        
        # Vérifier la clé OpenAI
        openai_api_key = os.getenv('OPENAI_API_KEY')
        if not openai_api_key:
            return jsonify({
                'error': 'Clé API OpenAI non configurée sur le serveur',
                'suggestion': 'Contactez l\'administrateur pour configurer OPENAI_API_KEY'
            }), 500
        
        # Créer les embeddings avec un modèle plus récent et performant
        logger.info("🧠 Création des embeddings avec OpenAI...")
        # Utiliser text-embedding-3-small ou text-embedding-ada-002 selon disponibilité
        try:
            # Essayer d'abord le modèle plus récent (meilleure qualité sémantique)
            embeddings = OpenAIEmbeddings(
                api_key=openai_api_key,
                model="text-embedding-3-large"  # Plus performant que ada-002
            )
            logger.info("✅ Utilisation de text-embedding-3-large pour meilleure qualité sémantique")
        except Exception as e:
            logger.warning(f"text-embedding-3-large non disponible, fallback vers ada-002: {str(e)}")
            embeddings = OpenAIEmbeddings(
                api_key=openai_api_key,
                model="text-embedding-ada-002"
            )
        
        # Test rapide de l'API
        try:
            test_embedding = await embeddings.aembed_query("test")
            logger.info(f"✅ API OpenAI fonctionnelle, dimension: {len(test_embedding)}")
        except Exception as e:
            logger.error(f"❌ Erreur de test API OpenAI: {str(e)}")
            return jsonify({
                'error': f'Erreur API OpenAI: {str(e)}',
                'suggestion': 'Vérifiez votre clé API et votre quota'
            }), 500
        
        # Inférer des actions suggérées pour ce corpus
        logger.info("🧠 Inférence des actions suggérées pour le corpus...")
        language = data.get('language', 'en')
        inferred_actions = await infer_corpus_actions(split_docs, language=language)
        
        # Créer le vector store
        logger.info("🗃️ Création du vector store...")
        vector_store = await FAISS.afrom_documents(split_docs, embeddings)
        
        # Stocker le vector store et les actions pour cette session
        vector_stores[session_id] = {
            'store': vector_store,
            'created_at': datetime.utcnow().isoformat(),
            'files_count': len(files),
            'chunks_count': len(split_docs),
            'files_indexed': [f['fileName'] for f in files],
            'auto_actions': inferred_actions
        }
        
        logger.info(f"✅ Indexation terminée pour la session {session_id}")
        if images_processed_count > 0:
            logger.info(f"🖼️ {images_processed_count} image(s) traitée(s) avec OCR")
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'indexed_files_count': len(files),
            'chunks_count': len(split_docs),
            'files_indexed': [f['fileName'] for f in files],
            'vector_store_ready': True,
            'suggested_actions': inferred_actions.get('suggested_actions', []),
            'corpus_domain': inferred_actions.get('domain', 'unknown'),
            'images_ocr_processed': images_processed_count,
            'ocr_enabled': ocr_enabled
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Erreur lors de l'indexation: {str(e)}")
        return jsonify({
            'error': f'Erreur lors de l\'indexation: {str(e)}',
            'success': False
        }), 500


@app.route('/infer-corpus-actions', methods=['POST'])
async def infer_corpus_actions_endpoint():
    """
    Endpoint pour inférer les actions suggérées à partir de documents
    """
    try:
        data = request.get_json()
        documents_data = data.get('documents', [])
        language = data.get('language', 'en')
        
        if not documents_data:
            return jsonify({
                "error": "No documents provided",
                'success': False
            }), 400
        
        # Convertir les documents en format Document
        documents = []
        for doc_data in documents_data:
            documents.append(Document(
                page_content=doc_data.get('content', ''),
                metadata={
                    'fileName': doc_data.get('file_name', 'document'),
                    'source': doc_data.get('file_name', 'document')
                }
            ))
        
        # Inférer les actions suggérées
        inferred_actions = await infer_corpus_actions(documents, language=language)
        
        return jsonify({
            'success': True,
            'suggested_actions': inferred_actions.get('suggested_actions', []),
            'domain': inferred_actions.get('domain', 'unknown')
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Erreur lors de l'inférence des actions: {str(e)}")
        return jsonify({
            'error': f'Error inferring actions: {str(e)}',
            'success': False
        }), 500


@app.route('/query', methods=['POST'])
async def handle_query():
    """
    Route de requête améliorée avec support des modes local et online
    """
    data = request.get_json()
    
    # Extraction des paramètres
    user_query = data.get('user_query')
    research_mode = data.get('research_mode', 'local')
    selected_model = data.get('selected_model', DEFAULT_MODEL)
    language = data.get('language', 'en')
    session_id = request.headers.get('Session-ID', 'default')  # NOUVEAU: récupérer session ID
    conversation_history = data.get('conversation_history', [])
    
    # Paramètres de contrôle des modes
    disable_online_search = data.get('disable_online_search', False)
    enable_online_search = data.get('enable_online_search', False)
    use_backend_vectorstore = data.get('use_backend_vectorstore', False)  # NOUVEAU
    
    # Validation des paramètres essentiels
    if not user_query:
        return jsonify({"error": "user_query est requis"}), 400
    
    # Validation et nettoyage du modèle
    if selected_model not in MODEL_CONFIG:
        logger.warning(f"Modèle inconnu {selected_model}, utilisation du modèle par défaut {DEFAULT_MODEL}")
        selected_model = DEFAULT_MODEL
    
    try:
        # Détermination du mode effectif
        effective_mode = research_mode
        
        # MODE LOCAL
        if effective_mode == 'local' or disable_online_search:
            logger.info(f"🔒 Mode LOCAL activé pour la requête: {user_query[:50]}...")
            
            # Validation des paramètres pour le mode local
            file_name = data.get('file_name')
            file_content = data.get('file_content', '')
            directory_content = data.get('directory_content', [])
            repo_structure = data.get('repo_structure', '')
            is_binary = data.get('is_binary', False)
            
            # ✅ Directory/corpus mode (répertoire importé)
            # Si un vector store existe pour cette session, autoriser les requêtes même sans file_name
            # Si use_backend_vectorstore est True OU si un vector store existe pour la session, on peut procéder
            has_vector_store = session_id in vector_stores
            should_use_vectorstore = use_backend_vectorstore or has_vector_store
            
            if not file_name:
                if should_use_vectorstore and has_vector_store:
                    file_name = "__DIRECTORY_CORPUS__"
                    file_content = ""
                    # on laisse directory_content se remplir depuis le vector store ci-dessous
                    is_binary = False
                    use_backend_vectorstore = True  # Forcer à True puisque le vector store existe
                    logger.info(f"📚 Mode CORPUS (répertoire) activé: query sur l'ensemble des documents, session={session_id}")
                else:
                    return jsonify({
                        "error": "Mode local nécessite un fichier sélectionné ou un répertoire indexé",
                        "mode": "local",
                        "session_id": session_id,
                        "has_vector_store": has_vector_store,
                        "use_backend_vectorstore": use_backend_vectorstore,
                        "suggestion": "Sélectionnez un fichier, ou importez un répertoire pour l'indexer"
                    }), 400
            
            # NOUVELLE LOGIQUE: Recherche sémantique améliorée avec recherche hybride (sémantique + mots-clés)
            relevant_docs: List[Document] = []
            # Utiliser le vector store si disponible (que use_backend_vectorstore soit True ou False, si le store existe on l'utilise)
            if should_use_vectorstore and has_vector_store:
                try:
                    session_store = vector_stores.get(session_id)
                    if not session_store or not isinstance(session_store, dict) or 'store' not in session_store or session_store.get('store') is None:
                        logger.warning(f"Vector store non disponible pour session {session_id}")
                    else:
                        session_vector_store = session_store['store']
                        logger.info(f"🔍 Hybrid retrieval (semantic + lexical) dans le vector store session {session_id}")

                        # Hierarchical retrieval strategy:
                        # 1) If a file is selected: search in that file first, then in whole repo if needed
                        # 2) If no file selected: search directly in whole repo until meaningful answer
                        file_first_docs: List[Document] = []
                        file_first_debug: Dict[str, Any] = {}
                        corpus_docs: List[Document] = []
                        corpus_debug: Dict[str, Any] = {}
                        
                        if file_name and file_name != "__DIRECTORY_CORPUS__":
                            # STEP 1: Search in selected file first
                            logger.info(f"📄 Recherche dans le fichier sélectionné '{file_name}'...")
                            file_first_docs, file_first_debug = hybrid_retrieve_documents(
                                vector_store=session_vector_store,
                                query=user_query,
                                k_candidates=80,  # Increased for better recall within file
                                k_final=10,  # Increased from 8 to 10
                                semantic_weight=0.55,
                                bm25_weight=0.35,
                                exact_weight=0.10,
                                preferred_sources=[file_name],
                            )
                            logger.info(f"📊 Fichier: {len(file_first_docs)} résultats trouvés")
                            
                            # STEP 2: Toujours chercher aussi dans le corpus pour compléter
                            # Si l'info n'est pas dans le fichier, elle sera dans le corpus
                            logger.info(f"📚 Recherche complémentaire dans le corpus du répertoire...")
                            corpus_result = hybrid_retrieve_documents(
                                vector_store=session_vector_store,
                                query=user_query,
                                k_candidates=120,  # Increased for better recall
                                k_final=15,  # Increased from 12 to 15
                                semantic_weight=0.55,
                                bm25_weight=0.35,
                                exact_weight=0.10,
                            )
                            if isinstance(corpus_result, tuple) and len(corpus_result) == 2:
                                corpus_docs, corpus_debug = corpus_result
                                if not isinstance(corpus_debug, dict):
                                    corpus_debug = {}
                            else:
                                corpus_docs = []
                                corpus_debug = {}
                            logger.info(f"📚 Corpus: {len(corpus_docs)} résultats trouvés")
                        else:
                            # No file selected: use the advanced search_semantic_documents_sync
                            # which includes person name filtering, function name search, filename matching, etc.
                            logger.info(f"📚 No file selected: Using advanced search in whole repository...")
                            directory_content = search_semantic_documents_sync(
                                session_vector_store, 
                                user_query, 
                                session_id, 
                                conversation_history
                            )
                            # Convert List[Dict] to List[Document] for consistency with file_first logic
                            corpus_docs = []
                            for doc_dict in directory_content:
                                # Create a Document object from the dict
                                from langchain.schema import Document
                                doc = Document(
                                    page_content=doc_dict.get('content', ''),
                                    metadata={
                                        'fileName': doc_dict.get('fileName', ''),
                                        'file_name': doc_dict.get('fileName', ''),
                                        'pageNumber': doc_dict.get('pageNumber'),
                                        'is_page_chunk': doc_dict.get('isPageChunk', False)
                                    }
                                )
                                corpus_docs.append(doc)
                            corpus_debug = {}
                            logger.info(f"📚 Advanced repository search: {len(corpus_docs)} docs found")

                        # Merge preserving order: file_first_docs first (if any), then corpus_docs
                        # Dedupe by (source, preview)
                        merged: List[Document] = []
                        seen = set()
                        
                        # Add file documents first (priority)
                        for d in file_first_docs:
                            src = (d.metadata.get("source") or d.metadata.get("fileName") or d.metadata.get("file_name") or "").strip().lower()
                            key = (src, (d.page_content or "")[:120])
                            if key not in seen:
                                seen.add(key)
                                merged.append(d)
                        
                        # Then add corpus documents (excluding duplicates)
                        for d in corpus_docs:
                            src = (d.metadata.get("source") or d.metadata.get("fileName") or d.metadata.get("file_name") or "").strip().lower()
                            key = (src, (d.page_content or "")[:120])
                            if key not in seen:
                                seen.add(key)
                                merged.append(d)
                                # Limit total results
                                if len(merged) >= 20:
                                    break

                        relevant_docs = merged
                        logger.info(
                            f"🎯 Hierarchical search completed: {len(relevant_docs)} total docs | "
                            f"from_file={len(file_first_docs)} | "
                            f"from_corpus={len([d for d in merged if d not in file_first_docs])}"
                        )
                        
                        # Ajout des documents pertinents au contexte avec métadonnées enrichies
                        # Éviter les doublons en utilisant un set de file_name + début du contenu
                        if not directory_content:  # Only populate if not already filled by search_semantic_documents_sync
                            seen_docs = set()
                            for doc in relevant_docs:
                                file_name_from_meta = doc.metadata.get("fileName") or doc.metadata.get("file_name") or "document_vectorstore"
                                content_preview = doc.page_content[:100]  # Premier 100 chars pour détecter les doublons
                                doc_key = f"{file_name_from_meta}:{content_preview}"
                                
                                if doc_key not in seen_docs:
                                    seen_docs.add(doc_key)
                                    page_number = doc.metadata.get("page_number")
                                    result_dict = {
                                        "fileName": file_name_from_meta, 
                                        "content": doc.page_content[:2500]  # Augmenté à 2500 chars par chunk
                                    }
                                    if page_number is not None:
                                        result_dict["pageNumber"] = page_number
                                        result_dict["isPageChunk"] = doc.metadata.get("is_page_chunk", False)
                                    directory_content.append(result_dict)
                            
                            logger.info(f"📚 {len(directory_content)} documents uniques ajoutés depuis le vector store de session")
                        else:
                            logger.info(f"📚 {len(directory_content)} documents ajoutés via search_semantic_documents_sync")
                    
                except Exception as e:
                    logger.warning(f"Erreur lors de la récupération depuis le vector store: {str(e)}")
            
            elif use_backend_vectorstore:
                logger.warning(f"Vector store demandé mais non trouvé pour la session {session_id}")
            
            # ANCIENNE LOGIQUE: Utiliser le vector store global (fallback)
            elif vector_store and user_query:
                try:
                    logger.info(f"🔍 Recherche sémantique améliorée dans le vector store global")
                    search_results_with_scores = vector_store.similarity_search_with_score(
                        user_query, 
                        k=15
                    )
                    
                    filtered_docs = []
                    for doc, distance in search_results_with_scores:
                        similarity_score = 1 - distance
                        if similarity_score >= 0.65:
                            filtered_docs.append((doc, similarity_score))
                    
                    filtered_docs.sort(key=lambda x: x[1], reverse=True)
                    top_docs = [doc for doc, score in filtered_docs[:10]]
                    
                    if len(top_docs) > 5:
                        try:
                            reranked_docs = await rerank_documents_with_llm(
                                user_query, 
                                top_docs, 
                                selected_model
                            )
                            relevant_docs = reranked_docs[:8]
                        except Exception as rerank_error:
                            logger.warning(f"Re-ranking échoué: {str(rerank_error)}")
                            relevant_docs = top_docs[:8]
                    else:
                        relevant_docs = top_docs
                    
                    for doc in relevant_docs:
                        directory_content.append({ 
                            "fileName": doc.metadata.get("file_name", "document_vectorstore"), 
                            "content": doc.page_content[:2000]
                        })
                    logger.info(f"📚 {len(relevant_docs)} documents ajoutés depuis le vector store global")
                except Exception as e:
                    logger.warning(f"Erreur lors de la récupération depuis le vector store global: {str(e)}")
            
            # Exécution de la requête en mode local - recherche en ligne DÉSACTIVÉE par défaut
            enable_auto_search = data.get('enable_auto_online_search', False)  # Désactivé par défaut en mode local
            response, actual_model_used = await query_model_local_mode(
                file_name=file_name,
                file_content=file_content,
                directory_content=directory_content,
                repo_structure=repo_structure,
                user_query=user_query,
                is_binary=is_binary,
                selected_model=selected_model,
                language=language,
                conversation_history=conversation_history,
                enable_auto_online_search=enable_auto_search
            )
            
            # Extraire les informations de page des documents utilisés
            pages_used = []
            for doc in directory_content:
                if isinstance(doc, dict) and doc.get('pageNumber') is not None:
                    pages_used.append({
                        "fileName": doc.get('fileName'),
                        "pageNumber": doc.get('pageNumber')
                    })
            
            return jsonify({
                "response": response,
                "mode": "local",
                "model_used": actual_model_used,
                "model_config": MODEL_CONFIG.get(selected_model, {}),
                "context_info": {
                    "file_name": file_name,
                    "directory_files_count": len(directory_content),
                    "vector_store_docs": len(relevant_docs),
                    "is_binary": is_binary,
                    "session_id": session_id,
                    "used_backend_vectorstore": use_backend_vectorstore,
                    "pages_referenced": pages_used if pages_used else None
                }
            }), 200
        
        # MODE ONLINE avec recherche automatique si l'info n'est pas trouvée
        elif effective_mode == 'online' or enable_online_search:
            logger.info(f"🌐 Mode ONLINE activé pour la requête: {user_query[:50]}...")
            
            enable_auto_search = data.get('enable_auto_online_search', True)  # Activé par défaut
            response = await query_model_online_mode(
                user_query=user_query,
                selected_model=selected_model,
                language=language,
                enable_auto_online_search=enable_auto_search
            )
            
            return jsonify({
                "response": response,
                "mode": "online",
                "model_used": selected_model,
                "model_config": MODEL_CONFIG.get(selected_model, {}),
                "search_info": {
                    "query": user_query,
                    "language": language,
                    "search_performed": "intelligent_enrichment",
                    "approach": "searchgpt_style"
                }
            }), 200
        
        else:
            return jsonify({
                "error": f"Mode de recherche non supporté: {effective_mode}",
                "supported_modes": ["local", "online"]
            }), 400
            
    except Exception as e:
        logger.error(f"Erreur lors du traitement de la requête: {str(e)}")
        return jsonify({
            "error": f"Erreur lors du traitement: {str(e)}",
            "mode": effective_mode,
            "model_used": selected_model
        }), 500

@app.route('/query-stream', methods=['POST'])
def handle_query_stream():
    """
    Version streaming de la route query pour réponses en temps réel
    """
    data = request.get_json()
    
    # Récupérer les paramètres
    user_query = data.get('user_query')
    research_mode = data.get('research_mode', 'local')
    selected_model = data.get('selected_model', DEFAULT_MODEL)
    language = data.get('language', 'en')
    session_id = request.headers.get('Session-ID', 'default')
    
    if not user_query:
        return jsonify({"error": "user_query est requis"}), 400
    
    # Validation du modèle
    if selected_model not in MODEL_CONFIG:
        logger.warning(f"Unknown model {selected_model}, using default {DEFAULT_MODEL}")
        selected_model = DEFAULT_MODEL
    
    logger.info(f"🚀 Streaming request - Model: {selected_model}, Mode: {research_mode}")
    
    try:
        # MODE ONLINE
        if research_mode == 'online' or data.get('enable_online_search'):
            t = translations.get(language, translations['en'])
            
            prompt = (
                f"{t['online_mode_title']}\n"
                f"{t['recent_info_mention']}\n\n"
                f"{t['question']}: {user_query}\n\n"
                f"{t['online_instructions_title']}\n"
                f"{t['give_best_answer']}\n"
                f"{t['be_precise_dates']}\n"
                f"{t['mention_recent_useful']}\n\n"
                f"{t['emoji']}"
            )
            
            return Response(
                stream_with_context(stream_response(prompt, selected_model)),
                mimetype='text/event-stream',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'X-Accel-Buffering': 'no'
                }
            )
        
        # MODE LOCAL
        elif research_mode == 'local':
            file_name = data.get('file_name')
            file_content = data.get('file_content', '')
            directory_content = data.get('directory_content', [])
            repo_structure = data.get('repo_structure', '')
            use_backend_vectorstore = data.get('use_backend_vectorstore', False)
            conversation_history = data.get('conversation_history', [])
            
            if not file_name:
                return jsonify({"error": "Mode local nécessite un file_name"}), 400
            
            # Si le vector store backend est disponible, utiliser la recherche sémantique améliorée avec stratégie hiérarchique
            if use_backend_vectorstore and session_id in vector_stores:
                try:
                    session_store = vector_stores.get(session_id)
                    if not session_store or not isinstance(session_store, dict) or 'store' not in session_store or session_store.get('store') is None:
                        logger.warning(f"Vector store non disponible pour session {session_id}")
                    else:
                        session_vector_store = session_store['store']
                        
                        # Hierarchical search strategy for streaming:
                        # If file_name is provided and not __DIRECTORY_CORPUS__, search file first, then corpus
                        file_first_results = []
                        corpus_results = []
                        
                        if file_name and file_name != "__DIRECTORY_CORPUS__":
                            # Step 1: Search in selected file first
                            logger.info(f"📄 [Stream] Recherche dans le fichier sélectionné '{file_name}'...")
                            file_first_docs, _ = hybrid_retrieve_documents(
                                vector_store=session_vector_store,
                                query=user_query,
                                k_candidates=80,
                                k_final=10,
                                semantic_weight=0.55,
                                bm25_weight=0.35,
                                exact_weight=0.10,
                                preferred_sources=[file_name],
                            )
                            
                            # Convert to directory_content format
                            for doc in file_first_docs:
                                file_name_from_meta = doc.metadata.get("fileName") or doc.metadata.get("file_name") or "document_vectorstore"
                                page_number = doc.metadata.get("page_number")
                                result_dict = {
                                    "fileName": file_name_from_meta,
                                    "content": doc.page_content[:2500]
                                }
                                if page_number is not None:
                                    result_dict["pageNumber"] = page_number
                                    result_dict["isPageChunk"] = doc.metadata.get("is_page_chunk", False)
                                file_first_results.append(result_dict)
                            
                            logger.info(f"📊 [Stream] Fichier: {len(file_first_results)} résultats")
                            
                            # Step 2: Always search in corpus to complement file results
                            # If info not in file, it will be in corpus
                            logger.info(f"📚 [Stream] Recherche complémentaire dans le corpus...")
                            corpus_docs, _ = hybrid_retrieve_documents(
                                vector_store=session_vector_store,
                                query=user_query,
                                k_candidates=120,
                                k_final=15,
                                semantic_weight=0.55,
                                bm25_weight=0.35,
                                exact_weight=0.10,
                            )
                            
                            # Convert and dedupe
                            seen_keys = set()
                            for doc in corpus_docs:
                                file_name_from_meta = doc.metadata.get("fileName") or doc.metadata.get("file_name") or "document_vectorstore"
                                content_preview = doc.page_content[:100]
                                key = f"{file_name_from_meta}:{content_preview}"
                                if key not in seen_keys:
                                    seen_keys.add(key)
                                    page_number = doc.metadata.get("page_number")
                                    result_dict = {
                                        "fileName": file_name_from_meta,
                                        "content": doc.page_content[:2500]
                                    }
                                    if page_number is not None:
                                        result_dict["pageNumber"] = page_number
                                        result_dict["isPageChunk"] = doc.metadata.get("is_page_chunk", False)
                                    corpus_results.append(result_dict)
                            
                            logger.info(f"📚 [Stream] Corpus: {len(corpus_results)} résultats")
                        else:
                            # No file selected: use the advanced search_semantic_documents_sync
                            # which includes person name filtering, function name search, filename matching, etc.
                            logger.info(f"📚 [Stream] No file selected: Using advanced search in whole repository...")
                            corpus_results = search_semantic_documents_sync(
                                session_vector_store, 
                                user_query, 
                                session_id, 
                                conversation_history
                            )
                        
                        # Combine results: file_first first, then corpus
                        directory_content = file_first_results + corpus_results
                        logger.info(f"✅ [Stream] {len(directory_content)} documents récupérés par recherche hiérarchique (file={len(file_first_results)}, corpus={len(corpus_results)})")
                except Exception as e:
                    logger.warning(f"Erreur lors de la recherche sémantique en streaming: {str(e)}")
            
            t = translations.get(language, translations['en'])
            
            # Construire le prompt local avec query_model_local_mode pour avoir le même prompt amélioré
            import asyncio
            try:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                enable_auto_search = data.get('enable_auto_online_search', False)  # Désactivé par défaut en mode local
                prompt = loop.run_until_complete(
                    query_model_local_mode(
                        file_name, file_content, directory_content, repo_structure,
                        user_query, False, selected_model, language, conversation_history,
                        enable_auto_online_search=enable_auto_search
                    )
                )
                loop.close()
            except Exception as e:
                logger.warning(f"Erreur lors de la construction du prompt amélioré: {str(e)}")
                # Fallback vers l'ancien prompt
            directory_content_summary = ' '.join(
                [f"{t['other_file']}: {doc['fileName']} : {doc['content']}" 
                 for doc in directory_content]
            ) if directory_content else t['no_other_files']
            
            prompt = (
                f"{t['local_analysis_mode']}\n"
                f"{t['no_external_search']}\n\n"
                f"{t['project_structure']}:\n{repo_structure}\n\n"
                f"{t['main_file']}: {file_name}\n\n"
                f"{t['file_content']}:\n{file_content}\n\n"
                f"{t['directory_context']}:\n{directory_content_summary}\n\n"
                f"{t['question']}: {user_query}\n\n"
                f"{t['instructions']}:\n"
                f"{t['base_response_only']}\n"
                f"{t['missing_info_clarify']}\n"
                f"{t['no_speculation']}\n"
                f"{t['focus_local_analysis']}\n\n"
                f"{t['emoji']}"
            )
            
            return Response(
                stream_with_context(stream_response(prompt, selected_model)),
                mimetype='text/event-stream',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'X-Accel-Buffering': 'no'
                }
            )
        
        else:
            return jsonify({"error": f"Mode non supporté: {research_mode}"}), 400
            
    except Exception as e:
        logger.error(f"Erreur dans handle_query_stream: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/vector-store-status', methods=['GET'])
async def get_vector_store_status():
    """
    Endpoint pour vérifier le statut des vector stores
    """
    session_id = request.headers.get('Session-ID', 'default')
    
    return jsonify({
        'session_id': session_id,
        'session_has_vectorstore': session_id in vector_stores,
        'global_vectorstore_exists': vector_store is not None,
        'total_sessions': len(vector_stores),
        'session_info': vector_stores.get(session_id, {}).get('files_indexed', []) if session_id in vector_stores else None
    }), 200

@app.route('/extract-structured', methods=['POST'])
async def extract_structured():
    """
    Endpoint pour extraire les données structurées d'un document
    Utilisé pour les avocats/notaires pour extraire automatiquement les informations importantes
    """
    try:
        data = request.get_json()
        
        # Validation des paramètres
        file_name = data.get('file_name')
        file_content = data.get('file_content', '')
        selected_model = data.get('selected_model', DEFAULT_MODEL)
        language = data.get('language', 'fr')
        
        if not file_name:
            return jsonify({
                "error": "file_name est requis",
                "success": False
            }), 400
        
        if not file_content:
            return jsonify({
                "error": "file_content est requis",
                "success": False
            }), 400
        
        # Validation du modèle
        if selected_model not in MODEL_CONFIG:
            logger.warning(f"Modèle inconnu {selected_model}, utilisation du modèle par défaut {DEFAULT_MODEL}")
            selected_model = DEFAULT_MODEL
        
        logger.info(f"📋 Extraction structurée demandée pour: {file_name}")
        
        # Détecter le type de document
        document_type = await detect_document_type(file_content, file_name)
        logger.info(f"🔍 Type de document détecté: {document_type}")
        
        # Extraire les données structurées
        structured_data = await extract_structured_data(
            file_content, 
            file_name, 
            document_type, 
            selected_model, 
            language
        )
        
        if "error" in structured_data:
            return jsonify({
                "success": False,
                "error": structured_data.get("error"),
                "data": structured_data
            }), 500
        
        return jsonify({
            "success": True,
            "data": structured_data,
            "document_type": document_type,
            "file_name": file_name
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Erreur lors de l'extraction structurée: {str(e)}")
        return jsonify({
            "success": False,
            "error": f"Erreur lors de l'extraction: {str(e)}"
        }), 500

@app.route('/summarize_file_stream', methods=['POST'])
def summarize_file_stream():
    """Génère un résumé d'un fichier avec streaming (4 lignes max)"""
    try:
        data = request.get_json()
        file_name = data.get('file_name', '')
        file_content = data.get('file_content', '')
        language = data.get('language', 'en')
        
        # Get translations for the selected language
        t = translations.get(language, translations['en'])
        
        if not file_content:
            error_msg = t.get('no_content_provided', 'No content provided')
            return jsonify({"error": error_msg}), 400
        
        # Build language-specific prompt
        lang_names = {'en': 'English', 'fr': 'French', 'es': 'Spanish'}
        lang_name = lang_names.get(language, language)
        prompt_text = t['summarize_file_prompt'].format(language=lang_name)
        summary_label = t['summarize_file_summary_label']
        system_content = t['summarize_file_system'].format(language=lang_name)
        
        prompt = f"""{prompt_text}

Document: {file_name}
Content:
{file_content[:2000]}

{summary_label}"""
        
        def generate():
            try:
                # Essayer Mistral d'abord, puis fallback vers OpenAI
                mistral_api_key = os.getenv("MISTRAL_API_KEY")
                openai_api_key = os.getenv("OPENAI_API_KEY")
                
                use_mistral = mistral_api_key and len(mistral_api_key.strip()) > 0
                use_openai = openai_api_key and len(openai_api_key.strip()) > 0
                
                if not use_mistral and not use_openai:
                    error_msg = "Aucune clé API configurée (ni Mistral ni OpenAI)"
                    logger.error(f"❌ {error_msg}")
                    yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
                    return
                
                # Essayer Mistral d'abord si disponible
                if use_mistral:
                    try:
                        mistral_client = OpenAI(base_url="https://api.mistral.ai/v1", api_key=mistral_api_key)
                        stream = mistral_client.chat.completions.create(
                    model="mistral-small",
                    messages=[
                        {"role": "system", "content": system_content},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=200,
                    temperature=0.3,
                    stream=True
                )
                        logger.info("✅ Utilisation de Mistral pour le résumé")
                    except Exception as mistral_error:
                        logger.warning(f"⚠️ Erreur Mistral ({mistral_error}), fallback vers OpenAI...")
                        use_mistral = False
                
                # Fallback vers OpenAI si Mistral n'est pas disponible ou a échoué
                if not use_mistral:
                    if not use_openai:
                        error_msg = "Mistral indisponible et OpenAI non configuré"
                        logger.error(f"❌ {error_msg}")
                        yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
                        return
                    
                    try:
                        from openai import OpenAI as OpenAIClient
                        openai_client = OpenAIClient(api_key=openai_api_key)
                        stream = openai_client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": system_content},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=200,
                            temperature=0.3,
                            stream=True
                        )
                        logger.info("✅ Utilisation d'OpenAI (fallback) pour le résumé")
                    except Exception as openai_error:
                        logger.error(f"❌ Erreur OpenAI: {openai_error}")
                        yield f"data: {json.dumps({'error': f'Erreur API: {openai_error}', 'done': True})}\n\n"
                        return
                
                accumulated_text = ""
                line_count = 0
                
                for chunk in stream:
                    if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        accumulated_text += content
                        
                        # Compter les lignes
                        current_lines = accumulated_text.split('\n')
                        if len(current_lines) > 4:
                            # Limiter à 4 lignes
                            accumulated_text = '\n'.join(current_lines[:4])
                            yield f"data: {json.dumps({'content': content, 'done': False})}\n\n"
                            yield f"data: {json.dumps({'done': True})}\n\n"
                            return
                        
                        yield f"data: {json.dumps({'content': content, 'done': False})}\n\n"
                
                # Limiter à 4 lignes à la fin
                final_lines = accumulated_text.split('\n')
                if len(final_lines) > 4:
                    accumulated_text = '\n'.join(final_lines[:4])
                
                yield f"data: {json.dumps({'done': True})}\n\n"
                
            except Exception as e:
                logger.error(f"❌ Erreur lors du streaming du résumé: {str(e)}")
                yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"
        
        return Response(generate(), mimetype='text/event-stream')
        
    except Exception as e:
        logger.error(f"❌ Erreur lors de la génération du résumé: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/summarize_repository_stream', methods=['POST'])
def summarize_repository_stream():
    """Génère un résumé d'un répertoire avec streaming"""
    try:
        data = request.get_json()
        files_info = data.get('files', [])
        language = data.get('language', 'en')
        
        # Get translations for the selected language
        t = translations.get(language, translations['en'])
        
        if not files_info or len(files_info) == 0:
            error_msg = t.get('no_files_provided', 'No files provided')
            return jsonify({"error": error_msg}), 400
        
        # Compter les sous-répertoires et fichiers
        file_count = len(files_info)
        subdirectories = set()
        for file_info in files_info:
            name = file_info.get('name', '')
            if '/' in name:
                parts = name.split('/')
                if len(parts) > 1:
                    subdirectories.add(parts[0])
        
        subdirectory_count = len(subdirectories)
        
        # Handle pluralization for different languages
        s_files = 's' if file_count > 1 else ''
        s_dirs = 's' if subdirectory_count > 1 else ''
        if language == 'fr':
            s_files = 's' if file_count > 1 else ''
            s_dirs = 's' if subdirectory_count > 1 else ''
        elif language == 'es':
            s_files = 's' if file_count > 1 else ''
            s_dirs = 's' if subdirectory_count > 1 else ''
        else:  # English
            s_files = 's' if file_count > 1 else ''
            s_dirs = 's' if subdirectory_count > 1 else ''
        
        def generate():
            try:
                # Essayer Mistral d'abord, puis fallback vers OpenAI
                mistral_api_key = os.getenv("MISTRAL_API_KEY")
                openai_api_key = os.getenv("OPENAI_API_KEY")
                
                use_mistral = mistral_api_key and len(mistral_api_key.strip()) > 0
                use_openai = openai_api_key and len(openai_api_key.strip()) > 0
                
                if not use_mistral and not use_openai:
                    error_msg = "Aucune clé API configurée (ni Mistral ni OpenAI)"
                    logger.error(f"❌ {error_msg}")
                    yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
                    return
                
                # Construire le prompt avec les informations du répertoire
                file_names_text = '\n'.join([f"- {f.get('display_name', f.get('name', ''))}" for f in files_info[:10]])
                
                # Get language-specific translations
                lang_names = {'en': 'English', 'fr': 'French', 'es': 'Spanish'}
                lang_name = lang_names.get(language, language)
                
                prompt_start = t['summarize_repo_prompt_start'].format(
                    file_count=file_count,
                    s_files=s_files,
                    subdir_count=subdirectory_count,
                    s_dirs=s_dirs
                )
                files_label = t['summarize_repo_files_label']
                instructions = t['summarize_repo_instructions'].format(language=lang_name)
                expected_format = t['summarize_repo_expected_format'].format(
                    file_count=file_count,
                    s_files=s_files,
                    subdir_count=subdirectory_count,
                    s_dirs=s_dirs
                )
                summary_label = t['summarize_repo_summary_label']
                system_content = t['summarize_repo_system'].format(language=lang_name)
                
                prompt = f"""{prompt_start}

{files_label}
{file_names_text}

{instructions}

{expected_format}

{summary_label}"""
                
                # Essayer Mistral d'abord si disponible
                if use_mistral:
                    try:
                        mistral_client = OpenAI(base_url="https://api.mistral.ai/v1", api_key=mistral_api_key)
                        stream = mistral_client.chat.completions.create(
                    model="mistral-small",
                    messages=[
                        {"role": "system", "content": system_content},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=400,
                    temperature=0.3,
                    stream=True
                )
                        logger.info("✅ Utilisation de Mistral pour le résumé du répertoire")
                    except Exception as mistral_error:
                        logger.warning(f"⚠️ Erreur Mistral ({mistral_error}), fallback vers OpenAI...")
                        use_mistral = False
                
                # Fallback vers OpenAI si Mistral n'est pas disponible ou a échoué
                if not use_mistral:
                    if not use_openai:
                        error_msg = "Mistral indisponible et OpenAI non configuré"
                        logger.error(f"❌ {error_msg}")
                        yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
                        return
                    
                    try:
                        from openai import OpenAI as OpenAIClient
                        openai_client = OpenAIClient(api_key=openai_api_key)
                        stream = openai_client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": system_content},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=400,
                            temperature=0.3,
                            stream=True
                        )
                        logger.info("✅ Utilisation d'OpenAI (fallback) pour le résumé du répertoire")
                    except Exception as openai_error:
                        logger.error(f"❌ Erreur OpenAI: {openai_error}")
                        yield f"data: {json.dumps({'error': f'Erreur API: {openai_error}', 'done': True})}\n\n"
                        return
                
                for chunk in stream:
                    if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        yield f"data: {json.dumps({'content': content, 'done': False})}\n\n"
                
                yield f"data: {json.dumps({'done': True})}\n\n"
                
            except Exception as e:
                logger.error(f"❌ Erreur lors du streaming du résumé de répertoire: {str(e)}")
                yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"
        
        return Response(generate(), mimetype='text/event-stream')
        
    except Exception as e:
        logger.error(f"❌ Erreur lors de la génération du résumé de répertoire: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    """Endpoint de santé pour vérifier que l'API fonctionne"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "1.0.0",
        "services": {
            "openai": bool(os.getenv("OPENAI_API_KEY")),
            "mistral": bool(os.getenv("MISTRAL_API_KEY")),
            "google_oauth": AuthConfig.is_google_configured(),
            "database": os.path.exists(AuthConfig.DATABASE_PATH)
        },
        "endpoints": {
            "extract_structured": "/extract-structured",
            "summarize_file_stream": "/summarize_file_stream",
            "summarize_repository_stream": "/summarize_repository_stream"
        }
    }), 200

@app.route('/test-endpoints', methods=['GET'])
def test_endpoints():
    """Endpoint pour tester la disponibilité des endpoints"""
    endpoints_status = {}
    
    # Liste des endpoints à tester
    endpoints_to_test = [
        '/extract-structured',
        '/summarize_file_stream',
        '/summarize_repository_stream',
        '/query',
        '/upload'
    ]
    
    # Vérifier si les routes existent dans l'application Flask
    for endpoint in endpoints_to_test:
        rule = None
        for rule in app.url_map.iter_rules():
            if rule.rule == endpoint:
                endpoints_status[endpoint] = {
                    "exists": True,
                    "methods": list(rule.methods - {'HEAD', 'OPTIONS'}),
                    "rule": rule.rule
                }
                break
        else:
            endpoints_status[endpoint] = {
                "exists": False,
                "error": "Route not found"
            }
    
    return jsonify({
        "status": "ok",
        "endpoints": endpoints_status,
        "total_routes": len(list(app.url_map.iter_rules()))
    }), 200

@app.route('/users/me', methods=['GET'])
def get_current_user():
    """Récupère les informations de l'utilisateur connecté"""
    from auth import verify_jwt_token
    
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return jsonify({'error': 'Token manquant'}), 401
    
    token = auth_header.replace('Bearer ', '')
    payload = verify_jwt_token(token)
    
    if not payload:
        return jsonify({'error': 'Token invalide'}), 401
    
    return jsonify({
        'user': {
            'id': payload['user_id'],
            'email': payload['email'],
            'name': payload['name']
        }
    }), 200

# Route pour tester la configuration Google (DEBUG)
@app.route('/debug/google-config', methods=['GET'])
def debug_google_config():
    google_client_id = os.getenv('GOOGLE_CLIENT_ID')
    return jsonify({
        'google_client_id_set': bool(google_client_id),
        'google_client_secret_set': bool(os.getenv('GOOGLE_CLIENT_SECRET')),
        'client_id_preview': google_client_id[:20] + '...' if google_client_id else 'Missing',
        'full_client_id': google_client_id  # Temporaire pour debug
    })

def main():
    print("🚀 Starting Enhanced AI backend on http://0.0.0.0:5000")
    print(f"📋 Available models: {', '.join(MODEL_CONFIG.keys())}")
    print(f"🎯 Default model: {DEFAULT_MODEL} (Mistral)")
    print("📡 Endpoints:")
    print("   GET  /models - List all available models")
    print("   POST /models/<model_id>/test - Test specific model")
    print("   POST /upload - Upload files for analysis")
    print("   POST /query - Query with selected model")
    print("   POST /index-directory - Index directory files")
    print("🔐 Authentication endpoints:")
    print("   POST /auth/login - Email/password login")
    print("   POST /auth/register - User registration")
    print("   GET  /auth/google - Google OAuth")
    print("   GET  /auth/verify - Verify JWT token")
    print("   POST /marketing/subscribe - Save marketing emails")
    print("   GET  /health - Health check")
    print("   GET  /users/me - Current user info")
    print("✨ New features:")
    print("   • Complete authentication system")
    print("   • Google OAuth integration")
    print("   • Marketing email capture")
    print("   • JWT-based sessions")
    print("   • SQLite database")
    if AuthConfig.is_google_configured():
        print("   • ✅ Google OAuth configured")
    else:
        print("   • ⚠️  Google OAuth NOT configured")
    app.run(debug=False, host='0.0.0.0', port=5000)

if __name__ == "__main__":
    main()