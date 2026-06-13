import os
import json
import base64
import requests
from flask import Flask, request, jsonify, Response, stream_with_context
from flask.helpers import make_response
from flask_cors import CORS
from dotenv import load_dotenv
from docx import Document as DocxDocument
import PyPDF2
from io import BytesIO
from openai import OpenAI
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from typing import Dict, List, Optional, Any
import aiocache
import asyncio
import logging
import time
import re
from datetime import datetime

# Load environment variables early (before importing modules that read env at import time)
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), '.env'))
load_dotenv()

from legacy_auth import register_auth_routes, init_database

# Imports des modules refactorisÃ©s
from config import (
    AuthConfig, FlaskConfig, AIConfig,
    MODEL_CONFIG, DEFAULT_MODEL, MISTRAL_MODEL, OLLAMA_MODELS, OLLAMA_API_URL
)
from utils.file_utils import (
    extract_text_from_docx, extract_text_from_pdf,
    extract_pages_from_pdf, extract_sections_from_docx,
)
from utils.docling_parser import parse_with_docling
from utils.translations import translations
from services.api_clients import (
    call_openai_api, call_mistral_api, call_ollama_api, call_gemini_api,
    stream_response, openai_client, get_model_config
)
from services.analysis_service import analyze_file_content, save_file_description
from services.search_service import perform_online_search, search_serpapi, rerank_documents_with_llm
from services.hybrid_retrieval import hybrid_retrieve_documents
from services.langgraph_citation_service import run_answer_graph
from services.vector_store_service import (
    vector_stores, embeddings, get_vector_store,
    create_vector_store, add_documents_to_vector_store,
    compute_document_hash, get_index_path, save_faiss_index, load_faiss_index,
)
from services.aws_persistence import aws_persistence_service

# Debug: VÃ©rifier le chargement des variables d'environnement
print(f"GOOGLE_CLIENT_ID chargÃ©: {os.getenv('GOOGLE_CLIENT_ID')[:20] + '...' if os.getenv('GOOGLE_CLIENT_ID') else 'NON CHARGÃ‰'}")
print(f"OPENAI_API_KEY chargÃ©: {'âœ…' if os.getenv('OPENAI_API_KEY') else 'âŒ'}")

# Variable globale pour stocker les vector stores par session
vector_stores = {}

# Logger setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _extract_user_email_from_auth_header() -> Optional[str]:
    """Resolve authenticated user email from a Firebase/JWT bearer token if present."""
    auth_header = request.headers.get('Authorization', '')
    if not auth_header.startswith('Bearer '):
        return None

    try:
        from legacy_auth import verify_auth_token

        token = auth_header.replace('Bearer ', '', 1)
        auth_result = verify_auth_token(token)
        if not auth_result:
            return None

        user = auth_result.get('user', {})
        return user.get('email')
    except Exception as exc:
        logger.warning(f"Impossible de rÃ©soudre l'utilisateur depuis le token: {str(exc)}")
        return None

# Configuration des modÃ¨les importÃ©e depuis config.models
# MODEL_CONFIG, DEFAULT_MODEL, MISTRAL_MODEL, OLLAMA_MODELS, OLLAMA_API_URL sont maintenant importÃ©s

# Flask setup
app = Flask(__name__)
CORS(app, resources={r"/*": {
    "origins": ["http://flexianalyse.com", "http://localhost:5173", "https://flexianalyse.com"],
    "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    "allow_headers": ["Content-Type", "Authorization", "Session-ID"]
}})

# Ajouter un handler pour les requÃªtes OPTIONS (CORS preflight)
@app.before_request
def handle_preflight():
    if request.method == "OPTIONS":
        response = make_response()
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add('Access-Control-Allow-Headers', "Content-Type,Authorization,Session-ID")
        response.headers.add('Access-Control-Allow-Methods', "GET,PUT,POST,DELETE,OPTIONS")
        return response

# Clients et composants initialisÃ©s dans les modules
# openai_client, embeddings sont maintenant importÃ©s depuis services.api_clients et services.vector_store_service
vector_store = None  # GardÃ© pour compatibilitÃ© avec le code existant

# Cache setup
cache = aiocache.SimpleMemoryCache()

# Description file path
DESCRIPTIONS_FILE = os.path.join(os.path.dirname(__file__), "file_descriptions.json")
if not os.path.exists(DESCRIPTIONS_FILE):
    with open(DESCRIPTIONS_FILE, 'w') as f:
        json.dump([], f)

# Traductions importÃ©es depuis utils.translations
# translations est maintenant importÃ©
_OLD_translations = {
    'en': {
        # Existing keys
        'analyze': 'Analyze',
        'content_of_file': "the content of the file",
        'and_provide': "and provide a description of its purpose.",
        'content': "Content",
        'description': "Description",
        'template': "I want an answer like: The file text.js aims to .....",
        'project_structure': 'Project structure',
        'file': 'File',
        'other_file': 'Other file',
        'question': 'Question',
        'emoji': 'Add some relevant emojis when it necessary to make it pleasant to read ðŸ˜ŠðŸ“„âœ¨',
        
        # New keys for local query mode
        'local_analysis_mode': 'ðŸ”’ LOCAL ANALYSIS MODE - Analyze ONLY the context provided below.',
        'no_external_search': 'Do NOT perform external searches or reference information not provided.',
        'main_file': 'Main file',
        'file_content': 'File content',
        'directory_context': 'Directory context',
        'no_other_files': 'No other files in context.',
        'instructions': 'Instructions',
        'base_response_only': '- Base your response ONLY on the code/content provided above',
        'missing_info_clarify': '- If information is not in the provided context, state it clearly',
        'no_speculation': '- Do not speculate on elements not present in the files',
        'focus_local_analysis': '- Focus on analyzing the local code/content',
        'error_local_analysis': 'Error during local analysis',
        'cached_result': 'Using cached result',
        'analyzing': 'Analyzing',
        'with_context': 'with context',
        'files_in_directory': 'files in directory',
        
        # Query analysis keys
        'analyze_query_prompt': 'Analyze this question and determine if it requires recent/current information that your model might not have in its training data.',
        'query_label': 'Question',
        'json_response_format': 'Respond in JSON format ONLY',
        'needs_search_field': 'needs_search',
        'reason_field': 'reason',
        'search_keywords_field': 'search_keywords',
        'cutoff_relevance_field': 'estimated_cutoff_relevance',
        'short_explanation': 'short explanation',
        'keyword_or_null': 'or null',
        'examples_need_search': 'Examples of questions requiring search',
        'current_prices_stocks': '- Current prices, recent stock prices',
        'recent_events_news': '- Recent events, news',
        'new_software_versions': '- New software/technology versions',
        'recent_statistics_data': '- Recent statistics, government data',
        'recent_people_companies': '- Information about recent people/companies',
        'examples_no_search': 'Examples of questions NOT requiring search',
        'general_concepts': '- General concepts, established theories',
        'programming_syntax': '- Programming, language syntax',
        'history_facts': '- History, historical facts',
        'math_science': '- Mathematics, fundamental sciences',
        'analysis_error': 'Analysis error',
        'query_analysis_error': 'Error during query analysis',
        'automatic_analysis_failed': 'Automatic analysis inconclusive',
        'error_analysis_fallback': 'Analysis error, searching for safety',
        'analyzing_query': 'Analyzing query',
        'search_needed': 'Search needed',
        'no_search_needed': 'No search needed',
        'json_parse_failed': 'Failed to parse JSON response',
        'fallback_analysis': 'Using fallback analysis',

        # NEW KEYS FOR ONLINE MODE
        'online_mode_title': 'ðŸ¤– Answer this question using your training knowledge.',
        'recent_info_mention': 'If you think more recent information could enrich your response, mention it at the end.',
        'online_instructions_title': 'Instructions:',
        'give_best_answer': '- Give your best answer based on your knowledge',
        'be_precise_dates': '- Be precise about dates/versions you know',
        'mention_recent_useful': '- Mention if more recent info would be useful',
        'enrichment_title': 'ðŸ”„ RESPONSE ENRICHMENT',
        'initial_response': 'Your initial response:',
        'new_info_found': 'New information found:',
        'enrichment_instructions': 'Instructions:',
        'combine_intelligently': '- Intelligently combine your initial response with new info',
        'update_obsolete': '- Update obsolete parts if necessary',
        'distinguish_info': '- Clearly distinguish basic info from recent info',
        'cite_sources': '- Cite sources for recent information',
        'keep_structure': '- Keep the structure and style of your initial response',
        'prioritize_recent': '- If new info contradicts your knowledge, prioritize recent sources',
        'enriched_response': 'Enriched response:',
        'info_enriched': 'ðŸ’¡ **Information enriched**: This response combines my basic knowledge with recent data found online.',
        'search_reason': 'ðŸ” **Search reason**: {reason}',
        'default_search_reason': 'Potentially outdated information',
        'source_training': 'ðŸ§  **Source**: Response based on my training knowledge. No recent search was deemed necessary.',
        'online_processing_error': 'Error during online mode processing: {error}',
        'search_analysis_log': 'ðŸ” Search analysis: {analysis}',
        'enriching_current_data': 'ðŸŒ Enriching with current data...',
        'no_search_necessary': 'âœ… No search necessary, response based on model knowledge',
        'error_online_mode': 'Error during online mode: {error}',
        'separator': '---'
    },
    'fr': {
        # Existing keys
        'analyze': 'Analysez',
        'content_of_file': "le contenu du fichier",
        'and_provide': "et fournissez une description de son objectif.",
        'content': "Contenu",
        'description': "Description",
        'template': "Je veux une rÃ©ponse du genre: Le Fichier text.js a pour Objectif : .....",
        'project_structure': 'Structure du projet',
        'file': 'Fichier',
        'other_file': 'Autre fichier',
        'question': 'Question',
        'emoji': 'Ajoute quelques emojis pertinents quand c\'est nÃ©cessaire pour rendre la lecture agrÃ©able ðŸ˜ŠðŸ“„âœ¨',
        
        # New keys for local query mode
        'local_analysis_mode': 'ðŸ”’ MODE ANALYSE LOCALE - Analysez UNIQUEMENT le contexte fourni ci-dessous.',
        'no_external_search': 'Ne faites PAS de recherche externe ou de rÃ©fÃ©rence Ã  des informations non fournies.',
        'main_file': 'Fichier principal',
        'file_content': 'Contenu du fichier',
        'directory_context': 'Contexte du rÃ©pertoire',
        'no_other_files': 'Aucun autre fichier dans le contexte.',
        'instructions': 'Instructions',
        'base_response_only': '- Basez votre rÃ©ponse UNIQUEMENT sur le code/contenu fourni ci-dessus',
        'missing_info_clarify': '- Si l\'information n\'est pas dans le contexte fourni, dites-le clairement',
        'no_speculation': '- Ne spÃ©culez pas sur des Ã©lÃ©ments non prÃ©sents dans les fichiers',
        'focus_local_analysis': '- Concentrez-vous sur l\'analyse du code/contenu local',
        'error_local_analysis': 'Erreur lors de l\'analyse locale',
        'cached_result': 'Utilisation du rÃ©sultat en cache',
        'analyzing': 'Analyse de',
        'with_context': 'avec contexte de',
        'files_in_directory': 'fichiers dans le rÃ©pertoire',
        
        # Query analysis keys
        'analyze_query_prompt': 'Analysez cette question et dÃ©terminez si elle nÃ©cessite des informations rÃ©centes/actuelles que votre modÃ¨le pourrait ne pas avoir dans ses donnÃ©es d\'entraÃ®nement.',
        'query_label': 'Question',
        'json_response_format': 'RÃ©pondez au format JSON UNIQUEMENT',
        'needs_search_field': 'needs_search',
        'reason_field': 'reason',
        'search_keywords_field': 'search_keywords',
        'cutoff_relevance_field': 'estimated_cutoff_relevance',
        'short_explanation': 'explication courte',
        'keyword_or_null': 'ou null',
        'examples_need_search': 'Exemples de questions nÃ©cessitant une recherche',
        'current_prices_stocks': '- Prix actuels, cours de bourse rÃ©cents',
        'recent_events_news': '- Ã‰vÃ©nements rÃ©cents, actualitÃ©s',
        'new_software_versions': '- Nouvelles versions de logiciels/technologies',
        'recent_statistics_data': '- Statistiques rÃ©centes, donnÃ©es gouvernementales',
        'recent_people_companies': '- Informations sur des personnes/entreprises rÃ©centes',
        'examples_no_search': 'Exemples de questions NE nÃ©cessitant PAS de recherche',
        'general_concepts': '- Concepts gÃ©nÃ©raux, thÃ©ories Ã©tablies',
        'programming_syntax': '- Programmation, syntaxe de langages',
        'history_facts': '- Histoire, faits historiques',
        'math_science': '- MathÃ©matiques, sciences fondamentales',
        'analysis_error': 'Erreur d\'analyse',
        'query_analysis_error': 'Erreur lors de l\'analyse de la requÃªte',
        'automatic_analysis_failed': 'Analyse automatique non concluante',
        'error_analysis_fallback': 'Erreur d\'analyse, recherche par sÃ©curitÃ©',
        'analyzing_query': 'Analyse de la requÃªte',
        'search_needed': 'Recherche nÃ©cessaire',
        'no_search_needed': 'Aucune recherche nÃ©cessaire',
        'json_parse_failed': 'Ã‰chec du parsing JSON',
        'fallback_analysis': 'Utilisation de l\'analyse de secours',

        # NEW KEYS FOR ONLINE MODE
        'online_mode_title': 'ðŸ¤– RÃ©pondez Ã  cette question en utilisant vos connaissances d\'entraÃ®nement.',
        'recent_info_mention': 'Si vous pensez que des informations plus rÃ©centes pourraient enrichir votre rÃ©ponse, mentionnez-le Ã  la fin.',
        'online_instructions_title': 'Instructions :',
        'give_best_answer': '- Donnez votre meilleure rÃ©ponse basÃ©e sur vos connaissances',
        'be_precise_dates': '- Soyez prÃ©cis sur les dates/versions que vous connaissez',
        'mention_recent_useful': '- Mentionnez si des infos plus rÃ©centes seraient utiles',
        'enrichment_title': 'ðŸ”„ ENRICHISSEMENT DE RÃ‰PONSE',
        'initial_response': 'Votre rÃ©ponse initiale :',
        'new_info_found': 'Nouvelles informations trouvÃ©es :',
        'enrichment_instructions': 'Instructions :',
        'combine_intelligently': '- Combinez intelligemment votre rÃ©ponse initiale avec les nouvelles infos',
        'update_obsolete': '- Mettez Ã  jour les parties obsolÃ¨tes si nÃ©cessaire',
        'distinguish_info': '- Distinguez clairement les infos de base des infos rÃ©centes',
        'cite_sources': '- Citez les sources pour les informations rÃ©centes',
        'keep_structure': '- Gardez la structure et le style de votre rÃ©ponse initiale',
        'prioritize_recent': '- Si les nouvelles infos contredisent vos connaissances, privilÃ©giez les sources rÃ©centes',
        'enriched_response': 'RÃ©ponse enrichie :',
        'info_enriched': 'ðŸ’¡ **Informations enrichies** : Cette rÃ©ponse combine mes connaissances de base avec des donnÃ©es rÃ©centes trouvÃ©es en ligne.',
        'search_reason': 'ðŸ” **Raison de la recherche** : {reason}',
        'default_search_reason': 'Information potentiellement obsolÃ¨te',
        'source_training': 'ðŸ§  **Source** : RÃ©ponse basÃ©e sur mes connaissances d\'entraÃ®nement. Aucune recherche rÃ©cente n\'a Ã©tÃ© jugÃ©e nÃ©cessaire.',
        'online_processing_error': 'Erreur lors du traitement en mode online : {error}',
        'search_analysis_log': 'ðŸ” Analyse de recherche : {analysis}',
        'enriching_current_data': 'ðŸŒ Enrichissement avec des donnÃ©es actuelles...',
        'no_search_necessary': 'âœ… Pas de recherche nÃ©cessaire, rÃ©ponse basÃ©e sur les connaissances du modÃ¨le',
        'error_online_mode': 'Erreur lors du mode online : {error}',
        'separator': '---'
    },
    'es': {
        # Existing keys
        'analyze': 'Analiza',
        'content_of_file': "el contenido del archivo",
        'and_provide': "y proporciona una descripciÃ³n de su propÃ³sito.",
        'content': "Contenido",
        'description': "DescripciÃ³n",
        'template': "Quiero una respuesta como: El archivo text.js tiene como objetivo .....",
        'project_structure': 'Estructura del proyecto',
        'file': 'Archivo',
        'other_file': 'Otro archivo',
        'question': 'Pregunta',
        'emoji': 'Agrega algunos emojis relevantes cuando sea necesario para que sea agradable de leer ðŸ˜ŠðŸ“„âœ¨',
        
        # New keys for local query mode
        'local_analysis_mode': 'ðŸ”’ MODO ANÃLISIS LOCAL - Analiza ÃšNICAMENTE el contexto proporcionado a continuaciÃ³n.',
        'no_external_search': 'NO realices bÃºsquedas externas o referencias a informaciÃ³n no proporcionada.',
        'main_file': 'Archivo principal',
        'file_content': 'Contenido del archivo',
        'directory_context': 'Contexto del directorio',
        'no_other_files': 'NingÃºn otro archivo en el contexto.',
        'instructions': 'Instrucciones',
        'base_response_only': '- Basa tu respuesta ÃšNICAMENTE en el cÃ³digo/contenido proporcionado arriba',
        'missing_info_clarify': '- Si la informaciÃ³n no estÃ¡ en el contexto proporcionado, indÃ­calo claramente',
        'no_speculation': '- No especules sobre elementos no presentes en los archivos',
        'focus_local_analysis': '- ConcÃ©ntrate en analizar el cÃ³digo/contenido local',
        'error_local_analysis': 'Error durante el anÃ¡lisis local',
        'cached_result': 'Usando resultado en cachÃ©',
        'analyzing': 'Analizando',
        'with_context': 'con contexto de',
        'files_in_directory': 'archivos en el directorio',
        
        # Query analysis keys
        'analyze_query_prompt': 'Analiza esta pregunta y determina si requiere informaciÃ³n reciente/actual que tu modelo podrÃ­a no tener en sus datos de entrenamiento.',
        'query_label': 'Pregunta',
        'json_response_format': 'Responde en formato JSON ÃšNICAMENTE',
        'needs_search_field': 'needs_search',
        'reason_field': 'reason',
        'search_keywords_field': 'search_keywords',
        'cutoff_relevance_field': 'estimated_cutoff_relevance',
        'short_explanation': 'explicaciÃ³n corta',
        'keyword_or_null': 'o null',
        'examples_need_search': 'Ejemplos de preguntas que requieren bÃºsqueda',
        'current_prices_stocks': '- Precios actuales, precios de acciones recientes',
        'recent_events_news': '- Eventos recientes, noticias',
        'new_software_versions': '- Nuevas versiones de software/tecnologÃ­as',
        'recent_statistics_data': '- EstadÃ­sticas recientes, datos gubernamentales',
        'recent_people_companies': '- InformaciÃ³n sobre personas/empresas recientes',
        'examples_no_search': 'Ejemplos de preguntas que NO requieren bÃºsqueda',
        'general_concepts': '- Conceptos generales, teorÃ­as establecidas',
        'programming_syntax': '- ProgramaciÃ³n, sintaxis de lenguajes',
        'history_facts': '- Historia, hechos histÃ³ricos',
        'math_science': '- MatemÃ¡ticas, ciencias fundamentales',
        'analysis_error': 'Error de anÃ¡lisis',
        'query_analysis_error': 'Error durante el anÃ¡lisis de la consulta',
        'automatic_analysis_failed': 'AnÃ¡lisis automÃ¡tico no concluyente',
        'error_analysis_fallback': 'Error de anÃ¡lisis, bÃºsqueda por seguridad',
        'analyzing_query': 'Analizando consulta',
        'search_needed': 'BÃºsqueda necesaria',
        'no_search_needed': 'No se necesita bÃºsqueda',
        'json_parse_failed': 'Error al analizar respuesta JSON',
        'fallback_analysis': 'Usando anÃ¡lisis de respaldo',

        # NEW KEYS FOR ONLINE MODE
        'online_mode_title': 'ðŸ¤– Responde esta pregunta usando tus conocimientos de entrenamiento.',
        'recent_info_mention': 'Si crees que informaciÃ³n mÃ¡s reciente podrÃ­a enriquecer tu respuesta, menciÃ³nalo al final.',
        'online_instructions_title': 'Instrucciones:',
        'give_best_answer': '- Da tu mejor respuesta basada en tus conocimientos',
        'be_precise_dates': '- SÃ© preciso sobre fechas/versiones que conoces',
        'mention_recent_useful': '- Menciona si informaciÃ³n mÃ¡s reciente serÃ­a Ãºtil',
        'enrichment_title': 'ðŸ”„ ENRIQUECIMIENTO DE RESPUESTA',
        'initial_response': 'Tu respuesta inicial:',
        'new_info_found': 'Nueva informaciÃ³n encontrada:',
        'enrichment_instructions': 'Instrucciones:',
        'combine_intelligently': '- Combina inteligentemente tu respuesta inicial con la nueva info',
        'update_obsolete': '- Actualiza las partes obsoletas si es necesario',
        'distinguish_info': '- Distingue claramente la info bÃ¡sica de la info reciente',
        'cite_sources': '- Cita fuentes para la informaciÃ³n reciente',
        'keep_structure': '- MantÃ©n la estructura y estilo de tu respuesta inicial',
        'prioritize_recent': '- Si la nueva info contradice tus conocimientos, prioriza fuentes recientes',
        'enriched_response': 'Respuesta enriquecida:',
        'info_enriched': 'ðŸ’¡ **InformaciÃ³n enriquecida**: Esta respuesta combina mis conocimientos bÃ¡sicos con datos recientes encontrados en lÃ­nea.',
        'search_reason': 'ðŸ” **RazÃ³n de bÃºsqueda**: {reason}',
        'default_search_reason': 'InformaciÃ³n potencialmente obsoleta',
        'source_training': 'ðŸ§  **Fuente**: Respuesta basada en mis conocimientos de entrenamiento. No se considerÃ³ necesaria una bÃºsqueda reciente.',
        'online_processing_error': 'Error durante el procesamiento en modo online: {error}',
        'search_analysis_log': 'ðŸ” AnÃ¡lisis de bÃºsqueda: {analysis}',
        'enriching_current_data': 'ðŸŒ Enriqueciendo con datos actuales...',
        'no_search_necessary': 'âœ… No es necesaria bÃºsqueda, respuesta basada en conocimientos del modelo',
        'error_online_mode': 'Error durante modo online: {error}',
        'separator': '---'
    }
}

# Description template
description_template = "Je veux une reponse du genre: Le Fichier text.js a pour Objectif : ....."

# Fonctions utilitaires importÃ©es depuis les modules
# extract_text_from_docx, extract_text_from_pdf sont maintenant importÃ©s depuis utils.file_utils
# get_model_config est maintenant importÃ© depuis services.api_clients

# call_openai_api est maintenant importÃ© depuis services.api_clients
def _OLD_call_openai_api(prompt, selected_model="gpt-3.5-turbo", max_retries=3, max_tokens_override=None):
    """
    Enhanced OpenAI API call supporting both Chat Completions and Responses API
    max_tokens_override: override la limite de tokens de la config du modÃ¨le
    """
    model_config = get_model_config(selected_model)
    model_id = model_config["model_id"]
    api_type = model_config.get("api_type", "chat")
    
    for attempt in range(max_retries):
        try:
            logger.info(f"[OpenAI] Using model: {model_id} with API: {api_type}")
            
            # GPT-5 variants use the new Responses API
            if api_type == "responses":
                request_params = {
                    "model": model_id,
                    "input": prompt
                }
                
                # Add reasoning parameters for GPT-5
                if "reasoning_effort" in model_config:
                    request_params["reasoning"] = {
                        "effort": model_config["reasoning_effort"]
                    }
                
                # Add verbosity parameters for GPT-5
                if "verbosity" in model_config:
                    request_params["text"] = {
                        "verbosity": model_config["verbosity"]
                    }
                
                logger.info(f"[OpenAI] GPT-5 request params: reasoning={model_config.get('reasoning_effort')}, verbosity={model_config.get('verbosity')}")
                response = openai_client.responses.create(**request_params)
                
                # Extract text from responses API - Updated extraction logic
                try:
                    # The actual content is in output array, look for message type with content
                    if hasattr(response, 'output') and response.output:
                        for item in response.output:
                            # Look for message type items
                            if hasattr(item, 'type') and item.type == 'message':
                                if hasattr(item, 'content') and item.content:
                                    for content_item in item.content:
                                        if hasattr(content_item, 'text') and content_item.text:
                                            logger.info(f"[OpenAI] Successfully extracted GPT-5 response")
                                            return content_item.text.strip()
                    
                    # Fallback extraction methods
                    if hasattr(response, 'text') and hasattr(response.text, 'content'):
                        return response.text.content.strip()
                    elif hasattr(response, 'content'):
                        return response.content.strip()
                    else:
                        # Final fallback: try to extract from string representation
                        response_str = str(response)
                        logger.warning(f"[OpenAI] Using fallback extraction for GPT-5 response")
                        
                        # Try to extract text from the string representation
                        import re
                        text_match = re.search(r"text='([^']*)'", response_str)
                        if text_match:
                            extracted_text = text_match.group(1)
                            # Decode escape sequences
                            extracted_text = extracted_text.replace('\\n', '\n').replace("\\'", "'")
                            return extracted_text.strip()
                        
                        return "Error: Could not extract response content from GPT-5"
                        
                except Exception as extraction_error:
                    logger.error(f"[OpenAI] Error extracting GPT-5 response: {extraction_error}")
                    return f"Error extracting response: {str(extraction_error)}"
            
            # Standard models use Chat Completions API
            else:
                # Utiliser max_tokens_override si fourni, sinon la config du modÃ¨le
                max_tokens = max_tokens_override if max_tokens_override is not None else model_config.get("max_tokens", 500)
                request_params = {
                    "model": model_id,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.5,
                    "max_tokens": max_tokens
                }
                
                response = openai_client.chat.completions.create(**request_params)
                return response.choices[0].message.content.strip()
            
        except Exception as e:
            logger.warning(f"[OpenAI] Attempt {attempt+1}/{max_retries} failed with {model_id}: {str(e)}")
            
            # Smart fallback strategy
            error_str = str(e).lower()
            if any(keyword in error_str for keyword in ["model", "not found", "unsupported", "invalid"]):
                # GPT-5 fallback chain
                if model_id == "gpt-5":
                    logger.info("[OpenAI] GPT-5 not available, falling back to GPT-5-mini")
                    model_id = "gpt-5-mini"
                    model_config = get_model_config("gpt-5-mini")
                    api_type = model_config.get("api_type", "responses")
                    continue
                elif model_id == "gpt-5-mini":
                    logger.info("[OpenAI] GPT-5-mini not available, falling back to GPT-5-nano")
                    model_id = "gpt-5-nano"
                    model_config = get_model_config("gpt-5-nano")
                    api_type = model_config.get("api_type", "responses")
                    continue
                elif model_id == "gpt-5-nano":
                    logger.info("[OpenAI] GPT-5-nano not available, falling back to GPT-4o")
                    model_id = "gpt-4o"
                    model_config = get_model_config("gpt-4o")
                    api_type = model_config.get("api_type", "chat")
                    continue
                elif model_id == "gpt-4o":
                    logger.info("[OpenAI] GPT-4o not available, falling back to GPT-3.5-turbo")
                    model_id = "gpt-3.5-turbo"
                    model_config = get_model_config("gpt-3.5-turbo")
                    api_type = model_config.get("api_type", "chat")
                    continue
            
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logger.warning(f"[OpenAI] Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                continue
            else:
                logger.error(f"[OpenAI] Max retries exceeded: {str(e)}")
                raise e
    
    raise RuntimeError("[OpenAI] Max retries exceeded")

# stream_response est maintenant importÃ© depuis services.api_clients
def _OLD_stream_response(prompt, selected_model="gpt-3.5-turbo"):
    """
    GÃ©nÃ¨re une rÃ©ponse en streaming pour n'importe quel modÃ¨le
    """
    model_config = get_model_config(selected_model)
    
    try:
        # Pour les modÃ¨les OpenAI avec support du streaming
        if selected_model.lower() in ["gpt-3.5-turbo", "gpt-4o", "openai"]:
            model_id = model_config["model_id"]
            
            stream = openai_client.chat.completions.create(
                model=model_id,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=model_config.get("max_tokens", 500),
                stream=True
            )
            
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
            
            yield f"data: {json.dumps({'done': True})}\n\n"
            
        # Pour Mistral - pas de streaming natif, on simule
        elif selected_model.lower() == "mistral":
            try:
                # Appel normal Ã  Mistral (sans streaming)
                response = call_mistral_api(prompt)
                
                # Simuler le streaming en envoyant par morceaux
                words = response.split(' ')
                for i in range(0, len(words), 3):  # Envoyer 3 mots Ã  la fois
                    chunk = ' '.join(words[i:i+3])
                    if i + 3 < len(words):
                        chunk += ' '
                    yield f"data: {json.dumps({'content': chunk})}\n\n"
                    time.sleep(0.05)  # Petit dÃ©lai pour simuler le streaming
                
                yield f"data: {json.dumps({'done': True})}\n\n"
                
            except Exception as mistral_error:
                # Si Mistral Ã©choue, fallback vers GPT-3.5
                logger.warning(f"Mistral failed: {str(mistral_error)}, falling back to GPT-3.5")
                yield f"data: {json.dumps({'warning': 'Mistral indisponible, utilisation de GPT-3.5'})}\n\n"
                
                # Utiliser GPT-3.5 en fallback
                stream = openai_client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5,
                    max_tokens=500,
                    stream=True
                )
                
                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
                
                yield f"data: {json.dumps({'done': True})}\n\n"
        
        # Pour Llama/Ollama - pas de streaming natif
        elif selected_model.lower() in ["llama3", "llama3.2"]:
            try:
                response = call_ollama_api(prompt, selected_model)
                
                # Simuler le streaming
                words = response.split(' ')
                for i in range(0, len(words), 3):
                    chunk = ' '.join(words[i:i+3])
                    if i + 3 < len(words):
                        chunk += ' '
                    yield f"data: {json.dumps({'content': chunk})}\n\n"
                    time.sleep(0.05)
                
                yield f"data: {json.dumps({'done': True})}\n\n"
                
            except Exception as ollama_error:
                logger.warning(f"Ollama failed: {str(ollama_error)}, falling back to GPT-3.5")
                yield f"data: {json.dumps({'warning': 'Ollama indisponible, utilisation de GPT-3.5'})}\n\n"
                
                # Fallback vers GPT-3.5
                stream = openai_client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5,
                    max_tokens=500,
                    stream=True
                )
                
                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
                
                yield f"data: {json.dumps({'done': True})}\n\n"
        
        else:
            # ModÃ¨le inconnu, utiliser GPT-3.5 par dÃ©faut
            logger.warning(f"Unknown model {selected_model}, using GPT-3.5")
            yield f"data: {json.dumps({'warning': f'ModÃ¨le {selected_model} inconnu, utilisation de GPT-3.5'})}\n\n"
            
            stream = openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=500,
                stream=True
            )
            
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
            
            yield f"data: {json.dumps({'done': True})}\n\n"
            
    except Exception as e:
        logger.error(f"Erreur streaming: {str(e)}")
        yield f"data: {json.dumps({'error': str(e)})}\n\n"

# call_mistral_api est maintenant importÃ© depuis services.api_clients
def _OLD_call_mistral_api(prompt, max_retries=3):
    mistral_api_key = os.getenv("MISTRAL_API_KEY")
    if not mistral_api_key:
        raise ValueError("MISTRAL_API_KEY environment variable not set")

    url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {mistral_api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": MISTRAL_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 500,
        "temperature": 0.5
    }

    for attempt in range(max_retries):
        try:
            response = requests.post(url, headers=headers, json=payload)
            if response.status_code == 429:
                retry_after = int(response.headers.get("Retry-After", 1))
                logger.warning(f"[Mistral] Rate limited. Retrying in {retry_after} sec (attempt {attempt+1}/{max_retries})...")
                time.sleep(retry_after)
                continue

            response.raise_for_status()
            result = response.json()
            return result["choices"][0]["message"]["content"].strip()

        except requests.exceptions.HTTPError as e:
            if response.status_code == 429 and attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logger.warning(f"[Mistral] 429 Too Many Requests. Retrying in {wait_time} sec (attempt {attempt+1}/{max_retries})...")
                time.sleep(wait_time)
                continue
            else:
                logger.error(f"[Mistral] HTTP Error: {e}")
                raise e

        except Exception as e:
            logger.error(f"[Mistral] Unexpected Error: {e}")
            raise e

    raise RuntimeError("[Mistral] Max retries exceeded")

# call_ollama_api est maintenant importÃ© depuis services.api_clients
def _OLD_call_ollama_api(prompt, selected_model="llama3", max_retries=3):
    """Enhanced Ollama API call"""
    model_config = get_model_config(selected_model)
    model_id = model_config["model_id"]
    
    for attempt in range(max_retries):
        try:
            response = requests.post(OLLAMA_API_URL, json={
                "model": model_id,
                "prompt": prompt,
                "stream": False
            })
            response.raise_for_status()
            data = response.json()
            return data.get("response") or data.get("message", {}).get("content", "No response").strip()
            
        except Exception as e:
            logger.warning(f"[Ollama] Attempt {attempt+1}/{max_retries} failed: {str(e)}")
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logger.warning(f"[Ollama] Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                continue
            else:
                logger.error(f"[Ollama] Max retries exceeded: {str(e)}")
                raise e
    
    raise RuntimeError("[Ollama] Max retries exceeded")

def _get_type_label(doc_type: str, language: str = 'en') -> str:
    """Retourne un label lisible pour un type de document dans la langue donnÃ©e."""
    labels = {
        'fr': {
            'cv_resume': 'CV / RÃ©sumÃ©',
            'facture_invoice': 'Facture',
            'contrat_location': 'Contrat de location',
            'contrat_travail': 'Contrat de travail',
            'contrat_vente': 'Contrat de vente',
            'contrat_generique': 'Contrat',
            'contrat_prenuptial': 'Contrat de mariage',
            'procuration_poa': 'Procuration',
            'accord_confidentialite_nda': 'Accord de confidentialitÃ©',
            'acte_propriete_immobiliere': 'Acte de propriÃ©tÃ©',
            'testament': 'Testament',
            'acte_notarie': 'Acte notariÃ©',
            'lettre': 'Lettre / Courrier',
            'document_financier': 'Document financier',
            'assurance_insurance': 'Assurance',
            'jugement_decision_justice': 'Jugement / DÃ©cision',
            'releve_bancaire': 'RelevÃ© bancaire',
            'certificat_attestation': 'Certificat / Attestation',
            'contrat_pret_loan': 'Contrat de prÃªt',
            'devis_estimation': 'Devis / Estimation',
            'bon_commande_purchase_order': 'Bon de commande',
            'proces_verbal': 'ProcÃ¨s-verbal',
            'rapport_expertise': "Rapport d'expertise",
            'permis_licence': 'Permis / Licence',
            'document_generique': 'Document',
        },
        'en': {
            'cv_resume': 'CV / Resume',
            'facture_invoice': 'Invoice',
            'contrat_location': 'Rental Contract',
            'contrat_travail': 'Employment Contract',
            'contrat_vente': 'Sale Contract',
            'contrat_generique': 'Contract',
            'contrat_prenuptial': 'Prenuptial Agreement',
            'procuration_poa': 'Power of Attorney',
            'accord_confidentialite_nda': 'NDA',
            'acte_propriete_immobiliere': 'Real Estate Deed',
            'testament': 'Will / Testament',
            'acte_notarie': 'Notarial Act',
            'lettre': 'Letter',
            'document_financier': 'Financial Document',
            'assurance_insurance': 'Insurance Policy',
            'jugement_decision_justice': 'Court Decision',
            'releve_bancaire': 'Bank Statement',
            'certificat_attestation': 'Certificate',
            'contrat_pret_loan': 'Loan Agreement',
            'devis_estimation': 'Quote / Estimate',
            'bon_commande_purchase_order': 'Purchase Order',
            'proces_verbal': 'Meeting Minutes',
            'rapport_expertise': 'Expert Report',
            'permis_licence': 'Permit / License',
            'document_generique': 'Document',
        },
        'es': {
            'cv_resume': 'CV / CurrÃ­culum',
            'facture_invoice': 'Factura',
            'contrat_location': 'Contrato de alquiler',
            'contrat_travail': 'Contrato de trabajo',
            'contrat_vente': 'Contrato de venta',
            'contrat_generique': 'Contrato',
            'contrat_prenuptial': 'Contrato prenupcial',
            'procuration_poa': 'Poder notarial',
            'accord_confidentialite_nda': 'Acuerdo de confidencialidad',
            'acte_propriete_immobiliere': 'Escritura de propiedad',
            'testament': 'Testamento',
            'acte_notarie': 'Acta notarial',
            'lettre': 'Carta',
            'document_financier': 'Documento financiero',
            'assurance_insurance': 'PÃ³liza de seguro',
            'jugement_decision_justice': 'Sentencia judicial',
            'releve_bancaire': 'Extracto bancario',
            'certificat_attestation': 'Certificado',
            'contrat_pret_loan': 'Contrato de prÃ©stamo',
            'devis_estimation': 'Presupuesto',
            'bon_commande_purchase_order': 'Orden de compra',
            'proces_verbal': 'Acta de reuniÃ³n',
            'rapport_expertise': 'Informe pericial',
            'permis_licence': 'Permiso / Licencia',
            'document_generique': 'Documento',
        }
    }
    lang_labels = labels.get(language, labels['en'])
    return lang_labels.get(doc_type, doc_type.replace('_', ' ').title())

async def infer_corpus_actions(documents: List[Document], language: str = 'en') -> Dict[str, Any]:
    """
    Utilise un petit appel modÃ¨le pour deviner le type de corpus (CV, rapports annuels, etc.)
    et proposer des actions suggÃ©rÃ©es (boutons) adaptÃ©es.
    """
    try:
        # D'abord, dÃ©tecter le type de document dominant dans le corpus (avec plus de contexte)
        document_types = {}
        document_details = {}  # Stocker plus d'infos par type
        detection_confidences = {}  # Stocker les confiances par type
        
        for doc in documents[:30]:  # Analyser plus de documents
            doc_content = doc.page_content[:4000] if len(doc.page_content) > 4000 else doc.page_content
            detection_result = await detect_document_type_detailed(doc_content, doc.metadata.get('fileName', ''))
            doc_type = detection_result['type']
            doc_confidence = detection_result.get('confidence', 0.5)
            
            document_types[doc_type] = document_types.get(doc_type, 0) + 1
            
            # Stocker la meilleure confiance par type
            if doc_type not in detection_confidences or doc_confidence > detection_confidences[doc_type]:
                detection_confidences[doc_type] = doc_confidence
            
            # Stocker les dÃ©tails pour enrichir le prompt
            if doc_type not in document_details:
                document_details[doc_type] = []
            meta_name = doc.metadata.get('fileName') or doc.metadata.get('source') or 'document'
            # Prendre un Ã©chantillon plus intelligent (dÃ©but + milieu si disponible)
            content_sample = doc_content[:1000]
            if len(doc_content) > 2000:
                mid_point = len(doc_content) // 2
                content_sample += " ... " + doc_content[mid_point:mid_point+500]
            document_details[doc_type].append({
                'name': meta_name,
                'snippet': content_sample.replace('\n', ' ').strip()[:1200]
            })
        
        # Trouver le type de document le plus frÃ©quent
        dominant_type = max(document_types.items(), key=lambda x: x[1])[0] if document_types else 'document_generique'
        dominant_confidence = detection_confidences.get(dominant_type, 0.5)
        
        # Construire un rÃ©sumÃ© enrichi du corpus pour le prompt
        sample_texts = []
        # Inclure tous les documents du type dominant (jusqu'Ã  10)
        dominant_docs = document_details.get(dominant_type, [])[:10]
        for doc_info in dominant_docs:
            sample_texts.append(f"- {doc_info['name']}: {doc_info['snippet']}")
        
        # Ajouter quelques exemples d'autres types si prÃ©sents
        other_types = [dt for dt in document_types.keys() if dt != dominant_type]
        for other_type in other_types[:2]:  # Max 2 autres types
            for doc_info in document_details.get(other_type, [])[:2]:  # 2 exemples par type
                sample_texts.append(f"- [{other_type}] {doc_info['name']}: {doc_info['snippet'][:600]}")
        
        corpus_preview = "\n".join(sample_texts)
        
        # Informations additionnelles pour amÃ©liorer la prÃ©cision
        type_distribution = ", ".join([f"{dt}: {count}" for dt, count in sorted(document_types.items(), key=lambda x: x[1], reverse=True)[:3]])

        # Actions spÃ©cifiques selon le type de document et la langue
        specific_actions_prompts = {
            'fr': {
                'contrat_location': """
Actions spÃ©cifiques pour un contrat de location (bail) :
1. "VÃ©rifier les parties" - Identifie et liste toutes les parties (locataire, bailleur) avec leurs coordonnÃ©es complÃ¨tes
2. "VÃ©rifier les dates" - Extrait toutes les dates importantes : signature, dÃ©but, fin, durÃ©e, prÃ©avis
3. "VÃ©rifier les montants" - Liste tous les montants : loyer, caution, charges, indexation
4. "Analyser les clauses Ã  risque" - Identifie les clauses potentiellement problÃ©matiques ou dÃ©savantageuses
5. "VÃ©rifier les obligations" - Liste les obligations du locataire et du bailleur
6. "VÃ©rifier le bien louÃ©" - DÃ©taille les caractÃ©ristiques du bien (adresse, superficie, type)
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'contrat_travail': """
Actions spÃ©cifiques pour un contrat de travail :
1. "VÃ©rifier les parties" - Identifie l'employeur et l'employÃ© avec leurs coordonnÃ©es
2. "VÃ©rifier les dates" - Extrait les dates : signature, dÃ©but, pÃ©riode d'essai, fin
3. "VÃ©rifier la rÃ©munÃ©ration" - DÃ©taille le salaire, primes, avantages, rÃ©visions
4. "VÃ©rifier les obligations" - Liste les obligations de l'employÃ© et de l'employeur
5. "Analyser les clauses Ã  risque" - Identifie les clauses restrictives ou problÃ©matiques
6. "VÃ©rifier les conditions" - DÃ©taille les conditions de travail, horaires, lieu
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'contrat_vente': """
Actions spÃ©cifiques pour un contrat de vente :
1. "VÃ©rifier les parties" - Identifie l'acheteur et le vendeur avec leurs coordonnÃ©es
2. "VÃ©rifier les dates" - Extrait les dates : signature, livraison, paiement
3. "VÃ©rifier les montants" - DÃ©taille le prix, acompte, modalitÃ©s de paiement
4. "VÃ©rifier l'objet" - DÃ©crit prÃ©cisÃ©ment l'objet de la vente
5. "Analyser les garanties" - Liste les garanties et conditions de garantie
6. "VÃ©rifier les conditions" - DÃ©taille les conditions de vente, dÃ©lais, pÃ©nalitÃ©s
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'contrat_generique': """
Actions spÃ©cifiques pour un contrat :
1. "VÃ©rifier les parties" - Identifie toutes les parties avec leurs coordonnÃ©es
2. "VÃ©rifier les dates importantes" - Extrait toutes les dates clÃ©s du contrat
3. "VÃ©rifier les montants" - Liste tous les montants et modalitÃ©s financiÃ¨res
4. "Analyser les clauses Ã  risque" - Identifie les clauses potentiellement problÃ©matiques
5. "VÃ©rifier les obligations" - Liste les obligations de chaque partie
6. "VÃ©rifier l'objet" - DÃ©crit prÃ©cisÃ©ment l'objet du contrat
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'testament': """
Actions spÃ©cifiques pour un testament :
1. "VÃ©rifier le testateur" - Identifie le testateur et ses coordonnÃ©es
2. "VÃ©rifier les bÃ©nÃ©ficiaires" - Liste tous les bÃ©nÃ©ficiaires et leurs parts
3. "VÃ©rifier les dates" - Extrait les dates : rÃ©daction, signature, modifications
4. "VÃ©rifier les legs" - DÃ©taille tous les legs et hÃ©ritages
5. "VÃ©rifier les conditions" - Liste les conditions et clauses particuliÃ¨res
6. "VÃ©rifier l'exÃ©cuteur" - Identifie l'exÃ©cuteur testamentaire
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'acte_notarie': """
Actions spÃ©cifiques pour un acte notariÃ© :
1. "VÃ©rifier les parties" - Identifie toutes les parties impliquÃ©es
2. "VÃ©rifier les dates" - Extrait toutes les dates importantes
3. "VÃ©rifier les montants" - Liste tous les montants et transactions
4. "VÃ©rifier l'objet" - DÃ©crit prÃ©cisÃ©ment l'objet de l'acte
5. "VÃ©rifier le notaire" - Identifie le notaire et son Ã©tude
6. "VÃ©rifier les conditions" - DÃ©taille les conditions et clauses
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'lettre': """
Actions spÃ©cifiques pour une lettre :
1. "VÃ©rifier l'expÃ©diteur" - Identifie l'expÃ©diteur avec ses coordonnÃ©es
2. "VÃ©rifier le destinataire" - Identifie le destinataire avec ses coordonnÃ©es
3. "VÃ©rifier la date" - Extrait la date de la lettre
4. "VÃ©rifier l'objet" - DÃ©crit l'objet et le but de la lettre
5. "VÃ©rifier infos clÃ©s" - Extrait les informations importantes (montants, rÃ©fÃ©rences, engagements)
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'document_financier': """
Actions spÃ©cifiques pour un document financier :
1. "VÃ©rifier les parties" - Identifie les parties concernÃ©es (employeur, employÃ©, institution)
2. "VÃ©rifier la pÃ©riode" - Extrait la pÃ©riode couverte par le document
3. "VÃ©rifier les montants" - Liste tous les montants (revenus, dÃ©ductions, impÃ´ts, totaux)
4. "VÃ©rifier dÃ©ductions" - DÃ©taille toutes les dÃ©ductions (impÃ´ts, cotisations)
5. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'cv_resume': """
Actions spÃ©cifiques pour un CV/Resume :
1. "VÃ©rifier identitÃ©" - Identifie le nom complet, coordonnÃ©es et informations de contact
2. "VÃ©rifier expÃ©rience" - Liste toutes les expÃ©riences professionnelles avec dates, postes et entreprises
3. "VÃ©rifier formation" - DÃ©taille les diplÃ´mes, formations et certifications avec dates et institutions
4. "VÃ©rifier compÃ©tences" - Liste les compÃ©tences techniques, linguistiques et autres compÃ©tences
5. "VÃ©rifier rÃ©alisations" - Extrait les rÃ©alisations, projets et accomplissements majeurs
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'facture_invoice': """
Actions spÃ©cifiques pour une facture/Invoice :
1. "VÃ©rifier les parties" - Identifie l'Ã©metteur (vendeur) et le client avec leurs coordonnÃ©es
2. "VÃ©rifier les dates" - Extrait la date d'Ã©mission, date d'Ã©chÃ©ance et date de paiement
3. "VÃ©rifier les montants" - Liste le montant HT, TVA, montant TTC et modalitÃ©s de paiement
4. "VÃ©rifier les articles" - DÃ©taille tous les articles/lignes de facturation avec quantitÃ©s et prix
5. "VÃ©rifier rÃ©fÃ©rences" - Extrait le numÃ©ro de facture, rÃ©fÃ©rences client et numÃ©ros de commande
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'contrat_prenuptial': """
Actions spÃ©cifiques pour un contrat de mariage/prÃ©nuptial :
1. "VÃ©rifier les parties" - Identifie les Ã©poux/futurs Ã©poux avec leurs coordonnÃ©es complÃ¨tes
2. "VÃ©rifier les dates" - Extrait la date de signature et la date de mariage prÃ©vue
3. "VÃ©rifier le rÃ©gime matrimonial" - DÃ©taille le rÃ©gime choisi (sÃ©paration de biens, communautÃ©, etc.)
4. "VÃ©rifier les biens" - Liste les biens propres et les biens communs avec leurs valeurs
5. "VÃ©rifier les clauses particuliÃ¨res" - Identifie les clauses spÃ©cifiques (hÃ©ritage, donation, etc.)
6. "Analyser les conditions" - DÃ©taille les conditions de modification et de dissolution
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'procuration_poa': """
Actions spÃ©cifiques pour une procuration/Power Of Attorney :
1. "VÃ©rifier le mandant" - Identifie le mandant (donneur de procuration) avec ses coordonnÃ©es
2. "VÃ©rifier le mandataire" - Identifie le mandataire (reprÃ©sentant) avec ses coordonnÃ©es
3. "VÃ©rifier les dates" - Extrait la date de signature, date de dÃ©but et date d'expiration
4. "VÃ©rifier les pouvoirs" - DÃ©taille tous les pouvoirs accordÃ©s (signature, gestion, dÃ©cisions)
5. "VÃ©rifier les limitations" - Liste les limitations et restrictions des pouvoirs
6. "VÃ©rifier les conditions" - DÃ©taille les conditions de rÃ©vocation et d'utilisation
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'accord_confidentialite_nda': """
Actions spÃ©cifiques pour un accord de confidentialitÃ©/NDA :
1. "VÃ©rifier les parties" - Identifie toutes les parties (divulgateur et bÃ©nÃ©ficiaire) avec leurs coordonnÃ©es
2. "VÃ©rifier les dates" - Extrait la date de signature et la durÃ©e de l'accord
3. "VÃ©rifier les informations confidentielles" - DÃ©taille la portÃ©e des informations couvertes
4. "VÃ©rifier les obligations" - Liste les obligations de confidentialitÃ© et de non-divulgation
5. "Analyser les exceptions" - Identifie les exceptions autorisÃ©es (loi, ordre judiciaire, etc.)
6. "VÃ©rifier les sanctions" - DÃ©taille les pÃ©nalitÃ©s et recours en cas de violation
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'acte_propriete_immobiliere': """
Actions spÃ©cifiques pour un acte de propriÃ©tÃ© immobiliÃ¨re/Real Estate Deed :
1. "VÃ©rifier le propriÃ©taire" - Identifie le propriÃ©taire actuel avec ses coordonnÃ©es
2. "VÃ©rifier le bien" - DÃ©taille les caractÃ©ristiques du bien (adresse, superficie, type, parcelle cadastrale)
3. "VÃ©rifier les dates" - Extrait la date d'acquisition, date de l'acte et historique des transactions
4. "VÃ©rifier les montants" - Liste le prix d'acquisition, taxes et frais associÃ©s
5. "VÃ©rifier les charges" - DÃ©taille les servitudes, hypothÃ¨ques et autres charges
6. "VÃ©rifier le bornage" - Identifie les limites et bornes de la propriÃ©tÃ©
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'assurance_insurance': """
Actions spÃ©cifiques pour un contrat/police d'assurance :
1. "VÃ©rifier l'assurÃ©" - Identifie l'assurÃ© et le souscripteur avec leurs coordonnÃ©es complÃ¨tes
2. "VÃ©rifier la couverture" - DÃ©taille les garanties, les risques couverts et les exclusions
3. "VÃ©rifier les montants" - Liste les primes, franchises, plafonds d'indemnisation
4. "VÃ©rifier les dates" - Extrait les dates de souscription, dÃ©but, fin, renouvellement
5. "Analyser les exclusions" - Identifie toutes les exclusions et limitations de garantie
6. "VÃ©rifier les sinistres" - DÃ©taille les procÃ©dures de dÃ©claration de sinistre et dÃ©lais
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'jugement_decision_justice': """
Actions spÃ©cifiques pour un jugement/dÃ©cision de justice :
1. "VÃ©rifier les parties" - Identifie le demandeur, le dÃ©fendeur et leurs avocats
2. "VÃ©rifier la juridiction" - Identifie le tribunal, la chambre, le(s) juge(s)
3. "VÃ©rifier les dates" - Extrait les dates d'audience, de dÃ©libÃ©rÃ© et de prononcÃ©
4. "Analyser le dispositif" - RÃ©sume les dÃ©cisions rendues (condamnation, dÃ©boutÃ©, etc.)
5. "Analyser les motifs" - RÃ©sume les arguments retenus par le tribunal
6. "VÃ©rifier les voies de recours" - Identifie les possibilitÃ©s d'appel et dÃ©lais
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'releve_bancaire': """
Actions spÃ©cifiques pour un relevÃ© bancaire :
1. "VÃ©rifier le titulaire" - Identifie le titulaire du compte avec ses coordonnÃ©es
2. "VÃ©rifier le compte" - Extrait le numÃ©ro de compte, IBAN, BIC et type de compte
3. "VÃ©rifier la pÃ©riode" - Extrait la pÃ©riode couverte par le relevÃ©
4. "VÃ©rifier les soldes" - Liste le solde initial, solde final et variations
5. "Analyser les opÃ©rations" - RÃ©sume les opÃ©rations principales (virements, prÃ©lÃ¨vements, etc.)
6. "VÃ©rifier les frais" - Identifie tous les frais bancaires et commissions
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'certificat_attestation': """
Actions spÃ©cifiques pour un certificat/attestation :
1. "VÃ©rifier l'Ã©metteur" - Identifie l'organisme ou la personne qui dÃ©livre le certificat
2. "VÃ©rifier le bÃ©nÃ©ficiaire" - Identifie la personne concernÃ©e par le certificat
3. "VÃ©rifier l'objet" - DÃ©crit prÃ©cisÃ©ment ce qui est certifiÃ© ou attestÃ©
4. "VÃ©rifier les dates" - Extrait la date de dÃ©livrance et la durÃ©e de validitÃ©
5. "VÃ©rifier l'authenticitÃ©" - Identifie les Ã©lÃ©ments d'authenticitÃ© (cachet, signature, numÃ©ro)
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'contrat_pret_loan': """
Actions spÃ©cifiques pour un contrat de prÃªt/crÃ©dit :
1. "VÃ©rifier les parties" - Identifie l'emprunteur et le prÃªteur avec leurs coordonnÃ©es
2. "VÃ©rifier les montants" - DÃ©taille le capital empruntÃ©, taux d'intÃ©rÃªt, TAEG, mensualitÃ©s
3. "VÃ©rifier les dates" - Extrait les dates de signature, dÃ©but, fin et Ã©chÃ©ances
4. "VÃ©rifier les garanties" - Liste les garanties exigÃ©es (hypothÃ¨que, caution, nantissement)
5. "Analyser les conditions" - DÃ©taille les conditions de remboursement anticipÃ© et pÃ©nalitÃ©s
6. "VÃ©rifier le tableau d'amortissement" - RÃ©sume le plan de remboursement
7. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'devis_estimation': """
Actions spÃ©cifiques pour un devis/estimation :
1. "VÃ©rifier les parties" - Identifie l'Ã©metteur du devis et le client
2. "VÃ©rifier les articles" - DÃ©taille tous les postes, quantitÃ©s et prix unitaires
3. "VÃ©rifier les montants" - Liste le total HT, TVA, total TTC, remises Ã©ventuelles
4. "VÃ©rifier la validitÃ©" - Extrait la durÃ©e de validitÃ© du devis et conditions
5. "VÃ©rifier les conditions" - DÃ©taille les conditions de paiement et de livraison
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'bon_commande_purchase_order': """
Actions spÃ©cifiques pour un bon de commande :
1. "VÃ©rifier les parties" - Identifie le client et le fournisseur avec leurs coordonnÃ©es
2. "VÃ©rifier les articles" - DÃ©taille les articles commandÃ©s, quantitÃ©s et prix
3. "VÃ©rifier les dates" - Extrait la date de commande et la date de livraison prÃ©vue
4. "VÃ©rifier les conditions" - DÃ©taille les conditions de livraison et de paiement
5. "VÃ©rifier les rÃ©fÃ©rences" - Extrait les numÃ©ros de commande et rÃ©fÃ©rences
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'proces_verbal': """
Actions spÃ©cifiques pour un procÃ¨s-verbal/compte-rendu :
1. "VÃ©rifier les participants" - Liste tous les participants, prÃ©sents et absents
2. "VÃ©rifier les dates" - Extrait la date, l'heure et le lieu de la rÃ©union
3. "VÃ©rifier l'ordre du jour" - Liste les points Ã  l'ordre du jour traitÃ©s
4. "Analyser les dÃ©cisions" - RÃ©sume toutes les dÃ©cisions prises et votes effectuÃ©s
5. "VÃ©rifier les actions" - Liste les actions dÃ©cidÃ©es avec responsables et Ã©chÃ©ances
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'rapport_expertise': """
Actions spÃ©cifiques pour un rapport d'expertise/audit :
1. "VÃ©rifier l'expert" - Identifie l'expert ou l'auditeur et ses qualifications
2. "VÃ©rifier l'objet" - DÃ©crit la mission et le pÃ©rimÃ¨tre de l'expertise
3. "VÃ©rifier la mÃ©thodologie" - RÃ©sume la mÃ©thodologie utilisÃ©e
4. "Analyser les conclusions" - Extrait les conclusions et constats principaux
5. "VÃ©rifier les recommandations" - Liste toutes les recommandations formulÃ©es
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON""",
                'permis_licence': """
Actions spÃ©cifiques pour un permis/licence/autorisation :
1. "VÃ©rifier le titulaire" - Identifie le titulaire du permis avec ses coordonnÃ©es
2. "VÃ©rifier l'autoritÃ©" - Identifie l'organisme Ã©metteur du permis
3. "VÃ©rifier l'objet" - DÃ©crit prÃ©cisÃ©ment ce qui est autorisÃ©
4. "VÃ©rifier les dates" - Extrait la date de dÃ©livrance, d'expiration et de renouvellement
5. "VÃ©rifier les conditions" - Liste les conditions, restrictions et obligations
6. "Extraire donnÃ©es structurÃ©es" - Extrait toutes les donnÃ©es dans un format structurÃ© JSON"""
            },
            'en': {
                'contrat_location': """
Specific actions for a rental contract (lease):
1. "Verify parties" - Identifies and lists all parties (tenant, landlord) with their full contact details
2. "Verify dates" - Extracts all important dates: signature, start, end, duration, notice period
3. "Verify amounts" - Lists all amounts: rent, deposit, charges, indexation
4. "Analyze risky clauses" - Identifies potentially problematic or disadvantageous clauses
5. "Verify obligations" - Lists the obligations of the tenant and landlord
6. "Verify rented property" - Details the property characteristics (address, area, type)
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_travail': """
Specific actions for an employment contract:
1. "Verify parties" - Identifies the employer and employee with their contact details
2. "Verify dates" - Extracts dates: signature, start, probation period, end
3. "Verify remuneration" - Details salary, bonuses, benefits, revisions
4. "Verify obligations" - Lists the obligations of the employee and employer
5. "Analyze risky clauses" - Identifies restrictive or problematic clauses
6. "Verify conditions" - Details working conditions, hours, location
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_vente': """
Specific actions for a sale contract:
1. "Verify parties" - Identifies the buyer and seller with their contact details
2. "Verify dates" - Extracts dates: signature, delivery, payment
3. "Verify amounts" - Details price, down payment, payment terms
4. "Verify subject" - Precisely describes the subject of the sale
5. "Analyze guarantees" - Lists guarantees and warranty conditions
6. "Verify conditions" - Details sale conditions, deadlines, penalties
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_generique': """
Specific actions for a contract:
1. "Verify parties" - Identifies all parties with their contact details
2. "Verify important dates" - Extracts all key dates in the contract
3. "Verify amounts" - Lists all amounts and financial terms
4. "Analyze risky clauses" - Identifies potentially problematic clauses
5. "Verify obligations" - Lists the obligations of each party
6. "Verify subject" - Precisely describes the subject of the contract
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'testament': """
Specific actions for a will:
1. "Verify testator" - Identifies the testator and their contact details
2. "Verify beneficiaries" - Lists all beneficiaries and their shares
3. "Verify dates" - Extracts dates: drafting, signature, modifications
4. "Verify bequests" - Details all bequests and inheritances
5. "Verify conditions" - Lists conditions and particular clauses
6. "Verify executor" - Identifies the executor
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'acte_notarie': """
Specific actions for a notarial act:
1. "Verify parties" - Identifies all parties involved
2. "Verify dates" - Extracts all important dates
3. "Verify amounts" - Lists all amounts and transactions
4. "Verify subject" - Precisely describes the subject of the act
5. "Verify notary" - Identifies the notary and their office
6. "Verify conditions" - Details conditions and clauses
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'lettre': """
Specific actions for a letter:
1. "Verify sender" - Identifies the sender with their contact details
2. "Verify recipient" - Identifies the recipient with their contact details
3. "Verify date" - Extracts the letter date
4. "Verify subject" - Describes the subject and purpose of the letter
5. "Verify key info" - Extracts important information (amounts, references, commitments)
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'document_financier': """
Specific actions for a financial document:
1. "Verify parties" - Identifies parties concerned (employer, employee, institution)
2. "Verify period" - Extracts the period covered by the document
3. "Verify amounts" - Lists all amounts (income, deductions, taxes, totals)
4. "Verify deductions" - Details all deductions (taxes, contributions)
5. "Extract structured data" - Extracts all data in a structured JSON format""",
                'cv_resume': """
Specific actions for a CV/Resume:
1. "Verify identity" - Identifies full name, contact information and personal details
2. "Verify experience" - Lists all professional experiences with dates, positions and companies
3. "Verify education" - Details degrees, training and certifications with dates and institutions
4. "Verify skills" - Lists technical, language and other skills
5. "Verify achievements" - Extracts major achievements, projects and accomplishments
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'facture_invoice': """
Specific actions for an invoice:
1. "Verify parties" - Identifies the issuer (seller) and customer with their contact details
2. "Verify dates" - Extracts issue date, due date and payment date
3. "Verify amounts" - Lists net amount, tax, total amount and payment terms
4. "Verify items" - Details all invoice items/lines with quantities and prices
5. "Verify references" - Extracts invoice number, customer references and order numbers
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_prenuptial': """
Specific actions for a prenuptial agreement:
1. "Verify parties" - Identifies the spouses/future spouses with their full contact details
2. "Verify dates" - Extracts signature date and planned marriage date
3. "Verify marital regime" - Details the chosen regime (separation of property, community, etc.)
4. "Verify assets" - Lists separate and community property with their values
5. "Verify special clauses" - Identifies specific clauses (inheritance, donation, etc.)
6. "Analyze conditions" - Details conditions for modification and dissolution
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'procuration_poa': """
Specific actions for a Power Of Attorney:
1. "Verify principal" - Identifies the principal (grantor) with contact details
2. "Verify agent" - Identifies the agent (attorney-in-fact) with contact details
3. "Verify dates" - Extracts signature date, start date and expiration date
4. "Verify powers" - Details all granted powers (signing, management, decisions)
5. "Verify limitations" - Lists limitations and restrictions of powers
6. "Verify conditions" - Details conditions for revocation and use
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'accord_confidentialite_nda': """
Specific actions for a Non-Disclosure Agreement (NDA):
1. "Verify parties" - Identifies all parties (discloser and recipient) with their contact details
2. "Verify dates" - Extracts signature date and duration of the agreement
3. "Verify confidential information" - Details the scope of covered information
4. "Verify obligations" - Lists confidentiality and non-disclosure obligations
5. "Analyze exceptions" - Identifies authorized exceptions (law, court order, etc.)
6. "Verify penalties" - Details penalties and remedies in case of breach
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'acte_propriete_immobiliere': """
Specific actions for a Real Estate Deed:
1. "Verify owner" - Identifies the current owner with contact details
2. "Verify property" - Details property characteristics (address, area, type, cadastral lot)
3. "Verify dates" - Extracts acquisition date, deed date and transaction history
4. "Verify amounts" - Lists acquisition price, taxes and associated fees
5. "Verify encumbrances" - Details easements, mortgages and other encumbrances
6. "Verify boundaries" - Identifies property limits and boundaries
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'assurance_insurance': """
Specific actions for an insurance policy/contract:
1. "Verify insured" - Identifies the insured and policyholder with their full contact details
2. "Verify coverage" - Details guarantees, covered risks and exclusions
3. "Verify amounts" - Lists premiums, deductibles, indemnity ceilings
4. "Verify dates" - Extracts subscription, start, end and renewal dates
5. "Analyze exclusions" - Identifies all exclusions and coverage limitations
6. "Verify claims" - Details claim declaration procedures and deadlines
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'jugement_decision_justice': """
Specific actions for a court judgment/decision:
1. "Verify parties" - Identifies plaintiff, defendant and their lawyers
2. "Verify jurisdiction" - Identifies the court, chamber, judge(s)
3. "Verify dates" - Extracts hearing, deliberation and ruling dates
4. "Analyze ruling" - Summarizes decisions rendered (conviction, dismissal, etc.)
5. "Analyze reasoning" - Summarizes arguments retained by the court
6. "Verify appeals" - Identifies appeal possibilities and deadlines
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'releve_bancaire': """
Specific actions for a bank statement:
1. "Verify account holder" - Identifies the account holder with contact details
2. "Verify account" - Extracts account number, IBAN, BIC and account type
3. "Verify period" - Extracts the period covered by the statement
4. "Verify balances" - Lists opening balance, closing balance and variations
5. "Analyze transactions" - Summarizes main transactions (transfers, debits, etc.)
6. "Verify fees" - Identifies all bank fees and commissions
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'certificat_attestation': """
Specific actions for a certificate/attestation:
1. "Verify issuer" - Identifies the organization or person issuing the certificate
2. "Verify beneficiary" - Identifies the person concerned by the certificate
3. "Verify subject" - Describes precisely what is certified or attested
4. "Verify dates" - Extracts issue date and validity period
5. "Verify authenticity" - Identifies authenticity elements (stamp, signature, number)
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'contrat_pret_loan': """
Specific actions for a loan/credit agreement:
1. "Verify parties" - Identifies the borrower and lender with their contact details
2. "Verify amounts" - Details borrowed capital, interest rate, APR, monthly payments
3. "Verify dates" - Extracts signing, start, end and maturity dates
4. "Verify collateral" - Lists required guarantees (mortgage, surety, pledge)
5. "Analyze conditions" - Details early repayment conditions and penalties
6. "Verify amortization" - Summarizes the repayment schedule
7. "Extract structured data" - Extracts all data in a structured JSON format""",
                'devis_estimation': """
Specific actions for a quote/estimate:
1. "Verify parties" - Identifies the quote issuer and the client
2. "Verify items" - Details all line items, quantities and unit prices
3. "Verify amounts" - Lists net total, tax, gross total, any discounts
4. "Verify validity" - Extracts quote validity period and conditions
5. "Verify conditions" - Details payment and delivery conditions
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'bon_commande_purchase_order': """
Specific actions for a purchase order:
1. "Verify parties" - Identifies the client and supplier with their contact details
2. "Verify items" - Details ordered items, quantities and prices
3. "Verify dates" - Extracts order date and expected delivery date
4. "Verify conditions" - Details delivery and payment conditions
5. "Verify references" - Extracts order numbers and references
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'proces_verbal': """
Specific actions for meeting minutes/official records:
1. "Verify participants" - Lists all participants, present and absent
2. "Verify dates" - Extracts date, time and location of the meeting
3. "Verify agenda" - Lists agenda items discussed
4. "Analyze decisions" - Summarizes all decisions made and votes taken
5. "Verify action items" - Lists decided actions with responsible parties and deadlines
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'rapport_expertise': """
Specific actions for an expert/audit report:
1. "Verify expert" - Identifies the expert or auditor and their qualifications
2. "Verify subject" - Describes the mission and scope of the expertise
3. "Verify methodology" - Summarizes the methodology used
4. "Analyze conclusions" - Extracts main conclusions and findings
5. "Verify recommendations" - Lists all recommendations made
6. "Extract structured data" - Extracts all data in a structured JSON format""",
                'permis_licence': """
Specific actions for a permit/license/authorization:
1. "Verify holder" - Identifies the permit holder with their contact details
2. "Verify authority" - Identifies the issuing authority
3. "Verify subject" - Describes precisely what is authorized
4. "Verify dates" - Extracts issue date, expiration date and renewal date
5. "Verify conditions" - Lists conditions, restrictions and obligations
6. "Extract structured data" - Extracts all data in a structured JSON format"""
            },
            'es': {
                'contrat_location': """
Acciones especÃ­ficas para un contrato de alquiler (arrendamiento):
1. "Verificar partes" - Identifica y enumera todas las partes (inquilino, arrendador) con sus datos de contacto completos
2. "Verificar fechas" - Extrae todas las fechas importantes: firma, inicio, fin, duraciÃ³n, preaviso
3. "Verificar montos" - Enumera todos los montos: alquiler, depÃ³sito, gastos, indexaciÃ³n
4. "Analizar clÃ¡usulas de riesgo" - Identifica clÃ¡usulas potencialmente problemÃ¡ticas o desventajosas
5. "Verificar obligaciones" - Enumera las obligaciones del inquilino y del arrendador
6. "Verificar propiedad alquilada" - Detalla las caracterÃ­sticas de la propiedad (direcciÃ³n, superficie, tipo)
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_travail': """
Acciones especÃ­ficas para un contrato de trabajo:
1. "Verificar partes" - Identifica al empleador y al empleado con sus datos de contacto
2. "Verificar fechas" - Extrae las fechas: firma, inicio, perÃ­odo de prueba, fin
3. "Verificar remuneraciÃ³n" - Detalla salario, bonos, beneficios, revisiones
4. "Verificar obligaciones" - Enumera las obligaciones del empleado y del empleador
5. "Analizar clÃ¡usulas de riesgo" - Identifica clÃ¡usulas restrictivas o problemÃ¡ticas
6. "Verificar condiciones" - Detalla las condiciones de trabajo, horarios, lugar
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_vente': """
Acciones especÃ­ficas para un contrato de venta:
1. "Verificar partes" - Identifica al comprador y al vendedor con sus datos de contacto
2. "Verificar fechas" - Extrae las fechas: firma, entrega, pago
3. "Verificar montos" - Detalla precio, anticipo, modalidades de pago
4. "Verificar objeto" - Describe precisamente el objeto de la venta
5. "Analizar garantÃ­as" - Enumera las garantÃ­as y condiciones de garantÃ­a
6. "Verificar condiciones" - Detalla las condiciones de venta, plazos, penalizaciones
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_generique': """
Acciones especÃ­ficas para un contrato:
1. "Verificar partes" - Identifica todas las partes con sus datos de contacto
2. "Verificar fechas importantes" - Extrae todas las fechas clave del contrato
3. "Verificar montos" - Enumera todos los montos y modalidades financieras
4. "Analizar clÃ¡usulas de riesgo" - Identifica clÃ¡usulas potencialmente problemÃ¡ticas
5. "Verificar obligaciones" - Enumera las obligaciones de cada parte
6. "Verificar objeto" - Describe precisamente el objeto del contrato
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'testament': """
Acciones especÃ­ficas para un testamento:
1. "Verificar testador" - Identifica al testador y sus datos de contacto
2. "Verificar beneficiarios" - Enumera todos los beneficiarios y sus partes
3. "Verificar fechas" - Extrae las fechas: redacciÃ³n, firma, modificaciones
4. "Verificar legados" - Detalla todos los legados y herencias
5. "Verificar condiciones" - Enumera las condiciones y clÃ¡usulas particulares
6. "Verificar ejecutor" - Identifica al albacea
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'acte_notarie': """
Acciones especÃ­ficas para un acta notarial:
1. "Verificar partes" - Identifica todas las partes involucradas
2. "Verificar fechas" - Extrae todas las fechas importantes
3. "Verificar montos" - Enumera todos los montos y transacciones
4. "Verificar objeto" - Describe precisamente el objeto del acta
5. "Verificar notario" - Identifica al notario y su estudio
6. "Verificar condiciones" - Detalla las condiciones y clÃ¡usulas
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'lettre': """
Acciones especÃ­ficas para una carta:
1. "Verificar remitente" - Identifica al remitente con sus datos de contacto
2. "Verificar destinatario" - Identifica al destinatario con sus datos de contacto
3. "Verificar fecha" - Extrae la fecha de la carta
4. "Verificar objeto" - Describe el objeto y el propÃ³sito de la carta
5. "Verificar informaciÃ³n clave" - Extrae informaciÃ³n importante (montos, referencias, compromisos)
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'document_financier': """
Acciones especÃ­ficas para un documento financiero:
1. "Verificar partes" - Identifica las partes concernidas (empleador, empleado, instituciÃ³n)
2. "Verificar perÃ­odo" - Extrae el perÃ­odo cubierto por el documento
3. "Verificar montos" - Enumera todos los montos (ingresos, deducciones, impuestos, totales)
4. "Verificar deducciones" - Detalla todas las deducciones (impuestos, cotizaciones)
5. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'cv_resume': """
Acciones especÃ­ficas para un CV/Resume:
1. "Verificar identidad" - Identifica el nombre completo, datos de contacto e informaciÃ³n personal
2. "Verificar experiencia" - Enumera todas las experiencias profesionales con fechas, puestos y empresas
3. "Verificar formaciÃ³n" - Detalla tÃ­tulos, formaciones y certificaciones con fechas e instituciones
4. "Verificar competencias" - Enumera las competencias tÃ©cnicas, lingÃ¼Ã­sticas y otras habilidades
5. "Verificar logros" - Extrae los logros, proyectos y realizaciones principales
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'facture_invoice': """
Acciones especÃ­ficas para una factura/Invoice:
1. "Verificar partes" - Identifica el emisor (vendedor) y el cliente con sus datos de contacto
2. "Verificar fechas" - Extrae la fecha de emisiÃ³n, fecha de vencimiento y fecha de pago
3. "Verificar montos" - Enumera el importe sin IVA, IVA, importe total y modalidades de pago
4. "Verificar artÃ­culos" - Detalla todos los artÃ­culos/lÃ­neas de facturaciÃ³n con cantidades y precios
5. "Verificar referencias" - Extrae el nÃºmero de factura, referencias del cliente y nÃºmeros de pedido
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_prenuptial': """
Acciones especÃ­ficas para un contrato de matrimonio/prenupcial:
1. "Verificar partes" - Identifica los cÃ³nyuges/futuros cÃ³nyuges con sus datos de contacto completos
2. "Verificar fechas" - Extrae la fecha de firma y la fecha de matrimonio prevista
3. "Verificar rÃ©gimen matrimonial" - Detalla el rÃ©gimen elegido (separaciÃ³n de bienes, comunidad, etc.)
4. "Verificar bienes" - Enumera los bienes propios y los bienes comunes con sus valores
5. "Verificar clÃ¡usulas particulares" - Identifica las clÃ¡usulas especÃ­ficas (herencia, donaciÃ³n, etc.)
6. "Analizar condiciones" - Detalla las condiciones de modificaciÃ³n y disoluciÃ³n
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'procuration_poa': """
Acciones especÃ­ficas para una procuraciÃ³n/Power Of Attorney:
1. "Verificar mandante" - Identifica al mandante (otorgante) con sus datos de contacto
2. "Verificar mandatario" - Identifica al mandatario (representante) con sus datos de contacto
3. "Verificar fechas" - Extrae la fecha de firma, fecha de inicio y fecha de expiraciÃ³n
4. "Verificar poderes" - Detalla todos los poderes concedidos (firma, gestiÃ³n, decisiones)
5. "Verificar limitaciones" - Enumera las limitaciones y restricciones de los poderes
6. "Verificar condiciones" - Detalla las condiciones de revocaciÃ³n y uso
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'accord_confidentialite_nda': """
Acciones especÃ­ficas para un acuerdo de confidencialidad/NDA:
1. "Verificar partes" - Identifica todas las partes (divulgador y beneficiario) con sus datos de contacto
2. "Verificar fechas" - Extrae la fecha de firma y la duraciÃ³n del acuerdo
3. "Verificar informaciÃ³n confidencial" - Detalla el alcance de la informaciÃ³n cubierta
4. "Verificar obligaciones" - Enumera las obligaciones de confidencialidad y no divulgaciÃ³n
5. "Analizar excepciones" - Identifica las excepciones autorizadas (ley, orden judicial, etc.)
6. "Verificar sanciones" - Detalla las penalizaciones y recursos en caso de violaciÃ³n
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructur ado""",
                'acte_propriete_immobiliere': """
Acciones especÃ­ficas para un acta de propiedad inmobiliaria/Real Estate Deed:
1. "Verificar propietario" - Identifica al propietario actual con sus datos de contacto
2. "Verificar propiedad" - Detalla las caracterÃ­sticas de la propiedad (direcciÃ³n, superficie, tipo, lote catastral)
3. "Verificar fechas" - Extrae la fecha de adquisiciÃ³n, fecha del acta e historial de transacciones
4. "Verificar montos" - Enumera el precio de adquisiciÃ³n, impuestos y honorarios asociados
5. "Verificar cargas" - Detalla las servidumbres, hipotecas y otras cargas
6. "Verificar lÃ­mites" - Identifica los lÃ­mites y linderos de la propiedad
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'assurance_insurance': """
Acciones especÃ­ficas para un contrato/pÃ³liza de seguro:
1. "Verificar asegurado" - Identifica al asegurado y al tomador con sus datos de contacto completos
2. "Verificar cobertura" - Detalla las garantÃ­as, riesgos cubiertos y exclusiones
3. "Verificar montos" - Enumera primas, franquicias, topes de indemnizaciÃ³n
4. "Verificar fechas" - Extrae las fechas de suscripciÃ³n, inicio, fin y renovaciÃ³n
5. "Analizar exclusiones" - Identifica todas las exclusiones y limitaciones de cobertura
6. "Verificar siniestros" - Detalla los procedimientos de declaraciÃ³n de siniestro y plazos
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'jugement_decision_justice': """
Acciones especÃ­ficas para un fallo/decisiÃ³n judicial:
1. "Verificar partes" - Identifica al demandante, demandado y sus abogados
2. "Verificar jurisdicciÃ³n" - Identifica el tribunal, sala, juez(es)
3. "Verificar fechas" - Extrae fechas de audiencia, deliberaciÃ³n y pronunciamiento
4. "Analizar fallo" - Resume las decisiones dictadas (condena, desestimaciÃ³n, etc.)
5. "Analizar fundamentos" - Resume los argumentos retenidos por el tribunal
6. "Verificar recursos" - Identifica las posibilidades de apelaciÃ³n y plazos
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'releve_bancaire': """
Acciones especÃ­ficas para un extracto bancario:
1. "Verificar titular" - Identifica al titular de la cuenta con sus datos de contacto
2. "Verificar cuenta" - Extrae nÃºmero de cuenta, IBAN, BIC y tipo de cuenta
3. "Verificar perÃ­odo" - Extrae el perÃ­odo cubierto por el extracto
4. "Verificar saldos" - Enumera saldo inicial, saldo final y variaciones
5. "Analizar operaciones" - Resume las operaciones principales (transferencias, dÃ©bitos, etc.)
6. "Verificar comisiones" - Identifica todas las comisiones y cargos bancarios
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'certificat_attestation': """
Acciones especÃ­ficas para un certificado/atestaciÃ³n:
1. "Verificar emisor" - Identifica el organismo o persona que emite el certificado
2. "Verificar beneficiario" - Identifica la persona concernida por el certificado
3. "Verificar objeto" - Describe precisamente lo que se certifica o atestigua
4. "Verificar fechas" - Extrae la fecha de emisiÃ³n y el perÃ­odo de validez
5. "Verificar autenticidad" - Identifica elementos de autenticidad (sello, firma, nÃºmero)
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'contrat_pret_loan': """
Acciones especÃ­ficas para un contrato de prÃ©stamo/crÃ©dito:
1. "Verificar partes" - Identifica al prestatario y al prestamista con sus datos de contacto
2. "Verificar montos" - Detalla capital prestado, tasa de interÃ©s, TAE, cuotas mensuales
3. "Verificar fechas" - Extrae fechas de firma, inicio, fin y vencimientos
4. "Verificar garantÃ­as" - Enumera las garantÃ­as exigidas (hipoteca, aval, prenda)
5. "Analizar condiciones" - Detalla condiciones de reembolso anticipado y penalizaciones
6. "Verificar amortizaciÃ³n" - Resume el plan de reembolso
7. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'devis_estimation': """
Acciones especÃ­ficas para un presupuesto/estimaciÃ³n:
1. "Verificar partes" - Identifica al emisor del presupuesto y al cliente
2. "Verificar artÃ­culos" - Detalla todas las partidas, cantidades y precios unitarios
3. "Verificar montos" - Enumera total sin IVA, IVA, total con IVA, descuentos
4. "Verificar validez" - Extrae el perÃ­odo de validez del presupuesto y condiciones
5. "Verificar condiciones" - Detalla condiciones de pago y entrega
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'bon_commande_purchase_order': """
Acciones especÃ­ficas para una orden de compra:
1. "Verificar partes" - Identifica al cliente y al proveedor con sus datos de contacto
2. "Verificar artÃ­culos" - Detalla los artÃ­culos pedidos, cantidades y precios
3. "Verificar fechas" - Extrae la fecha de pedido y la fecha de entrega prevista
4. "Verificar condiciones" - Detalla las condiciones de entrega y pago
5. "Verificar referencias" - Extrae nÃºmeros de pedido y referencias
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'proces_verbal': """
Acciones especÃ­ficas para un acta/minuta de reuniÃ³n:
1. "Verificar participantes" - Enumera todos los participantes, presentes y ausentes
2. "Verificar fechas" - Extrae la fecha, hora y lugar de la reuniÃ³n
3. "Verificar orden del dÃ­a" - Enumera los puntos del orden del dÃ­a tratados
4. "Analizar decisiones" - Resume todas las decisiones tomadas y votaciones realizadas
5. "Verificar acciones" - Enumera las acciones decididas con responsables y plazos
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'rapport_expertise': """
Acciones especÃ­ficas para un informe de peritaje/auditorÃ­a:
1. "Verificar perito" - Identifica al perito o auditor y sus cualificaciones
2. "Verificar objeto" - Describe la misiÃ³n y el perÃ­metro del peritaje
3. "Verificar metodologÃ­a" - Resume la metodologÃ­a utilizada
4. "Analizar conclusiones" - Extrae las conclusiones y hallazgos principales
5. "Verificar recomendaciones" - Enumera todas las recomendaciones formuladas
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado""",
                'permis_licence': """
Acciones especÃ­ficas para un permiso/licencia/autorizaciÃ³n:
1. "Verificar titular" - Identifica al titular del permiso con sus datos de contacto
2. "Verificar autoridad" - Identifica el organismo emisor del permiso
3. "Verificar objeto" - Describe precisamente lo que estÃ¡ autorizado
4. "Verificar fechas" - Extrae fecha de emisiÃ³n, caducidad y renovaciÃ³n
5. "Verificar condiciones" - Enumera condiciones, restricciones y obligaciones
6. "Extraer datos estructurados" - Extrae todos los datos en un formato JSON estructurado"""
            }
        }

        # SÃ©lectionner les prompts spÃ©cifiques selon la langue
        lang_prompts = specific_actions_prompts.get(language, specific_actions_prompts['en'])
        specific_prompt = lang_prompts.get(dominant_type, '')
        
        base_prompts = {
            'fr': f"""
Tu reÃ§ois un aperÃ§u de plusieurs documents importÃ©s par un utilisateur.
Le type de document dominant dÃ©tectÃ© est : {dominant_type}

{specific_prompt}

Ã€ partir de ces textes UNIQUEMENT, propose jusqu'Ã  7 actions suggÃ©rÃ©es au format JSON (elles seront affichÃ©es comme des boutons dans l'interface).
Les actions doivent Ãªtre PRATIQUES et correspondre Ã  ce que les utilisateurs vÃ©rifient habituellement pour ce type de document.

Retourne du JSON STRICT avec exactement cette forme :
{{
  "domain": "etiquette_courte_du_domaine",
  "suggested_actions": [
    {{
      "id": "identifiant_machine",
      "title": "Label court du bouton (max 25 caractÃ¨res)",
      "description": "Une phrase expliquant ce que fait cette action pour l'utilisateur.",
      "sample_prompt": "Prompt complet en langage naturel que l'app pourra envoyer Ã  l'assistant quand l'utilisateur clique sur ce bouton."
    }}
  ]
}}

IMPORTANT :
- Les titres doivent Ãªtre courts et clairs (max 25 caractÃ¨res)
- Les actions doivent Ãªtre spÃ©cifiques au type de document dÃ©tectÃ©
- Priorise les actions que les utilisateurs vÃ©rifient habituellement (parties, dates, montants, clauses, obligations)
- Inclus toujours "Extraire donnÃ©es structurÃ©es" comme derniÃ¨re action
""",
            'en': f"""
You are a document analysis expert. You receive a preview of documents uploaded by a user.

CONTEXTUAL INFORMATION:
- Dominant document type detected: {dominant_type}
- Type distribution in corpus: {type_distribution}
- Total documents analyzed: {len(documents)}

{specific_prompt}

ANALYZE THE PROVIDED DOCUMENTS and propose up to 7 PRECISE and USEFUL suggested actions in JSON format.
Actions must be:
1. SPECIFIC to the actual content of the documents (analyze the content to propose relevant actions)
2. PRACTICAL (what users actually check in this document type)
3. ACTIONABLE (each action should allow extracting or verifying precise information)

DETAILED INSTRUCTIONS:
- First analyze the document content to identify key elements present (dates, amounts, parties, clauses, etc.)
- Propose actions that exactly match what is present in the documents
- Titles must be short and clear (max 25 chars), actionable and direct
- sample_prompt must be precise and request specific information findable in the documents
- Always prioritize: parties/identities â†’ dates â†’ amounts/finances â†’ obligations/clauses â†’ other details
- Always include "Extract structured data" as the last action

Return STRICT valid JSON with this exact shape:
{{
  "domain": "short_domain_label",
  "suggested_actions": [
    {{
      "id": "machine_readable_id",
      "title": "Short button label (max 25 chars)",
      "description": "One sentence explaining what this action does for the user.",
      "sample_prompt": "Complete, precise and actionable natural language prompt the app will send to the assistant when the user clicks this action. The prompt must request specific information present in the documents."
    }}
  ]
}}
""",
            'es': f"""
Eres un experto en anÃ¡lisis documental. Recibes una vista previa de documentos cargados por un usuario.

INFORMACIÃ“N CONTEXTUAL:
- Tipo de documento dominante detectado: {dominant_type}
- DistribuciÃ³n de tipos en el corpus: {type_distribution}
- Total de documentos analizados: {len(documents)}

{specific_prompt}

ANALIZA LOS DOCUMENTOS PROPORCIONADOS y propone hasta 7 acciones sugeridas PRECISAS y ÃšTILES en formato JSON.
Las acciones deben ser:
1. ESPECÃFICAS al contenido real de los documentos (analiza el contenido para proponer acciones relevantes)
2. PRÃCTICAS (lo que los usuarios realmente verifican en este tipo de documento)
3. ACCIONABLES (cada acciÃ³n debe permitir extraer o verificar informaciÃ³n precisa)

INSTRUCCIONES DETALLADAS:
- Primero analiza el contenido del documento para identificar elementos clave presentes (fechas, montos, partes, clÃ¡usulas, etc.)
- Propone acciones que correspondan exactamente a lo que estÃ¡ presente en los documentos
- Los tÃ­tulos deben ser cortos y claros (mÃ¡x 25 caracteres), accionables y directos
- Los sample_prompt deben ser precisos y solicitar informaciÃ³n especÃ­fica encontrable en los documentos
- Prioriza siempre: partes/identidades â†’ fechas â†’ montos/finanzas â†’ obligaciones/clÃ¡usulas â†’ otros detalles
- Siempre incluye "Extraer datos estructurados" como Ãºltima acciÃ³n

Retorna JSON ESTRICTO con exactamente esta forma:
{{
  "domain": "etiqueta_corta_del_dominio",
  "suggested_actions": [
    {{
      "id": "identificador_legible_por_maquina",
      "title": "Etiqueta corta del botÃ³n (mÃ¡x 25 caracteres)",
      "description": "Una oraciÃ³n que explica lo que hace esta acciÃ³n para el usuario.",
      "sample_prompt": "Prompt completo, preciso y accionable en lenguaje natural que la app enviarÃ¡ al asistente cuando el usuario haga clic en esta acciÃ³n. El prompt debe solicitar informaciÃ³n especÃ­fica presente en los documentos."
    }}
  ]
}}
"""
        }
        
        prompt = base_prompts.get(language, base_prompts['en']) + "\n\nCORPUS PREVIEW:\n" + corpus_preview

        # Use Mistral instead of OpenAI for suggested actions generation
        raw_response = call_mistral_api(prompt)

        # Essayer d'extraire un JSON valide
        json_match = re.search(r'\{.*\}', raw_response, re.DOTALL)
        if json_match:
            actions = json.loads(json_match.group())
        else:
            actions = json.loads(raw_response)

        # Validation minimale
        if not isinstance(actions, dict) or "suggested_actions" not in actions:
            raise ValueError("Invalid actions format")

        # Enrichir avec les informations de dÃ©tection
        actions['detected_type'] = dominant_type
        actions['detected_type_confidence'] = round(dominant_confidence, 2)
        actions['type_distribution'] = document_types
        actions['detected_type_label'] = _get_type_label(dominant_type, language)
        
        return actions
    except Exception as e:
        logger.warning(f"âš ï¸  Unable to infer corpus actions: {e}")
        
        # DÃ©tecter le type de document pour le fallback
        try:
            # Utiliser le premier document pour dÃ©tecter le type
            if documents:
                first_doc = documents[0]
                doc_type = await detect_document_type(
                    first_doc.page_content[:2000] if first_doc.page_content else '',
                    first_doc.metadata.get('fileName', '')
                )
            else:
                doc_type = 'document_generique'
        except:
            doc_type = 'document_generique'
        
        # Fallback avec actions spÃ©cifiques selon le type de document et la langue
        fallback_actions = {
            'fr': {
                'contrat_location': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie toutes les parties (locataire, bailleur) avec leurs coordonnÃ©es", "sample_prompt": "Identifie toutes les parties de ce contrat de location : le locataire et le bailleur. Liste leurs noms complets, adresses, tÃ©lÃ©phones et emails si disponibles."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait toutes les dates importantes du contrat", "sample_prompt": "Extrais toutes les dates importantes de ce contrat de location : date de signature, date de dÃ©but, date de fin, durÃ©e, prÃ©avis."},
                    {"id": "verify_amounts", "title": "VÃ©rifier les montants", "description": "Liste tous les montants : loyer, caution, charges", "sample_prompt": "Liste tous les montants mentionnÃ©s dans ce contrat de location : le loyer mensuel, la caution, les charges, et toute indexation prÃ©vue."},
                    {"id": "analyze_risky_clauses", "title": "Analyser clauses Ã  risque", "description": "Identifie les clauses potentiellement problÃ©matiques", "sample_prompt": "Analyse ce contrat de location et identifie les clauses potentiellement problÃ©matiques ou dÃ©savantageuses pour le locataire ou le bailleur."},
                    {"id": "verify_obligations", "title": "VÃ©rifier les obligations", "description": "Liste les obligations du locataire et du bailleur", "sample_prompt": "Liste toutes les obligations du locataire et du bailleur mentionnÃ©es dans ce contrat de location."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce contrat de location : parties, dates, montants, bien louÃ©, clauses importantes."}
                ],
                'contrat_travail': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie l'employeur et l'employÃ©", "sample_prompt": "Identifie les parties de ce contrat de travail : l'employeur et l'employÃ©. Liste leurs noms complets et coordonnÃ©es."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait les dates importantes du contrat", "sample_prompt": "Extrais toutes les dates importantes de ce contrat de travail : date de signature, date de dÃ©but, pÃ©riode d'essai, date de fin si applicable."},
                    {"id": "verify_remuneration", "title": "VÃ©rifier rÃ©munÃ©ration", "description": "DÃ©taille le salaire, primes et avantages", "sample_prompt": "DÃ©taille la rÃ©munÃ©ration dans ce contrat de travail : salaire de base, primes, avantages, rÃ©visions salariales prÃ©vues."},
                    {"id": "verify_obligations", "title": "VÃ©rifier les obligations", "description": "Liste les obligations de l'employÃ© et de l'employeur", "sample_prompt": "Liste toutes les obligations de l'employÃ© et de l'employeur mentionnÃ©es dans ce contrat de travail."},
                    {"id": "analyze_risky_clauses", "title": "Analyser clauses Ã  risque", "description": "Identifie les clauses restrictives ou problÃ©matiques", "sample_prompt": "Analyse ce contrat de travail et identifie les clauses potentiellement restrictives ou problÃ©matiques (clause de non-concurrence, clause d'exclusivitÃ©, etc.)."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce contrat de travail : parties, dates, rÃ©munÃ©ration, obligations, conditions de travail."}
                ],
                'contrat_vente': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie l'acheteur et le vendeur", "sample_prompt": "Identifie les parties de ce contrat de vente : l'acheteur et le vendeur. Liste leurs noms complets et coordonnÃ©es."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait les dates importantes", "sample_prompt": "Extrais toutes les dates importantes de ce contrat de vente : date de signature, date de livraison, dates de paiement."},
                    {"id": "verify_amounts", "title": "VÃ©rifier les montants", "description": "DÃ©taille le prix et modalitÃ©s de paiement", "sample_prompt": "DÃ©taille tous les montants de ce contrat de vente : prix total, acompte, modalitÃ©s de paiement, Ã©chÃ©ances."},
                    {"id": "verify_object", "title": "VÃ©rifier l'objet", "description": "DÃ©crit prÃ©cisÃ©ment l'objet de la vente", "sample_prompt": "DÃ©cris prÃ©cisÃ©ment l'objet de cette vente : nature du bien, caractÃ©ristiques, quantitÃ©, Ã©tat."},
                    {"id": "verify_guarantees", "title": "VÃ©rifier garanties", "description": "Liste les garanties et conditions", "sample_prompt": "Liste toutes les garanties mentionnÃ©es dans ce contrat de vente et leurs conditions."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce contrat de vente : parties, dates, montants, objet, garanties, conditions."}
                ],
                'contrat_generique': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie toutes les parties", "sample_prompt": "Identifie toutes les parties de ce contrat avec leurs noms complets et coordonnÃ©es."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait toutes les dates importantes", "sample_prompt": "Extrais toutes les dates importantes de ce contrat : signature, Ã©chÃ©ances, dates de paiement."},
                    {"id": "verify_amounts", "title": "VÃ©rifier les montants", "description": "Liste tous les montants et modalitÃ©s", "sample_prompt": "Liste tous les montants mentionnÃ©s dans ce contrat et leurs modalitÃ©s de paiement."},
                    {"id": "verify_object", "title": "VÃ©rifier l'objet", "description": "DÃ©crit prÃ©cisÃ©ment l'objet du contrat", "sample_prompt": "DÃ©cris prÃ©cisÃ©ment l'objet de ce contrat en une phrase claire."},
                    {"id": "analyze_risky_clauses", "title": "Analyser clauses Ã  risque", "description": "Identifie les clauses problÃ©matiques", "sample_prompt": "Analyse ce contrat et identifie les clauses potentiellement problÃ©matiques ou dÃ©savantageuses."},
                    {"id": "verify_obligations", "title": "VÃ©rifier les obligations", "description": "Liste les obligations de chaque partie", "sample_prompt": "Liste toutes les obligations de chaque partie mentionnÃ©es dans ce contrat."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce contrat : parties, dates, montants, objet, obligations, clauses importantes."}
                ],
                'testament': [
                    {"id": "verify_testator", "title": "VÃ©rifier le testateur", "description": "Identifie le testateur", "sample_prompt": "Identifie le testateur de ce testament avec ses coordonnÃ©es complÃ¨tes."},
                    {"id": "verify_beneficiaries", "title": "VÃ©rifier bÃ©nÃ©ficiaires", "description": "Liste tous les bÃ©nÃ©ficiaires et leurs parts", "sample_prompt": "Liste tous les bÃ©nÃ©ficiaires de ce testament et leurs parts respectives."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait les dates importantes", "sample_prompt": "Extrais toutes les dates importantes de ce testament : date de rÃ©daction, signature, modifications Ã©ventuelles."},
                    {"id": "verify_legacies", "title": "VÃ©rifier les legs", "description": "DÃ©taille tous les legs et hÃ©ritages", "sample_prompt": "DÃ©taille tous les legs et hÃ©ritages mentionnÃ©s dans ce testament."},
                    {"id": "verify_executor", "title": "VÃ©rifier l'exÃ©cuteur", "description": "Identifie l'exÃ©cuteur testamentaire", "sample_prompt": "Identifie l'exÃ©cuteur testamentaire mentionnÃ© dans ce testament."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce testament : testateur, bÃ©nÃ©ficiaires, legs, dates, exÃ©cuteur."}
                ],
                'acte_notarie': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie toutes les parties", "sample_prompt": "Identifie toutes les parties impliquÃ©es dans cet acte notariÃ© avec leurs coordonnÃ©es."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait toutes les dates importantes", "sample_prompt": "Extrais toutes les dates importantes de cet acte notariÃ©."},
                    {"id": "verify_amounts", "title": "VÃ©rifier les montants", "description": "Liste tous les montants et transactions", "sample_prompt": "Liste tous les montants et transactions mentionnÃ©s dans cet acte notariÃ©."},
                    {"id": "verify_object", "title": "VÃ©rifier l'objet", "description": "DÃ©crit prÃ©cisÃ©ment l'objet de l'acte", "sample_prompt": "DÃ©cris prÃ©cisÃ©ment l'objet de cet acte notariÃ©."},
                    {"id": "verify_notary", "title": "VÃ©rifier le notaire", "description": "Identifie le notaire et son Ã©tude", "sample_prompt": "Identifie le notaire qui a rÃ©digÃ© cet acte et son Ã©tude notariale."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de cet acte notariÃ© : parties, dates, montants, objet, notaire."}
                ],
                'lettre': [
                    {"id": "verify_sender", "title": "VÃ©rifier l'expÃ©diteur", "description": "Identifie l'expÃ©diteur de la lettre", "sample_prompt": "Identifie l'expÃ©diteur de cette lettre : nom, fonction, organisation, coordonnÃ©es."},
                    {"id": "verify_recipient", "title": "VÃ©rifier le destinataire", "description": "Identifie le destinataire de la lettre", "sample_prompt": "Identifie le destinataire de cette lettre : nom, fonction, organisation, coordonnÃ©es."},
                    {"id": "verify_date", "title": "VÃ©rifier la date", "description": "Extrait la date de la lettre", "sample_prompt": "Extrais la date de cette lettre (date d'Ã©criture, date d'envoi si mentionnÃ©e)."},
                    {"id": "verify_object", "title": "VÃ©rifier l'objet", "description": "DÃ©crit l'objet et le but de la lettre", "sample_prompt": "DÃ©cris l'objet et le but principal de cette lettre en une phrase claire."},
                    {"id": "verify_key_information", "title": "VÃ©rifier infos clÃ©s", "description": "Extrait les informations importantes mentionnÃ©es", "sample_prompt": "Extrais toutes les informations importantes mentionnÃ©es dans cette lettre : montants, dates, rÃ©fÃ©rences, engagements."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de cette lettre : expÃ©diteur, destinataire, date, objet, informations clÃ©s."}
                ],
                'document_financier': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie les parties concernÃ©es", "sample_prompt": "Identifie toutes les parties concernÃ©es par ce document financier : employeur, employÃ©, institution, etc."},
                    {"id": "verify_period", "title": "VÃ©rifier la pÃ©riode", "description": "Extrait la pÃ©riode couverte", "sample_prompt": "Extrais la pÃ©riode couverte par ce document financier : dates de dÃ©but et de fin."},
                    {"id": "verify_amounts", "title": "VÃ©rifier les montants", "description": "Liste tous les montants et totaux", "sample_prompt": "Liste tous les montants mentionnÃ©s dans ce document financier : revenus, dÃ©ductions, impÃ´ts, totaux."},
                    {"id": "verify_deductions", "title": "VÃ©rifier dÃ©ductions", "description": "DÃ©taille toutes les dÃ©ductions", "sample_prompt": "DÃ©taille toutes les dÃ©ductions mentionnÃ©es dans ce document financier : impÃ´ts, cotisations, autres dÃ©ductions."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce document financier : parties, pÃ©riode, montants, dÃ©ductions, totaux."}
                ],
                'assurance_insurance': [
                    {"id": "verify_insured", "title": "VÃ©rifier l'assurÃ©", "description": "Identifie l'assurÃ© et le souscripteur", "sample_prompt": "Identifie l'assurÃ© et le souscripteur de cette police d'assurance avec leurs coordonnÃ©es complÃ¨tes."},
                    {"id": "verify_coverage", "title": "VÃ©rifier couverture", "description": "DÃ©taille les garanties et exclusions", "sample_prompt": "DÃ©taille les garanties couvertes, les risques couverts et les exclusions de cette police d'assurance."},
                    {"id": "verify_amounts", "title": "VÃ©rifier les montants", "description": "Liste primes, franchises, plafonds", "sample_prompt": "Liste les montants : prime annuelle, franchise, plafond d'indemnisation de cette police d'assurance."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait les dates de souscription et Ã©chÃ©ance", "sample_prompt": "Extrais les dates importantes : souscription, effet, Ã©chÃ©ance, renouvellement de cette police d'assurance."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de cette police d'assurance : assurÃ©, couverture, montants, dates."}
                ],
                'jugement_decision_justice': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie demandeur, dÃ©fendeur et avocats", "sample_prompt": "Identifie les parties de ce jugement : demandeur, dÃ©fendeur et leurs avocats respectifs."},
                    {"id": "verify_jurisdiction", "title": "VÃ©rifier juridiction", "description": "Identifie le tribunal et le juge", "sample_prompt": "Identifie le tribunal, la chambre et le(s) juge(s) ayant rendu cette dÃ©cision."},
                    {"id": "analyze_ruling", "title": "Analyser le dispositif", "description": "RÃ©sume la dÃ©cision rendue", "sample_prompt": "RÃ©sume la dÃ©cision rendue dans ce jugement : condamnations, montants, obligations."},
                    {"id": "verify_appeals", "title": "VÃ©rifier recours", "description": "Identifie les voies de recours", "sample_prompt": "Identifie les voies de recours possibles et les dÃ©lais d'appel mentionnÃ©s dans ce jugement."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce jugement : parties, juridiction, dÃ©cision, recours."}
                ],
                'releve_bancaire': [
                    {"id": "verify_holder", "title": "VÃ©rifier le titulaire", "description": "Identifie le titulaire du compte", "sample_prompt": "Identifie le titulaire du compte avec ses coordonnÃ©es complÃ¨tes."},
                    {"id": "verify_account", "title": "VÃ©rifier le compte", "description": "Extrait numÃ©ro de compte, IBAN, BIC", "sample_prompt": "Extrais le numÃ©ro de compte, IBAN, BIC et type de compte de ce relevÃ© bancaire."},
                    {"id": "verify_balances", "title": "VÃ©rifier les soldes", "description": "Liste solde initial, final et variations", "sample_prompt": "Liste le solde initial, le solde final, le total des dÃ©bits et crÃ©dits de ce relevÃ© bancaire."},
                    {"id": "analyze_transactions", "title": "Analyser opÃ©rations", "description": "RÃ©sume les opÃ©rations principales", "sample_prompt": "RÃ©sume les opÃ©rations principales de ce relevÃ© bancaire : virements, prÃ©lÃ¨vements, paiements."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce relevÃ© bancaire : titulaire, compte, soldes, opÃ©rations."}
                ],
                'certificat_attestation': [
                    {"id": "verify_issuer", "title": "VÃ©rifier l'Ã©metteur", "description": "Identifie l'organisme Ã©metteur", "sample_prompt": "Identifie l'organisme ou la personne qui a Ã©mis ce certificat/attestation."},
                    {"id": "verify_beneficiary", "title": "VÃ©rifier bÃ©nÃ©ficiaire", "description": "Identifie la personne concernÃ©e", "sample_prompt": "Identifie la personne concernÃ©e par ce certificat/attestation avec ses coordonnÃ©es."},
                    {"id": "verify_subject", "title": "VÃ©rifier l'objet", "description": "DÃ©crit ce qui est certifiÃ©", "sample_prompt": "DÃ©cris prÃ©cisÃ©ment ce qui est certifiÃ© ou attestÃ© dans ce document."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait date d'Ã©mission et validitÃ©", "sample_prompt": "Extrais la date d'Ã©mission et la pÃ©riode de validitÃ© de ce certificat/attestation."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce certificat : Ã©metteur, bÃ©nÃ©ficiaire, objet, dates."}
                ],
                'contrat_pret_loan': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie emprunteur et prÃªteur", "sample_prompt": "Identifie l'emprunteur et le prÃªteur de ce contrat de prÃªt avec leurs coordonnÃ©es."},
                    {"id": "verify_amounts", "title": "VÃ©rifier les montants", "description": "DÃ©taille capital, taux, mensualitÃ©s", "sample_prompt": "DÃ©taille les montants : capital empruntÃ©, taux d'intÃ©rÃªt, TAEG, mensualitÃ©s de ce contrat de prÃªt."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait dates de signature et Ã©chÃ©ances", "sample_prompt": "Extrais les dates : signature, dÃ©but, fin, Ã©chÃ©ances de ce contrat de prÃªt."},
                    {"id": "verify_guarantees", "title": "VÃ©rifier garanties", "description": "Liste les garanties exigÃ©es", "sample_prompt": "Liste les garanties exigÃ©es dans ce contrat de prÃªt : hypothÃ¨que, caution, nantissement."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce contrat de prÃªt : parties, montants, dates, garanties."}
                ],
                'devis_estimation': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie Ã©metteur et client", "sample_prompt": "Identifie l'Ã©metteur du devis et le client avec leurs coordonnÃ©es."},
                    {"id": "verify_items", "title": "VÃ©rifier les postes", "description": "DÃ©taille les postes et prix", "sample_prompt": "DÃ©taille tous les postes du devis : description, quantitÃ©s, prix unitaires et totaux."},
                    {"id": "verify_amounts", "title": "VÃ©rifier les montants", "description": "Liste total HT, TVA, TTC", "sample_prompt": "Liste les montants : total HT, TVA, total TTC, remises Ã©ventuelles de ce devis."},
                    {"id": "verify_validity", "title": "VÃ©rifier validitÃ©", "description": "Extrait la durÃ©e de validitÃ©", "sample_prompt": "Extrais la durÃ©e de validitÃ© et les conditions de ce devis."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce devis : parties, postes, montants, validitÃ©."}
                ],
                'bon_commande_purchase_order': [
                    {"id": "verify_parties", "title": "VÃ©rifier les parties", "description": "Identifie client et fournisseur", "sample_prompt": "Identifie le client et le fournisseur de ce bon de commande avec leurs coordonnÃ©es."},
                    {"id": "verify_items", "title": "VÃ©rifier les articles", "description": "DÃ©taille articles, quantitÃ©s, prix", "sample_prompt": "DÃ©taille les articles commandÃ©s : rÃ©fÃ©rences, descriptions, quantitÃ©s et prix."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait date de commande et livraison", "sample_prompt": "Extrais la date de commande et la date de livraison prÃ©vue de ce bon de commande."},
                    {"id": "verify_conditions", "title": "VÃ©rifier conditions", "description": "DÃ©taille conditions de livraison et paiement", "sample_prompt": "DÃ©taille les conditions de livraison et de paiement de ce bon de commande."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce bon de commande : parties, articles, dates, conditions."}
                ],
                'proces_verbal': [
                    {"id": "verify_participants", "title": "VÃ©rifier participants", "description": "Liste prÃ©sents et absents", "sample_prompt": "Liste tous les participants de ce procÃ¨s-verbal : prÃ©sents, absents et excusÃ©s."},
                    {"id": "verify_dates", "title": "VÃ©rifier date et lieu", "description": "Extrait date, heure et lieu", "sample_prompt": "Extrais la date, l'heure et le lieu de la rÃ©union dÃ©crite dans ce procÃ¨s-verbal."},
                    {"id": "analyze_decisions", "title": "Analyser dÃ©cisions", "description": "RÃ©sume les dÃ©cisions prises", "sample_prompt": "RÃ©sume toutes les dÃ©cisions prises et les votes rÃ©alisÃ©s dans ce procÃ¨s-verbal."},
                    {"id": "verify_actions", "title": "VÃ©rifier actions", "description": "Liste les actions dÃ©cidÃ©es", "sample_prompt": "Liste les actions dÃ©cidÃ©es avec les responsables et les Ã©chÃ©ances mentionnÃ©s dans ce procÃ¨s-verbal."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce procÃ¨s-verbal : participants, dÃ©cisions, actions, dates."}
                ],
                'rapport_expertise': [
                    {"id": "verify_expert", "title": "VÃ©rifier l'expert", "description": "Identifie l'expert et ses qualifications", "sample_prompt": "Identifie l'expert ou l'auditeur et ses qualifications dans ce rapport."},
                    {"id": "verify_subject", "title": "VÃ©rifier l'objet", "description": "DÃ©crit la mission et le pÃ©rimÃ¨tre", "sample_prompt": "DÃ©cris la mission et le pÃ©rimÃ¨tre de l'expertise de ce rapport."},
                    {"id": "analyze_conclusions", "title": "Analyser conclusions", "description": "Extrait les conclusions principales", "sample_prompt": "Extrais les conclusions et constats principaux de ce rapport d'expertise."},
                    {"id": "verify_recommendations", "title": "VÃ©rifier recommandations", "description": "Liste les recommandations", "sample_prompt": "Liste toutes les recommandations formulÃ©es dans ce rapport d'expertise."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce rapport : expert, objet, conclusions, recommandations."}
                ],
                'permis_licence': [
                    {"id": "verify_holder", "title": "VÃ©rifier le titulaire", "description": "Identifie le titulaire du permis", "sample_prompt": "Identifie le titulaire de ce permis/licence avec ses coordonnÃ©es complÃ¨tes."},
                    {"id": "verify_authority", "title": "VÃ©rifier l'autoritÃ©", "description": "Identifie l'organisme Ã©metteur", "sample_prompt": "Identifie l'autoritÃ© ou l'organisme qui a dÃ©livrÃ© ce permis/licence."},
                    {"id": "verify_subject", "title": "VÃ©rifier l'objet", "description": "DÃ©crit ce qui est autorisÃ©", "sample_prompt": "DÃ©cris prÃ©cisÃ©ment ce qui est autorisÃ© par ce permis/licence."},
                    {"id": "verify_dates", "title": "VÃ©rifier les dates", "description": "Extrait dates de dÃ©livrance et expiration", "sample_prompt": "Extrais les dates de dÃ©livrance, d'expiration et de renouvellement de ce permis/licence."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce permis : titulaire, autoritÃ©, objet, dates, conditions."}
                ],
                'default': [
                    {"id": "summarize_all", "title": "RÃ©sumer documents", "description": "GÃ©nÃ¨re un rÃ©sumÃ© global des documents", "sample_prompt": "Fournis un rÃ©sumÃ© clair et structurÃ© de tous les documents uploadÃ©s, en mettant en Ã©vidence les thÃ¨mes principaux et les informations importantes. Adapte le rÃ©sumÃ© au type de document : pour un CV, concentre-toi sur l'expÃ©rience professionnelle, les compÃ©tences et les rÃ©alisations ; pour un document financier, mentionne les montants et chiffres pertinents ; pour un contrat, mentionne les parties et dates importantes. Ne mentionne PAS d'informations qui ne sont pas prÃ©sentes dans les documents (par exemple, ne mentionne pas d'informations financiÃ¨res si le document est un CV)."},
                    {"id": "extract_key_points", "title": "Extraire points clÃ©s", "description": "Liste les points clÃ©s et entitÃ©s", "sample_prompt": "Extrais les points clÃ©s, dÃ©cisions importantes et entitÃ©s nommÃ©es (personnes, entreprises, lieux) de tous les documents uploadÃ©s et organise-les en puces."},
                    {"id": "extract_structured", "title": "Extraire donnÃ©es structurÃ©es", "description": "Extrait toutes les donnÃ©es dans un format structurÃ©", "sample_prompt": "Extrait toutes les donnÃ©es structurÃ©es de ce document : parties, dates, montants, informations clÃ©s."}
                ]
            },
            'en': {
                'contrat_location': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies all parties (tenant, landlord) with their contact details", "sample_prompt": "Identify all parties in this rental contract: the tenant and the landlord. List their full names, addresses, phone numbers and emails if available."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts all important dates from the contract", "sample_prompt": "Extract all important dates from this rental contract: signing date, start date, end date, duration, notice period."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists all amounts: rent, deposit, charges", "sample_prompt": "List all amounts mentioned in this rental contract: monthly rent, deposit, charges, and any planned indexation."},
                    {"id": "analyze_risky_clauses", "title": "Analyze risky clauses", "description": "Identifies potentially problematic clauses", "sample_prompt": "Analyze this rental contract and identify potentially problematic or disadvantageous clauses for the tenant or landlord."},
                    {"id": "verify_obligations", "title": "Verify obligations", "description": "Lists obligations of tenant and landlord", "sample_prompt": "List all obligations of the tenant and landlord mentioned in this rental contract."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this rental contract: parties, dates, amounts, rented property, important clauses."}
                ],
                'contrat_travail': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies employer and employee", "sample_prompt": "Identify the parties in this employment contract: the employer and the employee. List their full names and contact details."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts important dates from the contract", "sample_prompt": "Extract all important dates from this employment contract: signing date, start date, probation period, end date if applicable."},
                    {"id": "verify_remuneration", "title": "Verify remuneration", "description": "Details salary, bonuses and benefits", "sample_prompt": "Detail the remuneration in this employment contract: base salary, bonuses, benefits, planned salary revisions."},
                    {"id": "verify_obligations", "title": "Verify obligations", "description": "Lists obligations of employee and employer", "sample_prompt": "List all obligations of the employee and employer mentioned in this employment contract."},
                    {"id": "analyze_risky_clauses", "title": "Analyze risky clauses", "description": "Identifies restrictive or problematic clauses", "sample_prompt": "Analyze this employment contract and identify potentially restrictive or problematic clauses (non-compete clause, exclusivity clause, etc.)."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this employment contract: parties, dates, remuneration, obligations, working conditions."}
                ],
                'contrat_vente': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies buyer and seller", "sample_prompt": "Identify the parties in this sale contract: the buyer and the seller. List their full names and contact details."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts important dates", "sample_prompt": "Extract all important dates from this sale contract: signing date, delivery date, payment dates."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Details price and payment terms", "sample_prompt": "Detail all amounts in this sale contract: total price, down payment, payment terms, deadlines."},
                    {"id": "verify_object", "title": "Verify subject", "description": "Precisely describes the subject of the sale", "sample_prompt": "Precisely describe the subject of this sale: nature of the property, characteristics, quantity, condition."},
                    {"id": "verify_guarantees", "title": "Verify guarantees", "description": "Lists guarantees and conditions", "sample_prompt": "List all guarantees mentioned in this sale contract and their conditions."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this sale contract: parties, dates, amounts, subject, guarantees, conditions."}
                ],
                'contrat_generique': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies all parties", "sample_prompt": "Identify all parties in this contract with their full names and contact details."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts all important dates", "sample_prompt": "Extract all important dates from this contract: signing, deadlines, payment dates."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists all amounts and terms", "sample_prompt": "List all amounts mentioned in this contract and their payment terms."},
                    {"id": "verify_object", "title": "Verify subject", "description": "Precisely describes the subject of the contract", "sample_prompt": "Precisely describe the subject of this contract in a clear sentence."},
                    {"id": "analyze_risky_clauses", "title": "Analyze risky clauses", "description": "Identifies problematic clauses", "sample_prompt": "Analyze this contract and identify potentially problematic or disadvantageous clauses."},
                    {"id": "verify_obligations", "title": "Verify obligations", "description": "Lists obligations of each party", "sample_prompt": "List all obligations of each party mentioned in this contract."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this contract: parties, dates, amounts, subject, obligations, important clauses."}
                ],
                'testament': [
                    {"id": "verify_testator", "title": "Verify testator", "description": "Identifies the testator", "sample_prompt": "Identify the testator of this will with their complete contact details."},
                    {"id": "verify_beneficiaries", "title": "Verify beneficiaries", "description": "Lists all beneficiaries and their shares", "sample_prompt": "List all beneficiaries of this will and their respective shares."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts important dates", "sample_prompt": "Extract all important dates from this will: drafting date, signature, any modifications."},
                    {"id": "verify_legacies", "title": "Verify bequests", "description": "Details all bequests and inheritances", "sample_prompt": "Detail all bequests and inheritances mentioned in this will."},
                    {"id": "verify_executor", "title": "Verify executor", "description": "Identifies the executor", "sample_prompt": "Identify the executor mentioned in this will."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this will: testator, beneficiaries, bequests, dates, executor."}
                ],
                'acte_notarie': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies all parties", "sample_prompt": "Identify all parties involved in this notarial act with their contact details."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts all important dates", "sample_prompt": "Extract all important dates from this notarial act."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists all amounts and transactions", "sample_prompt": "List all amounts and transactions mentioned in this notarial act."},
                    {"id": "verify_object", "title": "Verify subject", "description": "Precisely describes the subject of the act", "sample_prompt": "Precisely describe the subject of this notarial act."},
                    {"id": "verify_notary", "title": "Verify notary", "description": "Identifies the notary and their office", "sample_prompt": "Identify the notary who drafted this act and their notarial office."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this notarial act: parties, dates, amounts, subject, notary."}
                ],
                'lettre': [
                    {"id": "verify_sender", "title": "Verify sender", "description": "Identifies the letter sender", "sample_prompt": "Identify the sender of this letter: name, function, organization, contact details."},
                    {"id": "verify_recipient", "title": "Verify recipient", "description": "Identifies the letter recipient", "sample_prompt": "Identify the recipient of this letter: name, function, organization, contact details."},
                    {"id": "verify_date", "title": "Verify date", "description": "Extracts the letter date", "sample_prompt": "Extract the date of this letter (writing date, sending date if mentioned)."},
                    {"id": "verify_object", "title": "Verify subject", "description": "Describes the subject and purpose of the letter", "sample_prompt": "Describe the subject and main purpose of this letter in a clear sentence."},
                    {"id": "verify_key_information", "title": "Verify key info", "description": "Extracts important information mentioned", "sample_prompt": "Extract all important information mentioned in this letter: amounts, dates, references, commitments."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this letter: sender, recipient, date, subject, key information."}
                ],
                'document_financier': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies concerned parties", "sample_prompt": "Identify all parties concerned by this financial document: employer, employee, institution, etc."},
                    {"id": "verify_period", "title": "Verify period", "description": "Extracts the covered period", "sample_prompt": "Extract the period covered by this financial document: start and end dates."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists all amounts and totals", "sample_prompt": "List all amounts mentioned in this financial document: income, deductions, taxes, totals."},
                    {"id": "verify_deductions", "title": "Verify deductions", "description": "Details all deductions", "sample_prompt": "Detail all deductions mentioned in this financial document: taxes, contributions, other deductions."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this financial document: parties, period, amounts, deductions, totals."}
                ],
                'assurance_insurance': [
                    {"id": "verify_insured", "title": "Verify insured", "description": "Identifies insured party and policyholder", "sample_prompt": "Identify the insured party and the policyholder of this insurance policy with their full contact details."},
                    {"id": "verify_coverage", "title": "Verify coverage", "description": "Details guarantees and exclusions", "sample_prompt": "Detail the guarantees covered, risks covered and exclusions of this insurance policy."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists premiums, deductibles, ceilings", "sample_prompt": "List the amounts: annual premium, deductible, compensation ceiling of this insurance policy."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts subscription and expiry dates", "sample_prompt": "Extract the important dates: subscription, effective date, expiry, renewal of this insurance policy."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this insurance policy: insured, coverage, amounts, dates."}
                ],
                'jugement_decision_justice': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies plaintiff, defendant and lawyers", "sample_prompt": "Identify the parties in this court ruling: plaintiff, defendant and their respective lawyers."},
                    {"id": "verify_jurisdiction", "title": "Verify jurisdiction", "description": "Identifies court and judge", "sample_prompt": "Identify the court, chamber and judge(s) who rendered this decision."},
                    {"id": "analyze_ruling", "title": "Analyze ruling", "description": "Summarizes the decision", "sample_prompt": "Summarize the decision rendered in this ruling: convictions, amounts, obligations."},
                    {"id": "verify_appeals", "title": "Verify appeals", "description": "Identifies appeal options", "sample_prompt": "Identify possible appeal options and deadlines mentioned in this ruling."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this ruling: parties, jurisdiction, decision, appeals."}
                ],
                'releve_bancaire': [
                    {"id": "verify_holder", "title": "Verify holder", "description": "Identifies account holder", "sample_prompt": "Identify the account holder with their full contact details."},
                    {"id": "verify_account", "title": "Verify account", "description": "Extracts account number, IBAN, BIC", "sample_prompt": "Extract the account number, IBAN, BIC and account type from this bank statement."},
                    {"id": "verify_balances", "title": "Verify balances", "description": "Lists opening, closing balances and variations", "sample_prompt": "List the opening balance, closing balance, total debits and credits of this bank statement."},
                    {"id": "analyze_transactions", "title": "Analyze transactions", "description": "Summarizes main transactions", "sample_prompt": "Summarize the main transactions in this bank statement: transfers, direct debits, payments."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this bank statement: holder, account, balances, transactions."}
                ],
                'certificat_attestation': [
                    {"id": "verify_issuer", "title": "Verify issuer", "description": "Identifies issuing organization", "sample_prompt": "Identify the organization or person that issued this certificate/attestation."},
                    {"id": "verify_beneficiary", "title": "Verify beneficiary", "description": "Identifies the person concerned", "sample_prompt": "Identify the person concerned by this certificate/attestation with their contact details."},
                    {"id": "verify_subject", "title": "Verify subject", "description": "Describes what is certified", "sample_prompt": "Precisely describe what is certified or attested in this document."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts issue date and validity", "sample_prompt": "Extract the issue date and validity period of this certificate/attestation."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this certificate: issuer, beneficiary, subject, dates."}
                ],
                'contrat_pret_loan': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies borrower and lender", "sample_prompt": "Identify the borrower and lender of this loan agreement with their contact details."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Details principal, rate, installments", "sample_prompt": "Detail the amounts: principal borrowed, interest rate, APR, monthly installments of this loan agreement."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts signing dates and deadlines", "sample_prompt": "Extract the dates: signing, start, end, deadlines of this loan agreement."},
                    {"id": "verify_guarantees", "title": "Verify guarantees", "description": "Lists required guarantees", "sample_prompt": "List the guarantees required in this loan agreement: mortgage, surety, pledge."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this loan agreement: parties, amounts, dates, guarantees."}
                ],
                'devis_estimation': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies issuer and client", "sample_prompt": "Identify the quote issuer and the client with their contact details."},
                    {"id": "verify_items", "title": "Verify line items", "description": "Details items and prices", "sample_prompt": "Detail all line items in this quote: description, quantities, unit prices and totals."},
                    {"id": "verify_amounts", "title": "Verify amounts", "description": "Lists subtotal, tax, total", "sample_prompt": "List the amounts: subtotal, tax, grand total, any discounts in this quote."},
                    {"id": "verify_validity", "title": "Verify validity", "description": "Extracts validity period", "sample_prompt": "Extract the validity period and conditions of this quote."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this quote: parties, line items, amounts, validity."}
                ],
                'bon_commande_purchase_order': [
                    {"id": "verify_parties", "title": "Verify parties", "description": "Identifies customer and supplier", "sample_prompt": "Identify the customer and supplier of this purchase order with their contact details."},
                    {"id": "verify_items", "title": "Verify items", "description": "Details articles, quantities, prices", "sample_prompt": "Detail the ordered items: references, descriptions, quantities and prices."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts order and delivery dates", "sample_prompt": "Extract the order date and expected delivery date of this purchase order."},
                    {"id": "verify_conditions", "title": "Verify conditions", "description": "Details delivery and payment conditions", "sample_prompt": "Detail the delivery and payment conditions of this purchase order."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this purchase order: parties, items, dates, conditions."}
                ],
                'proces_verbal': [
                    {"id": "verify_participants", "title": "Verify participants", "description": "Lists attendees and absentees", "sample_prompt": "List all participants in these meeting minutes: attendees, absentees and excused."},
                    {"id": "verify_dates", "title": "Verify date and venue", "description": "Extracts date, time and location", "sample_prompt": "Extract the date, time and location of the meeting described in these minutes."},
                    {"id": "analyze_decisions", "title": "Analyze decisions", "description": "Summarizes decisions taken", "sample_prompt": "Summarize all decisions taken and votes conducted in these meeting minutes."},
                    {"id": "verify_actions", "title": "Verify action items", "description": "Lists decided action items", "sample_prompt": "List the action items decided with owners and deadlines mentioned in these minutes."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from these minutes: participants, decisions, action items, dates."}
                ],
                'rapport_expertise': [
                    {"id": "verify_expert", "title": "Verify expert", "description": "Identifies expert and qualifications", "sample_prompt": "Identify the expert or auditor and their qualifications in this report."},
                    {"id": "verify_subject", "title": "Verify subject", "description": "Describes the mission and scope", "sample_prompt": "Describe the mission and scope of the expertise in this report."},
                    {"id": "analyze_conclusions", "title": "Analyze conclusions", "description": "Extracts main conclusions", "sample_prompt": "Extract the main conclusions and findings from this expert report."},
                    {"id": "verify_recommendations", "title": "Verify recommendations", "description": "Lists recommendations", "sample_prompt": "List all recommendations made in this expert report."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this report: expert, subject, conclusions, recommendations."}
                ],
                'permis_licence': [
                    {"id": "verify_holder", "title": "Verify holder", "description": "Identifies permit holder", "sample_prompt": "Identify the holder of this permit/license with their full contact details."},
                    {"id": "verify_authority", "title": "Verify authority", "description": "Identifies issuing body", "sample_prompt": "Identify the authority or body that issued this permit/license."},
                    {"id": "verify_subject", "title": "Verify subject", "description": "Describes what is authorized", "sample_prompt": "Precisely describe what is authorized by this permit/license."},
                    {"id": "verify_dates", "title": "Verify dates", "description": "Extracts issue and expiry dates", "sample_prompt": "Extract the issue date, expiry date and renewal of this permit/license."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this permit: holder, authority, subject, dates, conditions."}
                ],
                'default': [
                    {"id": "summarize_all", "title": "Summarize documents", "description": "Generates a global summary of documents", "sample_prompt": "Provide a clear and structured summary of all uploaded documents, highlighting the main themes and important information. Adapt the summary to the document type: for a CV, focus on professional experience, skills and achievements; for a financial document, mention relevant amounts and figures; for a contract, mention important parties and dates. Do NOT mention information that is not present in the documents (for example, do not mention financial information if the document is a CV)."},
                    {"id": "extract_key_points", "title": "Extract key points", "description": "Lists key points and entities", "sample_prompt": "Extract key points, important decisions and named entities (people, companies, places) from all uploaded documents and organize them in bullet points."},
                    {"id": "extract_structured", "title": "Extract structured data", "description": "Extracts all data in a structured format", "sample_prompt": "Extract all structured data from this document: parties, dates, amounts, key information."}
                ]
            },
            'es': {
                'contrat_location': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica todas las partes (inquilino, arrendador) con sus datos de contacto", "sample_prompt": "Identifica todas las partes de este contrato de alquiler: el inquilino y el arrendador. Enumera sus nombres completos, direcciones, telÃ©fonos y correos electrÃ³nicos si estÃ¡n disponibles."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae todas las fechas importantes del contrato", "sample_prompt": "Extrae todas las fechas importantes de este contrato de alquiler: fecha de firma, fecha de inicio, fecha de fin, duraciÃ³n, preaviso."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera todos los montos: alquiler, depÃ³sito, gastos", "sample_prompt": "Enumera todos los montos mencionados en este contrato de alquiler: el alquiler mensual, el depÃ³sito, los gastos y cualquier indexaciÃ³n prevista."},
                    {"id": "analyze_risky_clauses", "title": "Analizar clÃ¡usulas de riesgo", "description": "Identifica clÃ¡usulas potencialmente problemÃ¡ticas", "sample_prompt": "Analiza este contrato de alquiler e identifica las clÃ¡usulas potencialmente problemÃ¡ticas o desventajosas para el inquilino o el arrendador."},
                    {"id": "verify_obligations", "title": "Verificar obligaciones", "description": "Enumera las obligaciones del inquilino y del arrendador", "sample_prompt": "Enumera todas las obligaciones del inquilino y del arrendador mencionadas en este contrato de alquiler."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este contrato de alquiler: partes, fechas, montos, propiedad alquilada, clÃ¡usulas importantes."}
                ],
                'contrat_travail': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica al empleador y al empleado", "sample_prompt": "Identifica las partes de este contrato de trabajo: el empleador y el empleado. Enumera sus nombres completos y datos de contacto."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae las fechas importantes del contrato", "sample_prompt": "Extrae todas las fechas importantes de este contrato de trabajo: fecha de firma, fecha de inicio, perÃ­odo de prueba, fecha de fin si es aplicable."},
                    {"id": "verify_remuneration", "title": "Verificar remuneraciÃ³n", "description": "Detalla el salario, bonos y beneficios", "sample_prompt": "Detalla la remuneraciÃ³n en este contrato de trabajo: salario base, bonos, beneficios, revisiones salariales previstas."},
                    {"id": "verify_obligations", "title": "Verificar obligaciones", "description": "Enumera las obligaciones del empleado y del empleador", "sample_prompt": "Enumera todas las obligaciones del empleado y del empleador mencionadas en este contrato de trabajo."},
                    {"id": "analyze_risky_clauses", "title": "Analizar clÃ¡usulas de riesgo", "description": "Identifica clÃ¡usulas restrictivas o problemÃ¡ticas", "sample_prompt": "Analiza este contrato de trabajo e identifica las clÃ¡usulas potencialmente restrictivas o problemÃ¡ticas (clÃ¡usula de no competencia, clÃ¡usula de exclusividad, etc.)."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este contrato de trabajo: partes, fechas, remuneraciÃ³n, obligaciones, condiciones de trabajo."}
                ],
                'contrat_vente': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica al comprador y al vendedor", "sample_prompt": "Identifica las partes de este contrato de venta: el comprador y el vendedor. Enumera sus nombres completos y datos de contacto."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae las fechas importantes", "sample_prompt": "Extrae todas las fechas importantes de este contrato de venta: fecha de firma, fecha de entrega, fechas de pago."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Detalla el precio y modalidades de pago", "sample_prompt": "Detalla todos los montos de este contrato de venta: precio total, anticipo, modalidades de pago, plazos."},
                    {"id": "verify_object", "title": "Verificar objeto", "description": "Describe precisamente el objeto de la venta", "sample_prompt": "Describe precisamente el objeto de esta venta: naturaleza del bien, caracterÃ­sticas, cantidad, estado."},
                    {"id": "verify_guarantees", "title": "Verificar garantÃ­as", "description": "Enumera las garantÃ­as y condiciones", "sample_prompt": "Enumera todas las garantÃ­as mencionadas en este contrato de venta y sus condiciones."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este contrato de venta: partes, fechas, montos, objeto, garantÃ­as, condiciones."}
                ],
                'contrat_generique': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica todas las partes", "sample_prompt": "Identifica todas las partes de este contrato con sus nombres completos y datos de contacto."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae todas las fechas importantes", "sample_prompt": "Extrae todas las fechas importantes de este contrato: firma, plazos, fechas de pago."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera todos los montos y modalidades", "sample_prompt": "Enumera todos los montos mencionados en este contrato y sus modalidades de pago."},
                    {"id": "verify_object", "title": "Verificar objeto", "description": "Describe precisamente el objeto del contrato", "sample_prompt": "Describe precisamente el objeto de este contrato en una oraciÃ³n clara."},
                    {"id": "analyze_risky_clauses", "title": "Analizar clÃ¡usulas de riesgo", "description": "Identifica clÃ¡usulas problemÃ¡ticas", "sample_prompt": "Analiza este contrato e identifica las clÃ¡usulas potencialmente problemÃ¡ticas o desventajosas."},
                    {"id": "verify_obligations", "title": "Verificar obligaciones", "description": "Enumera las obligaciones de cada parte", "sample_prompt": "Enumera todas las obligaciones de cada parte mencionadas en este contrato."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurudos de este contrato: partes, fechas, montos, objeto, obligaciones, clÃ¡usulas importantes."}
                ],
                'testament': [
                    {"id": "verify_testator", "title": "Verificar testador", "description": "Identifica al testador", "sample_prompt": "Identifica al testador de este testamento con sus datos de contacto completos."},
                    {"id": "verify_beneficiaries", "title": "Verificar beneficiarios", "description": "Enumera todos los beneficiarios y sus partes", "sample_prompt": "Enumera todos los beneficiarios de este testamento y sus partes respectivas."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae las fechas importantes", "sample_prompt": "Extrae todas las fechas importantes de este testamento: fecha de redacciÃ³n, firma, modificaciones eventuales."},
                    {"id": "verify_legacies", "title": "Verificar legados", "description": "Detalla todos los legados y herencias", "sample_prompt": "Detalla todos los legados y herencias mencionados en este testamento."},
                    {"id": "verify_executor", "title": "Verificar ejecutor", "description": "Identifica al albacea", "sample_prompt": "Identifica al albacea mencionado en este testamento."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este testamento: testador, beneficiarios, legados, fechas, ejecutor."}
                ],
                'acte_notarie': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica todas las partes", "sample_prompt": "Identifica todas las partes involucradas en este acta notarial con sus datos de contacto."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae todas las fechas importantes", "sample_prompt": "Extrae todas las fechas importantes de este acta notarial."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera todos los montos y transacciones", "sample_prompt": "Enumera todos los montos y transacciones mencionados en este acta notarial."},
                    {"id": "verify_object", "title": "Verificar objeto", "description": "Describe precisamente el objeto del acta", "sample_prompt": "Describe precisamente el objeto de este acta notarial."},
                    {"id": "verify_notary", "title": "Verificar notario", "description": "Identifica al notario y su estudio", "sample_prompt": "Identifica al notario que redactÃ³ este acta y su estudio notarial."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este acta notarial: partes, fechas, montos, objeto, notario."}
                ],
                'lettre': [
                    {"id": "verify_sender", "title": "Verificar remitente", "description": "Identifica al remitente de la carta", "sample_prompt": "Identifica al remitente de esta carta: nombre, funciÃ³n, organizaciÃ³n, datos de contacto."},
                    {"id": "verify_recipient", "title": "Verificar destinatario", "description": "Identifica al destinatario de la carta", "sample_prompt": "Identifica al destinatario de esta carta: nombre, funciÃ³n, organizaciÃ³n, datos de contacto."},
                    {"id": "verify_date", "title": "Verificar fecha", "description": "Extrae la fecha de la carta", "sample_prompt": "Extrae la fecha de esta carta (fecha de escritura, fecha de envÃ­o si se menciona)."},
                    {"id": "verify_object", "title": "Verificar objeto", "description": "Describe el objeto y el propÃ³sito de la carta", "sample_prompt": "Describe el objeto y el propÃ³sito principal de esta carta en una oraciÃ³n clara."},
                    {"id": "verify_key_information", "title": "Verificar info clave", "description": "Extrae informaciÃ³n importante mencionada", "sample_prompt": "Extrae toda la informaciÃ³n importante mencionada en esta carta: montos, fechas, referencias, compromisos."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de esta carta: remitente, destinatario, fecha, objeto, informaciÃ³n clave."}
                ],
                'document_financier': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica las partes concernidas", "sample_prompt": "Identifica todas las partes concernidas por este documento financiero: empleador, empleado, instituciÃ³n, etc."},
                    {"id": "verify_period", "title": "Verificar perÃ­odo", "description": "Extrae el perÃ­odo cubierto", "sample_prompt": "Extrae el perÃ­odo cubierto por este documento financiero: fechas de inicio y de fin."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera todos los montos y totales", "sample_prompt": "Enumera todos los montos mencionados en este documento financiero: ingresos, deducciones, impuestos, totales."},
                    {"id": "verify_deductions", "title": "Verificar deducciones", "description": "Detalla todas las deducciones", "sample_prompt": "Detalla todas las deducciones mencionadas en este documento financiero: impuestos, cotizaciones, otras deducciones."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este documento financiero: partes, perÃ­odo, montos, deducciones, totales."}
                ],
                'assurance_insurance': [
                    {"id": "verify_insured", "title": "Verificar asegurado", "description": "Identifica al asegurado y al tomador", "sample_prompt": "Identifica al asegurado y al tomador de esta pÃ³liza de seguro con sus datos de contacto completos."},
                    {"id": "verify_coverage", "title": "Verificar cobertura", "description": "Detalla garantÃ­as y exclusiones", "sample_prompt": "Detalla las garantÃ­as cubiertas, los riesgos cubiertos y las exclusiones de esta pÃ³liza de seguro."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera primas, franquicias, topes", "sample_prompt": "Enumera los montos: prima anual, franquicia, tope de indemnizaciÃ³n de esta pÃ³liza de seguro."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae fechas de suscripciÃ³n y vencimiento", "sample_prompt": "Extrae las fechas importantes: suscripciÃ³n, efecto, vencimiento, renovaciÃ³n de esta pÃ³liza de seguro."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de esta pÃ³liza de seguro: asegurado, cobertura, montos, fechas."}
                ],
                'jugement_decision_justice': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica demandante, demandado y abogados", "sample_prompt": "Identifica las partes de esta sentencia: demandante, demandado y sus respectivos abogados."},
                    {"id": "verify_jurisdiction", "title": "Verificar jurisdicciÃ³n", "description": "Identifica el tribunal y el juez", "sample_prompt": "Identifica el tribunal, la sala y el/los juez/jueces que dictaron esta decisiÃ³n."},
                    {"id": "analyze_ruling", "title": "Analizar fallo", "description": "Resume la decisiÃ³n dictada", "sample_prompt": "Resume la decisiÃ³n dictada en esta sentencia: condenas, montos, obligaciones."},
                    {"id": "verify_appeals", "title": "Verificar recursos", "description": "Identifica vÃ­as de recurso", "sample_prompt": "Identifica las vÃ­as de recurso posibles y los plazos de apelaciÃ³n mencionados en esta sentencia."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de esta sentencia: partes, jurisdicciÃ³n, decisiÃ³n, recursos."}
                ],
                'releve_bancaire': [
                    {"id": "verify_holder", "title": "Verificar titular", "description": "Identifica al titular de la cuenta", "sample_prompt": "Identifica al titular de la cuenta con sus datos de contacto completos."},
                    {"id": "verify_account", "title": "Verificar cuenta", "description": "Extrae nÃºmero de cuenta, IBAN, BIC", "sample_prompt": "Extrae el nÃºmero de cuenta, IBAN, BIC y tipo de cuenta de este extracto bancario."},
                    {"id": "verify_balances", "title": "Verificar saldos", "description": "Enumera saldo inicial, final y variaciones", "sample_prompt": "Enumera el saldo inicial, el saldo final, el total de dÃ©bitos y crÃ©ditos de este extracto bancario."},
                    {"id": "analyze_transactions", "title": "Analizar operaciones", "description": "Resume las operaciones principales", "sample_prompt": "Resume las operaciones principales de este extracto bancario: transferencias, domiciliaciones, pagos."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este extracto bancario: titular, cuenta, saldos, operaciones."}
                ],
                'certificat_attestation': [
                    {"id": "verify_issuer", "title": "Verificar emisor", "description": "Identifica al organismo emisor", "sample_prompt": "Identifica al organismo o persona que emitiÃ³ este certificado/atestaciÃ³n."},
                    {"id": "verify_beneficiary", "title": "Verificar beneficiario", "description": "Identifica a la persona concernida", "sample_prompt": "Identifica a la persona concernida por este certificado/atestaciÃ³n con sus datos de contacto."},
                    {"id": "verify_subject", "title": "Verificar objeto", "description": "Describe lo que se certifica", "sample_prompt": "Describe precisamente lo que se certifica o atestigua en este documento."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae fecha de emisiÃ³n y validez", "sample_prompt": "Extrae la fecha de emisiÃ³n y el perÃ­odo de validez de este certificado/atestaciÃ³n."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este certificado: emisor, beneficiario, objeto, fechas."}
                ],
                'contrat_pret_loan': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica prestatario y prestamista", "sample_prompt": "Identifica al prestatario y al prestamista de este contrato de prÃ©stamo con sus datos de contacto."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Detalla capital, tasa, cuotas", "sample_prompt": "Detalla los montos: capital prestado, tasa de interÃ©s, TAE, cuotas mensuales de este contrato de prÃ©stamo."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae fechas de firma y vencimientos", "sample_prompt": "Extrae las fechas: firma, inicio, fin, vencimientos de este contrato de prÃ©stamo."},
                    {"id": "verify_guarantees", "title": "Verificar garantÃ­as", "description": "Enumera las garantÃ­as exigidas", "sample_prompt": "Enumera las garantÃ­as exigidas en este contrato de prÃ©stamo: hipoteca, aval, prenda."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este contrato de prÃ©stamo: partes, montos, fechas, garantÃ­as."}
                ],
                'devis_estimation': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica emisor y cliente", "sample_prompt": "Identifica al emisor del presupuesto y al cliente con sus datos de contacto."},
                    {"id": "verify_items", "title": "Verificar partidas", "description": "Detalla partidas y precios", "sample_prompt": "Detalla todas las partidas del presupuesto: descripciÃ³n, cantidades, precios unitarios y totales."},
                    {"id": "verify_amounts", "title": "Verificar montos", "description": "Enumera subtotal, IVA, total", "sample_prompt": "Enumera los montos: subtotal, IVA, total, descuentos eventuales de este presupuesto."},
                    {"id": "verify_validity", "title": "Verificar validez", "description": "Extrae el perÃ­odo de validez", "sample_prompt": "Extrae el perÃ­odo de validez y las condiciones de este presupuesto."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este presupuesto: partes, partidas, montos, validez."}
                ],
                'bon_commande_purchase_order': [
                    {"id": "verify_parties", "title": "Verificar partes", "description": "Identifica cliente y proveedor", "sample_prompt": "Identifica al cliente y al proveedor de esta orden de compra con sus datos de contacto."},
                    {"id": "verify_items", "title": "Verificar artÃ­culos", "description": "Detalla artÃ­culos, cantidades, precios", "sample_prompt": "Detalla los artÃ­culos pedidos: referencias, descripciones, cantidades y precios."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae fecha de pedido y entrega", "sample_prompt": "Extrae la fecha de pedido y la fecha de entrega prevista de esta orden de compra."},
                    {"id": "verify_conditions", "title": "Verificar condiciones", "description": "Detalla condiciones de entrega y pago", "sample_prompt": "Detalla las condiciones de entrega y de pago de esta orden de compra."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de esta orden de compra: partes, artÃ­culos, fechas, condiciones."}
                ],
                'proces_verbal': [
                    {"id": "verify_participants", "title": "Verificar participantes", "description": "Enumera presentes y ausentes", "sample_prompt": "Enumera todos los participantes de este acta: presentes, ausentes y excusados."},
                    {"id": "verify_dates", "title": "Verificar fecha y lugar", "description": "Extrae fecha, hora y lugar", "sample_prompt": "Extrae la fecha, hora y lugar de la reuniÃ³n descrita en este acta."},
                    {"id": "analyze_decisions", "title": "Analizar decisiones", "description": "Resume las decisiones tomadas", "sample_prompt": "Resume todas las decisiones tomadas y las votaciones realizadas en este acta."},
                    {"id": "verify_actions", "title": "Verificar acciones", "description": "Enumera las acciones decididas", "sample_prompt": "Enumera las acciones decididas con los responsables y los plazos mencionados en este acta."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este acta: participantes, decisiones, acciones, fechas."}
                ],
                'rapport_expertise': [
                    {"id": "verify_expert", "title": "Verificar experto", "description": "Identifica al experto y sus cualificaciones", "sample_prompt": "Identifica al experto o auditor y sus cualificaciones en este informe."},
                    {"id": "verify_subject", "title": "Verificar objeto", "description": "Describe la misiÃ³n y el alcance", "sample_prompt": "Describe la misiÃ³n y el alcance de la peritaciÃ³n de este informe."},
                    {"id": "analyze_conclusions", "title": "Analizar conclusiones", "description": "Extrae las conclusiones principales", "sample_prompt": "Extrae las conclusiones y hallazgos principales de este informe pericial."},
                    {"id": "verify_recommendations", "title": "Verificar recomendaciones", "description": "Enumera las recomendaciones", "sample_prompt": "Enumera todas las recomendaciones formuladas en este informe pericial."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este informe: experto, objeto, conclusiones, recomendaciones."}
                ],
                'permis_licence': [
                    {"id": "verify_holder", "title": "Verificar titular", "description": "Identifica al titular del permiso", "sample_prompt": "Identifica al titular de este permiso/licencia con sus datos de contacto completos."},
                    {"id": "verify_authority", "title": "Verificar autoridad", "description": "Identifica al organismo emisor", "sample_prompt": "Identifica la autoridad u organismo que otorgÃ³ este permiso/licencia."},
                    {"id": "verify_subject", "title": "Verificar objeto", "description": "Describe lo que se autoriza", "sample_prompt": "Describe precisamente lo que se autoriza mediante este permiso/licencia."},
                    {"id": "verify_dates", "title": "Verificar fechas", "description": "Extrae fechas de emisiÃ³n y expiraciÃ³n", "sample_prompt": "Extrae las fechas de emisiÃ³n, expiraciÃ³n y renovaciÃ³n de este permiso/licencia."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este permiso: titular, autoridad, objeto, fechas, condiciones."}
                ],
                'default': [
                    {"id": "summarize_all", "title": "Resumir documentos", "description": "Genera un resumen global de los documentos", "sample_prompt": "Proporciona un resumen claro y estructurado de todos los documentos cargados, destacando los temas principales y la informaciÃ³n importante. Adapta el resumen al tipo de documento: para un CV, concÃ©ntrate en la experiencia profesional, habilidades y logros; para un documento financiero, menciona montos y cifras relevantes; para un contrato, menciona partes y fechas importantes. NO menciones informaciÃ³n que no estÃ© presente en los documentos (por ejemplo, no menciones informaciÃ³n financiera si el documento es un CV)."},
                    {"id": "extract_key_points", "title": "Extraer puntos clave", "description": "Enumera los puntos clave y entidades", "sample_prompt": "Extrae los puntos clave, decisiones importantes y entidades nombradas (personas, empresas, lugares) de todos los documentos cargados y organÃ­zalos en viÃ±etas."},
                    {"id": "extract_structured", "title": "Extraer datos estructurados", "description": "Extrae todos los datos en un formato estructurado", "sample_prompt": "Extrae todos los datos estructurados de este documento: partes, fechas, montos, informaciÃ³n clave."}
                ]
            }
        }
        
        # SÃ©lectionner les actions selon la langue et le type de document
        lang_actions = fallback_actions.get(language, fallback_actions['en'])
        actions_list = lang_actions.get(doc_type, lang_actions.get('default', []))
        
        return {
            "domain": doc_type,
            "suggested_actions": actions_list,
            "detected_type": doc_type,
            "detected_type_label": _get_type_label(doc_type, language),
            "detected_type_confidence": 0.5,
            "type_distribution": {doc_type: 1}
        }

# save_file_description est maintenant importÃ© depuis services.analysis_service
async def _OLD_save_file_description(file_name, description):
    with open(DESCRIPTIONS_FILE, 'r') as f:
        descriptions = json.load(f)
    descriptions.append({"file_name": file_name, "description": description})
    with open(DESCRIPTIONS_FILE, 'w') as f:
        json.dump(descriptions, f, indent=4)

async def query_model(file_name, file_content, directory_content, repo_structure, user_query, is_binary=False, selected_model=DEFAULT_MODEL, language='en'):
    cache_key = f"query_{file_name}_{user_query}_{selected_model}_{language}"
    cached = await cache.get(cache_key)
    if cached:
        return cached

    # fallback to English if unknown language
    t = translations.get(language, translations['en'])

    # Construire le rÃ©sumÃ© des autres fichiers
    directory_content_summary = ' '.join(
        [f"{t['other_file']}: {doc['fileName']} : {doc['content']}" for doc in directory_content]
    )

    # Construire le prompt multilangue
    prompt = (
        f"{t['project_structure']}:\n{repo_structure}\n\n"
        f"{t['file']}: {file_name}\n\n"
        f"{t['content']}:\n{file_content}\n\n"
        f"{directory_content_summary}\n\n"
        f"{t['question']}: {user_query}\n\n"
        f"{t['emoji']}"
    )

    try:
        # Enhanced model routing with GPT-5 variants
        if selected_model.lower() in ["gpt-3.5-turbo", "gpt-4o"]:
            result = call_openai_api(prompt, selected_model)
        elif selected_model.lower() in ["gpt-5", "gpt-5-mini", "gpt-5-nano"]:
            result = call_openai_api(prompt, selected_model)
        elif selected_model.lower() == "openai":
            # Legacy compatibility
            result = call_openai_api(prompt, "openai")
        elif selected_model.lower() == "mistral":
            result = call_mistral_api(prompt)
        elif selected_model.lower().startswith("gemini") or selected_model.lower() in ["gemini-3-flash", "gemini-pro"]:
            result = call_gemini_api(prompt, selected_model)
        elif selected_model.lower() in OLLAMA_MODELS or selected_model.lower() == "llama3":
            result = call_ollama_api(prompt, selected_model)
        else:
            logger.warning(f"Unknown model {selected_model}, falling back to default {DEFAULT_MODEL}")
            result = call_openai_api(prompt, DEFAULT_MODEL)
            
    except Exception as e:
        result = f"Error querying model: {str(e)}"

    await cache.set(cache_key, result, ttl=3600)
    return result

# Validation de la configuration
try:
    FlaskConfig.validate_env_vars()
    logger.info("âœ… Configuration validÃ©e")
except ValueError as e:
    logger.error(f"âŒ Erreur de configuration: {e}")

# Initialisation de la base de donnÃ©es d'authentification
init_database()

# Enregistrement des routes d'authentification
register_auth_routes(app)

# New endpoint to get available models with detailed info
@app.route('/models', methods=['GET'])
def get_available_models():
    """
    Returns a list of available models with detailed configuration
    """
    models = []
    for model_id, config in MODEL_CONFIG.items():
        model_info = {
            "id": model_id,
            "name": config["name"],
            "provider": config["provider"],
            "description": config["description"],
            "cost_tier": config["cost_tier"],
            "max_tokens": config["max_tokens"],
            "is_default": model_id == DEFAULT_MODEL
        }
        models.append(model_info)
    
    return jsonify({
        "models": models,
        "default_model": DEFAULT_MODEL,
        "total_count": len(models)
    }), 200

@app.route('/models/<model_id>/test', methods=['POST'])
def test_model(model_id):
    """
    Test if a specific model is available and working
    """
    if model_id not in MODEL_CONFIG:
        return jsonify({"error": "Model not found"}), 404
    
    test_prompt = "Hello, please respond with 'Model is working correctly.'"
    
    try:
        if model_id in ["gpt-3.5-turbo", "gpt-4o", "gpt-5", "gpt-5-mini", "gpt-5-nano", "openai"]:
            response = call_openai_api(test_prompt, model_id)
        elif model_id == "mistral":
            response = call_mistral_api(test_prompt)
        elif model_id.startswith("gemini") or model_id in ["gemini-3-flash", "gemini-pro"]:
            response = call_gemini_api(test_prompt, model_id)
        elif model_id in OLLAMA_MODELS or model_id == "llama3":
            response = call_ollama_api(test_prompt, model_id)
        else:
            return jsonify({"error": "Model not supported"}), 400
            
        return jsonify({
            "model_id": model_id,
            "status": "available",
            "test_response": response,
            "config": MODEL_CONFIG[model_id]
        }), 200
        
    except Exception as e:
        logger.error(f"Model test failed for {model_id}: {str(e)}")
        return jsonify({
            "model_id": model_id,
            "status": "unavailable",
            "error": str(e),
            "config": MODEL_CONFIG[model_id]
        }), 503

# Enhanced Flask Routes
@app.route('/upload', methods=['POST'])
async def upload_files():
    if 'files' not in request.files:
        return jsonify({"error": "No files provided"}), 400

    files = request.files.getlist('files')
    selected_model = request.form.get('model', DEFAULT_MODEL)  # Use GPT-3.5-Turbo as default
    language = request.form.get('language', 'en')
    session_id = request.headers.get('Session-ID', 'default')
    user_email = _extract_user_email_from_auth_header()
    results, texts = [], []

    # Validate model
    if selected_model not in MODEL_CONFIG:
        logger.warning(f"Unknown model {selected_model}, using default {DEFAULT_MODEL}")
        selected_model = DEFAULT_MODEL

    for file in files:
        file_name = file.filename
        extension = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''
        is_binary = extension in ['docx', 'pdf']
        chunk_store_result = None

        raw_bytes = file.read()
        doc_hash = compute_document_hash(raw_bytes)

        # --- Docling structured extraction (PDF / DOCX) with basic fallback ---
        chunk_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
        chunks_with_meta: list = []
        doc_meta: dict = {"parser": "basic", "num_pages": 1}

        if extension in ('pdf', 'docx'):
            docling_result = parse_with_docling(raw_bytes, extension, file_name)
        else:
            docling_result = None

        if docling_result is not None:
            # --- Docling path: structure-aware chunks ---
            chunks_with_meta, file_content, doc_meta = docling_result
            logger.info(
                "\u2705 Docling parsed %s: %d chunks, parser=%s",
                file_name, len(chunks_with_meta), doc_meta.get("parser"),
            )
        else:
            # --- Fallback: basic page-level extraction ---
            if extension == 'pdf':
                page_sections = extract_pages_from_pdf(raw_bytes)
                file_content = "\n".join(p["text"] for p in page_sections)
                doc_meta["num_pages"] = len(page_sections)
            elif extension == 'docx':
                page_sections = extract_sections_from_docx(raw_bytes)
                file_content = extract_text_from_docx(BytesIO(raw_bytes))
                doc_meta["num_pages"] = len(page_sections)
            else:
                file_content = raw_bytes.decode('utf-8', errors='ignore')
                page_sections = [{"page": 1, "text": file_content}]

            global_chunk_idx = 0
            for section in page_sections:
                for chunk_doc in chunk_splitter.create_documents([section["text"]]):
                    if chunk_doc.page_content.strip():
                        chunks_with_meta.append({
                            "text":         chunk_doc.page_content,
                            "page":         section["page"],
                            "chunk_id":     f"chunk_{global_chunk_idx}",
                            "chunk_index":  global_chunk_idx,
                            "element_type": "text",
                            "headings":     [],
                            "section":      "",
                        })
                        global_chunk_idx += 1
            logger.info(
                "\u26a0\ufe0f Basic extraction for %s: %d chunks",
                file_name, len(chunks_with_meta),
            )

        description = await analyze_file_content(file_content, file_name, is_binary, extension, selected_model, language)
        await save_file_description(file_name, description)
        if not description.startswith("Error"):
            texts.append(description)
        persist_result = aws_persistence_service.persist_document(
            file_name=file_name,
            content=file_content,
            raw_bytes=raw_bytes,
            mime_type=file.mimetype,
            metadata={
                'extension': extension,
                'is_binary': is_binary,
                'language': language,
                'model_used': selected_model,
                'description_excerpt': (description or '')[:500],
            },
            user_email=user_email,
            session_id=session_id,
            source='upload',
        )

        # Ensure globally unique chunk IDs across multiple uploaded documents.
        # Preferred prefix: persisted document UUID. Fallback: content hash.
        doc_id_prefix = persist_result.get('document_id') or doc_hash
        for idx, chunk in enumerate(chunks_with_meta):
            chunk['chunk_id'] = f"{doc_id_prefix}_chunk_{idx}"
            chunk['chunk_index'] = idx

        if persist_result.get('stored') and persist_result.get('document_id') and persist_result.get('organization_id'):
            chunk_texts = [c["text"] for c in chunks_with_meta]
            per_chunk_meta = [
                {
                    "document_id":  persist_result['document_id'],
                    "file_name":    file_name,
                    "page":         c["page"],
                    "chunk_id":     c["chunk_id"],
                    "element_type": c.get("element_type", "text"),
                    "headings":     c.get("headings", []),
                    "section":      c.get("section", ""),
                    "session_id":   session_id,
                    "model_used":   selected_model,
                    "source":       "upload",
                    "parser":       doc_meta.get("parser", "basic"),
                }
                for c in chunks_with_meta
            ]

            if chunk_texts:
                try:
                    loop = asyncio.get_event_loop()
                    chunk_embeddings = await loop.run_in_executor(None, embeddings.embed_documents, chunk_texts)
                    chunk_store_result = aws_persistence_service.persist_document_chunks(
                        document_id=persist_result['document_id'],
                        organization_id=persist_result['organization_id'],
                        chunks=chunk_texts,
                        embeddings=chunk_embeddings,
                        per_chunk_metadata=per_chunk_meta,
                    )
                except Exception as exc:
                    logger.error("Failed to generate/store pgvector chunks for %s: %s", file_name, exc)
                    chunk_store_result = {"stored": False, "reason": "embedding_error", "error": str(exc)}
        elif not persist_result.get('stored'):
            logger.warning("AWS persistence skipped for %s: %s", file_name, persist_result)

        # --- FAISS index: load from cache or build from chunks ---
        index_path = get_index_path(doc_hash)
        cached_vs = load_faiss_index(index_path)
        faiss_docs = [
            Document(
                page_content=c["text"],
                metadata={
                    "document_id":  file_name,
                    "file_name":    file_name,
                    "page":         c["page"],
                    "chunk_id":     c["chunk_id"],
                    "element_type": c.get("element_type", "text"),
                    "section":      c.get("section", ""),
                    "headings":     c.get("headings", []),
                    "session_id":   session_id,
                }
            )
            for c in chunks_with_meta
        ]

        global vector_store
        if cached_vs is not None:
            logger.info("ðŸ“‚ Reusing cached FAISS index for %s (hash=%s)", file_name, doc_hash[:12])
            if vector_store is None:
                vector_store = cached_vs
            else:
                loop = asyncio.get_event_loop()
                await loop.run_in_executor(None, vector_store.merge_from, cached_vs)
        elif faiss_docs:
            if vector_store is None:
                vector_store = FAISS.from_documents(faiss_docs, embeddings)
            else:
                loop = asyncio.get_event_loop()
                await loop.run_in_executor(None, vector_store.add_documents, faiss_docs)
            per_doc_vs = FAISS.from_documents(faiss_docs, embeddings)
            save_faiss_index(per_doc_vs, index_path)
            logger.info("âœ… FAISS index built and saved for %s (%d chunks)", file_name, len(faiss_docs))

        results.append({
            "file_name":     file_name,
            "description":   description,
            "model_used":    selected_model,
            "storage":       persist_result,
            "chunks":        chunk_store_result,
            "pages_indexed": doc_meta.get("num_pages", len(set(c["page"] for c in chunks_with_meta))),
            "chunks_indexed":len(chunks_with_meta),
            "cache_hit":     cached_vs is not None,
            "parser":        doc_meta.get("parser", "basic"),
        })

    if texts:
        splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
        docs = [doc for text in texts for doc in splitter.create_documents([text])]
        if vector_store is None:
            vector_store = FAISS.from_documents(docs, embeddings)
        else:
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, vector_store.add_documents, docs)

    return jsonify({
        "message": "Files processed successfully", 
        "results": results,
        "model_used": selected_model,
        "model_config": MODEL_CONFIG.get(selected_model, {})
    }), 200

# perform_online_search et search_serpapi sont maintenant importÃ©s depuis services.search_service
def _OLD_perform_online_search(query: str, language: str = 'en') -> str:
    return _OLD_search_serpapi(query, language)

def _OLD_search_serpapi(query: str, language: str = 'en') -> str:
    """
    SerpAPI - 100 recherches gratuites/mois
    Inscription: https://serpapi.com/
    """
    api_key = os.getenv("SERPAPI_KEY")
    if not api_key:
        return "ClÃ© API manquante"
    
    try:
        url = "https://serpapi.com/search"
        params = {
            'api_key': api_key,
            'engine': 'google',
            'q': query,
            'hl': language,
            'num': 5,
            'no_cache': 'true'  # Force fresh results
        }
        
        response = requests.get(url, params=params, timeout=15)
        response.raise_for_status()
        data = response.json()
        
        results = []
        
        # RÃ©ponse directe (featured snippet)
        if data.get('answer_box'):
            answer = data['answer_box']
            if answer.get('answer'):
                results.append(f"**RÃ©ponse directe**: {answer['answer']}")
            if answer.get('snippet'):
                results.append(f"**Information**: {answer['snippet']}")
        
        # RÃ©sultats organiques
        if data.get('organic_results'):
            for result in data['organic_results'][:5]:
                title = result.get('title', '')
                snippet = result.get('snippet', '')
                link = result.get('link', '')
                
                if title and snippet:
                    results.append(f"**{title}**\n{snippet}\nSource: {link}")
        
        return "\n\n".join(results) if results else "Aucun rÃ©sultat trouvÃ©."
        
    except Exception as e:
        raise Exception(f"Erreur SerpAPI: {str(e)}")

# rerank_documents_with_llm est maintenant importÃ© depuis services.search_service
async def _OLD_rerank_documents_with_llm(query: str, documents: List[Document], model: str = DEFAULT_MODEL) -> List[Document]:
    """
    Re-rank les documents avec un LLM pour amÃ©liorer la pertinence sÃ©mantique.
    Le LLM Ã©value chaque document par rapport Ã  la requÃªte et les trie par pertinence.
    """
    try:
        if len(documents) <= 5:
            return documents  # Pas besoin de re-ranking pour peu de documents
        
        # Construire le prompt de re-ranking
        docs_summary = []
        for i, doc in enumerate(documents):
            file_name = doc.metadata.get("fileName") or doc.metadata.get("file_name") or f"document_{i}"
            content_snippet = doc.page_content[:500]  # Premier 500 chars pour le prompt
            docs_summary.append(f"[{i}] Fichier: {file_name}\nContenu: {content_snippet}...")
        
        rerank_prompt = f"""Tu dois classer ces documents par ordre de pertinence pour rÃ©pondre Ã  cette question: "{query}"

Documents Ã  classer:
{chr(10).join(docs_summary)}

Retourne UNIQUEMENT une liste JSON des indices (nombres entre crochets) dans l'ordre de pertinence dÃ©croissante.
Format: [3, 1, 5, 2, 0, 4, ...]

RÃ©ponds UNIQUEMENT avec le JSON, sans explication."""
        
        # Appel au modÃ¨le pour re-ranking
        rerank_response = call_openai_api(rerank_prompt, model)
        
        # Extraire les indices du JSON
        import json
        json_match = re.search(r'\[.*?\]', rerank_response)
        if json_match:
            ranked_indices = json.loads(json_match.group())
            # VÃ©rifier que tous les indices sont valides
            valid_indices = [idx for idx in ranked_indices if 0 <= idx < len(documents)]
            if len(valid_indices) == len(documents):
                reranked = [documents[idx] for idx in valid_indices]
                logger.info(f"âœ… Re-ranking rÃ©ussi: {len(reranked)} documents reclassÃ©s")
                return reranked
        
        # Fallback: retourner l'ordre original si le parsing Ã©choue
        logger.warning("Re-ranking JSON invalide, utilisation de l'ordre original")
        return documents
        
    except Exception as e:
        logger.warning(f"Erreur lors du re-ranking: {str(e)}, utilisation de l'ordre original")
        return documents

def search_semantic_documents_sync(vector_store, user_query: str, session_id: str = 'default', conversation_history: Optional[List[Dict]] = None) -> List[Dict]:
    """
    Fonction helper synchrone pour rechercher des documents pertinents dans le vector store.
    Retourne une liste de dictionnaires avec 'fileName' et 'content'.
    AMÃ‰LIORATION: DÃ©tecte les noms de personnes et filtre strictement par nom.
    Utilise l'historique de conversation pour dÃ©tecter les pronoms (he, she, his, her).
    """
    try:
        # Enterprise-style hybrid retrieval (deterministic): semantic + lexical reranking.
        # This is more precise than raw similarity_search, and avoids LLM reranking drift.
        result = hybrid_retrieve_documents(
            vector_store=vector_store,
            query=user_query,
            k_candidates=100,  # Increased for better recall
            k_final=25,  # Increased from 20 to 25 for more comprehensive results
            semantic_weight=0.55,  # Adjusted weights for better precision
            bm25_weight=0.35,
            exact_weight=0.10,
        )
        if isinstance(result, tuple) and len(result) == 2:
            docs, debug = result
            if not isinstance(debug, dict):
                debug = {}
        else:
            docs = []
            debug = {}
        if docs:
            logger.info(f"ðŸ”Ž search_semantic_documents_sync hybrid selected {len(docs)} docs | top={debug.get('top', []) if isinstance(debug, dict) else []}")
            out: List[Dict] = []
            seen = set()
            for d in docs:
                fn = d.metadata.get("fileName") or d.metadata.get("file_name") or d.metadata.get("source") or "document_vectorstore"
                key = (str(fn).lower(), (d.page_content or "")[:120])
                if key in seen:
                    continue
                seen.add(key)
                
                # Extraire les informations de page/section depuis les mÃ©tadonnÃ©es Docling
                page_number = d.metadata.get("page_number")
                is_page_chunk = d.metadata.get("is_page_chunk", False)
                section = d.metadata.get("section") or ""
                chunk_id = d.metadata.get("chunk_id")

                result_dict = {
                    "fileName": fn,
                    "content": (d.page_content or "")[:2500]
                }

                if page_number is not None:
                    result_dict["pageNumber"] = page_number
                    result_dict["isPageChunk"] = is_page_chunk

                # Si c'est un chunk de page, extraire le numÃ©ro de page du contenu ou des mÃ©tadonnÃ©es
                if is_page_chunk and page_number is None:
                    content_preview = d.page_content or ""
                    if "[Page" in content_preview and "de" in content_preview:
                        import re
                        match = re.search(r'\[Page\s+(\d+)', content_preview)
                        if match:
                            result_dict["pageNumber"] = int(match.group(1))
                            result_dict["isPageChunk"] = True

                if section:
                    result_dict["section"] = section
                if chunk_id is not None:
                    result_dict["chunkId"] = chunk_id

                out.append(result_dict)
            return out

        # DÃ‰TECTION DES NOMS DE PERSONNES dans la requÃªte ET l'historique
        person_names = set()
        words = user_query.split()
        
        # DÃ©tecter les pronoms et chercher le nom dans l'historique
        query_lower = user_query.lower()
        has_pronoun = any(pronoun in query_lower for pronoun in ['he', 'she', 'his', 'her', 'him', 'they', 'their'])
        
        if has_pronoun and conversation_history:
            # Chercher le dernier nom mentionnÃ© dans l'historique
            for turn in reversed(conversation_history):
                content = (turn.get("content") or "").lower()
                # Chercher les noms dans l'historique
                for name in ['karim', 'dominique', 'essome', 'ngami']:
                    if name in content:
                        person_names.add(name)
                        logger.info(f"ðŸ‘¤ Nom dÃ©tectÃ© depuis l'historique (pronoun dÃ©tectÃ©): {name}")
                        break
                if person_names:
                    break
        
        # DÃ©tecter les noms dans la requÃªte actuelle
        for i, word in enumerate(words):
            clean_word = word.strip('.,!?;:()[]{}"\'').strip()
            # DÃ©tecter les noms propres (capitalisÃ©s) ou les mots aprÃ¨s "his", "her", "their", etc.
            if clean_word and clean_word[0].isupper() and len(clean_word) > 2:
                person_names.add(clean_word.lower())
            # DÃ©tecter aprÃ¨s des mots indicateurs
            if i > 0 and words[i-1].lower() in ['his', 'her', 'their', 'karim', 'dominique', 'about', 'for']:
                if len(clean_word) > 2:
                    person_names.add(clean_word.lower())
        
        # Ajouter aussi les noms communs dans les requÃªtes (karim, dominique, etc.)
        common_names = ['karim', 'dominique', 'essome', 'ngami']
        for name in common_names:
            if name in query_lower:
                person_names.add(name)
        
        logger.info(f"ðŸ‘¤ Noms de personnes dÃ©tectÃ©s (requÃªte + historique): {person_names}")
        
        # Recherche sÃ©mantique
        # AMÃ‰LIORATION: Recherche avec plus de rÃ©sultats et recherche multiple pour meilleure prÃ©cision
        # Recherche principale avec la requÃªte complÃ¨te
        search_results_with_scores = vector_store.similarity_search_with_score(user_query, k=50)
        
        # Recherche supplÃ©mentaire avec les mots-clÃ©s individuels pour capturer les noms de fonctions/mÃ©thodes
        query_words = [word.strip('.,!?;:()[]{}"\'').lower() for word in user_query.split() if len(word.strip('.,!?;:()[]{}"\'')) > 2]
        additional_results = []
        for word in query_words[:5]:  # Limiter Ã  5 mots pour Ã©viter trop de rÃ©sultats
            try:
                word_results = vector_store.similarity_search_with_score(word, k=10)
                additional_results.extend(word_results)
            except:
                pass
        
        # Combiner et dÃ©dupliquer les rÃ©sultats
        all_search_results = {}
        for doc, score in search_results_with_scores:
            doc_id = id(doc)
            if doc_id not in all_search_results or score < all_search_results[doc_id][1]:
                all_search_results[doc_id] = (doc, score)
        
        for doc, score in additional_results:
            doc_id = id(doc)
            if doc_id not in all_search_results or score < all_search_results[doc_id][1]:
                all_search_results[doc_id] = (doc, score)
        
        search_results_with_scores = list(all_search_results.values())
        
        # Extraire les mots-clÃ©s (y compris noms propres courts)
        query_keywords = set()
        for word in user_query.split():
            clean_word = word.strip('.,!?;:()[]{}"\'').lower()
            if len(clean_word) > 2:
                query_keywords.add(clean_word)
        
        logger.info(f"ðŸ”‘ Mots-clÃ©s extraits: {query_keywords}")
        
        # Recherche par nom de fichier - Essayer de rÃ©cupÃ©rer tous les documents
        filename_matches = []
        try:
            # MÃ©thode 1: Recherche avec une requÃªte trÃ¨s large pour rÃ©cupÃ©rer beaucoup de documents
            all_docs_from_store = vector_store.similarity_search("document file content", k=1000)
            
            # MÃ©thode 2: Si on peut accÃ©der au docstore directement, l'utiliser
            if hasattr(vector_store, 'docstore') and hasattr(vector_store.docstore, '_dict'):
                all_docs_dict = vector_store.docstore._dict
                logger.info(f"ðŸ“¦ AccÃ¨s direct au docstore: {len(all_docs_dict)} documents disponibles")
                # Convertir les valeurs du dict en documents
                all_docs_from_store = list(all_docs_dict.values()) if all_docs_dict else all_docs_from_store
            
            logger.info(f"ðŸ” Recherche dans {len(all_docs_from_store)} documents pour correspondance de nom")
            
            for doc in all_docs_from_store:
                file_name_from_meta = (doc.metadata.get("fileName") or doc.metadata.get("file_name") or "").lower()
                for keyword in query_keywords:
                    if keyword in file_name_from_meta:
                        filename_matches.append(doc)
                        logger.info(f"ðŸ“ Fichier trouvÃ© par nom: {doc.metadata.get('fileName', 'N/A')} (mot-clÃ©: '{keyword}')")
                        break
        except Exception as e:
            logger.warning(f"Erreur lors de la rÃ©cupÃ©ration de tous les documents: {str(e)}")
        
        logger.info(f"ðŸ“‚ {len(filename_matches)} fichiers trouvÃ©s par correspondance de nom")
        
        # AMÃ‰LIORATION: Recherche explicite par nom de fonction/mÃ©thode dans le contenu
        function_name_matches = []
        # Extraire les noms potentiels de fonctions/mÃ©thodes de la requÃªte (mots avec underscore ou camelCase)
        potential_function_names = []
        query_words_for_func = user_query.split()
        for word in query_words_for_func:
            clean_word = word.strip('.,!?;:()[]{}"\'').strip()
            # DÃ©tecter les noms de fonctions (avec underscore ou camelCase)
            if '_' in clean_word or (clean_word and clean_word[0].islower() and any(c.isupper() for c in clean_word[1:])):
                potential_function_names.append(clean_word.lower())
                potential_function_names.append(clean_word)  # Garder aussi la version originale
        
        # Rechercher explicitement ces noms dans le contenu
        if potential_function_names:
            try:
                all_docs_for_function_search = vector_store.similarity_search("function method def class", k=500)
                for doc in all_docs_for_function_search:
                    content_lower = doc.page_content.lower()
                    for func_name in potential_function_names:
                        if func_name.lower() in content_lower:
                            function_name_matches.append(doc)
                            logger.info(f"ðŸ”§ Fonction/mÃ©thode trouvÃ©e par nom: '{func_name}' dans {doc.metadata.get('fileName', 'N/A')}")
                            break
            except Exception as e:
                logger.warning(f"Erreur lors de la recherche par nom de fonction: {str(e)}")
        
        logger.info(f"ðŸ”§ {len(function_name_matches)} documents trouvÃ©s par correspondance de nom de fonction/mÃ©thode")
        
        # Combiner les rÃ©sultats avec prioritÃ© aux correspondances exactes
        all_candidate_docs = {}
        # PrioritÃ© 1: Correspondances de noms de fonctions (score trÃ¨s Ã©levÃ©)
        for doc in function_name_matches:
            all_candidate_docs[id(doc)] = (doc, 0.95)  # Score trÃ¨s Ã©levÃ© pour correspondances exactes
        
        # PrioritÃ© 2: Correspondances de noms de fichiers (score Ã©levÃ©)
        for doc in filename_matches:
            doc_id = id(doc)
            if doc_id not in all_candidate_docs:
                all_candidate_docs[doc_id] = (doc, 0.85)
        
        # PrioritÃ© 3: RÃ©sultats de recherche sÃ©mantique (score normalisÃ©)
        for doc, score in search_results_with_scores:
            doc_id = id(doc)
            # Si dÃ©jÃ  prÃ©sent avec un meilleur score, garder le meilleur
            if doc_id not in all_candidate_docs:
                all_candidate_docs[doc_id] = (doc, 1 - score)
            elif all_candidate_docs[doc_id][1] < (1 - score):
                # Si le nouveau score est meilleur, le remplacer
                all_candidate_docs[doc_id] = (doc, 1 - score)
        
        # FILTRAGE STRICT PAR NOM DE PERSONNE (CRITIQUE)
        # Si un nom de personne est dÃ©tectÃ©, ne garder QUE les fichiers correspondant Ã  ce nom
        query_words = [w.strip('.,!?;:()[]{}"\'').lower() for w in user_query.split() if len(w.strip('.,!?;:()[]{}"\'')) > 2]
        filename_match_ids = {id(doc) for doc in filename_matches}
        filtered_docs = []
        
        # Liste des noms de personnes connus pour exclusion
        known_person_names = {
            'karim': ['dominique', 'essome'],
            'dominique': ['karim', 'ngami'],
            'essome': ['karim', 'ngami'],
            'ngami': ['dominique', 'essome']
        }
        
        # DÃ©terminer les noms Ã  exclure
        names_to_exclude = set()
        for detected_name in person_names:
            if detected_name in known_person_names:
                names_to_exclude.update(known_person_names[detected_name])
        
        logger.info(f"ðŸš« Noms Ã  exclure (autres personnes): {names_to_exclude}")
        
        for doc, similarity_score in all_candidate_docs.values():
            doc_id = id(doc)
            file_name_lower = (doc.metadata.get("fileName") or doc.metadata.get("file_name") or "").lower()
            doc_content_lower = doc.page_content[:500].lower()  # VÃ©rifier aussi dans le contenu
            
            # EXCLUSION STRICTE: Si un nom de personne est dÃ©tectÃ©, exclure les fichiers d'autres personnes
            should_exclude = False
            if person_names:
                # RÃˆGLE 1: Exclure TOUJOURS les fichiers contenant un nom Ã  exclure (mÃªme avec score Ã©levÃ©)
                for exclude_name in names_to_exclude:
                    if exclude_name in file_name_lower:
                        should_exclude = True
                        logger.info(f"âŒ Fichier EXCLU (nom de fichier contient '{exclude_name}'): {doc.metadata.get('fileName', 'N/A')}")
                        break
                    # VÃ©rifier aussi dans le contenu (premiers 200 caractÃ¨res pour Ãªtre plus strict)
                    if exclude_name in doc_content_lower[:200]:
                        should_exclude = True
                        logger.info(f"âŒ Fichier EXCLU (contenu contient '{exclude_name}'): {doc.metadata.get('fileName', 'N/A')}")
                        break
                
                # RÃˆGLE 2: Si un nom de personne est dÃ©tectÃ©, ne garder QUE les fichiers qui contiennent ce nom
                if not should_exclude:
                    file_contains_person_name = any(name in file_name_lower for name in person_names)
                    # VÃ©rifier aussi dans le contenu si pas dans le nom
                    if not file_contains_person_name:
                        file_contains_person_name = any(name in doc_content_lower[:300] for name in person_names)
                    
                    if not file_contains_person_name:
                        # Si aucun nom de personne n'est dans le fichier, EXCLURE TOUJOURS (mÃªme avec score Ã©levÃ©)
                        should_exclude = True
                        logger.info(f"âš ï¸ Fichier EXCLU (ne contient AUCUN nom de la personne recherchÃ©e): {doc.metadata.get('fileName', 'N/A')}")
            
            if should_exclude:
                continue  # Exclure ce document
            
            filename_keyword_matches = sum(1 for word in query_words if word in file_name_lower)
            
            # Score Ã©levÃ© pour les fichiers trouvÃ©s par nom de personne
            if doc_id in filename_match_ids:
                # Bonus supplÃ©mentaire si le fichier contient le nom de la personne dÃ©tectÃ©e
                person_name_bonus = 0.0
                if person_names:
                    for name in person_names:
                        if name in file_name_lower:
                            person_name_bonus = 0.15
                            break
                filtered_docs.append((doc, 0.9 + person_name_bonus))
                logger.info(f"âœ… Fichier inclus (correspond au nom): {doc.metadata.get('fileName', 'N/A')}")
            elif similarity_score >= 0.45 or filename_keyword_matches >= 1:
                filtered_docs.append((doc, similarity_score))
        
        filtered_docs.sort(key=lambda x: x[1], reverse=True)
        # AMÃ‰LIORATION: Augmenter le nombre de documents rÃ©cupÃ©rÃ©s pour amÃ©liorer la prÃ©cision
        # Si un nom de personne est dÃ©tectÃ©, limiter aux top 20 fichiers les plus pertinents
        # Sinon, prendre jusqu'Ã  30 documents pour avoir plus de contexte
        max_docs = 20 if person_names else 30
        top_docs = [doc for doc, score in filtered_docs[:max_docs]]
        
        logger.info(f"ðŸ“Š {len(top_docs)} documents finaux sÃ©lectionnÃ©s aprÃ¨s filtrage strict par nom")
        
        # Construire la liste de rÃ©sultats avec amÃ©lioration du contenu
        seen_docs = set()
        results = []
        for doc in top_docs:
            file_name_from_meta = doc.metadata.get("fileName") or doc.metadata.get("file_name") or "document_vectorstore"
            content_preview = doc.page_content[:100]
            doc_key = f"{file_name_from_meta}:{content_preview}"
            
            if doc_key not in seen_docs:
                seen_docs.add(doc_key)
                # AMÃ‰LIORATION: Augmenter la taille du contenu et chercher le contexte autour des mots-clÃ©s
                content = doc.page_content
                query_lower = user_query.lower()
                query_words = [w.strip('.,!?;:()[]{}"\'').lower() for w in user_query.split() if len(w.strip('.,!?;:()[]{}"\'')) > 2]
                
                # Si on trouve un mot-clÃ© dans le contenu, essayer d'inclure plus de contexte autour
                content_lower = content.lower()
                for word in query_words:
                    if word in content_lower:
                        # Trouver la position du mot et inclure plus de contexte
                        idx = content_lower.find(word)
                        if idx > 0:
                            # Inclure 500 caractÃ¨res avant et aprÃ¨s pour capturer toute la fonction
                            start = max(0, idx - 500)
                            end = min(len(content), idx + len(word) + 500)
                            # Si le contenu extrait est plus pertinent, l'utiliser
                            if end - start > len(content[:3000]):
                                content = content[start:end]
                                break
                
                results.append({
                    "fileName": file_name_from_meta,
                    "content": content[:3000]  # AugmentÃ© de 2500 Ã  3000 pour plus de contexte
                })
        
        logger.info(f"ðŸ“š {len(results)} documents uniques rÃ©cupÃ©rÃ©s pour la requÃªte")
        return results
        
    except Exception as e:
        logger.error(f"Erreur lors de la recherche sÃ©mantique synchrone: {str(e)}")
        return []

async def detect_missing_information(response: str, user_query: str, language: str = 'en') -> bool:
    """
    DÃ©tecte si la rÃ©ponse du modÃ¨le indique que l'information n'est pas disponible dans les documents
    ou si une recherche en ligne serait utile pour obtenir des informations en temps rÃ©el.
    Retourne True si l'information semble absente ou si une recherche en ligne serait bÃ©nÃ©fique.
    """
    response_lower = response.lower()
    query_lower = user_query.lower()
    
    # Phrases indicatrices que l'information n'est pas disponible
    missing_indicators = [
        "n'apparaÃ®t pas",
        "n'est pas disponible",
        "n'est pas prÃ©sent",
        "n'est pas trouvÃ©",
        "n'est pas mentionnÃ©",
        "n'est pas indiquÃ©",
        "n'est pas fourni",
        "n'est pas dans",
        "ne figure pas",
        "ne contient pas",
        "aucune information",
        "pas d'information",
        "information absente",
        "information non disponible",
        "je ne trouve pas",
        "je n'ai pas trouvÃ©",
        "impossible de trouver",
        "ne peut pas Ãªtre trouvÃ©",
        "does not appear",
        "is not available",
        "is not present",
        "is not found",
        "is not mentioned",
        "is not indicated",
        "is not provided",
        "is not in",
        "no information",
        "information not available",
        "i cannot find",
        "i did not find",
        "unable to find",
        "cannot be found"
    ]
    
    # Indicateurs que des informations plus rÃ©centes seraient utiles
    needs_recent_info_indicators = [
        "informations plus rÃ©centes",
        "information plus rÃ©cente",
        "donnÃ©es plus rÃ©centes",
        "vÃ©rifier les changements",
        "confirmer que",
        "pourrait Ãªtre utile",
        "serait utile",
        "pourrait nÃ©cessiter",
        "nÃ©cessiterait",
        "more recent information",
        "recent information",
        "more recent data",
        "check for changes",
        "confirm that",
        "might be useful",
        "would be useful",
        "might require",
        "would require",
        "should verify",
        "devrait vÃ©rifier",
        "il est important de vÃ©rifier",
        "it is important to check"
    ]
    
    # Mots-clÃ©s dans la requÃªte qui indiquent un besoin d'informations en temps rÃ©el
    real_time_query_keywords = [
        "heure actuelle",
        "heure maintenant",
        "quelle heure",
        "what time",
        "current time",
        "time now",
        "prix actuel",
        "prix maintenant",
        "current price",
        "price now",
        "taux actuel",
        "current rate",
        "taux de change",
        "exchange rate",
        "cours actuel",
        "current rate",
        "Ã©vÃ©nements rÃ©cents",
        "recent events",
        "nouvelles rÃ©centes",
        "recent news",
        "maintenant",
        "now",
        "actuel",
        "current",
        "aujourd'hui",
        "today",
        "en ce moment",
        "right now",
        "Ã  l'instant",
        "at the moment"
    ]
    
    # VÃ©rifier si la rÃ©ponse contient des indicateurs d'absence
    has_missing_indicator = any(indicator in response_lower for indicator in missing_indicators)
    
    # VÃ©rifier si la rÃ©ponse suggÃ¨re qu'une information plus rÃ©cente serait utile
    needs_recent_info = any(indicator in response_lower for indicator in needs_recent_info_indicators)
    
    # VÃ©rifier si la requÃªte demande des informations en temps rÃ©el
    needs_realtime_info = any(keyword in query_lower for keyword in real_time_query_keywords)
    
    # VÃ©rifier aussi la longueur de la rÃ©ponse - si elle est trÃ¨s courte, c'est probablement qu'il n'y a pas d'info
    is_very_short = len(response.strip()) < 100
    
    # VÃ©rifier si la rÃ©ponse contient des phrases comme "d'aprÃ¨s les documents" ou "dans les documents fournis"
    # Si oui, c'est probablement qu'il n'y a pas d'info ailleurs
    mentions_documents_only = any(phrase in response_lower for phrase in [
        "dans les documents fournis",
        "dans les documents",
        "dans le document",
        "d'aprÃ¨s les documents",
        "selon les documents",
        "in the provided documents",
        "in the documents",
        "in the document",
        "according to the documents"
    ])
    
    # DÃ©clencher une recherche en ligne si :
    # 1. La rÃ©ponse indique que l'info est absente
    # 2. La rÃ©ponse suggÃ¨re qu'une info plus rÃ©cente serait utile
    # 3. La requÃªte demande des informations en temps rÃ©el
    # 4. La rÃ©ponse est trÃ¨s courte ET mentionne seulement les documents
    should_search_online = (
        has_missing_indicator or 
        needs_recent_info or 
        needs_realtime_info or 
        (is_very_short and mentions_documents_only)
    )
    
    if should_search_online:
        reason = []
        if has_missing_indicator:
            reason.append("information absente")
        if needs_recent_info:
            reason.append("besoin d'informations rÃ©centes")
        if needs_realtime_info:
            reason.append("demande d'informations en temps rÃ©el")
        logger.info(f"ðŸ” Recherche en ligne dÃ©clenchÃ©e: {', '.join(reason)}")
    
    return should_search_online

async def query_model_local_mode(file_name: str, file_content: str, directory_content: List[Dict], 
                                repo_structure: str, user_query: str, is_binary: bool = False, 
                                selected_model: str = DEFAULT_MODEL, language: str = 'en',
                                conversation_history: Optional[List[Dict]] = None,
                                enable_auto_online_search: bool = True) -> tuple[str, str]:
    """
    Mode LOCAL: analyse STRICTEMENT les documents fournis (fichier principal + contexte).
    On rÃ©duit et structure le contexte pour amÃ©liorer la prÃ©cision et limiter les hallucinations.
    IMPORTANT: pas de cache sur ce mode, pour toujours prendre en compte le contexte le plus rÃ©cent
    (nouveaux fichiers indexÃ©s, corrections, etc.).
    """
    t = translations.get(language, translations['en'])
    
    # AmÃ©lioration: Organiser le contexte par pertinence et prÃ©server les informations clÃ©s
    max_chunk_len = 2000  # AugmentÃ© encore plus pour prÃ©server le maximum de contexte
    contextual_docs: List[str] = []
    
    # Grouper les documents par fichier source pour Ã©viter la duplication et prÃ©server le contexte complet
    # Structure: {file_label: [(content, page_number), ...]}
    docs_by_file = {}
    for doc in directory_content or []:
        if not isinstance(doc, dict):
            continue
        file_label = doc.get('fileName') or doc.get('file_name') or "document_contextuel"
        raw_content = doc.get('content', '') or ''
        if raw_content is None:
            raw_content = ''
        page_number = doc.get('pageNumber')  # page number from Docling/basic chunking
        chunk_id = doc.get('chunkId')
        section = doc.get('section') or ''   # heading hierarchy e.g. "Chapter 2 > Safety"
        
        if file_label not in docs_by_file:
            docs_by_file[file_label] = []
        docs_by_file[file_label].append((raw_content, page_number, chunk_id, section))
    
    # Construire un rÃ©sumÃ© structurÃ© par fichier (les premiers sont les plus pertinents)
    # IMPORTANT: PrÃ©server le maximum de contenu pour capturer les informations comme les adresses
    for idx, (file_label, contents_with_pages) in enumerate(docs_by_file.items(), 1):
        # Annoter chaque chunk avec son numÃ©ro de page AVANT de les combiner
        # Ainsi le modÃ¨le peut identifier exactement quelle page contient quelle information
        annotated_chunks = []
        for raw_content, page_number, chunk_id, section in contents_with_pages:
            parts = []
            if page_number is not None:
                parts.append(f"p.{page_number}")
            if section:
                parts.append(section)
            prefix = f"[{' | '.join(parts)}] " if parts else ""
            annotated_chunks.append(f"{prefix}{raw_content}")

        pages = [c[1] for c in contents_with_pages if c[1] is not None]
        
        # Combiner les chunks du mÃªme fichier avec un sÃ©parateur clair
        combined_content = "\n---\n".join(annotated_chunks)
        
        # CrÃ©er un label avec les informations de page
        page_info = ""
        if pages:
            unique_pages = sorted(set(pages))
            if len(unique_pages) == 1:
                page_info = f" ðŸ“„ Page {unique_pages[0]}"
            else:
                page_info = f" ðŸ“„ Pages {', '.join(map(str, unique_pages))}"
        
        # Pour les fichiers pertinents (top 5), prÃ©server encore plus de contenu
        extended_max_len = max_chunk_len * 1.5 if idx <= 5 else max_chunk_len
        
        # Tronquer intelligemment en prÃ©servant le dÃ©but et la fin
        if len(combined_content) > extended_max_len:
            # Garder le dÃ©but (souvent le plus important) et un peu de la fin
            # Augmenter la partie finale pour capturer les informations en fin de document
            content = combined_content[:int(extended_max_len - 300)] + "\n[... section tronquÃ©e ...]\n" + combined_content[-300:]
        else:
            content = combined_content
        
        # Marquer les documents les plus pertinents (premiers dans la liste)
        relevance_marker = "â­" if idx <= 3 else "ðŸ”" if idx <= 8 else ""
        # Header: page range + unique sections
        page_info = ""
        if pages:
            unique_pages = sorted(set(pages))
            if len(unique_pages) == 1:
                page_info = f" ðŸ“„ Page {unique_pages[0]}"
            else:
                page_info = f" ðŸ“„ Pages {', '.join(map(str, unique_pages))}"
        unique_sections = list(dict.fromkeys(
            c[3] for c in contents_with_pages if c[3]
        ))
        if unique_sections:
            page_info += f" ðŸ—‚ {' / '.join(unique_sections[:3])}"
        
        contextual_docs.append(f"{relevance_marker} [{idx}] {file_label}{page_info}:\n{content}")

    directory_content_summary = "\n\n" + "="*80 + "\n\n".join(contextual_docs) + "\n\n" + "="*80 if contextual_docs else t.get(
        'no_other_files', "Aucun autre fichier dans le contexte."
    )
    
    # Ajouter un header explicatif si on a du contexte
    if contextual_docs:
        directory_content_summary = (
            f"âš ï¸ ATTENTION: {len(docs_by_file)} fichiers ont Ã©tÃ© rÃ©cupÃ©rÃ©s par recherche sÃ©mantique "
            f"car ils sont potentiellement pertinents pour ta question.\n"
            f"Les documents sont classÃ©s par pertinence:\n"
            f"  â­ = TrÃ¨s pertinent (prioritÃ© haute)\n"
            f"  ðŸ” = Pertinent (prioritÃ© moyenne)\n"
            f"  [numÃ©ro] = Autre document Ã  analyser\n\n"
            f"TU DOIS ANALYSER TOUS CES FICHIERS, mÃªme si l'information semble absente du document principal.\n\n"
            f"{directory_content_summary}"
        )

    # On tronque aussi lÃ©gÃ¨rement le contenu du fichier principal si nÃ©cessaire
    main_max_len = 4000
    if file_content is None:
        file_content = ""
    trimmed_main_content = file_content[:main_max_len] + ("..." if len(file_content) > main_max_len else "")

    # DÃ©terminer si le contenu vient entiÃ¨rement des chunks (PDF/DOCX indexÃ© via Docling)
    content_is_chunked_only = not trimmed_main_content.strip() and bool(contextual_docs)

    # Construire un rÃ©sumÃ© compact de l'historique de conversation (si fourni)
    history_section = ""
    if conversation_history:
      history_lines: List[str] = []
      for turn in conversation_history[-10:]:
          role = turn.get("role", "user")
          content = (turn.get("content") or "")[:600]
          prefix = "Utilisateur" if role == "user" else "Assistant"
          history_lines.append(f"{prefix}: {content}")
      if history_lines:
          history_section = "=== HISTORIQUE DE LA CONVERSATION ===\n" + "\n".join(history_lines) + "\n\n"

    # DÃ©tection des noms de personnes dans la requÃªte pour instructions strictes
    person_names_in_query = set()
    words = user_query.split()
    for i, word in enumerate(words):
        clean_word = word.strip('.,!?;:()[]{}"\'').strip()
        if clean_word and clean_word[0].isupper() and len(clean_word) > 2:
            person_names_in_query.add(clean_word.lower())
        if i > 0 and words[i-1].lower() in ['his', 'her', 'their', 'karim', 'dominique', 'about', 'for']:
            if len(clean_word) > 2:
                person_names_in_query.add(clean_word.lower())
    
    query_lower = user_query.lower()
    common_names = ['karim', 'dominique', 'essome', 'ngami']
    for name in common_names:
        if name in query_lower:
            person_names_in_query.add(name)
    
    # Instructions strictes si un nom de personne est dÃ©tectÃ©
    person_filter_instructions = ""
    if person_names_in_query:
        person_names_str = ", ".join([name.capitalize() for name in person_names_in_query])
        person_filter_instructions = (
            f"\nðŸš¨ INSTRUCTION CRITIQUE - FILTRAGE PAR NOM DE PERSONNE:\n"
            f"La question concerne {person_names_str}. TU DOIS UNIQUEMENT utiliser les documents qui contiennent "
            f"le(s) nom(s) '{person_names_str}' dans leur nom de fichier OU dans leur contenu.\n"
            f"âŒ EXCLUSION STRICTE: N'utilise JAMAIS les documents d'autres personnes (comme Dominique si on demande Karim, "
            f"ou Karim si on demande Dominique).\n"
            f"âœ… PRIORITÃ‰: Les fichiers dont le nom contient '{person_names_str}' sont les SEULS documents pertinents.\n"
            f"âš ï¸ Si un document ne contient pas le nom '{person_names_str}', IGNORE-LE COMPLÃˆTEMENT, mÃªme s'il semble "
            f"sÃ©mantiquement similaire Ã  la question.\n\n"
        )
    
    # Prompt structurÃ© et trÃ¨s explicite avec instructions amÃ©liorÃ©es pour la recherche sÃ©mantique
    # IMPORTANT: Le modÃ¨le doit analyser TOUS les documents, pas seulement le fichier sÃ©lectionnÃ©
    num_other_docs = len(contextual_docs) if contextual_docs else 0
    total_docs_note = f"\nâš ï¸ ATTENTION: Tu as accÃ¨s Ã  {num_other_docs} autres documents en plus du document principal. " \
                      f"L'information demandÃ©e peut Ãªtre dans N'IMPORTE LEQUEL de ces documents. " \
                      f"ANALYSE TOUS LES DOCUMENTS avant de conclure qu'une information est absente.\n\n" if num_other_docs > 0 else ""

    # â”€â”€ Build the main-doc and chunks sections depending on whether we have raw content â”€â”€
    if content_is_chunked_only:
        # Binary file (PDF/DOCX): full content is in Docling chunks â€” make this crystal-clear to the model
        main_doc_section = (
            f"=== DOCUMENT ANALYSÃ‰: {file_name} ===\n"
            f"âš ï¸ Ce document (PDF/DOCX) a Ã©tÃ© parsÃ© par Docling et dÃ©coupÃ© en chunks ci-dessous.\n"
            f"ðŸš¨ RÃˆGLE ABSOLUE: L'intÃ©gralitÃ© du contenu de '{file_name}' EST dans les chunks ci-dessous. "
            f"NE DIS JAMAIS 'je ne trouve pas l'information' avant d'avoir lu CHAQUE chunk. "
            f"Si l'information est dans les chunks, elle EST dans le document.\n\n"
        )
        chunks_section_header = (
            f"=== CONTENU COMPLET DE '{file_name}' (CHUNKS DOCLING, classÃ©s par pertinence) ===\n"
            f"ðŸ”´ CRITIQUE: Ces chunks CONSTITUENT le document '{file_name}'. Ce sont les pages du document, pas des fichiers externes.\n"
            f"Chaque chunk est prÃ©fixÃ© [p.X | Section] indiquant sa page et section d'origine.\n"
            f"ANALYSE CHAQUE CHUNK pour rÃ©pondre Ã  la question. La rÃ©ponse EST lÃ .\n"
        )
    else:
        main_doc_section = (
            f"=== DOCUMENT PRINCIPAL (FICHIER SÃ‰LECTIONNÃ‰) ===\n"
            f"{t['main_file']}: {file_name}\n"
            f"{t['file_content']}:\n{trimmed_main_content}\n\n"
        )
        chunks_section_header = (
            f"=== AUTRES DOCUMENTS DU DOSSIER (RÃ‰CUPÃ‰RÃ‰S PAR RECHERCHE SÃ‰MANTIQUE) ===\n"
            f"âš ï¸ IMPORTANT: Ces documents ont Ã©tÃ© rÃ©cupÃ©rÃ©s car ils sont potentiellement pertinents. "
            f"ANALYSE-LES MÃ‰TICULEUSEMENT.\n"
        )

    prompt = (
        f"{t['local_analysis_mode']}\n"
        f"{t['no_external_search']}\n\n"
        f"{history_section}"
        f"{total_docs_note}"
        f"{person_filter_instructions}"
        f"=== CONTEXTE DU PROJET ===\n"
        f"{t['project_structure']}:\n{repo_structure}\n\n"
        f"{main_doc_section}"
        f"{chunks_section_header}"
        f"{directory_content_summary}\n\n"
        f"=== QUESTION Ã€ TRAITER ===\n"
        f"{t['question']}: {user_query}\n\n"
        f"{t['instructions']}:\n"
        f"{t['base_response_only']}\n"
        f"ðŸ” PROCÃ‰DURE DE RECHERCHE OBLIGATOIRE:\n"
    )

    if person_names_in_query:
        step1_text = "Si la question concerne une personne spÃ©cifique, vÃ©rifie d'abord que le document principal contient le nom de cette personne. Sinon, cherche dans les autres documents."
    else:
        step1_text = f"Lis TOUS les chunks de '{file_name}' listÃ©s ci-dessus." if content_is_chunked_only else f"Commence par analyser le document principal ({file_name})."

    prompt += (
        f"1. {step1_text}\n"
        f"2. {'Ne passe pas Ã  une conclusion sans avoir parcouru CHAQUE chunk.' if content_is_chunked_only else 'Si l information n est pas trouvÃ©e, PARCOURS SYSTÃ‰MATIQUEMENT TOUS les autres documents listÃ©s.'}\n"
        f"3. Les chunks marquÃ©s â­ sont les plus pertinents - commence par ceux-lÃ , mais lis aussi les autres.\n"
        f"4. Ne conclus JAMAIS qu'une information est absente avant d'avoir analysÃ© TOUS les chunks fournis.\n"
        f"5. ðŸ“Œ CITATION INLINE OBLIGATOIRE: Place la source IMMÃ‰DIATEMENT aprÃ¨s chaque fait, dans le mÃªme paragraphe.\n"
        f"   Format: **[{file_name}, p.X]** ou **[{file_name}, p.X | Section]** si la section est disponible.\n"
        f"   Exemple: Le salaire brut mensuel est de 3 500 â‚¬ **[contrat.pdf, p.4 | Article 6 > RÃ©munÃ©ration]**.\n"
        f"   âŒ NE regroupe PAS toutes les sources Ã  la fin â€” chaque fait doit avoir sa source INLINE.\n"
        f"6. CITATION VERBATIM: Pour les donnÃ©es clÃ©s (montants, dates, noms), entoure le texte exact du document avec ã€ ã€‘.\n"
        f"   Exemple: Le contrat stipule que ã€le salaire mensuel brut est fixÃ© Ã  3 500 eurosã€‘ **[contrat.pdf, p.4]**.\n"
        f"   N'invente JAMAIS de citation â€” copie exactement le texte du chunk.\n"
        f"7. Si une information ne peut ni Ãªtre lue ni dÃ©duite des chunks ci-dessus APRÃˆS AVOIR TOUT ANALYSÃ‰, "
        f"rÃ©pond explicitement qu'elle n'apparaÃ®t pas dans les documents fournis.\n"
        f"- Tu peux faire des calculs (totaux, salaire mensuel/annuel, etc.) Ã  partir des montants prÃ©sents dans les chunks.\n"
        f"{t['missing_info_clarify']}\n"
        f"{t['no_speculation']}\n"
        f"{t['focus_local_analysis']}\n"
        f"{t['emoji']}"
    )

    try:
        # LangGraph pipeline:
        # 1) generate answer with model fallback
        # 2) enforce source block from directory_content metadata (document/page/section)
        result, model_used = await run_answer_graph(
            prompt=prompt,
            selected_model=selected_model,
            user_query=user_query,
            directory_content=directory_content,
            model_executor=execute_model_query_with_fallback,
        )
        logger.info(f"âœ… ModÃ¨le utilisÃ© pour la rÃ©ponse: {model_used}")
        
        # DÃ©tection automatique si l'information n'est pas disponible et recherche en ligne si activÃ©e
        if enable_auto_online_search:
            missing_info = await detect_missing_information(result, user_query, language)
            
            if missing_info:
                logger.info("ðŸ” Information absente dÃ©tectÃ©e dans les documents. Recherche en ligne automatique...")
                t = translations.get(language, translations['en'])
                
                try:
                    # Effectuer une recherche en ligne
                    search_query = user_query
                    search_results = perform_online_search(search_query, language)
                    
                    if search_results and search_results != "Aucun rÃ©sultat trouvÃ©." and search_results != "ClÃ© API manquante":
                        # Fusionner les rÃ©sultats locaux avec les rÃ©sultats en ligne
                        enrichment_prompt = (
                            f"{t.get('enrichment_title', 'ðŸŒ Enrichissement avec recherche en ligne')}\n\n"
                            f"**RÃ©ponse initiale basÃ©e sur les documents locaux:**\n{result}\n\n"
                            f"**RÃ©sultats de recherche en ligne:**\n{search_results}\n\n"
                            f"{t.get('enrichment_instructions', 'Instructions:')}\n"
                            f"- Combine intelligemment les informations des documents locaux avec les donnÃ©es trouvÃ©es en ligne.\n"
                            f"- Si l'information n'est pas dans les documents locaux mais est trouvÃ©e en ligne, utilise les donnÃ©es en ligne.\n"
                            f"- Distingue clairement les sources: indique ce qui vient des documents locaux vs. ce qui vient de la recherche en ligne.\n"
                            f"- Cite les sources en ligne si disponibles.\n"
                            f"- Garde la structure et le format de la rÃ©ponse initiale si possible.\n"
                            f"- Priorise les informations rÃ©centes trouvÃ©es en ligne pour les donnÃ©es qui peuvent Ãªtre obsolÃ¨tes.\n\n"
                            f"**RÃ©ponse enrichie:**"
                        )
                        
                        enriched_result, _ = await execute_model_query_with_fallback(enrichment_prompt, selected_model, user_query)
                        
                        # Ajouter une note indiquant que la recherche en ligne a Ã©tÃ© utilisÃ©e
                        result = (
                            f"{enriched_result}\n\n"
                            f"---\n"
                            f"ðŸ’¡ **Note**: Cette rÃ©ponse combine les informations des documents locaux avec des donnÃ©es trouvÃ©es en ligne, "
                            f"car certaines informations n'Ã©taient pas disponibles dans les documents fournis.\n"
                        )
                        logger.info("âœ… RÃ©ponse enrichie avec des donnÃ©es en ligne")
                    else:
                        logger.info("âš ï¸ Aucun rÃ©sultat trouvÃ© en ligne ou clÃ© API manquante")
                        # Garder la rÃ©ponse originale si la recherche en ligne n'a rien donnÃ©
                except Exception as search_error:
                    logger.warning(f"âš ï¸ Erreur lors de la recherche en ligne automatique: {str(search_error)}")
                    # En cas d'erreur, garder la rÃ©ponse originale
                    pass
    except Exception as e:
        result = f"Erreur lors de l'analyse locale: {str(e)}"
        model_used = selected_model

    return result, model_used

def is_response_relevant(response: str, user_query: str) -> bool:
    """
    DÃ©termine si une rÃ©ponse est pertinente en vÃ©rifiant:
    1. Si la rÃ©ponse n'est pas vide ou trop courte
    2. Si la rÃ©ponse ne contient pas de phrases indiquant que l'information est absente
    3. Si la rÃ©ponse semble contenir une information rÃ©elle (pas juste "je ne sais pas")
    """
    if not response or len(response.strip()) < 20:
        logger.warning(f"âŒ RÃ©ponse trop courte ou vide: {len(response.strip() if response else 0)} chars")
        return False
    
    response_lower = response.lower()
    user_query_lower = user_query.lower()
    
    # Phrases indiquant que l'information n'est pas trouvÃ©e (plus complÃ¨tes)
    missing_phrases = [
        "n'apparaÃ®t pas dans les documents",
        "n'est pas prÃ©sente dans",
        "n'est pas disponible dans",
        "n'est pas trouvÃ©",
        "n'est pas mentionnÃ©",
        "n'est pas indiquÃ©",
        "ne trouve pas",
        "ne peut pas trouver",
        "ne peut pas Ãªtre trouvÃ©",
        "impossible de trouver",
        "information absente",
        "donnÃ©e absente",
        "non disponible",
        "pas d'information",
        "aucune information",
        "je ne peux pas",
        "je ne trouve pas",
        "je n'ai pas trouvÃ©",
        "dans les documents fournis, l'information",
        "l'information sur",
        "n'est pas directement mentionnÃ©",
        "does not appear in",
        "is not present in",
        "is not available in",
        "is not found",
        "is not mentioned",
        "cannot find",
        "unable to find",
        "information missing",
        "data not available",
        "no information",
        "i cannot",
        "i cannot find",
        "i did not find",
        "in the provided documents, the information",
        "the information about",
        "is not directly mentioned"
    ]
    
    # DÃ©tection stricte: si une phrase d'absence apparaÃ®t dans la rÃ©ponse, c'est non pertinent
    # VÃ©rifier particuliÃ¨rement dans les premiers 300 caractÃ¨res (dÃ©but de rÃ©ponse)
    response_start = response_lower[:300]
    for phrase in missing_phrases:
        if phrase in response_lower:
            # Si la phrase d'absence apparaÃ®t dans les 300 premiers caractÃ¨res, c'est dÃ©finitivement non pertinent
            if phrase in response_start:
                logger.warning(f"âŒ Phrase d'absence dÃ©tectÃ©e dans les premiers 300 chars: '{phrase}'")
                return False
            # Si la phrase apparaÃ®t plus loin mais que la rÃ©ponse contient principalement cette information, c'est aussi non pertinent
            # Compter combien de fois les phrases d'absence apparaissent
            missing_count = sum(1 for p in missing_phrases if p in response_lower)
            if missing_count >= 2:  # Si 2+ phrases d'absence, c'est probablement non pertinent
                logger.warning(f"âŒ Plusieurs phrases d'absence dÃ©tectÃ©es ({missing_count})")
                return False
    
    # VÃ©rifier que la rÃ©ponse contient au moins quelques mots de la requÃªte (pertinence sÃ©mantique)
    query_words = set(word.lower().strip('.,!?;:()[]{}"\'') for word in user_query_lower.split() if len(word) > 2)
    response_words = set(word.lower().strip('.,!?;:()[]{}"\'') for word in response_lower.split() if len(word) > 2)
    
    # Extraire les mots-clÃ©s importants de la requÃªte (noms propres, mots significatifs)
    important_query_words = [w for w in query_words if w not in ['est', 'sont', 'Ãªtre', 'avoir', 'was', 'is', 'are', 'the', 'a', 'an', 'le', 'la', 'les', 'un', 'une', 'des']]
    
    if important_query_words:
        # VÃ©rifier si au moins un mot important de la requÃªte est dans la rÃ©ponse
        common_words = query_words & response_words
        if len(common_words) == 0:
            logger.warning(f"âŒ Aucun mot en commun entre requÃªte et rÃ©ponse. RequÃªte: {important_query_words[:5]}")
            return False
        
        # VÃ©rifier que la rÃ©ponse contient au moins 30% des mots importants
        important_common = [w for w in important_query_words if w in response_words]
        if len(important_common) / len(important_query_words) < 0.3:
            logger.warning(f"âŒ Trop peu de mots importants trouvÃ©s ({len(important_common)}/{len(important_query_words)})")
            # Mais accepter si la rÃ©ponse est trÃ¨s dÃ©taillÃ©e et contient au moins un mot important
            if len(response.strip()) < 400 or len(important_common) == 0:
                return False
    
    # VÃ©rifier si la rÃ©ponse est trop gÃ©nÃ©rique (ex: "veuillez consulter les documents")
    generic_phrases = [
        "consultez les documents",
        "veuillez consulter",
        "please consult",
        "refer to the documents",
        "see the documents"
    ]
    if any(phrase in response_lower for phrase in generic_phrases) and len(response.strip()) < 150:
        logger.warning(f"âŒ RÃ©ponse trop gÃ©nÃ©rique")
        return False
    
    logger.info(f"âœ… RÃ©ponse considÃ©rÃ©e comme pertinente ({len(response.strip())} chars, {len(query_words & response_words)} mots en commun)")
    return True

async def execute_model_query_with_fallback(prompt: str, selected_model: str, user_query: str = "") -> tuple[str, str]:
    """
    ExÃ©cute la requÃªte sur le modÃ¨le sÃ©lectionnÃ© avec fallback en cascade si la rÃ©ponse n'est pas pertinente.
    Retourne (response, model_used)
    
    Ordre de fallback si GPT-3.5-turbo ne donne pas de rÃ©ponse pertinente:
    1. Mistral
    2. GPT-4o
    3. GPT-5-Nano
    4. GPT-5-Mini
    5. GPT-5
    """
    # Liste complÃ¨te des modÃ¨les de fallback (dans l'ordre de qualitÃ© croissante)
    all_fallback_models = ["gpt-3.5-turbo", "mistral", "gpt-4o", "gpt-5-nano", "gpt-5-mini", "gpt-5"]
    
    # Construire la liste des modÃ¨les Ã  essayer en commenÃ§ant par le modÃ¨le sÃ©lectionnÃ©
    selected_lower = selected_model.lower()
    
    if selected_lower in all_fallback_models:
        # Trouver l'index du modÃ¨le sÃ©lectionnÃ© et prendre tous les modÃ¨les suivants
        model_idx = all_fallback_models.index(selected_lower)
        models_to_try = all_fallback_models[model_idx:]
        logger.info(f"ðŸ“‹ Ordre de fallback: {models_to_try}")
    else:
        # ModÃ¨le non dans la liste, utiliser tel quel puis fallback complet
        models_to_try = [selected_lower] + all_fallback_models
        logger.info(f"ðŸ“‹ ModÃ¨le inconnu '{selected_lower}', utilisation avec fallback complet: {models_to_try}")
    
    last_exception = None
    
    for model in models_to_try:
        try:
            logger.info(f"ðŸ”„ Tentative avec le modÃ¨le: {model}")
            
            if model in ["gpt-3.5-turbo", "gpt-4o", "gpt-5", "gpt-5-mini", "gpt-5-nano"]:
                result = call_openai_api(prompt, model)
            elif model == "mistral":
                result = call_mistral_api(prompt)
            elif model.startswith("gemini") or model in ["gemini-3-flash", "gemini-pro"]:
                result = call_gemini_api(prompt, model)
            elif model in OLLAMA_MODELS or model == "llama3":
                result = call_ollama_api(prompt, model)
            else:
                # ModÃ¨le inconnu, utiliser OpenAI par dÃ©faut
                result = call_openai_api(prompt, DEFAULT_MODEL)
                model = DEFAULT_MODEL
            
            # VÃ©rifier la pertinence si user_query est fourni
            # Pour GPT-3.5-turbo, on vÃ©rifie toujours (plus susceptible d'Ã©chouer)
            # Pour les autres modÃ¨les moins puissants (mistral), on vÃ©rifie aussi
            # Pour les modÃ¨les plus puissants (gpt-4o, gpt-5*), on fait confiance mais on peut quand mÃªme vÃ©rifier
            should_check_relevance = user_query and (
                model == "gpt-3.5-turbo" or 
                model == "mistral" or
                model in ["gpt-5-nano", "gpt-5-mini"]  # VÃ©rifier aussi pour les modÃ¨les moins puissants
            )
            
            if should_check_relevance:
                if is_response_relevant(result, user_query):
                    logger.info(f"âœ… RÃ©ponse pertinente obtenue avec {model}")
                    return result, model
                else:
                    logger.warning(f"âš ï¸ RÃ©ponse non pertinente avec {model}, passage au modÃ¨le suivant...")
                    continue
            else:
                # Pour les modÃ¨les plus puissants (gpt-4o, gpt-5) ou si pas de user_query, accepter la rÃ©ponse
                logger.info(f"âœ… RÃ©ponse obtenue avec {model} (vÃ©rification de pertinence sautÃ©e)")
                return result, model
                
        except Exception as e:
            logger.warning(f"âŒ Erreur avec le modÃ¨le {model}: {str(e)}, passage au modÃ¨le suivant...")
            last_exception = e
            continue
    
    # Si tous les modÃ¨les ont Ã©chouÃ©, lever l'exception du dernier
    if last_exception:
        raise last_exception
    else:
        raise RuntimeError("Tous les modÃ¨les de fallback ont Ã©chouÃ©")

async def execute_model_query(prompt: str, selected_model: str, conversation_history: list = None) -> str:
    """
    ExÃ©cute la requÃªte sur le modÃ¨le sÃ©lectionnÃ© (version simple, sans fallback)
    """
    try:
        if selected_model.lower() in ["gpt-3.5-turbo", "gpt-4o"]:
            result = call_openai_api(prompt, selected_model, conversation_history=conversation_history)
        elif selected_model.lower() in ["gpt-5", "gpt-5-mini", "gpt-5-nano"]:
            result = call_openai_api(prompt, selected_model, conversation_history=conversation_history)
        elif selected_model.lower() == "openai":
            result = call_openai_api(prompt, "openai", conversation_history=conversation_history)
        elif selected_model.lower() == "mistral":
            # Pour Mistral, intÃ©grer l'historique dans le prompt si disponible
            if conversation_history and isinstance(conversation_history, list):
                history_text = "\n\n".join([
                    f"{'User' if msg.get('role') == 'user' else 'Assistant'}: {msg.get('content', '')}"
                    for msg in conversation_history
                ])
                full_prompt = f"{history_text}\n\nUser: {prompt}\n\nAssistant:"
            else:
                full_prompt = prompt
            result = call_mistral_api(full_prompt)
        elif selected_model.lower().startswith("gemini") or selected_model.lower() in ["gemini-3-flash", "gemini-pro"]:
            result = call_gemini_api(prompt, selected_model)
        elif selected_model.lower() in OLLAMA_MODELS or selected_model.lower() == "llama3":
            result = call_ollama_api(prompt, selected_model)
        else:
            logger.warning(f"ModÃ¨le inconnu {selected_model}, utilisation du modÃ¨le par dÃ©faut {DEFAULT_MODEL}")
            result = call_openai_api(prompt, DEFAULT_MODEL, conversation_history=conversation_history)
        
        return result
    except Exception as e:
        raise e

def _score_document_type(content_lower: str, file_name_lower: str, keywords: list, indicators: list,
                         keyword_weight: float = 3.0, indicator_weight: float = 1.0,
                         filename_weight: float = 5.0, header_boost: float = 2.0) -> float:
    """
    Calcule un score pondÃ©rÃ© pour un type de document.
    - Les mots-clÃ©s dans le nom du fichier ont le poids le plus Ã©levÃ©
    - Les mots-clÃ©s dans les premiers 800 chars du contenu ont un boost
    - Les indicateurs dans le contenu complet contribuent proportionnellement
    """
    score = 0.0
    header = content_lower[:800]
    
    # Score des mots-clÃ©s dans le nom du fichier (trÃ¨s fort signal)
    for kw in keywords:
        if kw in file_name_lower:
            score += filename_weight
    
    # Score des mots-clÃ©s dans l'en-tÃªte du document (fort signal)
    for kw in keywords:
        if kw in header:
            score += keyword_weight * header_boost
        elif kw in content_lower:
            score += keyword_weight
    
    # Score des indicateurs contextuels
    for ind in indicators:
        if ind in content_lower:
            score += indicator_weight
    
    return score

# Liste complÃ¨te de tous les types de documents valides
ALL_VALID_DOCUMENT_TYPES = [
    'cv_resume', 'facture_invoice', 'contrat_location', 'contrat_travail',
    'contrat_vente', 'contrat_generique', 'contrat_prenuptial', 'procuration_poa',
    'accord_confidentialite_nda', 'acte_propriete_immobiliere', 'testament',
    'acte_notarie', 'lettre', 'document_financier',
    'assurance_insurance', 'jugement_decision_justice', 'releve_bancaire',
    'certificat_attestation', 'contrat_pret_loan', 'devis_estimation',
    'bon_commande_purchase_order', 'proces_verbal', 'rapport_expertise',
    'permis_licence',
    'document_generique'
]

async def detect_document_type(file_content: str, file_name: str) -> str:
    """
    DÃ©tecte le type de document de maniÃ¨re prÃ©cise en utilisant une approche hybride:
    1. Scoring pondÃ©rÃ© basÃ© sur des patterns (rÃ¨gles) â€” rapide et fiable
    2. Analyse LLM pour les cas complexes (fallback)
    Retourne le type de document (str) pour compatibilitÃ© ascendante.
    Utiliser detect_document_type_detailed() pour obtenir le score de confiance.
    """
    result = await detect_document_type_detailed(file_content, file_name)
    return result['type']

async def detect_document_type_detailed(file_content: str, file_name: str) -> Dict[str, Any]:
    """
    DÃ©tecte le type de document avec scoring de confiance.
    Retourne: { 'type': str, 'confidence': float (0-1), 'method': 'pattern'|'llm'|'fallback', 'scores': dict }
    """
    content_lower = file_content[:5000].lower()
    file_name_lower = file_name.lower()
    content_sample = file_content[:1500]
    
    # ============ DÃ‰FINITION DES TYPES AVEC PATTERNS PONDÃ‰RÃ‰S ============
    
    type_definitions = {
        'cv_resume': {
            'keywords': ['curriculum vitae', 'cv', 'resume', 'rÃ©sumÃ© professionnel', 'professional summary'],
            'indicators': ['expÃ©rience professionnelle', 'professional experience', 'compÃ©tences', 'skills',
                          'formation', 'education', 'emploi', 'job', 'poste', 'position',
                          'langues', 'languages', 'certifications', 'rÃ©alisations', 'achievements',
                          'rÃ©fÃ©rences', 'references', 'profil', 'profile', 'objectif', 'objective'],
            'threshold': 8.0
        },
        'facture_invoice': {
            'keywords': ['facture', 'invoice', 'bill', 'reÃ§u', 'receipt', 'quittance', 'nota fiscal'],
            'indicators': ['montant total', 'total amount', 'tva', 'tax', 'date d\'Ã©mission', 'issue date',
                          'numÃ©ro de facture', 'invoice number', 'client', 'customer', 'paiement', 'payment',
                          'montant ht', 'net amount', 'montant ttc', 'gross amount', 'Ã©chÃ©ance', 'due date',
                          'bon de commande', 'purchase order', 'rÃ©fÃ©rence', 'reference'],
            'threshold': 7.0
        },
        'devis_estimation': {
            'keywords': ['devis', 'quote', 'quotation', 'estimation', 'estimate', 'proforma', 'pro forma',
                        'offre de prix', 'price offer', 'proposition commerciale', 'commercial proposal'],
            'indicators': ['prix unitaire', 'unit price', 'quantitÃ©', 'quantity', 'total ht', 'total hors taxe',
                          'validitÃ©', 'validity', 'durÃ©e de validitÃ©', 'conditions de paiement', 'payment terms',
                          'remise', 'discount', 'livraison', 'delivery', 'dÃ©lai', 'deadline',
                          'sous-total', 'subtotal', 'main d\'oeuvre', 'labor'],
            'threshold': 7.0
        },
        'bon_commande_purchase_order': {
            'keywords': ['bon de commande', 'purchase order', 'order form', 'commande', 'order',
                        'bon de livraison', 'delivery note', 'bordereau'],
            'indicators': ['numÃ©ro de commande', 'order number', 'fournisseur', 'supplier', 'vendor',
                          'quantitÃ© commandÃ©e', 'ordered quantity', 'date de livraison', 'delivery date',
                          'rÃ©fÃ©rence article', 'article reference', 'prix unitaire', 'unit price',
                          'conditions de livraison', 'delivery terms', 'adresse de livraison', 'shipping address'],
            'threshold': 7.0
        },
        'assurance_insurance': {
            'keywords': ['assurance', 'insurance', 'police d\'assurance', 'insurance policy', 'avenant',
                        'contrat d\'assurance', 'insurance contract', 'assureur', 'insurer',
                        'garantie d\'assurance', 'insurance coverage', 'souscription'],
            'indicators': ['prime', 'premium', 'franchise', 'deductible', 'sinistre', 'claim',
                          'couverture', 'coverage', 'assurÃ©', 'insured', 'bÃ©nÃ©ficiaire', 'beneficiary',
                          'indemnisation', 'indemnity', 'risque', 'risk', 'exclusion', 'exclusion',
                          'dÃ©claration de sinistre', 'claim declaration', 'plafond', 'ceiling',
                          'responsabilitÃ© civile', 'civil liability', 'tiers', 'third party'],
            'threshold': 8.0
        },
        'jugement_decision_justice': {
            'keywords': ['jugement', 'judgment', 'arrÃªt', 'ruling', 'ordonnance', 'order',
                        'dÃ©cision de justice', 'court decision', 'sentence', 'verdict',
                        'tribunal', 'court', 'cour d\'appel', 'court of appeal'],
            'indicators': ['attendu que', 'whereas', 'par ces motifs', 'for these reasons',
                          'condamne', 'condemns', 'dÃ©boute', 'dismisses', 'juge', 'judge',
                          'demandeur', 'plaintiff', 'dÃ©fendeur', 'defendant', 'partie civile',
                          'greffier', 'clerk', 'audience', 'hearing', 'chambre', 'chamber',
                          'code civil', 'civil code', 'code pÃ©nal', 'penal code',
                          'article', 'pourvoi', 'appeal', 'cassation'],
            'threshold': 8.0
        },
        'releve_bancaire': {
            'keywords': ['relevÃ© bancaire', 'bank statement', 'relevÃ© de compte', 'account statement',
                        'extrait de compte', 'account extract', 'relevÃ© mensuel', 'monthly statement'],
            'indicators': ['solde', 'balance', 'dÃ©bit', 'debit', 'crÃ©dit', 'credit',
                          'virement', 'transfer', 'retrait', 'withdrawal', 'dÃ©pÃ´t', 'deposit',
                          'numÃ©ro de compte', 'account number', 'iban', 'bic', 'swift',
                          'intÃ©rÃªts', 'interest', 'frais bancaires', 'bank fees',
                          'solde initial', 'opening balance', 'solde final', 'closing balance',
                          'date valeur', 'value date', 'libellÃ©', 'description'],
            'threshold': 7.0
        },
        'certificat_attestation': {
            'keywords': ['certificat', 'certificate', 'attestation', 'certification',
                        'acte de naissance', 'birth certificate', 'acte de mariage', 'marriage certificate',
                        'acte de dÃ©cÃ¨s', 'death certificate', 'attestation de domicile', 'proof of residence',
                        'attestation d\'emploi', 'employment certificate', 'attestation de scolaritÃ©'],
            'indicators': ['certifie', 'certify', 'atteste', 'attest', 'dÃ©clare', 'declare',
                          'fait Ã ', 'done at', 'le soussignÃ©', 'the undersigned', 'en foi de quoi',
                          'in witness whereof', 'cachet', 'stamp', 'sceau', 'seal',
                          'dÃ©livrÃ© Ã ', 'issued to', 'valable', 'valid', 'authentique', 'authentic'],
            'threshold': 7.0
        },
        'contrat_pret_loan': {
            'keywords': ['contrat de prÃªt', 'loan agreement', 'prÃªt', 'loan', 'crÃ©dit', 'credit',
                        'emprunt', 'borrowing', 'financement', 'financing',
                        'prÃªt immobilier', 'mortgage loan', 'prÃªt personnel', 'personal loan',
                        'contrat de crÃ©dit', 'credit agreement'],
            'indicators': ['taux d\'intÃ©rÃªt', 'interest rate', 'mensualitÃ©', 'monthly payment',
                          'amortissement', 'amortization', 'capital', 'principal',
                          'emprunteur', 'borrower', 'prÃªteur', 'lender', 'Ã©chÃ©ance', 'maturity',
                          'remboursement', 'repayment', 'garantie', 'collateral',
                          'taux annuel', 'annual rate', 'taeg', 'apr', 'durÃ©e du prÃªt', 'loan term',
                          'hypothÃ¨que', 'mortgage', 'cautionnement', 'surety'],
            'threshold': 8.0
        },
        'proces_verbal': {
            'keywords': ['procÃ¨s-verbal', 'procÃ¨s verbal', 'pv', 'minutes', 'compte-rendu', 'compte rendu',
                        'minutes of meeting', 'meeting minutes', 'rapport de rÃ©union', 'meeting report'],
            'indicators': ['rÃ©union', 'meeting', 'assemblÃ©e', 'assembly', 'sÃ©ance', 'session',
                          'ordre du jour', 'agenda', 'participants', 'attendees', 'prÃ©sents', 'present',
                          'absents', 'absent', 'dÃ©cision', 'decision', 'vote', 'rÃ©solution', 'resolution',
                          'secrÃ©taire', 'secretary', 'prÃ©sident', 'chairman', 'quorum',
                          'approbation', 'approval', 'adoptÃ©', 'adopted', 'unanimitÃ©', 'unanimity'],
            'threshold': 7.0
        },
        'rapport_expertise': {
            'keywords': ['rapport', 'report', 'expertise', 'audit', 'diagnostic',
                        'rapport d\'expertise', 'expert report', 'rapport d\'audit', 'audit report',
                        'rapport mÃ©dical', 'medical report', 'bilan', 'assessment'],
            'indicators': ['conclusion', 'findings', 'recommandation', 'recommendation',
                          'analyse', 'analysis', 'observation', 'observation', 'constat', 'finding',
                          'expert', 'Ã©valuation', 'evaluation', 'mÃ©thodologie', 'methodology',
                          'rÃ©sultat', 'result', 'annexe', 'appendix', 'synthÃ¨se', 'summary',
                          'pÃ©rimÃ¨tre', 'scope', 'mission', 'objectif', 'objective'],
            'threshold': 7.0
        },
        'permis_licence': {
            'keywords': ['permis', 'permit', 'licence', 'license', 'autorisation', 'authorization',
                        'agrÃ©ment', 'approval', 'habilitation', 'accreditation',
                        'permis de construire', 'building permit', 'permis de conduire', 'driving license',
                        'licence professionnelle', 'professional license'],
            'indicators': ['autorisÃ©', 'authorized', 'titulaire', 'holder', 'dÃ©livrÃ© par', 'issued by',
                          'date d\'expiration', 'expiration date', 'date de dÃ©livrance', 'issue date',
                          'numÃ©ro de permis', 'permit number', 'numÃ©ro de licence', 'license number',
                          'conditions', 'restrictions', 'catÃ©gorie', 'category',
                          'renouvellement', 'renewal', 'validitÃ©', 'validity'],
            'threshold': 7.0
        },
        'lettre': {
            'keywords': ['lettre', 'letter', 'correspondance', 'correspondence', 'courrier', 'mail'],
            'indicators': ['soutien', 'support', 'recommandation', 'recommendation', 'demande', 'request',
                          'cher monsieur', 'dear sir', 'madame', 'madam', 'monsieur', 'sir',
                          'veuillez agrÃ©er', 'yours sincerely', 'cordialement', 'best regards',
                          'objet:', 'subject:', 'rÃ©fÃ©rence:', 'reference:',
                          'piÃ¨ce jointe', 'attachment', 'ci-joint', 'enclosed'],
            'threshold': 7.0
        },
        'document_financier': {
            'keywords': ['t4', 't-4', 't4a', 't4a-', 'relevÃ©', 'statement', 'payroll', 'paie',
                        'salaire', 'salary', 'revenu', 'income', 'impÃ´t', 'tax', 'dÃ©duction', 'deduction',
                        'feuillet', 'slip', 'relevÃ© fiscal', 'tax statement', 'avis de cotisation', 'notice of assessment'],
            'indicators': ['revenus bruts', 'gross income', 'revenus nets', 'net income', 'impÃ´t retenu', 'tax withheld',
                          'annÃ©e', 'year', 'pÃ©riode', 'period', 'numÃ©ro d\'assurance sociale', 'social insurance number',
                          'cotisation', 'contribution', 'dÃ©claration', 'declaration', 'fiscal', 'fiscal'],
            'threshold': 7.0
        },
        'contrat_location': {
            'keywords': ['contrat de location', 'rental contract', 'bail', 'lease', 'contrat de bail'],
            'indicators': ['location', 'rental', 'loyer', 'rent', 'locataire', 'tenant',
                          'bailleur', 'landlord', 'propriÃ©taire', 'owner', 'garant', 'guarantor',
                          'charges locatives', 'maintenance fees', 'caution', 'deposit',
                          'Ã©tat des lieux', 'inventory', 'prÃ©avis', 'notice period'],
            'threshold': 7.0
        },
        'contrat_travail': {
            'keywords': ['contrat de travail', 'employment contract', 'employment agreement'],
            'indicators': ['travail', 'employment', 'employÃ©', 'employee', 'employeur', 'employer',
                          'salaire', 'salary', 'rÃ©munÃ©ration', 'remuneration', 'poste', 'position',
                          'pÃ©riode d\'essai', 'probation', 'cdi', 'cdd', 'contract duration',
                          'congÃ©', 'leave', 'licenciement', 'termination', 'dÃ©mission', 'resignation'],
            'threshold': 7.0
        },
        'contrat_vente': {
            'keywords': ['contrat de vente', 'sale contract', 'sale agreement', 'compromis de vente'],
            'indicators': ['vente', 'sale', 'achat', 'purchase', 'acheteur', 'buyer', 'vendeur', 'seller',
                          'prix', 'price', 'livraison', 'delivery', 'garantie', 'warranty',
                          'transfert de propriÃ©tÃ©', 'transfer of ownership', 'clause rÃ©solutoire'],
            'threshold': 7.0
        },
        'contrat_generique': {
            'keywords': ['contrat', 'contract', 'agreement', 'convention', 'accord'],
            'indicators': ['parties', 'obligation', 'clause', 'signature', 'durÃ©e', 'duration',
                          'rÃ©siliation', 'termination', 'litige', 'dispute', 'juridiction', 'jurisdiction'],
            'threshold': 5.0
        },
        'testament': {
            'keywords': ['testament', 'will', 'testamentaire', 'testamentary', 'derniÃ¨res volontÃ©s', 'last will'],
            'indicators': ['hÃ©ritier', 'heir', 'bÃ©nÃ©ficiaire', 'beneficiary', 'legs', 'bequest', 'legataire', 'legatee',
                          'exÃ©cuteur', 'executor', 'succession', 'inheritance', 'rÃ©vocation', 'revocation',
                          'olographe', 'holographic', 'codicille', 'codicil'],
            'threshold': 7.0
        },
        'contrat_prenuptial': {
            'keywords': ['contrat de mariage', 'prenuptial agreement', 'prenup', 'contrat prÃ©nuptial',
                        'marriage contract', 'convention matrimoniale', 'marital agreement'],
            'indicators': ['rÃ©gime matrimonial', 'matrimonial regime', 'biens', 'property', 'sÃ©paration de biens',
                          'separation of property', 'communautÃ©', 'community', 'Ã©poux', 'spouse', 'mariage', 'marriage',
                          'dot', 'dowry', 'donation entre Ã©poux', 'spousal donation'],
            'threshold': 7.0
        },
        'procuration_poa': {
            'keywords': ['procuration', 'power of attorney', 'power-of-attorney', 'mandat', 'mandate', 'poa'],
            'indicators': ['mandant', 'principal', 'mandataire', 'agent', 'attorney-in-fact', 'pouvoir', 'authority',
                          'reprÃ©senter', 'represent', 'agir au nom', 'act on behalf', 'signer', 'sign',
                          'dÃ©lÃ©gation', 'delegation', 'irrÃ©vocable', 'irrevocable'],
            'threshold': 7.0
        },
        'accord_confidentialite_nda': {
            'keywords': ['accord de confidentialitÃ©', 'non-disclosure agreement', 'nda', 'n.d.a.',
                        'confidentiality agreement', 'accord de non-divulgation'],
            'indicators': ['confidentiel', 'confidential', 'secret', 'proprietary', 'propriÃ©taire',
                          'divulgation', 'disclosure', 'rÃ©vÃ©ler', 'reveal', 'informations confidentielles',
                          'pÃ©nalitÃ©', 'penalty', 'durÃ©e de confidentialitÃ©', 'confidentiality period'],
            'threshold': 7.0
        },
        'acte_propriete_immobiliere': {
            'keywords': ['acte de propriÃ©tÃ©', 'real estate deed', 'property deed', 'acte de vente immobiliÃ¨re',
                        'deed of sale', 'title deed', 'titre de propriÃ©tÃ©'],
            'indicators': ['propriÃ©taire', 'owner', 'propriÃ©tÃ© immobiliÃ¨re', 'real estate', 'bien immobilier',
                          'property', 'parcelle', 'lot', 'cadastre', 'cadastral', 'superficie', 'area',
                          'adresse', 'address', 'bornage', 'boundary', 'hypothÃ¨que', 'mortgage'],
            'threshold': 7.0
        },
        'acte_notarie': {
            'keywords': ['acte notariÃ©', 'notarial act', 'acte authentique', 'authentic act', 'acte sous seing privÃ©'],
            'indicators': ['notaire', 'notary', 'Ã©tude notariale', 'notary office', 'minute',
                          'authentification', 'signature authentique', 'rÃ©pertoire des minutes',
                          'en prÃ©sence de', 'in the presence of'],
            'threshold': 7.0
        },
    }
    
    # ============ SCORING DE TOUS LES TYPES ============
    scores = {}
    for doc_type, definition in type_definitions.items():
        score = _score_document_type(
            content_lower, file_name_lower,
            definition['keywords'], definition['indicators']
        )
        if score > 0:
            scores[doc_type] = score
    
    # Trier par score dÃ©croissant
    sorted_scores = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    
    if sorted_scores:
        best_type, best_score = sorted_scores[0]
        threshold = type_definitions[best_type]['threshold']
        
        # Calculer la confiance (0-1) basÃ©e sur le score vs threshold et la marge avec le 2nd
        raw_confidence = min(best_score / (threshold * 2), 1.0)  # Normaliser
        
        # Bonus si la marge avec le 2Ã¨me est grande
        if len(sorted_scores) >= 2:
            second_score = sorted_scores[1][1]
            margin = (best_score - second_score) / best_score if best_score > 0 else 0
            raw_confidence = min(raw_confidence + margin * 0.2, 1.0)
        
        if best_score >= threshold:
            # RÃ©soudre les ambiguÃ¯tÃ©s contrat spÃ©cifique vs contrat_generique
            if best_type == 'contrat_generique' and len(sorted_scores) >= 2:
                second_type, second_score = sorted_scores[1]
                if second_type.startswith('contrat_') and second_score >= type_definitions[second_type]['threshold'] * 0.8:
                    return {
                        'type': second_type,
                        'confidence': raw_confidence,
                        'method': 'pattern',
                        'scores': dict(sorted_scores[:5])
                    }
            
            logger.info(f"ðŸ“‹ Document type detected: {best_type} (score={best_score:.1f}, confidence={raw_confidence:.2f})")
            return {
                'type': best_type,
                'confidence': raw_confidence,
                'method': 'pattern',
                'scores': dict(sorted_scores[:5])
            }
    
    # ============ DÃ‰TECTION PAR LLM POUR CAS COMPLEXES ============
    try:
        types_list = '\n'.join(f'- {t}' for t in ALL_VALID_DOCUMENT_TYPES)
        detection_prompt = f"""Analyze this document sample and identify its type. Return ONLY one of these exact types:
{types_list}

File name: {file_name}
Content sample (first 1500 chars): {content_sample}

Return ONLY the type identifier, nothing else:"""
        
        detected_type = call_mistral_api(detection_prompt).strip().lower()
        
        # Nettoyer le rÃ©sultat (enlever ponctuation, espaces, markdown)
        detected_type = re.sub(r'[^a-z0-9_]', '', detected_type.replace('-', '_'))
        
        if detected_type in ALL_VALID_DOCUMENT_TYPES:
            logger.info(f"ðŸ“‹ Document type detected via LLM: {detected_type}")
            return {
                'type': detected_type,
                'confidence': 0.6,
                'method': 'llm',
                'scores': {}
            }
    except Exception as e:
        logger.warning(f"LLM document type detection failed: {e}")
    
    # Fallback
    return {
        'type': 'document_generique',
        'confidence': 0.1,
        'method': 'fallback',
        'scores': dict(sorted_scores[:5]) if sorted_scores else {}
    }

async def extract_structured_data(file_content: str, file_name: str, document_type: str, 
                                  selected_model: str = DEFAULT_MODEL, language: str = 'fr') -> Dict[str, Any]:
    """
    Extrait les donnÃ©es structurÃ©es d'un document selon son type
    Retourne un dictionnaire JSON structurÃ©
    """
    # Prompts spÃ©cialisÃ©s par type de document
    prompts = {
        'contrat_location': """Tu dois extraire toutes les informations structurÃ©es de ce contrat de location.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Contrat de location",
  "parties": [
    {
      "nom": "nom complet",
      "role": "LOCATAIRE ou BAILLEUR",
      "adresse": "adresse complÃ¨te",
      "telephone": "numÃ©ro si disponible",
      "email": "email si disponible"
    }
  ],
  "dates_importantes": [
    {
      "type": "Date de signature",
      "valeur": "date au format JJ/MM/AAAA"
    },
    {
      "type": "Date de dÃ©but",
      "valeur": "date"
    },
    {
      "type": "Date de fin",
      "valeur": "date si disponible"
    },
    {
      "type": "DurÃ©e",
      "valeur": "durÃ©e (ex: 3 ans)"
    }
  ],
  "montants": [
    {
      "type": "Loyer mensuel",
      "valeur": nombre,
      "devise": "EUR"
    },
    {
      "type": "Caution",
      "valeur": nombre,
      "devise": "EUR"
    },
    {
      "type": "Charges",
      "valeur": nombre si disponible,
      "devise": "EUR"
    }
  ],
  "bien_loue": {
    "adresse": "adresse complÃ¨te du bien",
    "superficie": "superficie si disponible",
    "type": "appartement, maison, etc."
  },
  "clauses_cles": [
    "liste des clauses importantes (renouvellement, indexation, etc.)"
  ],
  "conditions": [
    "conditions particuliÃ¨res mentionnÃ©es"
  ]
}
Extrait TOUTES les informations disponibles. Si une information n'est pas prÃ©sente, utilise null ou omets le champ.""",

        'contrat_travail': """Tu dois extraire toutes les informations structurÃ©es de ce contrat de travail.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Contrat de travail",
  "employeur": {
    "nom": "nom de l'entreprise",
    "adresse": "adresse complÃ¨te",
    "siret": "SIRET si disponible"
  },
  "employe": {
    "nom": "nom complet",
    "adresse": "adresse complÃ¨te",
    "poste": "intitulÃ© du poste",
    "fonctions": "description des fonctions"
  },
  "dates_importantes": [
    {
      "type": "Date de signature",
      "valeur": "date"
    },
    {
      "type": "Date d'embauche",
      "valeur": "date"
    },
    {
      "type": "PÃ©riode d'essai",
      "valeur": "durÃ©e"
    }
  ],
  "remuneration": {
    "salaire_brut_mensuel": nombre,
    "salaire_net_mensuel": nombre si disponible,
    "devise": "EUR",
    "avantages": ["liste des avantages"]
  },
  "duree": {
    "type": "CDI, CDD, etc.",
    "duree": "durÃ©e si CDD"
  },
  "clauses_cles": [
    "clause de non-concurrence si prÃ©sente",
    "clause de confidentialitÃ© si prÃ©sente",
    "autres clauses importantes"
  ]
}""",

        'testament': """Tu dois extraire toutes les informations structurÃ©es de ce testament.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Testament",
  "testateur": {
    "nom": "nom complet",
    "adresse": "adresse complÃ¨te",
    "date_naissance": "date si disponible"
  },
  "beneficiaires": [
    {
      "nom": "nom complet",
      "relation": "relation avec le testateur",
      "legs": "description du legs",
      "conditions": "conditions Ã©ventuelles"
    }
  ],
  "executeur_testamentaire": {
    "nom": "nom si mentionnÃ©",
    "fonctions": "fonctions"
  },
  "biens_legues": [
    {
      "description": "description du bien",
      "beneficiaire": "nom du bÃ©nÃ©ficiaire",
      "valeur_estimee": "valeur si mentionnÃ©e"
    }
  ],
  "dates_importantes": [
    {
      "type": "Date de rÃ©daction",
      "valeur": "date"
    },
    {
      "type": "Date de signature",
      "valeur": "date"
    }
  ],
  "conditions_particulieres": [
    "conditions ou clauses particuliÃ¨res"
  ]
}""",

        'contrat_generique': """Tu dois extraire toutes les informations structurÃ©es de ce contrat.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Contrat",
  "parties": [
    {
      "nom": "nom complet",
      "role": "rÃ´le dans le contrat",
      "adresse": "adresse si disponible"
    }
  ],
  "objet": "objet du contrat en une phrase",
  "dates_importantes": [
    {
      "type": "type de date (signature, Ã©chÃ©ance, etc.)",
      "valeur": "date"
    }
  ],
  "montants": [
    {
      "type": "type de montant",
      "valeur": nombre,
      "devise": "EUR"
    }
  ],
  "clauses_cles": [
    "liste des clauses importantes"
  ],
  "obligations": [
    "obligations de chaque partie"
  ],
  "conditions": [
    "conditions particuliÃ¨res"
  ]
}""",

        'assurance_insurance': """Tu dois extraire toutes les informations structurÃ©es de ce contrat/police d'assurance.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Contrat d'assurance",
  "assureur": {
    "nom": "nom de la compagnie d'assurance",
    "adresse": "adresse complÃ¨te",
    "numero_police": "numÃ©ro de police"
  },
  "assure": {
    "nom": "nom complet de l'assurÃ©",
    "adresse": "adresse complÃ¨te",
    "qualite": "souscripteur, bÃ©nÃ©ficiaire, etc."
  },
  "couverture": {
    "type_assurance": "auto, habitation, vie, santÃ©, etc.",
    "garanties": ["liste des garanties couvertes"],
    "exclusions": ["liste des exclusions"]
  },
  "montants": [
    {"type": "Prime annuelle", "valeur": null, "devise": "EUR"},
    {"type": "Franchise", "valeur": null, "devise": "EUR"},
    {"type": "Plafond d'indemnisation", "valeur": null, "devise": "EUR"}
  ],
  "dates_importantes": [
    {"type": "Date de souscription", "valeur": "date"},
    {"type": "Date d'effet", "valeur": "date"},
    {"type": "Date d'Ã©chÃ©ance", "valeur": "date"},
    {"type": "Date de renouvellement", "valeur": "date"}
  ],
  "conditions_particulieres": ["conditions ou clauses spÃ©cifiques"]
}""",

        'jugement_decision_justice': """Tu dois extraire toutes les informations structurÃ©es de ce jugement/dÃ©cision de justice.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Jugement/DÃ©cision de justice",
  "juridiction": {
    "tribunal": "nom du tribunal",
    "chambre": "chambre si mentionnÃ©e",
    "juge": "nom du/des juge(s)",
    "greffier": "nom du greffier si mentionnÃ©"
  },
  "parties": [
    {"nom": "nom complet", "role": "demandeur/dÃ©fendeur/partie civile", "avocat": "nom de l'avocat si mentionnÃ©"}
  ],
  "dates_importantes": [
    {"type": "Date d'audience", "valeur": "date"},
    {"type": "Date de dÃ©libÃ©rÃ©", "valeur": "date"},
    {"type": "Date de prononcÃ©", "valeur": "date"}
  ],
  "dispositif": {
    "decision": "rÃ©sumÃ© de la dÃ©cision rendue",
    "condamnations": ["liste des condamnations"],
    "montants": [{"type": "type", "valeur": null, "devise": "EUR"}]
  },
  "voies_de_recours": {
    "appel_possible": true,
    "delai": "dÃ©lai d'appel",
    "juridiction_appel": "cour d'appel compÃ©tente"
  },
  "references_juridiques": ["articles de loi citÃ©s"]
}""",

        'releve_bancaire': """Tu dois extraire toutes les informations structurÃ©es de ce relevÃ© bancaire.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "RelevÃ© bancaire",
  "banque": {
    "nom": "nom de la banque",
    "agence": "agence si mentionnÃ©e"
  },
  "titulaire": {
    "nom": "nom complet du titulaire",
    "adresse": "adresse si disponible"
  },
  "compte": {
    "numero": "numÃ©ro de compte",
    "iban": "IBAN si disponible",
    "bic": "BIC si disponible",
    "type": "courant, Ã©pargne, etc."
  },
  "periode": {
    "debut": "date de dÃ©but",
    "fin": "date de fin"
  },
  "soldes": {
    "solde_initial": {"valeur": null, "devise": "EUR"},
    "solde_final": {"valeur": null, "devise": "EUR"},
    "total_debits": {"valeur": null, "devise": "EUR"},
    "total_credits": {"valeur": null, "devise": "EUR"}
  },
  "operations_principales": [
    {"date": "date", "libelle": "description", "montant": null, "type": "dÃ©bit/crÃ©dit"}
  ],
  "frais_bancaires": [{"type": "type de frais", "montant": null}]
}""",

        'certificat_attestation': """Tu dois extraire toutes les informations structurÃ©es de ce certificat/attestation.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Certificat/Attestation",
  "type_certificat": "type prÃ©cis (naissance, mariage, domicile, emploi, scolaritÃ©, etc.)",
  "emetteur": {
    "nom": "organisme ou personne Ã©mettrice",
    "qualite": "qualitÃ©/fonction",
    "adresse": "adresse si disponible"
  },
  "beneficiaire": {
    "nom": "nom complet de la personne concernÃ©e",
    "adresse": "adresse si disponible",
    "date_naissance": "date si disponible"
  },
  "objet": "ce qui est certifiÃ© ou attestÃ©",
  "dates_importantes": [
    {"type": "Date de dÃ©livrance", "valeur": "date"},
    {"type": "Date de validitÃ©", "valeur": "date si disponible"}
  ],
  "elements_authenticite": {
    "numero": "numÃ©ro de certificat si disponible",
    "cachet": true,
    "signature": true
  },
  "informations_complementaires": ["informations supplÃ©mentaires"]
}""",

        'contrat_pret_loan': """Tu dois extraire toutes les informations structurÃ©es de ce contrat de prÃªt/crÃ©dit.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Contrat de prÃªt",
  "type_pret": "immobilier, personnel, auto, professionnel, etc.",
  "preteur": {
    "nom": "nom de l'Ã©tablissement prÃªteur",
    "adresse": "adresse complÃ¨te"
  },
  "emprunteur": {
    "nom": "nom complet de l'emprunteur",
    "adresse": "adresse complÃ¨te"
  },
  "montants": {
    "capital_emprunte": {"valeur": null, "devise": "EUR"},
    "taux_interet": "taux en %",
    "taeg": "TAEG en % si disponible",
    "mensualite": {"valeur": null, "devise": "EUR"},
    "cout_total_credit": {"valeur": null, "devise": "EUR"}
  },
  "dates_importantes": [
    {"type": "Date de signature", "valeur": "date"},
    {"type": "Date de dÃ©but", "valeur": "date"},
    {"type": "Date de fin", "valeur": "date"},
    {"type": "DurÃ©e", "valeur": "durÃ©e en mois/annÃ©es"}
  ],
  "garanties": ["liste des garanties (hypothÃ¨que, caution, etc.)"],
  "conditions_remboursement_anticipe": "conditions si mentionnÃ©es",
  "assurance_emprunteur": "dÃ©tails si mentionnÃ©s"
}""",

        'devis_estimation': """Tu dois extraire toutes les informations structurÃ©es de ce devis/estimation.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Devis/Estimation",
  "emetteur": {
    "nom": "nom de l'entreprise",
    "adresse": "adresse complÃ¨te",
    "siret": "SIRET si disponible"
  },
  "client": {
    "nom": "nom du client",
    "adresse": "adresse si disponible"
  },
  "numero_devis": "numÃ©ro du devis",
  "articles": [
    {"description": "description", "quantite": null, "prix_unitaire": null, "total": null}
  ],
  "montants": {
    "total_ht": {"valeur": null, "devise": "EUR"},
    "tva": {"valeur": null, "taux": "taux en %"},
    "total_ttc": {"valeur": null, "devise": "EUR"},
    "remise": {"valeur": null, "devise": "EUR"}
  },
  "dates_importantes": [
    {"type": "Date d'Ã©mission", "valeur": "date"},
    {"type": "Date de validitÃ©", "valeur": "date"}
  ],
  "conditions": {
    "paiement": "conditions de paiement",
    "livraison": "conditions de livraison",
    "delai": "dÃ©lai d'exÃ©cution"
  }
}""",

        'bon_commande_purchase_order': """Tu dois extraire toutes les informations structurÃ©es de ce bon de commande.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Bon de commande",
  "client": {
    "nom": "nom du client",
    "adresse": "adresse complÃ¨te",
    "reference_client": "rÃ©fÃ©rence si disponible"
  },
  "fournisseur": {
    "nom": "nom du fournisseur",
    "adresse": "adresse complÃ¨te"
  },
  "numero_commande": "numÃ©ro de commande",
  "articles": [
    {"reference": "ref", "description": "description", "quantite": null, "prix_unitaire": null, "total": null}
  ],
  "montants": {
    "total_ht": {"valeur": null, "devise": "EUR"},
    "tva": {"valeur": null},
    "total_ttc": {"valeur": null, "devise": "EUR"}
  },
  "dates_importantes": [
    {"type": "Date de commande", "valeur": "date"},
    {"type": "Date de livraison prÃ©vue", "valeur": "date"}
  ],
  "conditions_livraison": "conditions de livraison",
  "conditions_paiement": "conditions de paiement"
}""",

        'proces_verbal': """Tu dois extraire toutes les informations structurÃ©es de ce procÃ¨s-verbal/compte-rendu.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "ProcÃ¨s-verbal/Compte-rendu",
  "type_reunion": "assemblÃ©e gÃ©nÃ©rale, conseil d'administration, rÃ©union de travail, etc.",
  "lieu": "lieu de la rÃ©union",
  "date_heure": "date et heure de la rÃ©union",
  "participants": {
    "presents": [{"nom": "nom", "qualite": "rÃ´le/fonction"}],
    "absents": [{"nom": "nom", "qualite": "rÃ´le/fonction"}],
    "excuses": [{"nom": "nom"}]
  },
  "president_seance": "nom du prÃ©sident de sÃ©ance",
  "secretaire": "nom du secrÃ©taire",
  "ordre_du_jour": ["liste des points Ã  l'ordre du jour"],
  "decisions": [
    {"sujet": "sujet de la dÃ©cision", "decision": "dÃ©cision prise", "vote": "rÃ©sultat du vote si applicable"}
  ],
  "actions_decidees": [
    {"action": "description", "responsable": "nom", "echeance": "date"}
  ],
  "prochaine_reunion": "date si mentionnÃ©e"
}""",

        'rapport_expertise': """Tu dois extraire toutes les informations structurÃ©es de ce rapport d'expertise/audit.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Rapport d'expertise/Audit",
  "type_expertise": "type prÃ©cis (immobilier, mÃ©dical, technique, financier, etc.)",
  "expert": {
    "nom": "nom de l'expert/auditeur",
    "qualifications": "qualifications/agrÃ©ments",
    "organisme": "organisme si applicable"
  },
  "commanditaire": {
    "nom": "nom du commanditaire",
    "qualite": "qualitÃ©"
  },
  "objet": "objet et pÃ©rimÃ¨tre de l'expertise",
  "dates_importantes": [
    {"type": "Date de la mission", "valeur": "date"},
    {"type": "Date du rapport", "valeur": "date"}
  ],
  "methodologie": "rÃ©sumÃ© de la mÃ©thodologie utilisÃ©e",
  "constats": ["liste des constats principaux"],
  "conclusions": ["liste des conclusions"],
  "recommandations": ["liste des recommandations"],
  "annexes": ["liste des annexes si mentionnÃ©es"]
}""",

        'permis_licence': """Tu dois extraire toutes les informations structurÃ©es de ce permis/licence/autorisation.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "type_document": "Permis/Licence/Autorisation",
  "type_permis": "type prÃ©cis (construire, conduire, exploitation, etc.)",
  "titulaire": {
    "nom": "nom complet du titulaire",
    "adresse": "adresse complÃ¨te"
  },
  "autorite_emettrice": {
    "nom": "organisme Ã©metteur",
    "adresse": "adresse si disponible"
  },
  "numero_permis": "numÃ©ro du permis/licence",
  "objet": "ce qui est autorisÃ© prÃ©cisÃ©ment",
  "dates_importantes": [
    {"type": "Date de dÃ©livrance", "valeur": "date"},
    {"type": "Date d'expiration", "valeur": "date"},
    {"type": "Date de renouvellement", "valeur": "date si applicable"}
  ],
  "conditions": ["liste des conditions et restrictions"],
  "categorie": "catÃ©gorie si applicable",
  "elements_authenticite": {
    "cachet": true,
    "signature": true,
    "numero_enregistrement": "numÃ©ro si disponible"
  }
}""",

        'document_generique': """Tu dois extraire les informations structurÃ©es de ce document.
Retourne UNIQUEMENT un JSON valide avec cette structure:
{
  "type_document": "Type de document dÃ©tectÃ©",
  "parties_mentionnees": [
    {
      "nom": "nom",
      "role": "rÃ´le si identifiable"
    }
  ],
  "dates_importantes": [
    {
      "type": "type de date",
      "valeur": "date"
    }
  ],
  "montants": [
    {
      "type": "type de montant",
      "valeur": nombre,
      "devise": "EUR"
    }
  ],
  "informations_cles": [
    "points clÃ©s du document"
  ]
}"""
    }
    
    # SÃ©lectionner le prompt appropriÃ©
    prompt_template = prompts.get(document_type, prompts['document_generique'])
    
    # Construire le prompt final
    extraction_prompt = f"""{prompt_template}

DOCUMENT Ã€ ANALYSER:
Nom du fichier: {file_name}

Contenu:
{file_content[:8000]}

IMPORTANT: 
- Retourne UNIQUEMENT le JSON, sans texte avant ou aprÃ¨s
- Le JSON DOIT Ãªtre complet et valide (toutes les accolades et crochets doivent Ãªtre fermÃ©s)
- Utilise des valeurs null pour les champs manquants
- Sois prÃ©cis et exhaustif
- Extrais TOUTES les informations disponibles dans le document
- Assure-toi que le JSON est bien formÃ© et complet avant de rÃ©pondre"""
    
    try:
        # Appeler le modÃ¨le pour l'extraction avec une limite de tokens plus Ã©levÃ©e (2000 tokens pour les JSON structurÃ©s)
        # On utilise directement call_openai_api avec max_tokens_override
        raw_response = call_openai_api(extraction_prompt, selected_model, max_retries=3, max_tokens_override=2000)
        
        # Nettoyer la rÃ©ponse : enlever les markdown code blocks si prÃ©sents
        cleaned_response = raw_response.strip()
        if cleaned_response.startswith('```json'):
            cleaned_response = cleaned_response[7:]  # Enlever ```json
        if cleaned_response.startswith('```'):
            cleaned_response = cleaned_response[3:]  # Enlever ```
        if cleaned_response.endswith('```'):
            cleaned_response = cleaned_response[:-3]  # Enlever ``` Ã  la fin
        cleaned_response = cleaned_response.strip()
        
        # Fonction pour rÃ©parer un JSON incomplet
        def repair_json(json_str: str) -> str:
            """Tente de rÃ©parer un JSON incomplet en fermant les structures ouvertes et les chaÃ®nes"""
            json_str = json_str.strip()
            if not json_str.startswith('{'):
                return json_str
            
            result = json_str
            
            # RÃ©parer les chaÃ®nes tronquÃ©es (ex: "devise": "E -> "devise": "EUR")
            # Chercher les guillemets non fermÃ©s Ã  la fin
            # On cherche le dernier guillemet ouvrant qui n'est pas suivi d'un guillemet fermant avant la fin
            # Pattern: trouver "..." qui n'est pas fermÃ©
            # On va simplement fermer les chaÃ®nes ouvertes Ã  la fin
            lines = result.split('\n')
            if lines:
                last_line = lines[-1].strip()
                # Si la derniÃ¨re ligne se termine par un guillemet ouvrant ou une chaÃ®ne incomplÃ¨te
                if last_line.count('"') % 2 != 0:  # Nombre impair de guillemets = chaÃ®ne non fermÃ©e
                    # Trouver le dernier guillemet ouvrant
                    last_quote_pos = last_line.rfind('"')
                    if last_quote_pos >= 0:
                        # VÃ©rifier s'il y a un guillemet fermant aprÃ¨s
                        after_quote = last_line[last_quote_pos + 1:]
                        if '"' not in after_quote:
                            # La chaÃ®ne est tronquÃ©e, on la ferme avec une valeur par dÃ©faut ou null
                            # On va plutÃ´t essayer de complÃ©ter intelligemment
                            # Pour l'instant, on ferme simplement la chaÃ®ne
                            if last_line.endswith('"'):
                                # DÃ©jÃ  un guillemet, on ajoute juste la fermeture
                                pass
                            else:
                                # TronquÃ© au milieu, on ferme avec null
                                # Trouver oÃ¹ commence la valeur
                                if ':' in last_line:
                                    key_part = last_line[:last_line.rfind(':') + 1]
                                    # Remplacer la valeur tronquÃ©e par null
                                    lines[-1] = key_part + ' null'
                                else:
                                    # Juste fermer la chaÃ®ne
                                    lines[-1] = last_line + '"'
                        result = '\n'.join(lines)
            
            # Compter les accolades et crochets ouverts
            open_braces = result.count('{')
            close_braces = result.count('}')
            open_brackets = result.count('[')
            close_brackets = result.count(']')
            
            # Fermer les structures ouvertes
            if open_braces > close_braces:
                # Fermer les objets
                for _ in range(open_braces - close_braces):
                    # Trouver la derniÃ¨re virgule ou le dernier caractÃ¨re et fermer
                    if result.rstrip().endswith(','):
                        result = result.rstrip()[:-1]  # Enlever la virgule
                    # Nettoyer les virgules avant de fermer
                    result = re.sub(r',(\s*)$', r'\1', result)
                    result += '\n}'
            
            if open_brackets > close_brackets:
                # Fermer les tableaux
                for _ in range(open_brackets - close_brackets):
                    if result.rstrip().endswith(','):
                        result = result.rstrip()[:-1]  # Enlever la virgule
                    # Nettoyer les virgules avant de fermer
                    result = re.sub(r',(\s*)$', r'\1', result)
                    result += '\n]'
            
            # Nettoyer les virgules finales avant les fermetures
            result = re.sub(r',(\s*[}\]])', r'\1', result)
            
            # Nettoyer les chaÃ®nes mal fermÃ©es restantes
            # Si on a encore des problÃ¨mes, on remplace les valeurs tronquÃ©es par null
            # Chercher les patterns comme "key": "incomplete_string (sans guillemet fermant)
            result = re.sub(r':\s*"([^"]*?)(?<!")\s*$', r': null', result, flags=re.MULTILINE)
            # Aussi pour les derniÃ¨res lignes qui se terminent par une chaÃ®ne incomplÃ¨te
            lines = result.split('\n')
            if lines:
                last_line = lines[-1].strip()
                # Si la ligne se termine par un guillemet ouvrant sans fermeture
                if last_line.count('"') % 2 != 0 and ':' in last_line:
                    # Extraire la clÃ© et remplacer la valeur par null
                    key_match = re.search(r'(\s*"[^"]+"\s*):\s*"[^"]*$', last_line)
                    if key_match:
                        lines[-1] = key_match.group(1) + ': null'
                    result = '\n'.join(lines)
            
            return result
        
        # Extraire le JSON de la rÃ©ponse
        json_match = re.search(r'\{.*', cleaned_response, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            
            # Essayer de parser directement
            try:
                structured_data = json.loads(json_str)
            except json.JSONDecodeError:
                # Essayer de rÃ©parer le JSON
                logger.warning("âš ï¸ JSON incomplet dÃ©tectÃ©, tentative de rÃ©paration...")
                repaired_json = repair_json(json_str)
                try:
                    structured_data = json.loads(repaired_json)
                    logger.info("âœ… JSON rÃ©parÃ© avec succÃ¨s")
                except json.JSONDecodeError as repair_error:
                    logger.error(f"âŒ Impossible de rÃ©parer le JSON: {str(repair_error)}")
                    # Essayer d'extraire au moins les parties valides
                    try:
                        # Extraire juste le dÃ©but du JSON jusqu'Ã  la premiÃ¨re erreur
                        # Trouver la position de l'erreur
                        error_pos = int(str(repair_error).split('char ')[1].split(')')[0]) if 'char' in str(repair_error) else len(json_str)
                        partial_json = json_str[:error_pos]
                        # Fermer proprement
                        partial_json = repair_json(partial_json)
                        structured_data = json.loads(partial_json)
                        logger.warning("âš ï¸ JSON partiel extrait (certaines donnÃ©es peuvent Ãªtre manquantes)")
                    except:
                        raise repair_error
        else:
            # Essayer de parser directement
            structured_data = json.loads(cleaned_response)
        
        # Ajouter des mÃ©tadonnÃ©es
        structured_data['metadata'] = {
            'file_name': file_name,
            'document_type': document_type,
            'extraction_date': datetime.utcnow().isoformat()
        }
        
        logger.info(f"âœ… Extraction structurÃ©e rÃ©ussie pour {file_name} (type: {document_type})")
        return structured_data
        
    except json.JSONDecodeError as e:
        logger.error(f"âŒ Erreur de parsing JSON lors de l'extraction: {str(e)}")
        logger.error(f"RÃ©ponse brute (premiers 2000 caractÃ¨res): {raw_response[:2000]}")
        logger.error(f"Longueur totale de la rÃ©ponse: {len(raw_response)} caractÃ¨res")
        # Retourner une structure minimale avec l'erreur
        return {
            "error": "Erreur lors de l'extraction des donnÃ©es structurÃ©es - JSON invalide ou incomplet",
            "error_details": str(e),
            "raw_response_preview": raw_response[:2000],
            "response_length": len(raw_response),
            "metadata": {
                "file_name": file_name,
                "document_type": document_type
            }
        }
    except Exception as e:
        logger.error(f"âŒ Erreur lors de l'extraction structurÃ©e: {str(e)}")
        return {
            "error": f"Erreur lors de l'extraction: {str(e)}",
            "metadata": {
                "file_name": file_name,
                "document_type": document_type
            }
        }
    
async def analyze_query_need_for_search(user_query: str, selected_model: str, language: str = 'en') -> Dict:
    """
    Analyse si la query nÃ©cessite des informations actuelles
    AmÃ©liorÃ© pour dÃ©tecter les questions nÃ©cessitant des informations en temps rÃ©el
    """
    t = translations.get(language, translations['en'])
    
    # DÃ©tection rapide basÃ©e sur des mots-clÃ©s pour les questions en temps rÃ©el
    query_lower = user_query.lower()
    realtime_keywords = [
        "heure actuelle", "heure maintenant", "quelle heure", "what time", "current time", "time now",
        "prix actuel", "prix maintenant", "current price", "price now",
        "taux actuel", "current rate", "taux de change", "exchange rate",
        "cours actuel", "current rate",
        "maintenant", "now", "actuel", "current", "aujourd'hui", "today",
        "en ce moment", "right now", "Ã  l'instant", "at the moment",
        "mÃ©tÃ©o", "weather", "tempÃ©rature", "temperature",
        "actualitÃ©", "news", "Ã©vÃ©nements rÃ©cents", "recent events"
    ]
    
    # Si la question contient des mots-clÃ©s de temps rÃ©el, dÃ©clencher directement la recherche
    needs_realtime = any(keyword in query_lower for keyword in realtime_keywords)
    
    if needs_realtime:
        logger.info(f"ðŸ” Question en temps rÃ©el dÃ©tectÃ©e: {user_query[:50]}...")
        return {
            "needs_search": True,
            "reason": "Question nÃ©cessitant des informations en temps rÃ©el",
            "search_keywords": [user_query],
            "estimated_cutoff_relevance": "high"
        }
    
    needs_search_field = t['needs_search_field']
    reason_field = t['reason_field']
    search_keywords_field = t['search_keywords_field']
    cutoff_relevance_field = t['cutoff_relevance_field']
    
    analysis_prompt = (
        f"{t['analyze_query_prompt']}\n\n"
        f"{t['query_label']}: {user_query}\n\n"
        f"{t['json_response_format']}:\n"
        f"{{\n"
        f'  "{needs_search_field}": true/false,\n'
        f'  "{reason_field}": "{t["short_explanation"]}",\n'
        f'  "{search_keywords_field}": ["mot-clÃ©1", "mot-clÃ©2"] {t["keyword_or_null"]},\n'
        f'  "{cutoff_relevance_field}": "high/medium/low"\n'
        f"}}\n\n"
        f"{t['examples_need_search']}:\n"
        f"{t['current_prices_stocks']}\n"
        f"{t['recent_events_news']}\n"
        f"{t['new_software_versions']}\n"
        f"{t['recent_statistics_data']}\n"
        f"{t['recent_people_companies']}\n"
        f"- Questions sur l'heure actuelle dans un pays ou une ville\n"
        f"- Questions sur les prix actuels, taux de change, cours boursiers\n"
        f"- Questions sur la mÃ©tÃ©o actuelle\n"
        f"- Questions sur les Ã©vÃ©nements rÃ©cents ou l'actualitÃ©\n\n"
        f"{t['examples_no_search']}:\n"
        f"{t['general_concepts']}\n"
        f"{t['programming_syntax']}\n"
        f"{t['history_facts']}\n"
        f"{t['math_science']}"
    )
    
    try:
        analysis_result = await execute_model_query(analysis_prompt, selected_model)
        # Nettoyer la rÃ©ponse pour extraire le JSON
        json_match = re.search(r'\{.*\}', analysis_result, re.DOTALL)
        if json_match:
            analysis_data = json.loads(json_match.group())
            return analysis_data
        else:
            # Fallback si le JSON n'est pas parsable
            return {
                "needs_search": True,
                "reason": "Analyse automatique non concluante",
                "search_keywords": [user_query],
                "estimated_cutoff_relevance": "medium"
            }
    except Exception as e:
        logger.warning(f"Erreur lors de l'analyse de la query: {str(e)}")
        return {
            "needs_search": True,
            "reason": "Erreur d'analyse, recherche par sÃ©curitÃ©",
            "search_keywords": [user_query],
            "estimated_cutoff_relevance": "medium"
        }
    
async def query_model_online_mode(user_query: str, selected_model: str = DEFAULT_MODEL, 
                                 language: str = 'en', enable_auto_online_search: bool = True,
                                 conversation_history: list = None) -> str:
    """
    Mode ONLINE: RÃ©ponse du modÃ¨le enrichie avec des donnÃ©es actuelles si nÃ©cessaire (comme SearchGPT)
    Avec recherche automatique si l'information n'est pas trouvÃ©e dans les connaissances du modÃ¨le
    """
    cache_key = f"online_query_{user_query}_{selected_model}_{language}"
    cached = await cache.get(cache_key)
    if cached:
        return cached

    t = translations.get(language, translations['en'])

    # Ã‰TAPE 1: RÃ©ponse initiale du modÃ¨le avec ses connaissances
    initial_prompt = (
        f"{t['online_mode_title']}\n"
        f"{t['recent_info_mention']}\n\n"
        f"{t['question']}: {user_query}\n\n"
        f"{t['online_instructions_title']}\n"
        f"{t['give_best_answer']}\n"
        f"{t['be_precise_dates']}\n"
        f"{t['mention_recent_useful']}\n\n"
        f"{t['emoji']}"
    )

    try:
        # RÃ©ponse initiale du modÃ¨le avec historique de conversation
        initial_response = await execute_model_query(initial_prompt, selected_model, conversation_history)
        
        # Ã‰TAPE 2: Analyse automatique du besoin de recherche
        search_analysis = await analyze_query_need_for_search(user_query, selected_model, language)
        
        logger.info(f"ðŸ” Analyse de recherche: {search_analysis}")
        
        # Ã‰TAPE 3: Recherche et enrichissement si nÃ©cessaire
        should_search = search_analysis.get("needs_search", False)
        
        # NOUVEAU: VÃ©rifier aussi si la rÃ©ponse indique que l'information n'est pas disponible
        if enable_auto_online_search and not should_search:
            missing_info = await detect_missing_information(initial_response, user_query, language)
            if missing_info:
                logger.info("ðŸ” Information absente dÃ©tectÃ©e dans les connaissances du modÃ¨le. Recherche en ligne automatique...")
                should_search = True
        
        if should_search:
            logger.info("ðŸŒ Enrichissement avec des donnÃ©es actuelles...")
            
            # Utiliser les mots-clÃ©s optimisÃ©s ou la query originale
            search_keywords = search_analysis.get("search_keywords", [user_query])
            search_query = " ".join(search_keywords) if isinstance(search_keywords, list) else user_query
            
            # Recherche en ligne
            try:
                search_results = perform_online_search(search_query, language)
                
                if search_results and search_results != "Aucun rÃ©sultat trouvÃ©." and search_results != "ClÃ© API manquante":
                    # Ã‰TAPE 4: Fusion intelligente des informations avec formatage amÃ©liorÃ©
                    enrichment_prompt = (
                        f"{t['enrichment_title']}\n\n"
                        f"{t['initial_response']}\n{initial_response}\n\n"
                        f"{t['new_info_found']}\n{search_results}\n\n"
                        f"{t['enrichment_instructions']}\n"
                        f"{t['combine_intelligently']}\n"
                        f"{t['update_obsolete']}\n"
                        f"{t['distinguish_info']}\n"
                        f"{t['cite_sources']}\n"
                        f"{t['keep_structure']}\n"
                        f"{t['prioritize_recent']}\n\n"
                        f"**FORMATAGE SPÃ‰CIAL POUR CERTAINES RÃ‰PONSES:**\n"
                        f"- Pour les questions sur l'heure: RÃ©ponds directement avec l'heure actuelle au format 'Il est XXhXX Ã  [lieu]' ou 'It is XX:XX in [place]'\n"
                        f"- Pour les prix (cryptomonnaies, actions, devises): Formate comme une carte avec le prix en gras, le changement en pourcentage, et une description claire\n"
                        f"- Pour les taux de change: Affiche le taux actuel de maniÃ¨re claire et structurÃ©e\n"
                        f"- Pour la mÃ©tÃ©o: Formate avec tempÃ©rature, conditions, et lieu\n"
                        f"- Utilise des emojis appropriÃ©s (ðŸ• pour l'heure, ðŸ’° pour les prix, ðŸŒ¤ï¸ pour la mÃ©tÃ©o, etc.)\n"
                        f"- Sois concis et prÃ©cis, va droit au but\n"
                        f"- Si les donnÃ©es de recherche sont claires, utilise-les directement sans trop d'explications\n\n"
                        f"{t['enriched_response']}"
                    )
                    
                    enriched_response, _ = await execute_model_query_with_fallback(enrichment_prompt, selected_model, user_query)
                    
                    # Ajout des mÃ©tadonnÃ©es de recherche
                    final_response = (
                        f"{enriched_response}\n\n"
                        f"---\n"
                        f"ðŸ’¡ **Informations enrichies**: Cette rÃ©ponse combine mes connaissances de base "
                        f"avec des donnÃ©es rÃ©centes trouvÃ©es en ligne.\n \n"
                        f"ðŸ” **Model**: {selected_model}"
                    )
                else:
                    # Pas de rÃ©sultats en ligne, garder la rÃ©ponse initiale
                    logger.info("âš ï¸ Aucun rÃ©sultat trouvÃ© en ligne ou clÃ© API manquante")
                    final_response = (
                        f"{initial_response}\n\n"
                        f"{t['separator']}\n"
                        f"{t['source_training']}"
                    )
            except Exception as search_error:
                logger.warning(f"âš ï¸ Erreur lors de la recherche en ligne: {str(search_error)}")
                # En cas d'erreur, garder la rÃ©ponse initiale
                final_response = (
                    f"{initial_response}\n\n"
                    f"{t['separator']}\n"
                    f"{t['source_training']}"
                )
            
        else:
            logger.info("âœ… Pas de recherche nÃ©cessaire, rÃ©ponse basÃ©e sur les connaissances du modÃ¨le")
            final_response = (
                f"{initial_response}\n\n"
                f"{t['separator']}\n"
                f"{t['source_training']}"
            )

    except Exception as e:
        logger.error(f"Erreur lors du mode online: {str(e)}")
        final_response = f"Erreur lors du traitement en mode online: {str(e)}"

    await cache.set(cache_key, final_response, ttl=1800)
    return final_response

@app.route('/index-directory', methods=['POST'])
async def index_directory():
    """
    Indexe les fichiers d'un rÃ©pertoire avec Docling + OpenAI embeddings.
    Accepte multipart/form-data: champ 'files' (multiple), champ 'language'.
    """
    try:
        session_id = request.headers.get('Session-ID', 'default')
        user_email = _extract_user_email_from_auth_header()
        language = request.form.get('language', 'en')

        uploaded_files = request.files.getlist('files')
        if not uploaded_files:
            return jsonify({
                "error": "Aucun fichier fourni pour l'indexation",
                'files_received': 0
            }), 400

        logger.info(f"ðŸ“‚ Indexation Docling de {len(uploaded_files)} fichier(s) pour la session {session_id}")
        documents = []
        files_indexed_names = []

        for uploaded_file in uploaded_files:
            filename = uploaded_file.filename or 'document'
            ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
            raw_bytes = uploaded_file.read()

            if not raw_bytes:
                logger.warning(f"Fichier ignorÃ© (vide): {filename}")
                continue

            # â€” Docling parse + chunk avec mÃ©tadonnÃ©es riches â€”
            parse_result = parse_with_docling(raw_bytes, ext, filename)

            if parse_result:
                chunks_with_meta, full_text, docling_meta = parse_result
                num_pages = docling_meta.get('num_pages', 0)
                logger.info(f"âœ… Docling: {filename} â†’ {num_pages} page(s), {len(chunks_with_meta)} chunk(s)")

                for chunk in chunks_with_meta:
                    if not chunk.get('text', '').strip():
                        continue
                    doc_metadata = {
                        'fileName': filename,
                        'page_number': chunk.get('page'),
                        'section': chunk.get('section', ''),
                        'chunk_id': chunk.get('chunk_id', ''),
                        'element_type': chunk.get('element_type', 'text'),
                        'headings': chunk.get('headings', []),
                        'indexed_at': datetime.utcnow().isoformat(),
                    }
                    documents.append(Document(
                        page_content=chunk['text'],
                        metadata=doc_metadata
                    ))

                files_indexed_names.append(filename)
                aws_persistence_service.persist_document(
                    file_name=filename,
                    content=full_text,
                    metadata={'num_pages': num_pages, 'parser': 'docling'},
                    user_email=user_email,
                    session_id=session_id,
                    source='index-directory',
                )
            else:
                # â€” Fallback: extraction basique si Docling indisponible â€”
                logger.warning(f"âš ï¸ Docling indisponible pour {filename}, fallback extraction basique")
                try:
                    if ext == 'pdf':
                        fallback_text = extract_text_from_pdf(raw_bytes)
                    elif ext == 'docx':
                        from io import BytesIO as _BytesIO
                        fallback_text = extract_text_from_docx(type('_F', (), {'read': lambda self: raw_bytes})())
                    else:
                        fallback_text = raw_bytes.decode('utf-8', errors='replace')

                    if fallback_text and fallback_text.strip():
                        documents.append(Document(
                            page_content=fallback_text,
                            metadata={
                                'fileName': filename,
                                'indexed_at': datetime.utcnow().isoformat(),
                                'parser': 'fallback',
                            }
                        ))
                        files_indexed_names.append(filename)
                except Exception as fb_err:
                    logger.error(f"âŒ Fallback extraction Ã©chouÃ©e pour {filename}: {fb_err}")

        if not documents:
            return jsonify({
                "error": "Aucun fichier valide Ã  indexer",
                'files_processed': 0
            }), 400

        logger.info(f"ðŸ—‚ï¸ {len(documents)} chunks Docling prÃªts pour l'embedding")

        # â€” Embeddings OpenAI â€”
        openai_api_key = os.getenv('OPENAI_API_KEY')
        if not openai_api_key:
            return jsonify({
                'error': 'ClÃ© API OpenAI non configurÃ©e sur le serveur',
                'suggestion': 'Contactez l\'administrateur pour configurer OPENAI_API_KEY'
            }), 500

        try:
            embeddings_model = OpenAIEmbeddings(
                api_key=openai_api_key,
                model="text-embedding-3-large"
            )
            logger.info("âœ… Utilisation de text-embedding-3-large")
        except Exception:
            embeddings_model = OpenAIEmbeddings(
                api_key=openai_api_key,
                model="text-embedding-ada-002"
            )

        try:
            test_embedding = await embeddings_model.aembed_query("test")
            logger.info(f"âœ… API OpenAI fonctionnelle, dimension: {len(test_embedding)}")
        except Exception as e:
            logger.error(f"âŒ Erreur de test API OpenAI: {str(e)}")
            return jsonify({
                'error': f'Erreur API OpenAI: {str(e)}',
                'suggestion': 'VÃ©rifiez votre clÃ© API et votre quota'
            }), 500

        # â€” InfÃ©rence des actions suggÃ©rÃ©es â€”
        logger.info("ðŸ§  InfÃ©rence des actions suggÃ©rÃ©es pour le corpus...")
        inferred_actions = await infer_corpus_actions(documents, language=language)

        # â€” CrÃ©ation du vector store FAISS â€”
        logger.info("ðŸ—ƒï¸ CrÃ©ation du vector store FAISS...")
        vector_store_obj = await FAISS.afrom_documents(documents, embeddings_model)

        vector_stores[session_id] = {
            'store': vector_store_obj,
            'created_at': datetime.utcnow().isoformat(),
            'files_count': len(files_indexed_names),
            'chunks_count': len(documents),
            'files_indexed': files_indexed_names,
            'auto_actions': inferred_actions
        }

        logger.info(
            f"âœ… Indexation Docling terminÃ©e â€” session {session_id}: "
            f"{len(files_indexed_names)} fichier(s), {len(documents)} chunks"
        )

        return jsonify({
            'success': True,
            'session_id': session_id,
            'indexed_files_count': len(files_indexed_names),
            'chunks_count': len(documents),
            'files_indexed': files_indexed_names,
            'vector_store_ready': True,
            'suggested_actions': inferred_actions.get('suggested_actions', []),
            'corpus_domain': inferred_actions.get('domain', 'unknown'),
            'detected_type': inferred_actions.get('detected_type', 'document_generique'),
            'detected_type_label': inferred_actions.get('detected_type_label', 'Document'),
            'detected_type_confidence': inferred_actions.get('detected_type_confidence', 0),
            'type_distribution': inferred_actions.get('type_distribution', {}),
        }), 200

    except Exception as e:
        logger.error(f"âŒ Erreur lors de l'indexation: {str(e)}")
        return jsonify({
            'error': f'Erreur lors de l\'indexation: {str(e)}',
            'success': False
        }), 500


@app.route('/infer-corpus-actions', methods=['POST'])
async def infer_corpus_actions_endpoint():
    """
    Endpoint pour infÃ©rer les actions suggÃ©rÃ©es Ã  partir de documents
    """
    try:
        data = request.get_json()
        documents_data = data.get('documents', [])
        language = data.get('language', 'en')
        
        if not documents_data:
            return jsonify({
                "error": "No documents provided",
                'success': False
            }), 400
        
        # Convertir les documents en format Document
        documents = []
        for doc_data in documents_data:
            documents.append(Document(
                page_content=doc_data.get('content', ''),
                metadata={
                    'fileName': doc_data.get('file_name', 'document'),
                    'source': doc_data.get('file_name', 'document')
                }
            ))
        
        # InfÃ©rer les actions suggÃ©rÃ©es
        inferred_actions = await infer_corpus_actions(documents, language=language)
        
        return jsonify({
            'success': True,
            'suggested_actions': inferred_actions.get('suggested_actions', []),
            'domain': inferred_actions.get('domain', 'unknown'),
            'detected_type': inferred_actions.get('detected_type', 'document_generique'),
            'detected_type_label': inferred_actions.get('detected_type_label', 'Document'),
            'detected_type_confidence': inferred_actions.get('detected_type_confidence', 0),
        }), 200
        
    except Exception as e:
        logger.error(f"âŒ Erreur lors de l'infÃ©rence des actions: {str(e)}")
        return jsonify({
            'error': f'Error inferring actions: {str(e)}',
            'success': False
        }), 500


@app.route('/query', methods=['POST'])
async def handle_query():
    """
    Route de requÃªte amÃ©liorÃ©e avec support des modes local et online
    """
    data = request.get_json()
    
    # Extraction des paramÃ¨tres
    user_query = data.get('user_query')
    research_mode = data.get('research_mode', 'local')
    selected_model = data.get('selected_model', DEFAULT_MODEL)
    language = data.get('language', 'en')
    session_id = request.headers.get('Session-ID', 'default')  # NOUVEAU: rÃ©cupÃ©rer session ID
    conversation_history = data.get('conversation_history', [])
    
    # ParamÃ¨tres de contrÃ´le des modes
    disable_online_search = data.get('disable_online_search', False)
    enable_online_search = data.get('enable_online_search', False)
    use_backend_vectorstore = data.get('use_backend_vectorstore', False)  # NOUVEAU
    
    # Validation des paramÃ¨tres essentiels
    if not user_query:
        return jsonify({"error": "user_query est requis"}), 400
    
    # Validation et nettoyage du modÃ¨le
    if selected_model not in MODEL_CONFIG:
        logger.warning(f"ModÃ¨le inconnu {selected_model}, utilisation du modÃ¨le par dÃ©faut {DEFAULT_MODEL}")
        selected_model = DEFAULT_MODEL
    
    try:
        # DÃ©termination du mode effectif
        effective_mode = research_mode
        
        # MODE LOCAL
        if effective_mode == 'local' or disable_online_search:
            logger.info(f"ðŸ”’ Mode LOCAL activÃ© pour la requÃªte: {user_query[:50]}...")
            
            # Validation des paramÃ¨tres pour le mode local
            file_name = data.get('file_name')
            file_content = data.get('file_content', '')
            directory_content = data.get('directory_content', [])
            repo_structure = data.get('repo_structure', '')
            is_binary = data.get('is_binary', False)
            
            # âœ… Directory/corpus mode (rÃ©pertoire importÃ©)
            # Si un vector store existe pour cette session, autoriser les requÃªtes mÃªme sans file_name
            # Si use_backend_vectorstore est True OU si un vector store existe pour la session, on peut procÃ©der
            has_vector_store = session_id in vector_stores
            should_use_vectorstore = use_backend_vectorstore or has_vector_store
            
            if not file_name:
                if should_use_vectorstore and has_vector_store:
                    file_name = "__DIRECTORY_CORPUS__"
                    file_content = ""
                    # on laisse directory_content se remplir depuis le vector store ci-dessous
                    is_binary = False
                    use_backend_vectorstore = True  # Forcer Ã  True puisque le vector store existe
                    logger.info(f"ðŸ“š Mode CORPUS (rÃ©pertoire) activÃ©: query sur l'ensemble des documents, session={session_id}")
                else:
                    return jsonify({
                        "error": "Mode local nÃ©cessite un fichier sÃ©lectionnÃ© ou un rÃ©pertoire indexÃ©",
                        "mode": "local",
                        "session_id": session_id,
                        "has_vector_store": has_vector_store,
                        "use_backend_vectorstore": use_backend_vectorstore,
                        "suggestion": "SÃ©lectionnez un fichier, ou importez un rÃ©pertoire pour l'indexer"
                    }), 400
            
            # VÃ©rifier si le fichier demandÃ© dÃ©passe 100 pages (il ne sera pas dans le vector store)
            # Si un fichier spÃ©cifique est demandÃ© et qu'il n'est pas dans le vector store, 
            # c'est probablement qu'il dÃ©passe 100 pages
            if file_name and file_name != "__DIRECTORY_CORPUS__" and should_use_vectorstore and has_vector_store:
                session_store = vector_stores.get(session_id)
                if session_store and isinstance(session_store, dict) and 'store' in session_store:
                    # VÃ©rifier si le fichier est dans les fichiers indexÃ©s
                    files_indexed = session_store.get('files_indexed', [])
                    if files_indexed and file_name not in files_indexed:
                        logger.warning(f"âš ï¸ Fichier {file_name} non trouvÃ© dans les fichiers indexÃ©s (probablement >100 pages). La requÃªte sera ignorÃ©e.")
                        return jsonify({
                            "error": f"Le fichier '{file_name}' dÃ©passe la limite de 100 pages et ne peut pas Ãªtre utilisÃ© pour les questions. Veuillez sÃ©lectionner un autre fichier.",
                            "mode": "local",
                            "file_name": file_name,
                            "suggestion": "SÃ©lectionnez un autre fichier du rÃ©pertoire qui ne dÃ©passe pas 100 pages"
                        }), 400
            
            # NOUVELLE LOGIQUE: Recherche sÃ©mantique amÃ©liorÃ©e avec recherche hybride (sÃ©mantique + mots-clÃ©s)
            relevant_docs: List[Document] = []
            # Utiliser le vector store si disponible (que use_backend_vectorstore soit True ou False, si le store existe on l'utilise)
            if should_use_vectorstore and has_vector_store:
                try:
                    session_store = vector_stores.get(session_id)
                    if not session_store or not isinstance(session_store, dict) or 'store' not in session_store or session_store.get('store') is None:
                        logger.warning(f"Vector store non disponible pour session {session_id}")
                    else:
                        session_vector_store = session_store['store']
                        logger.info(f"ðŸ” Hybrid retrieval (semantic + lexical) dans le vector store session {session_id}")

                        # Hierarchical retrieval strategy:
                        # 1) If a file is selected: search in that file first, then in whole repo if needed
                        # 2) If no file selected: search directly in whole repo until meaningful answer
                        file_first_docs: List[Document] = []
                        file_first_debug: Dict[str, Any] = {}
                        corpus_docs: List[Document] = []
                        corpus_debug: Dict[str, Any] = {}
                        
                        if file_name and file_name != "__DIRECTORY_CORPUS__":
                            # STEP 1: Search in selected file first
                            logger.info(f"ðŸ“„ Recherche dans le fichier sÃ©lectionnÃ© '{file_name}'...")
                            file_first_docs, file_first_debug = hybrid_retrieve_documents(
                                vector_store=session_vector_store,
                                query=user_query,
                                k_candidates=80,  # Increased for better recall within file
                                k_final=10,  # Increased from 8 to 10
                                semantic_weight=0.55,
                                bm25_weight=0.35,
                                exact_weight=0.10,
                                preferred_sources=[file_name],
                            )
                            logger.info(f"ðŸ“Š Fichier: {len(file_first_docs)} rÃ©sultats trouvÃ©s")
                            
                            # STEP 2: Toujours chercher aussi dans le corpus pour complÃ©ter
                            # Si l'info n'est pas dans le fichier, elle sera dans le corpus
                            logger.info(f"ðŸ“š Recherche complÃ©mentaire dans le corpus du rÃ©pertoire...")
                            corpus_result = hybrid_retrieve_documents(
                                vector_store=session_vector_store,
                                query=user_query,
                                k_candidates=120,  # Increased for better recall
                                k_final=15,  # Increased from 12 to 15
                                semantic_weight=0.55,
                                bm25_weight=0.35,
                                exact_weight=0.10,
                            )
                            if isinstance(corpus_result, tuple) and len(corpus_result) == 2:
                                corpus_docs, corpus_debug = corpus_result
                                if not isinstance(corpus_debug, dict):
                                    corpus_debug = {}
                            else:
                                corpus_docs = []
                                corpus_debug = {}
                            logger.info(f"ðŸ“š Corpus: {len(corpus_docs)} rÃ©sultats trouvÃ©s")
                        else:
                            # No file selected: use the advanced search_semantic_documents_sync
                            # which includes person name filtering, function name search, filename matching, etc.
                            logger.info(f"ðŸ“š No file selected: Using advanced search in whole repository...")
                            directory_content = search_semantic_documents_sync(
                                session_vector_store, 
                                user_query, 
                                session_id, 
                                conversation_history
                            )
                            # Convert List[Dict] to List[Document] for consistency with file_first logic
                            corpus_docs = []
                            for doc_dict in directory_content:
                                # Create a Document object from the dict
                                from langchain_core.documents import Document
                                doc = Document(
                                    page_content=doc_dict.get('content', ''),
                                    metadata={
                                        'fileName': doc_dict.get('fileName', ''),
                                        'file_name': doc_dict.get('fileName', ''),
                                        'pageNumber': doc_dict.get('pageNumber'),
                                        'is_page_chunk': doc_dict.get('isPageChunk', False)
                                    }
                                )
                                corpus_docs.append(doc)
                            corpus_debug = {}
                            logger.info(f"ðŸ“š Advanced repository search: {len(corpus_docs)} docs found")

                        # Merge preserving order: file_first_docs first (if any), then corpus_docs
                        # Dedupe by (source, preview)
                        merged: List[Document] = []
                        seen = set()
                        
                        # Add file documents first (priority)
                        for d in file_first_docs:
                            src = (d.metadata.get("source") or d.metadata.get("fileName") or d.metadata.get("file_name") or "").strip().lower()
                            key = (src, (d.page_content or "")[:120])
                            if key not in seen:
                                seen.add(key)
                                merged.append(d)
                        
                        # Then add corpus documents (excluding duplicates)
                        for d in corpus_docs:
                            src = (d.metadata.get("source") or d.metadata.get("fileName") or d.metadata.get("file_name") or "").strip().lower()
                            key = (src, (d.page_content or "")[:120])
                            if key not in seen:
                                seen.add(key)
                                merged.append(d)
                                # Limit total results
                                if len(merged) >= 20:
                                    break

                        relevant_docs = merged
                        logger.info(
                            f"ðŸŽ¯ Hierarchical search completed: {len(relevant_docs)} total docs | "
                            f"from_file={len(file_first_docs)} | "
                            f"from_corpus={len([d for d in merged if d not in file_first_docs])}"
                        )
                        
                        # Ajout des documents pertinents au contexte avec mÃ©tadonnÃ©es enrichies
                        # Ã‰viter les doublons en utilisant un set de file_name + dÃ©but du contenu
                        if not directory_content:  # Only populate if not already filled by search_semantic_documents_sync
                            seen_docs = set()
                            for doc in relevant_docs:
                                file_name_from_meta = doc.metadata.get("fileName") or doc.metadata.get("file_name") or "document_vectorstore"
                                content_preview = doc.page_content[:100]  # Premier 100 chars pour dÃ©tecter les doublons
                                doc_key = f"{file_name_from_meta}:{content_preview}"
                                
                                if doc_key not in seen_docs:
                                    seen_docs.add(doc_key)
                                    page_number = doc.metadata.get("page") or doc.metadata.get("page_number")
                                    chunk_id = doc.metadata.get("chunk_id")
                                    section = doc.metadata.get("section")
                                    result_dict = {
                                        "fileName": file_name_from_meta,
                                        "content": doc.page_content[:2500]
                                    }
                                    if page_number is not None:
                                        result_dict["pageNumber"] = page_number
                                    if chunk_id is not None:
                                        result_dict["chunkId"] = chunk_id
                                    if section:
                                        result_dict["section"] = section
                                    directory_content.append(result_dict)
                            
                            logger.info(f"ðŸ“š {len(directory_content)} documents uniques ajoutÃ©s depuis le vector store de session")
                        else:
                            logger.info(f"ðŸ“š {len(directory_content)} documents ajoutÃ©s via search_semantic_documents_sync")
                    
                except Exception as e:
                    logger.warning(f"Erreur lors de la rÃ©cupÃ©ration depuis le vector store: {str(e)}")
            
            elif use_backend_vectorstore:
                logger.warning(f"Vector store demandÃ© mais non trouvÃ© pour la session {session_id}")
            
            # ANCIENNE LOGIQUE: Utiliser le vector store global (fallback)
            elif vector_store and user_query:
                try:
                    logger.info(f"ðŸ” Recherche sÃ©mantique amÃ©liorÃ©e dans le vector store global")
                    search_results_with_scores = vector_store.similarity_search_with_score(
                        user_query, 
                        k=15
                    )
                    
                    filtered_docs = []
                    for doc, distance in search_results_with_scores:
                        similarity_score = 1 - distance
                        if similarity_score >= 0.65:
                            filtered_docs.append((doc, similarity_score))
                    
                    filtered_docs.sort(key=lambda x: x[1], reverse=True)
                    top_docs = [doc for doc, score in filtered_docs[:10]]
                    
                    if len(top_docs) > 5:
                        try:
                            reranked_docs = await rerank_documents_with_llm(
                                user_query, 
                                top_docs, 
                                selected_model
                            )
                            relevant_docs = reranked_docs[:8]
                        except Exception as rerank_error:
                            logger.warning(f"Re-ranking Ã©chouÃ©: {str(rerank_error)}")
                            relevant_docs = top_docs[:8]
                    else:
                        relevant_docs = top_docs
                    
                    for doc in relevant_docs:
                        result_dict = {
                            "fileName": doc.metadata.get("file_name", "document_vectorstore"),
                            "content": doc.page_content[:2000]
                        }
                        page_number = doc.metadata.get("page") or doc.metadata.get("page_number")
                        chunk_id = doc.metadata.get("chunk_id")
                        section = doc.metadata.get("section")
                        if page_number is not None:
                            result_dict["pageNumber"] = page_number
                        if chunk_id is not None:
                            result_dict["chunkId"] = chunk_id
                        if section:
                            result_dict["section"] = section
                        directory_content.append(result_dict)
                    logger.info(f"ðŸ“š {len(relevant_docs)} documents ajoutÃ©s depuis le vector store global")
                except Exception as e:
                    logger.warning(f"Erreur lors de la rÃ©cupÃ©ration depuis le vector store global: {str(e)}")
            
            # ExÃ©cution de la requÃªte en mode local - recherche en ligne DÃ‰SACTIVÃ‰E par dÃ©faut
            enable_auto_search = data.get('enable_auto_online_search', False)  # DÃ©sactivÃ© par dÃ©faut en mode local
            response, actual_model_used = await query_model_local_mode(
                file_name=file_name,
                file_content=file_content,
                directory_content=directory_content,
                repo_structure=repo_structure,
                user_query=user_query,
                is_binary=is_binary,
                selected_model=selected_model,
                language=language,
                conversation_history=conversation_history,
                enable_auto_online_search=enable_auto_search
            )
            
            # Extraire les informations de page des documents utilisÃ©s
            pages_used = []
            for doc in directory_content:
                if isinstance(doc, dict) and doc.get('pageNumber') is not None:
                    pages_used.append({
                        "fileName": doc.get('fileName'),
                        "pageNumber": doc.get('pageNumber')
                    })
            
            return jsonify({
                "response": response,
                "mode": "local",
                "model_used": actual_model_used,
                "model_config": MODEL_CONFIG.get(selected_model, {}),
                "context_info": {
                    "file_name": file_name,
                    "directory_files_count": len(directory_content),
                    "vector_store_docs": len(relevant_docs),
                    "is_binary": is_binary,
                    "session_id": session_id,
                    "used_backend_vectorstore": use_backend_vectorstore,
                    "pages_referenced": pages_used if pages_used else None
                }
            }), 200
        
        # MODE ONLINE avec recherche automatique si l'info n'est pas trouvÃ©e
        elif effective_mode == 'online' or enable_online_search:
            logger.info(f"ðŸŒ Mode ONLINE activÃ© pour la requÃªte: {user_query[:50]}...")
            
            enable_auto_search = data.get('enable_auto_online_search', True)  # ActivÃ© par dÃ©faut
            conversation_history = data.get('conversation_history', [])
            response = await query_model_online_mode(
                user_query=user_query,
                selected_model=selected_model,
                language=language,
                enable_auto_online_search=enable_auto_search,
                conversation_history=conversation_history
            )
            
            return jsonify({
                "response": response,
                "mode": "online",
                "model_used": selected_model,
                "model_config": MODEL_CONFIG.get(selected_model, {}),
                "search_info": {
                    "query": user_query,
                    "language": language,
                    "search_performed": "intelligent_enrichment",
                    "approach": "searchgpt_style"
                }
            }), 200
        
        else:
            return jsonify({
                "error": f"Mode de recherche non supportÃ©: {effective_mode}",
                "supported_modes": ["local", "online"]
            }), 400
            
    except Exception as e:
        logger.error(f"Erreur lors du traitement de la requÃªte: {str(e)}")
        return jsonify({
            "error": f"Erreur lors du traitement: {str(e)}",
            "mode": effective_mode,
            "model_used": selected_model
        }), 500

@app.route('/query-stream', methods=['POST'])
def handle_query_stream():
    """
    Version streaming de la route query pour rÃ©ponses en temps rÃ©el
    """
    data = request.get_json()
    
    # RÃ©cupÃ©rer les paramÃ¨tres
    user_query = data.get('user_query')
    research_mode = data.get('research_mode', 'local')
    selected_model = data.get('selected_model', DEFAULT_MODEL)
    language = data.get('language', 'en')
    session_id = request.headers.get('Session-ID', 'default')
    
    if not user_query:
        return jsonify({"error": "user_query est requis"}), 400
    
    # Validation du modÃ¨le
    if selected_model not in MODEL_CONFIG:
        logger.warning(f"Unknown model {selected_model}, using default {DEFAULT_MODEL}")
        selected_model = DEFAULT_MODEL
    
    logger.info(f"ðŸš€ Streaming request - Model: {selected_model}, Mode: {research_mode}")
    
    try:
        # MODE ONLINE
        if research_mode == 'online' or data.get('enable_online_search'):
            t = translations.get(language, translations['en'])
            conversation_history = data.get('conversation_history', [])
            
            prompt = (
                f"{t['online_mode_title']}\n"
                f"{t['recent_info_mention']}\n\n"
                f"{t['question']}: {user_query}\n\n"
                f"{t['online_instructions_title']}\n"
                f"{t['give_best_answer']}\n"
                f"{t['be_precise_dates']}\n"
                f"{t['mention_recent_useful']}\n\n"
                f"{t['emoji']}"
            )
            
            return Response(
                stream_with_context(stream_response(prompt, selected_model, conversation_history)),
                mimetype='text/event-stream',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'X-Accel-Buffering': 'no'
                }
            )
        
        # MODE LOCAL
        elif research_mode == 'local':
            file_name = data.get('file_name')
            file_content = data.get('file_content', '')
            directory_content = data.get('directory_content', [])
            repo_structure = data.get('repo_structure', '')
            use_backend_vectorstore = data.get('use_backend_vectorstore', False)
            conversation_history = data.get('conversation_history', [])
            
            if not file_name:
                return jsonify({"error": "Mode local nÃ©cessite un file_name"}), 400
            
            # Si le vector store backend est disponible, utiliser la recherche sÃ©mantique amÃ©liorÃ©e avec stratÃ©gie hiÃ©rarchique
            if use_backend_vectorstore and session_id in vector_stores:
                try:
                    session_store = vector_stores.get(session_id)
                    if not session_store or not isinstance(session_store, dict) or 'store' not in session_store or session_store.get('store') is None:
                        logger.warning(f"Vector store non disponible pour session {session_id}")
                    else:
                        session_vector_store = session_store['store']
                        
                        # Hierarchical search strategy for streaming:
                        # If file_name is provided and not __DIRECTORY_CORPUS__, search file first, then corpus
                        file_first_results = []
                        corpus_results = []
                        
                        if file_name and file_name != "__DIRECTORY_CORPUS__":
                            # Step 1: Search in selected file first
                            logger.info(f"ðŸ“„ [Stream] Recherche dans le fichier sÃ©lectionnÃ© '{file_name}'...")
                            file_first_docs, _ = hybrid_retrieve_documents(
                                vector_store=session_vector_store,
                                query=user_query,
                                k_candidates=80,
                                k_final=10,
                                semantic_weight=0.55,
                                bm25_weight=0.35,
                                exact_weight=0.10,
                                preferred_sources=[file_name],
                            )
                            
                            # Convert to directory_content format
                            for doc in file_first_docs:
                                file_name_from_meta = doc.metadata.get("fileName") or doc.metadata.get("file_name") or "document_vectorstore"
                                page_number = doc.metadata.get("page_number")
                                result_dict = {
                                    "fileName": file_name_from_meta,
                                    "content": doc.page_content[:2500]
                                }
                                if page_number is not None:
                                    result_dict["pageNumber"] = page_number
                                    result_dict["isPageChunk"] = doc.metadata.get("is_page_chunk", False)
                                file_first_results.append(result_dict)
                            
                            logger.info(f"ðŸ“Š [Stream] Fichier: {len(file_first_results)} rÃ©sultats")
                            
                            # Step 2: Always search in corpus to complement file results
                            # If info not in file, it will be in corpus
                            logger.info(f"ðŸ“š [Stream] Recherche complÃ©mentaire dans le corpus...")
                            corpus_docs, _ = hybrid_retrieve_documents(
                                vector_store=session_vector_store,
                                query=user_query,
                                k_candidates=120,
                                k_final=15,
                                semantic_weight=0.55,
                                bm25_weight=0.35,
                                exact_weight=0.10,
                            )
                            
                            # Convert and dedupe
                            seen_keys = set()
                            for doc in corpus_docs:
                                file_name_from_meta = doc.metadata.get("fileName") or doc.metadata.get("file_name") or "document_vectorstore"
                                content_preview = doc.page_content[:100]
                                key = f"{file_name_from_meta}:{content_preview}"
                                if key not in seen_keys:
                                    seen_keys.add(key)
                                    page_number = doc.metadata.get("page_number")
                                    result_dict = {
                                        "fileName": file_name_from_meta,
                                        "content": doc.page_content[:2500]
                                    }
                                    if page_number is not None:
                                        result_dict["pageNumber"] = page_number
                                        result_dict["isPageChunk"] = doc.metadata.get("is_page_chunk", False)
                                    corpus_results.append(result_dict)
                            
                            logger.info(f"ðŸ“š [Stream] Corpus: {len(corpus_results)} rÃ©sultats")
                        else:
                            # No file selected: use the advanced search_semantic_documents_sync
                            # which includes person name filtering, function name search, filename matching, etc.
                            logger.info(f"ðŸ“š [Stream] No file selected: Using advanced search in whole repository...")
                            corpus_results = search_semantic_documents_sync(
                                session_vector_store, 
                                user_query, 
                                session_id, 
                                conversation_history
                            )
                        
                        # Combine results: file_first first, then corpus
                        directory_content = file_first_results + corpus_results
                        logger.info(f"âœ… [Stream] {len(directory_content)} documents rÃ©cupÃ©rÃ©s par recherche hiÃ©rarchique (file={len(file_first_results)}, corpus={len(corpus_results)})")
                except Exception as e:
                    logger.warning(f"Erreur lors de la recherche sÃ©mantique en streaming: {str(e)}")
            
            t = translations.get(language, translations['en'])
            
            # Construire le prompt local avec query_model_local_mode pour avoir le mÃªme prompt amÃ©liorÃ©
            import asyncio
            try:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                enable_auto_search = data.get('enable_auto_online_search', False)  # DÃ©sactivÃ© par dÃ©faut en mode local
                prompt = loop.run_until_complete(
                    query_model_local_mode(
                        file_name, file_content, directory_content, repo_structure,
                        user_query, False, selected_model, language, conversation_history,
                        enable_auto_online_search=enable_auto_search
                    )
                )
                loop.close()
            except Exception as e:
                logger.warning(f"Erreur lors de la construction du prompt amÃ©liorÃ©: {str(e)}")
                # Fallback vers l'ancien prompt
            directory_content_summary = ' '.join(
                [f"{t['other_file']}: {doc['fileName']} : {doc['content']}" 
                 for doc in directory_content]
            ) if directory_content else t['no_other_files']
            
            prompt = (
                f"{t['local_analysis_mode']}\n"
                f"{t['no_external_search']}\n\n"
                f"{t['project_structure']}:\n{repo_structure}\n\n"
                f"{t['main_file']}: {file_name}\n\n"
                f"{t['file_content']}:\n{file_content}\n\n"
                f"{t['directory_context']}:\n{directory_content_summary}\n\n"
                f"{t['question']}: {user_query}\n\n"
                f"{t['instructions']}:\n"
                f"{t['base_response_only']}\n"
                f"{t['missing_info_clarify']}\n"
                f"{t['no_speculation']}\n"
                f"{t['focus_local_analysis']}\n\n"
                f"{t['emoji']}"
            )
            
            return Response(
                stream_with_context(stream_response(prompt, selected_model)),
                mimetype='text/event-stream',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'X-Accel-Buffering': 'no'
                }
            )
        
        else:
            return jsonify({"error": f"Mode non supportÃ©: {research_mode}"}), 400
            
    except Exception as e:
        logger.error(f"Erreur dans handle_query_stream: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/vector-store-status', methods=['GET'])
async def get_vector_store_status():
    """
    Endpoint pour vÃ©rifier le statut des vector stores
    """
    session_id = request.headers.get('Session-ID', 'default')
    
    return jsonify({
        'session_id': session_id,
        'session_has_vectorstore': session_id in vector_stores,
        'global_vectorstore_exists': vector_store is not None,
        'total_sessions': len(vector_stores),
        'session_info': vector_stores.get(session_id, {}).get('files_indexed', []) if session_id in vector_stores else None
    }), 200

@app.route('/extract-structured', methods=['POST'])
async def extract_structured():
    """
    Endpoint pour extraire les donnÃ©es structurÃ©es d'un document
    UtilisÃ© pour les avocats/notaires pour extraire automatiquement les informations importantes
    """
    try:
        data = request.get_json()
        
        # Validation des paramÃ¨tres
        file_name = data.get('file_name')
        file_content = data.get('file_content', '')
        selected_model = data.get('selected_model', DEFAULT_MODEL)
        language = data.get('language', 'fr')
        
        if not file_name:
            return jsonify({
                "error": "file_name est requis",
                "success": False
            }), 400
        
        if not file_content:
            return jsonify({
                "error": "file_content est requis",
                "success": False
            }), 400
        
        # Validation du modÃ¨le
        if selected_model not in MODEL_CONFIG:
            logger.warning(f"ModÃ¨le inconnu {selected_model}, utilisation du modÃ¨le par dÃ©faut {DEFAULT_MODEL}")
            selected_model = DEFAULT_MODEL
        
        logger.info(f"ðŸ“‹ Extraction structurÃ©e demandÃ©e pour: {file_name}")
        
        # DÃ©tecter le type de document
        document_type = await detect_document_type(file_content, file_name)
        logger.info(f"ðŸ” Type de document dÃ©tectÃ©: {document_type}")
        
        # Extraire les donnÃ©es structurÃ©es
        structured_data = await extract_structured_data(
            file_content, 
            file_name, 
            document_type, 
            selected_model, 
            language
        )
        
        if "error" in structured_data:
            return jsonify({
                "success": False,
                "error": structured_data.get("error"),
                "data": structured_data
            }), 500
        
        return jsonify({
            "success": True,
            "data": structured_data,
            "document_type": document_type,
            "file_name": file_name
        }), 200
        
    except Exception as e:
        logger.error(f"âŒ Erreur lors de l'extraction structurÃ©e: {str(e)}")
        return jsonify({
            "success": False,
            "error": f"Erreur lors de l'extraction: {str(e)}"
        }), 500

@app.route('/summarize_file_stream', methods=['POST'])
def summarize_file_stream():
    """GÃ©nÃ¨re un rÃ©sumÃ© d'un fichier avec streaming (4 lignes max).
    Accepte multipart/form-data: champ 'file', champ 'language'.
    """
    try:
        uploaded_file = request.files.get('file')
        language = request.form.get('language', 'en')

        if not uploaded_file:
            return jsonify({"error": "Aucun fichier fourni"}), 400

        file_name = uploaded_file.filename or 'document'
        ext = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''
        raw_bytes = uploaded_file.read()

        # Get translations for the selected language
        t = translations.get(language, translations['en'])

        # â€” Docling parse to extract text â€”
        file_content = ""
        parse_result = parse_with_docling(raw_bytes, ext, file_name)
        if parse_result:
            _, full_text, _ = parse_result
            file_content = full_text
        else:
            try:
                if ext == 'pdf':
                    file_content = extract_text_from_pdf(raw_bytes)
                elif ext == 'docx':
                    file_content = extract_text_from_docx(type('_F', (), {'read': lambda self: raw_bytes})())
                else:
                    file_content = raw_bytes.decode('utf-8', errors='replace')
            except Exception:
                file_content = ""

        if not file_content or not file_content.strip():
            error_msg = t.get('no_content_provided', 'No content provided')
            return jsonify({"error": error_msg}), 400

        # Build language-specific prompt
        lang_names = {'en': 'English', 'fr': 'French', 'es': 'Spanish'}
        lang_name = lang_names.get(language, language)
        prompt_text = t['summarize_file_prompt'].format(language=lang_name)
        summary_label = t['summarize_file_summary_label']
        system_content = t['summarize_file_system'].format(language=lang_name)

        prompt = f"""{prompt_text}

Document: {file_name}
Content:
{file_content[:2000]}

{summary_label}"""
        
        def generate():
            try:
                # Essayer Mistral d'abord, puis fallback vers OpenAI
                mistral_api_key = os.getenv("MISTRAL_API_KEY")
                openai_api_key = os.getenv("OPENAI_API_KEY")
                
                use_mistral = mistral_api_key and len(mistral_api_key.strip()) > 0
                use_openai = openai_api_key and len(openai_api_key.strip()) > 0
                
                if not use_mistral and not use_openai:
                    error_msg = "Aucune clÃ© API configurÃ©e (ni Mistral ni OpenAI)"
                    logger.error(f"âŒ {error_msg}")
                    yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
                    return
                
                # Essayer Mistral d'abord si disponible
                if use_mistral:
                    try:
                        mistral_client = OpenAI(base_url="https://api.mistral.ai/v1", api_key=mistral_api_key)
                        stream = mistral_client.chat.completions.create(
                    model="mistral-small",
                    messages=[
                        {"role": "system", "content": system_content},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=200,
                    temperature=0.3,
                    stream=True
                )
                        logger.info("âœ… Utilisation de Mistral pour le rÃ©sumÃ©")
                    except Exception as mistral_error:
                        logger.warning(f"âš ï¸ Erreur Mistral ({mistral_error}), fallback vers OpenAI...")
                        use_mistral = False
                
                # Fallback vers OpenAI si Mistral n'est pas disponible ou a Ã©chouÃ©
                if not use_mistral:
                    if not use_openai:
                        error_msg = "Mistral indisponible et OpenAI non configurÃ©"
                        logger.error(f"âŒ {error_msg}")
                        yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
                        return
                    
                    try:
                        from openai import OpenAI as OpenAIClient
                        openai_client = OpenAIClient(api_key=openai_api_key)
                        stream = openai_client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": system_content},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=200,
                            temperature=0.3,
                            stream=True
                        )
                        logger.info("âœ… Utilisation d'OpenAI (fallback) pour le rÃ©sumÃ©")
                    except Exception as openai_error:
                        logger.error(f"âŒ Erreur OpenAI: {openai_error}")
                        yield f"data: {json.dumps({'error': f'Erreur API: {openai_error}', 'done': True})}\n\n"
                        return
                
                accumulated_text = ""
                line_count = 0
                
                for chunk in stream:
                    if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        accumulated_text += content
                        
                        # Compter les lignes
                        current_lines = accumulated_text.split('\n')
                        if len(current_lines) > 4:
                            # Limiter Ã  4 lignes
                            accumulated_text = '\n'.join(current_lines[:4])
                            yield f"data: {json.dumps({'content': content, 'done': False})}\n\n"
                            yield f"data: {json.dumps({'done': True})}\n\n"
                            return
                        
                        yield f"data: {json.dumps({'content': content, 'done': False})}\n\n"
                
                # Limiter Ã  4 lignes Ã  la fin
                final_lines = accumulated_text.split('\n')
                if len(final_lines) > 4:
                    accumulated_text = '\n'.join(final_lines[:4])
                
                yield f"data: {json.dumps({'done': True})}\n\n"
                
            except Exception as e:
                logger.error(f"âŒ Erreur lors du streaming du rÃ©sumÃ©: {str(e)}")
                yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"
        
        return Response(generate(), mimetype='text/event-stream')
        
    except Exception as e:
        logger.error(f"âŒ Erreur lors de la gÃ©nÃ©ration du rÃ©sumÃ©: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/summarize_repository_stream', methods=['POST'])
def summarize_repository_stream():
    """GÃ©nÃ¨re un rÃ©sumÃ© d'un rÃ©pertoire avec streaming"""
    try:
        data = request.get_json()
        files_info = data.get('files', [])
        language = data.get('language', 'en')
        
        # Get translations for the selected language
        t = translations.get(language, translations['en'])
        
        if not files_info or len(files_info) == 0:
            error_msg = t.get('no_files_provided', 'No files provided')
            return jsonify({"error": error_msg}), 400
        
        # Compter les sous-rÃ©pertoires et fichiers
        file_count = len(files_info)
        subdirectories = set()
        for file_info in files_info:
            name = file_info.get('name', '')
            if '/' in name:
                parts = name.split('/')
                if len(parts) > 1:
                    subdirectories.add(parts[0])
        
        subdirectory_count = len(subdirectories)
        
        # Handle pluralization for different languages
        s_files = 's' if file_count > 1 else ''
        s_dirs = 's' if subdirectory_count > 1 else ''
        if language == 'fr':
            s_files = 's' if file_count > 1 else ''
            s_dirs = 's' if subdirectory_count > 1 else ''
        elif language == 'es':
            s_files = 's' if file_count > 1 else ''
            s_dirs = 's' if subdirectory_count > 1 else ''
        else:  # English
            s_files = 's' if file_count > 1 else ''
            s_dirs = 's' if subdirectory_count > 1 else ''
        
        def generate():
            try:
                # Essayer Mistral d'abord, puis fallback vers OpenAI
                mistral_api_key = os.getenv("MISTRAL_API_KEY")
                openai_api_key = os.getenv("OPENAI_API_KEY")
                
                use_mistral = mistral_api_key and len(mistral_api_key.strip()) > 0
                use_openai = openai_api_key and len(openai_api_key.strip()) > 0
                
                if not use_mistral and not use_openai:
                    error_msg = "Aucune clÃ© API configurÃ©e (ni Mistral ni OpenAI)"
                    logger.error(f"âŒ {error_msg}")
                    yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
                    return
                
                # Construire le prompt avec les informations du rÃ©pertoire
                file_names_text = '\n'.join([f"- {f.get('display_name', f.get('name', ''))}" for f in files_info[:10]])
                
                # Get language-specific translations
                lang_names = {'en': 'English', 'fr': 'French', 'es': 'Spanish'}
                lang_name = lang_names.get(language, language)
                
                prompt_start = t['summarize_repo_prompt_start'].format(
                    file_count=file_count,
                    s_files=s_files,
                    subdir_count=subdirectory_count,
                    s_dirs=s_dirs
                )
                files_label = t['summarize_repo_files_label']
                instructions = t['summarize_repo_instructions'].format(language=lang_name)
                expected_format = t['summarize_repo_expected_format'].format(
                    file_count=file_count,
                    s_files=s_files,
                    subdir_count=subdirectory_count,
                    s_dirs=s_dirs
                )
                summary_label = t['summarize_repo_summary_label']
                system_content = t['summarize_repo_system'].format(language=lang_name)
                
                prompt = f"""{prompt_start}

{files_label}
{file_names_text}

{instructions}

{expected_format}

{summary_label}"""
                
                # Essayer Mistral d'abord si disponible
                if use_mistral:
                    try:
                        mistral_client = OpenAI(base_url="https://api.mistral.ai/v1", api_key=mistral_api_key)
                        stream = mistral_client.chat.completions.create(
                    model="mistral-small",
                    messages=[
                        {"role": "system", "content": system_content},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=400,
                    temperature=0.3,
                    stream=True
                )
                        logger.info("âœ… Utilisation de Mistral pour le rÃ©sumÃ© du rÃ©pertoire")
                    except Exception as mistral_error:
                        logger.warning(f"âš ï¸ Erreur Mistral ({mistral_error}), fallback vers OpenAI...")
                        use_mistral = False
                
                # Fallback vers OpenAI si Mistral n'est pas disponible ou a Ã©chouÃ©
                if not use_mistral:
                    if not use_openai:
                        error_msg = "Mistral indisponible et OpenAI non configurÃ©"
                        logger.error(f"âŒ {error_msg}")
                        yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
                        return
                    
                    try:
                        from openai import OpenAI as OpenAIClient
                        openai_client = OpenAIClient(api_key=openai_api_key)
                        stream = openai_client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": system_content},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=400,
                            temperature=0.3,
                            stream=True
                        )
                        logger.info("âœ… Utilisation d'OpenAI (fallback) pour le rÃ©sumÃ© du rÃ©pertoire")
                    except Exception as openai_error:
                        logger.error(f"âŒ Erreur OpenAI: {openai_error}")
                        yield f"data: {json.dumps({'error': f'Erreur API: {openai_error}', 'done': True})}\n\n"
                        return
                
                for chunk in stream:
                    if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        yield f"data: {json.dumps({'content': content, 'done': False})}\n\n"
                
                yield f"data: {json.dumps({'done': True})}\n\n"
                
            except Exception as e:
                logger.error(f"âŒ Erreur lors du streaming du rÃ©sumÃ© de rÃ©pertoire: {str(e)}")
                yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"
        
        return Response(generate(), mimetype='text/event-stream')
        
    except Exception as e:
        logger.error(f"âŒ Erreur lors de la gÃ©nÃ©ration du rÃ©sumÃ© de rÃ©pertoire: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    """Endpoint de santÃ© pour vÃ©rifier que l'API fonctionne"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "1.0.0",
        "services": {
            "openai": bool(os.getenv("OPENAI_API_KEY")),
            "mistral": bool(os.getenv("MISTRAL_API_KEY")),
            "google_oauth": AuthConfig.is_google_configured(),
            "database": os.path.exists(AuthConfig.DATABASE_PATH)
        },
        "endpoints": {
            "extract_structured": "/extract-structured",
            "summarize_file_stream": "/summarize_file_stream",
            "summarize_repository_stream": "/summarize_repository_stream"
        }
    }), 200

@app.route('/test-endpoints', methods=['GET'])
def test_endpoints():
    """Endpoint pour tester la disponibilitÃ© des endpoints"""
    endpoints_status = {}
    
    # Liste des endpoints Ã  tester
    endpoints_to_test = [
        '/extract-structured',
        '/summarize_file_stream',
        '/summarize_repository_stream',
        '/query',
        '/upload'
    ]
    
    # VÃ©rifier si les routes existent dans l'application Flask
    for endpoint in endpoints_to_test:
        rule = None
        for rule in app.url_map.iter_rules():
            if rule.rule == endpoint:
                endpoints_status[endpoint] = {
                    "exists": True,
                    "methods": list(rule.methods - {'HEAD', 'OPTIONS'}),
                    "rule": rule.rule
                }
                break
        else:
            endpoints_status[endpoint] = {
                "exists": False,
                "error": "Route not found"
            }
    
    return jsonify({
        "status": "ok",
        "endpoints": endpoints_status,
        "total_routes": len(list(app.url_map.iter_rules()))
    }), 200

@app.route('/users/me', methods=['GET'])
def get_current_user():
    """RÃ©cupÃ¨re les informations de l'utilisateur connectÃ©"""
    from legacy_auth import verify_auth_token
    
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return jsonify({'error': 'Token manquant'}), 401
    
    token = auth_header.replace('Bearer ', '')
    auth_result = verify_auth_token(token)
    
    if not auth_result:
        return jsonify({'error': 'Token invalide'}), 401

    if auth_result['auth_type'] == 'firebase':
        return jsonify({
            'user': auth_result['user']
        }), 200
    
    return jsonify({
        'user': {
            'id': auth_result['payload']['user_id'],
            'email': auth_result['payload']['email'],
            'name': auth_result['payload']['name']
        }
    }), 200


@app.route('/users/me/recent-documents', methods=['GET'])
def get_recent_documents_for_user():
    """Return the most recent persisted documents for the authenticated user."""
    from legacy_auth import verify_auth_token, _get_pg_connection

    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return jsonify({'error': 'Token manquant'}), 401

    token = auth_header.replace('Bearer ', '')
    auth_result = verify_auth_token(token)

    if not auth_result or not auth_result.get('user'):
        return jsonify({'error': 'Token invalide'}), 401

    user_id = auth_result['user'].get('id')
    if not user_id:
        return jsonify({'error': 'Utilisateur introuvable'}), 404

    try:
        requested_limit = int(request.args.get('limit', 5))
    except (TypeError, ValueError):
        requested_limit = 5
    limit = max(1, min(requested_limit, 20))

    try:
        with _get_pg_connection() as conn:
            with conn.cursor() as cursor:
                cursor.execute(
                    '''
                    SELECT
                        id::text,
                        file_name,
                        mime_type,
                        size_bytes,
                        status,
                        created_at,
                        processed_at
                    FROM documents
                    WHERE uploaded_by_user_id = %s::uuid
                    ORDER BY COALESCE(processed_at, created_at) DESC
                    LIMIT %s
                    ''',
                    (user_id, limit)
                )
                rows = cursor.fetchall()

        recent_documents = []
        for row in rows:
            recent_documents.append({
                'id': row[0],
                'file_name': row[1],
                'mime_type': row[2],
                'size_bytes': row[3],
                'status': row[4],
                'created_at': row[5].isoformat() if row[5] else None,
                'processed_at': row[6].isoformat() if row[6] else None,
            })

        return jsonify({'documents': recent_documents}), 200
    except Exception as exc:
        logger.error(f"Erreur rÃ©cupÃ©ration documents rÃ©cents: {str(exc)}")
        return jsonify({'error': 'Impossible de charger les documents rÃ©cents'}), 500


@app.route('/users/me/documents/<document_id>/content', methods=['GET'])
def get_document_content_for_user(document_id):
    """Return the raw content of a user's document stored in S3 as base64."""
    from legacy_auth import verify_auth_token, _get_pg_connection

    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return jsonify({'error': 'Token manquant'}), 401

    token = auth_header.replace('Bearer ', '')
    auth_result = verify_auth_token(token)

    if not auth_result or not auth_result.get('user'):
        return jsonify({'error': 'Token invalide'}), 401

    user_id = auth_result['user'].get('id')
    if not user_id:
        return jsonify({'error': 'Utilisateur introuvable'}), 404

    try:
        with _get_pg_connection() as conn:
            with conn.cursor() as cursor:
                cursor.execute(
                    '''
                    SELECT file_name, mime_type, s3_bucket, s3_key
                    FROM documents
                    WHERE id = %s::uuid
                      AND uploaded_by_user_id = %s::uuid
                    LIMIT 1
                    ''',
                    (document_id, user_id)
                )
                row = cursor.fetchone()

        if not row:
            return jsonify({'error': 'Document introuvable'}), 404

        file_name, mime_type, s3_bucket, s3_key = row
        if not s3_bucket or not s3_key:
            return jsonify({'error': 'Contenu du document indisponible'}), 404

        if not aws_persistence_service.enabled or not aws_persistence_service._s3:
            return jsonify({'error': 'Stockage AWS indisponible'}), 503

        s3_object = aws_persistence_service._s3.get_object(Bucket=s3_bucket, Key=s3_key)
        file_bytes = s3_object['Body'].read()
        content_base64 = base64.b64encode(file_bytes).decode('ascii')

        return jsonify({
            'id': document_id,
            'file_name': file_name,
            'mime_type': mime_type,
            'content_base64': content_base64,
        }), 200
    except Exception as exc:
        logger.error(f"Erreur rÃ©cupÃ©ration contenu document: {str(exc)}")
        return jsonify({'error': 'Impossible de charger le contenu du document'}), 500

# Route pour tester la configuration Google (DEBUG)
@app.route('/debug/google-config', methods=['GET'])
def debug_google_config():
    google_client_id = os.getenv('GOOGLE_CLIENT_ID')
    return jsonify({
        'google_client_id_set': bool(google_client_id),
        'google_client_secret_set': bool(os.getenv('GOOGLE_CLIENT_SECRET')),
        'client_id_preview': google_client_id[:20] + '...' if google_client_id else 'Missing',
        'full_client_id': google_client_id  # Temporaire pour debug
    })

def main():
    print("ðŸš€ Starting Enhanced AI backend on http://0.0.0.0:5000")
    print(f"ðŸ“‹ Available models: {', '.join(MODEL_CONFIG.keys())}")
    print(f"ðŸŽ¯ Default model: {DEFAULT_MODEL} (Mistral)")
    print("ðŸ“¡ Endpoints:")
    print("   GET  /models - List all available models")
    print("   POST /models/<model_id>/test - Test specific model")
    print("   POST /upload - Upload files for analysis")
    print("   POST /query - Query with selected model")
    print("   POST /index-directory - Index directory files")
    print("ðŸ” Authentication endpoints:")
    print("   POST /auth/login - Email/password login")
    print("   POST /auth/register - User registration")
    print("   GET  /auth/google - Google OAuth")
    print("   GET  /auth/verify - Verify JWT token")
    print("   POST /marketing/subscribe - Save marketing emails")
    print("   GET  /health - Health check")
    print("   GET  /users/me - Current user info")
    print("âœ¨ New features:")
    print("   â€¢ Complete authentication system")
    print("   â€¢ Google OAuth integration")
    print("   â€¢ Marketing email capture")
    print("   â€¢ JWT-based sessions")
    print("   â€¢ SQLite database")
    if AuthConfig.is_google_configured():
        print("   â€¢ âœ… Google OAuth configured")
    else:
        print("   â€¢ âš ï¸  Google OAuth NOT configured")
    app.run(debug=False, host='0.0.0.0', port=5000)

if __name__ == "__main__":
    main()
