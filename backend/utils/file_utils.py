"""
Utilitaires pour l'extraction de texte depuis différents formats de fichiers
"""
from docx import Document as DocxDocument
import PyPDF2
from io import BytesIO
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file):
    """Extrait le texte d'un fichier DOCX"""
    try:
        doc = DocxDocument(BytesIO(file.read()))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return f"Error: Could not extract text from DOCX file. {str(e)}"


def extract_pages_from_pdf(file_bytes: bytes) -> List[Dict[str, Any]]:
    """Extrait le texte page par page d'un PDF.
    Retourne une liste de dicts: [{"page": int (1-indexed), "text": str}, ...]
    Seules les pages ayant du texte non vide sont incluses.
    """
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(pdf_reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({"page": i + 1, "text": text})
        return pages
    except Exception as e:
        logger.error(f"Error extracting pages from PDF: {e}")
        return []


def extract_text_from_pdf(file_bytes_or_file) -> str:
    """Extrait le texte complet d'un fichier PDF (toutes pages concaténées).
    Accepte bytes ou un objet fichier.
    """
    try:
        if isinstance(file_bytes_or_file, (bytes, bytearray)):
            raw = file_bytes_or_file
        else:
            raw = file_bytes_or_file.read()
        pages = extract_pages_from_pdf(raw)
        return "\n".join(p["text"] for p in pages)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return f"Error: Could not extract text from PDF file. {str(e)}"


def extract_sections_from_docx(file_bytes: bytes) -> List[Dict[str, Any]]:
    """Extrait le texte d'un DOCX en sections numérotées (simule des 'pages').
    Retourne une liste de dicts: [{"page": int, "text": str}, ...]
    Les paragraphes sont regroupés par blocs de ~50 paragraphes.
    """
    try:
        doc = DocxDocument(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        block_size = 50
        sections = []
        for i in range(0, len(paragraphs), block_size):
            block = "\n".join(paragraphs[i:i + block_size])
            if block.strip():
                sections.append({"page": (i // block_size) + 1, "text": block})
        return sections
    except Exception as e:
        logger.error(f"Error extracting sections from DOCX: {e}")
        return []


